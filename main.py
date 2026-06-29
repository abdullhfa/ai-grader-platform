# ─────────────────────────────────────────────────────────────────────
# CRITICAL: Force stdout/stderr to UTF-8 BEFORE any imports.
# Windows defaults to cp1252 which can't encode Arabic letters or emojis.
# ─────────────────────────────────────────────────────────────────────
import sys as _sys_init
import os as _os_init
_os_init.environ.setdefault("PYTHONIOENCODING", "utf-8")
try:
    if hasattr(_sys_init.stdout, "reconfigure"):
        _sys_init.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(_sys_init.stderr, "reconfigure"):
        _sys_init.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    try:
        import io as _io_init
        _sys_init.stdout = _io_init.TextIOWrapper(_sys_init.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
        _sys_init.stderr = _io_init.TextIOWrapper(_sys_init.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    except Exception:
        pass

# ── Global safe print (Windows consoles often remain cp1256/cp1252; emoji + Arabic break print()) ──
import builtins as _builtins_safe  # noqa: E402

_original_builtin_print = _builtins_safe.print


def _safe_builtin_print(*args, sep=" ", end="\n", file=None, flush=False):
    """Never raise UnicodeEncodeError: fall back to UTF-8 bytes on the underlying binary stream."""
    out = file if file is not None else _sys_init.stdout
    try:
        _original_builtin_print(*args, sep=sep, end=end, file=out, flush=flush)
    except UnicodeEncodeError:
        try:
            line = sep.join(str(a) for a in args) + (end or "")
            buf = getattr(out, "buffer", None)
            if buf is not None:
                buf.write(line.encode("utf-8", errors="replace"))
                if flush:
                    buf.flush()
            else:
                _original_builtin_print(
                    line.encode("ascii", errors="replace").decode("ascii"),
                    end="",
                    file=out,
                    flush=flush,
                )
        except Exception:
            pass


_builtins_safe.print = _safe_builtin_print


from typing import Any, Dict, List, Optional
import os
import json
import shutil
import time
import hmac
import hashlib
import re
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from dotenv import load_dotenv  # type: ignore
from fastapi import (  # type: ignore
    FastAPI,
    Request,
    Depends,
    HTTPException,
    UploadFile,
    File,
    Form,
    BackgroundTasks,
)
from fastapi.responses import (  # type: ignore
    HTMLResponse, JSONResponse, FileResponse, RedirectResponse
)
from fastapi.staticfiles import StaticFiles  # type: ignore
from fastapi.templating import Jinja2Templates  # type: ignore
from sqlalchemy.orm import Session  # type: ignore

from app.database import get_db, init_db  # type: ignore
from app import models  # type: ignore
from app.models import (  # type: ignore
    Submission,
    SubmissionStatus,
    GradingResult,
    GradingSummary,
    Assignment,
    AssignmentStatus,
    Textbook,
    BatchGrading,
    BatchStatus,
    StudentReport,
    ActivityLog,
)
from app.textbook_analyzer import (  # type: ignore
    SLIM_V3_SECTION_TITLES,
    apply_slim_v3_section_titles,
    process_textbook_and_assignment,
    strip_excluded_guide_sections,
)
from app.batch_grader import (  # type: ignore
    grade_batch_async,
    extract_text_from_file,
    check_plagiarism_for_submission,
)
from app.project_intelligence.submission_intake import INTAKE_IGNORE_DIR_NAMES  # type: ignore
from app.grading_mode_policy import (  # type: ignore
    grading_mode_display_label,
    normalize_grading_mode_choice,
    resolve_grading_policy,
)

INTAKE_IGNORE_DIRS_LIST = sorted(INTAKE_IGNORE_DIR_NAMES)
from app.document_processor import extract_student_name_from_file  # type: ignore
from app.report_generator import (  # type: ignore
    generate_student_report_pdf,
    generate_batch_summary_report,
)
from app.ai_provider import AIProvider, reset_global_provider  # type: ignore
from app.auth import (  # type: ignore
    create_session,
    delete_session,
    get_current_user,
    get_current_user_id,
    user_is_admin,
    register_user,
    authenticate_user,
    google_login_or_register,
)

load_dotenv(override=True)


def _orm_str(value: object) -> str:
    """Coerce SQLAlchemy Column / ORM attribute to plain str for typing + APIs."""
    return "" if value is None else str(value)


def _orm_int(value: object, default: int = 0) -> int:
    if value is None:
        return default
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _orm_float(value: object, default: float = 0.0) -> float:
    if value is None:
        return default
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _orm_set(instance: object, name: str, value: object) -> None:
    """Assign ORM attribute without Column[...] type-checker false positives."""
    setattr(instance, name, value)


def _patch_starlette_multipart_limits() -> None:
    """
    Starlette parses multipart with max_files=max_fields=1000 by default.
    Batch folder uploads (many small files) exceed that and show:
    "Too many files. Maximum number of files is 1000".
    """
    import starlette.requests as _sr

    _orig = _sr.Request._get_form

    try:
        cap_files = int(os.getenv("MULTIPART_MAX_FILES", "8000"))
        cap_fields = int(os.getenv("MULTIPART_MAX_FIELDS", "8000"))
    except ValueError:
        cap_files, cap_fields = 8000, 8000
    try:
        _multipart_ceiling = int(os.getenv("MULTIPART_HARD_CEILING", "12000"))
    except ValueError:
        _multipart_ceiling = 32000
    _multipart_ceiling = max(1000, _multipart_ceiling)
    cap_files = max(1000, min(cap_files, _multipart_ceiling))
    cap_fields = max(1000, min(cap_fields, _multipart_ceiling))

    async def _get_form(self, *, max_files=1000, max_fields=1000):  # type: ignore[no-untyped-def]
        mf = cap_files if max_files == 1000 else max_files
        fld = cap_fields if max_fields == 1000 else max_fields
        return await _orig(self, max_files=mf, max_fields=fld)

    _sr.Request._get_form = _get_form  # type: ignore[method-assign]


_patch_starlette_multipart_limits()


def _stderr_line_utf8(line: str) -> None:
    """Write one line to stderr as UTF-8 bytes (avoids Windows charmap crashes on emoji/Arabic)."""
    try:
        _sys_init.stderr.buffer.write((line.rstrip("\n") + "\n").encode("utf-8", errors="replace"))
    except Exception:
        pass


def _traceback_print_safe() -> None:
    """Log traceback via UTF-8 bytes (avoid Windows charmap failures on traceback source lines)."""
    import traceback as _tb

    try:
        _stderr_line_utf8(_tb.format_exc())
    except Exception:
        pass


# Patch traceback.print_exc — on Windows a broken stderr handle raises OSError [Errno 22]
# and masks the real exception (seen during /api/create-assignment error handling).
import traceback as _traceback_module  # noqa: E402

_original_traceback_print_exc = _traceback_module.print_exc


def _traceback_print_exc_safe(*args, **kwargs):
    try:
        _original_traceback_print_exc(*args, **kwargs)
    except OSError:
        _traceback_print_safe()
    except Exception:
        _traceback_print_safe()


_traceback_module.print_exc = _traceback_print_exc_safe


def _resolve_btec_unit_key(unit_number: str | None, units: dict) -> str | None:
    """Resolve API/URL unit id (e.g. L3_8 or bare '8') to a key in btec_units.json."""
    if not unit_number or not units:
        return None
    key = unit_number.strip()
    if key in units:
        return key
    if key.isdigit():
        for prefix in ("L3_", "L2_"):
            cand = f"{prefix}{key}"
            if cand in units:
                return cand
    return None


# Google OAuth config
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET", "")
GOOGLE_REDIRECT_URI = os.getenv("GOOGLE_REDIRECT_URI", "http://localhost:8800/auth/google/callback")

# Initialize database
init_db()

# Create FastAPI app
@asynccontextmanager
async def _app_lifespan(_app: FastAPI):
    """Startup/shutdown (replaces deprecated @app.on_event)."""
    os.environ.setdefault("PYTHONIOENCODING", "utf-8")
    try:
        if hasattr(_sys_init.stdout, "reconfigure"):
            _sys_init.stdout.reconfigure(encoding="utf-8", errors="replace")
        if hasattr(_sys_init.stderr, "reconfigure"):
            _sys_init.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    import asyncio

    from app.whatsapp_manager import start_whatsapp_service, stop_whatsapp_service  # type: ignore
    from app.core.logging_setup import configure_production_logging

    configure_production_logging()
    from app.security.app_secrets import get_app_secret

    get_app_secret()
    from app.batch_progress_store import hydrate_batch_progress
    from app.batch_checkpoint import resume_all_checkpoints_on_startup
    from app.batch_grade_worker import recover_orphaned_batches_on_startup

    hydrate_batch_progress(_app.state.batch_progress)
    start_whatsapp_service()

    async def _resume_batches_after_bind() -> None:
        await asyncio.sleep(2)
        try:
            queued = await resume_all_checkpoints_on_startup(_app.state.batch_progress)
            if queued:
                print(
                    f"✅ [BATCH-RESUME] queued {queued} batch job(s) to resume after startup"
                )
            await asyncio.sleep(1)
            recover_orphaned_batches_on_startup()
        except Exception as exc:
            print(f"⚠️ [BATCH-RESUME] startup resume failed: {exc}")

    asyncio.create_task(_resume_batches_after_bind())
    try:
        yield
    finally:
        stop_whatsapp_service()


app = FastAPI(
    title="أداة تصحيح واجبات  - نسخة متقدمة",
    description="AI-powered assignment grading tool with batch processing",
    version="2.0.0",
    lifespan=_app_lifespan,
)

from app.production.hardening import RateLimitMiddleware  # noqa: E402

app.add_middleware(RateLimitMiddleware)


from fastapi.exceptions import RequestValidationError  # noqa: E402


@app.exception_handler(RequestValidationError)
async def _api_validation_exception_handler(request: Request, exc: RequestValidationError):
    return JSONResponse(
        {"success": False, "detail": exc.errors(), "error": "validation_error"},
        status_code=422,
    )


@app.exception_handler(Exception)
async def _api_generic_exception_handler(request: Request, exc: Exception):
    _traceback_print_safe()
    _stderr_line_utf8(
        "[GLOBAL EXC HANDLER] " + request.method + " " + request.url.path + ": " + repr(exc)
    )
    return JSONResponse({"success": False,
        "detail": f"خطأ غير متوقع في الخادم: {type(exc).__name__}: {str(exc) or '(بدون رسالة)'}",
        "error": type(exc).__name__,
        "status_code": 500}, status_code=500)


# Create upload directories
UPLOAD_DIR = Path("uploads")
TEXTBOOKS_DIR = UPLOAD_DIR / "textbooks"
ASSIGNMENTS_DIR = UPLOAD_DIR / "assignments"
STUDENTS_DIR = UPLOAD_DIR / "students"
REPORTS_DIR = UPLOAD_DIR / "reports"

for dir_path in [
    UPLOAD_DIR,
    TEXTBOOKS_DIR,
    ASSIGNMENTS_DIR,
    STUDENTS_DIR,
    REPORTS_DIR,
]:
    dir_path.mkdir(parents=True, exist_ok=True)


def _safe_upload_basename(filename: Optional[str], *, default: str = "upload.bin") -> str:
    """Strip browser path components and illegal Windows filename characters."""
    raw = Path((filename or "").replace("\\", "/")).name
    unsafe = set('<>:"/\\|?*') | {chr(i) for i in range(32)}
    safe = "".join("_" if ch in unsafe else ch for ch in raw).strip()
    while "__" in safe:
        safe = safe.replace("__", "_")
    if not safe or safe in (".", ".."):
        safe = default
    return safe


def _resolve_uploaded_path(stored_path: str) -> Path:
    """Resolve DB-stored relative paths against project root (Windows-safe)."""
    p = Path(stored_path)
    if not p.is_absolute():
        p = (Path.cwd() / p).resolve()
    else:
        p = p.resolve()
    return p

# BTEC assessment module removed from routes

# Mount static files
app.mount("/static", StaticFiles(directory="app/static"), name="static")


@app.get("/uploads/{file_path:path}", name="serve_upload")
async def serve_upload_route(
    file_path: str,
    request: Request,
    db: Session = Depends(get_db),
):
    from app.security.upload_serve import serve_upload_file

    return await serve_upload_file(file_path, request, db)

# Setup templates
templates = Jinja2Templates(directory="app/templates")
templates.env.globals["intake_ignore_dirs"] = INTAKE_IGNORE_DIRS_LIST
from app.package_catalog import (  # type: ignore
    assignment_subtitle,
    package_feature_lines,
    student_package_feature_lines,
)

templates.env.filters["assignment_subtitle"] = assignment_subtitle
templates.env.filters["package_feature_lines"] = package_feature_lines
templates.env.filters["student_package_feature_lines"] = student_package_feature_lines

# ═══════════════════════════════════════════════════════
# Batch Grading Progress Tracker (in-memory)
# ═══════════════════════════════════════════════════════
from app.batch_progress_store import BatchProgressDict  # noqa: E402

batch_progress: BatchProgressDict = BatchProgressDict()

app.state.templates = templates
app.state.batch_progress = batch_progress

from app.routes.register import register_production_routers  # noqa: E402

register_production_routers(app)


# ═══════════════════════════════════════════════════════
# Activity Log Helper
# ═══════════════════════════════════════════════════════


def log_activity(
    db: Session,
    action: str,
    category: str,
    details: str = "",
    user_id: int | None = None,
    user_name: str = "",
    user_email: str = "",
    ip_address: str = "",
    user_agent: str = "",
    level: str = "info",
):
    """Log an activity to the database."""
    try:
        entry = ActivityLog(
            user_id=user_id,
            user_name=user_name or "",
            user_email=user_email or "",
            action=action,
            category=category,
            details=details[:2000] if details else "",
            ip_address=ip_address,
            user_agent=user_agent[:500] if user_agent else "",
            level=level,
        )
        db.add(entry)
        db.commit()
    except Exception as e:
        db.rollback()
        print(f"[LOG ERROR] {e}")


def _get_client_ip(request: Request) -> str:
    """Extract client IP from request."""
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.client.host if request.client else ""


def _get_user_display(db: Session, user_id: int | None) -> str:
    """Get user display name."""
    if not user_id:
        return "زائر"
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        return f"مستخدم #{user_id}"
    name = f"{user.first_name or ''} {user.last_name or ''}".strip()
    return name or str(user.email or '') or f"مستخدم #{user_id}"


def _get_user_email(db: Session, user_id: int | None) -> str:
    """Get user email."""
    if not user_id:
        return ""
    user = db.query(models.User).filter(models.User.id == user_id).first()
    return str(user.email) if user else ""


def _get_user_agent(request: Request) -> str:
    """Extract user agent from request."""
    return request.headers.get("user-agent", "")


# ═══════════════════════════════════════════════════════
# Subscription Helper
# ═══════════════════════════════════════════════════════


def get_active_subscription(db: Session, user_id: int):
    """Get the active subscription for a user, or None."""
    return (
        db.query(models.Subscription)
        .filter(
            models.Subscription.user_id == user_id,
            models.Subscription.status.in_(
                [models.SubscriptionStatus.ACTIVE, "active", "ACTIVE"]
            ),
        )
        .order_by(models.Subscription.created_at.desc())
        .first()
    )


def get_subscription_info(db: Session, user_id: int) -> dict:
    """Return a dict with subscription details for templates/API."""
    # Admin bypass - unlimited access
    user = db.query(models.User).filter(models.User.id == user_id).first()
    # pyright: ignore
    if user is not None and user.role == models.UserRole.ADMIN:
        return {
            "has_subscription": True,
            "remaining": 9999,
            "package_name": "Admin (غير محدود)",
            "assignments_limit": 9999,
            "assignments_used": 0,
            "subscription_id": 0,
            "grading_mode": "deep",
            "grading_profile_label": "Admin — Full verification",
            "grading_profile_description": "صلاحيات إدارية — تحقق كامل بدون قيود الباقة.",
        }

    # --- Per-subject balances (new system) ---
    balances = (
        db.query(models.SubjectBalance)
        .filter(models.SubjectBalance.user_id == user_id)
        .all()
    )
    if balances:
        total_limit = sum(b.assignments_limit or 0 for b in balances)
        total_used = sum(b.assignments_used or 0 for b in balances)
        total_remaining = max(0, total_limit - total_used)
        subjects_str = "|".join(str(b.subject) for b in balances if b.subject)
        subject_details = [
            {
                "subject": b.subject,
                "limit": b.assignments_limit or 0,
                "used": b.assignments_used or 0,
                "remaining": max(0, (b.assignments_limit or 0) - (b.assignments_used or 0)),
            }
            for b in balances
        ]
        sub = get_active_subscription(db, user_id)
        pkg_name = ""
        if sub:
            pkg = db.query(models.Package).filter(models.Package.id == sub.package_id).first()
            pkg_name = pkg.name if pkg else ""
        policy = resolve_grading_policy(pkg_name)
        return {
            "has_subscription": total_remaining > 0,
            "remaining": total_remaining,
            "package_name": pkg_name,
            "assignments_limit": total_limit,
            "assignments_used": total_used,
            "subscription_id": sub.id if sub else 0,
            "subjects": subjects_str,
            "subject_details": subject_details,
            "grading_mode": policy["grading_mode"],
            "grading_profile_label": policy["label_ar"],
            "grading_profile_description": policy.get("description_ar") or policy["label_ar"],
        }

    # --- Fall back to global subscription (old system) ---
    sub = get_active_subscription(db, user_id)
    if not sub:
        return {
            "has_subscription": False,
            "remaining": 0,
            "package_name": "",
            "assignments_limit": 0,
            "assignments_used": 0,
            "grading_mode": "deep",
            "grading_profile_label": "افتراضي — تحقق كامل",
            "grading_profile_description": "تحقق كامل — اشترِ باقة Basic أو Pro لمعرفة الفرق بالتفصيل.",
        }
    remaining = (sub.assignments_limit or 0) - (sub.assignments_used or 0)
    pkg = (
        db.query(models.Package)
        .filter(models.Package.id == sub.package_id)
        .first()
    )
    # Build subject_details from subscription.subjects for old system
    old_subject_details = []
    if sub.subjects:
        for subj in re.split(r"[,|]", str(sub.subjects)):
            subj = subj.strip()
            if subj:
                old_subject_details.append({
                    "subject": subj,
                    "limit": sub.assignments_limit or 0,
                    "used": sub.assignments_used or 0,
                    "remaining": max(remaining, 0),
                })
    policy = resolve_grading_policy(pkg.name if pkg else "")
    return {
        "has_subscription": True,
        "remaining": max(remaining, 0),
        "package_name": pkg.name if pkg else "",
        "assignments_limit": sub.assignments_limit or 0,
        "assignments_used": sub.assignments_used or 0,
        "subscription_id": sub.id,
        "subjects": sub.subjects or "",
        "subject_details": old_subject_details,
        "grading_mode": policy["grading_mode"],
        "grading_profile_label": policy["label_ar"],
        "grading_profile_description": policy.get("description_ar") or policy["label_ar"],
    }


def build_package_lock_token(user_id: int, package_id: int) -> str:
    """Create an HMAC-signed token to lock package selection per user."""
    from app.security.app_secrets import get_app_secret

    secret = get_app_secret()
    payload = f"{user_id}:{package_id}"
    signature = hmac.new(
        secret.encode("utf-8"), payload.encode("utf-8"), hashlib.sha256
    ).hexdigest()
    return f"{payload}:{signature}"


def verify_package_lock_token(token: str, user_id: int, package_id: int) -> bool:
    """Validate the signed package lock token for this user/package pair."""
    if not token:
        return False
    expected = build_package_lock_token(user_id, package_id)
    return hmac.compare_digest(token, expected)


def normalize_whatsapp_phone(phone: str) -> str:
    """Convert local Jordan phone numbers to WhatsApp-compatible E.164 digits."""
    from app.whatsapp_config import normalize_whatsapp_phone as _norm  # type: ignore
    return _norm(phone)


WHATSAPP_SERVICE_URL = os.getenv("WHATSAPP_SERVICE_URL", "http://localhost:3001")


async def send_whatsapp_activation_message(phone: str, amount: float) -> tuple[bool, str]:
    """Send activation confirmation to the teacher via whatsapp-web.js service."""
    to_number = normalize_whatsapp_phone(phone)
    if not to_number:
        return False, "invalid_phone"

    amount_text = f"{int(amount)}" if amount.is_integer() else f"{amount:.2f}"
    body_text = (
        f"مبروك تم تفعيل حسابك بقيمة ({amount_text})\n"
        "في حال وجود مشكله الرجاء التواصل معنا"
    )

    import httpx
    last_error = ""
    for attempt in range(1, 4):
        try:
            async with httpx.AsyncClient(timeout=30.0) as hc:
                response = await hc.post(
                    f"{WHATSAPP_SERVICE_URL}/send",
                    json={"phone": to_number, "message": body_text},
                )
            data = response.json()
            if response.status_code == 200 and data.get("success"):
                return True, "sent"
            last_error = data.get("error", f"http_{response.status_code}")
            if response.status_code == 503:
                break  # WhatsApp not connected, no point retrying
        except Exception as exc:
            last_error = f"whatsapp_exception_{exc}"
        if attempt < 3:
            import asyncio
            await asyncio.sleep(2)
    return False, last_error


@app.get("/login", response_class=HTMLResponse, name="login")
async def login_page(request: Request):
    """Login page"""
    from app.security.csrf import issue_csrf_token, set_csrf_cookie

    err = request.query_params.get("error", "")
    messages = {
        "invalid_credentials": "البريد الإلكتروني أو كلمة المرور غير صحيحة",
        "google_failed": "فشل تسجيل الدخول عبر Google",
        "locked": "تم قفل المحاولة مؤقتاً بسبب محاولات دخول فاشلة متكررة",
        "csrf": "انتهت صلاحية النموذج — أعد تحميل الصفحة وحاول مجدداً",
    }
    err_key = err.split("_")[0] if err.startswith("locked_") else err
    token = issue_csrf_token()
    response = templates.TemplateResponse(
        "login.html",
        {"request": request, "error": messages.get(err_key, err), "csrf_token": token},
    )
    set_csrf_cookie(response, token)
    return response


def _set_session_cookie(response, session_id: str):
    from app.security.app_secrets import session_cookie_kwargs

    response.set_cookie(key="session_id", value=session_id, **session_cookie_kwargs())
    return response


@app.post("/login", name="login_submit")
async def login_submit(request: Request, db: Session = Depends(get_db)):
    """Process login form"""
    from fastapi.responses import RedirectResponse  # type: ignore

    from app.security.csrf import validate_csrf_request
    from app.security.login_lockout import is_locked, record_failure, record_success

    form = await request.form()
    try:
        validate_csrf_request(request, str(form.get("csrf_token", "")))
    except HTTPException:
        return RedirectResponse(url="/login?error=csrf", status_code=302)

    email = str(form.get("username", "")).strip()  # type: ignore
    password = str(form.get("password", ""))  # type: ignore
    ip = _get_client_ip(request)
    locked, _secs = is_locked(email, ip)
    if locked:
        return RedirectResponse(url="/login?error=locked", status_code=302)

    user = authenticate_user(db, email, password)
    if not user:
        record_failure(email, ip)
        log_activity(db, "login_failed", "auth", f"محاولة دخول فاشلة: {email}", user_name=email, user_email=email, ip_address=ip, user_agent=_get_user_agent(request), level="warning")
        return RedirectResponse(
            url="/login?error=invalid_credentials", status_code=302
        )

    record_success(email, ip)
    # Create session
    session_id = create_session(int(user.id))  # type: ignore
    uname = f"{user.first_name or ''} {user.last_name or ''}".strip() or user.email
    log_activity(db, "login", "auth", f"تسجيل دخول: {uname}", user_id=int(user.id), user_name=uname, user_email=user.email, ip_address=ip, user_agent=_get_user_agent(request), level="success")

    # Redirect to dashboard with session cookie
    response = RedirectResponse(url="/dashboard", status_code=302)
    _set_session_cookie(response, session_id)
    return response


@app.get("/auth/google", name="google_login")
async def google_login():
    """Redirect to Google OAuth consent screen"""
    from urllib.parse import urlencode
    params = urlencode({
        "client_id": GOOGLE_CLIENT_ID,
        "redirect_uri": GOOGLE_REDIRECT_URI,
        "response_type": "code",
        "scope": "openid email profile",
        "access_type": "offline",
        "prompt": "select_account",
    })
    return RedirectResponse(
        url=f"https://accounts.google.com/o/oauth2/v2/auth?{params}",
        status_code=302,
    )


@app.get("/auth/google/callback", name="google_callback")
async def google_callback(request: Request, db: Session = Depends(get_db)):
    """Handle Google OAuth callback"""
    import httpx

    code = request.query_params.get("code")
    if not code:
        return RedirectResponse(
            url="/login?error=google_failed", status_code=302
        )

    # Exchange code for tokens
    async with httpx.AsyncClient() as client:
        token_resp = await client.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": GOOGLE_CLIENT_ID,
                "client_secret": GOOGLE_CLIENT_SECRET,
                "redirect_uri": GOOGLE_REDIRECT_URI,
                "grant_type": "authorization_code",
            },
        )
        if token_resp.status_code != 200:
            return RedirectResponse(
                url="/login?error=google_failed", status_code=302
            )
        tokens = token_resp.json()

        # Get user info
        userinfo_resp = await client.get(
            "https://www.googleapis.com/oauth2/v2/userinfo",
            headers={"Authorization": f"Bearer {tokens['access_token']}"},
        )
        if userinfo_resp.status_code != 200:
            return RedirectResponse(
                url="/login?error=google_failed", status_code=302
            )
        userinfo = userinfo_resp.json()

    # Login or register user
    user = google_login_or_register(
        db,
        google_id=userinfo.get("id", ""),
        email=userinfo.get("email", ""),
        first_name=userinfo.get("given_name", ""),
        last_name=userinfo.get("family_name", ""),
    )

    if not user:
        return RedirectResponse(
            url="/login?error=google_failed", status_code=302
        )

    # Create session
    session_id = create_session(int(user.id))  # type: ignore
    response = RedirectResponse(url="/dashboard", status_code=302)
    _set_session_cookie(response, session_id)
    return response
async def services_page(request: Request, db: Session = Depends(get_db)):
    """Services page"""
    user = get_current_user(request, db)
    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "services.html",
        {"request": request, "user": user, "subscription": sub_info},
    )


@app.get("/contact", response_class=HTMLResponse, name="contact")
async def contact_page(request: Request, db: Session = Depends(get_db)):
    """Contact us page"""
    user = get_current_user(request, db)
    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "contact.html",
        {"request": request, "user": user, "subscription": sub_info},
    )


@app.get("/subscribe", response_class=HTMLResponse)
async def subscribe_page(request: Request, db: Session = Depends(get_db)):
    """Subscription request page (package + phone + transfer + receipt)."""
    user_id = get_current_user_id(request)
    if not user_id:
        return RedirectResponse(url="/login", status_code=302)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    packages = (
        db.query(models.Package)
        .filter(models.Package.is_active.is_(True))
        .order_by(models.Package.price.asc())
        .all()
    )

    selected_package_id = request.query_params.get("package_id", "")
    if not selected_package_id:
        return RedirectResponse(url="/dashboard", status_code=302)

    selected_package = None
    package_lock_token = ""
    if selected_package_id:
        try:
            selected_package = (
                db.query(models.Package)
                .filter(
                    models.Package.id == int(selected_package_id),
                    models.Package.is_active.is_(True),
                )
                .first()
            )
            if selected_package is None:
                return RedirectResponse(url="/dashboard", status_code=302)
            else:
                package_lock_token = build_package_lock_token(
                    user_id, int(selected_package.id)
                )
        except ValueError:
            return RedirectResponse(url="/dashboard", status_code=302)

    return templates.TemplateResponse(
        "subscription_request.html",
        {
            "request": request,
            "user": user,
            "packages": packages,
            "selected_package_id": selected_package_id,
            "selected_package": selected_package,
            "package_lock_token": package_lock_token,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
        },
    )


@app.post("/api/subscribe-request")
async def subscribe_request_api(
    request: Request,
    phone: str = Form(...),
    transaction_id: str = Form(...),
    package_id: int = Form(...),
    subjects: str = Form(""),
    receipt_image: UploadFile | None = File(None),
    db: Session = Depends(get_db),
):
    """Create subscription verification request from subscription page."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse(
            {"status": "error", "message": "يجب تسجيل الدخول أولاً"}, status_code=401
        )

    try:
        import re

        phone = phone.strip()
        transaction_id = transaction_id.strip()
        subjects = subjects.strip()
        form_data = await request.form()
        package_lock_token = str(form_data.get("package_lock", "")).strip()

        if not package_lock_token:
            return JSONResponse(
                {
                    "status": "error",
                    "message": "يجب اختيار الباقة من الصفحة الرئيسية أولاً.",
                },
                status_code=400,
            )

        if not verify_package_lock_token(package_lock_token, user_id, package_id):
            return JSONResponse(
                {
                    "status": "error",
                    "message": "لا يمكن تعديل الباقة المحددة. يرجى إعادة الاختيار من صفحة الباقات.",
                },
                status_code=400,
            )

        if not phone or not transaction_id:
            return JSONResponse(
                {"status": "error", "message": "رقم الهاتف ورقم الحوالة مطلوبان"},
                status_code=400,
            )

        if not re.match(r"^07\d{8}$", phone):
            return JSONResponse(
                {"status": "error", "message": "رقم هاتف غير صحيح"},
                status_code=400,
            )

        package = (
            db.query(models.Package)
            .filter(models.Package.id == package_id, models.Package.is_active.is_(True))
            .first()
        )
        if not package:
            return JSONResponse(
                {"status": "error", "message": "الباقة غير موجودة"},
                status_code=404,
            )

        existing = (
            db.query(models.VerificationRequest)
            .filter(
                models.VerificationRequest.user_id == user_id,
                models.VerificationRequest.package_id == package_id,
                models.VerificationRequest.transaction_id == transaction_id,
                models.VerificationRequest.status == models.VerificationStatus.PENDING,
            )
            .first()
        )
        if existing is not None:  # type: ignore
            return JSONResponse(
                {
                    "status": "info",
                    "message": "لديك طلب اشتراك معلق بالفعل لنفس الباقة ورقم الحوالة.",
                },
                status_code=200,
            )

        note = ""
        if receipt_image is not None and receipt_image.filename:
            receipts_dir = Path("uploads") / "verification_receipts"
            receipts_dir.mkdir(parents=True, exist_ok=True)
            safe_name = Path(receipt_image.filename).name
            saved_name = f"{user_id}_{int(time.time())}_{safe_name}"
            saved_path = receipts_dir / saved_name
            content = await receipt_image.read()
            with open(saved_path, "wb") as f:
                f.write(content)
            note = f"receipt_image={saved_path.as_posix()}"

        verification = models.VerificationRequest(
            user_id=user_id,
            package_id=package_id,
            transaction_id=transaction_id,
            phone=phone,
            subjects=subjects,
            amount=float(package.price or 0),
            status=models.VerificationStatus.PENDING,
            admin_note=note,
        )
        db.add(verification)
        db.commit()
        log_activity(db, "subscribe_request", "subscription", f"طلب اشتراك جديد - باقة: {package.name}, رقم الحوالة: {transaction_id}", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="info")

        # Send WhatsApp notification to admin about new subscription request
        try:
            from datetime import datetime as _dt
            import httpx

            from app.whatsapp_config import get_whatsapp_sender_phone_e164  # type: ignore

            admin_phone = get_whatsapp_sender_phone_e164()
            amount_text = f"{int(package.price)}" if float(package.price or 0).is_integer() else f"{package.price:.2f}"
            today = _dt.now().strftime("%Y-%m-%d")
            wa_message = (
                f"تم اشتراك ({phone}) بمبلغ {amount_text} د.أ\n"
                f"الباقة: {package.name}\n"
                f"الوحدة: {subjects if subjects else 'غير محدد'}\n"
                f"بتاريخ: {today}\n"
                f"بانتظار الموافقه"
            )
            async with httpx.AsyncClient(timeout=15.0) as hc:
                await hc.post(
                    f"{WHATSAPP_SERVICE_URL}/send",
                    json={"phone": admin_phone, "message": wa_message},
                )
        except Exception as wa_err:
            print(f"WhatsApp subscription notify error: {wa_err}")

        return JSONResponse(
            {
                "status": "success",
                "message": "تم إرسال طلب الاشتراك بنجاح. بانتظار مراجعة الإدارة.",
            }
        )
    except Exception as e:
        print(f"Subscribe Request Error: {e}")
        return JSONResponse(
            {"status": "error", "message": "حدث خطأ أثناء إرسال طلب الاشتراك"},
            status_code=500,
        )


@app.post("/api/contact")
async def submit_contact(request: Request, db: Session = Depends(get_db)):
    """Handle contact form submission"""
    form = await request.form()
    name = str(form.get("name", "")).strip()
    email = str(form.get("email", "")).strip()
    msg_type = str(form.get("type", "")).strip()
    subject = str(form.get("subject", "")).strip()
    message = str(form.get("message", "")).strip()

    # Store in database
    try:
        contact = models.ContactMessage(
            name=name,
            email=email,
            message_type=msg_type,
            subject=subject,
            message=message,
        )
        db.add(contact)
        db.commit()
    except Exception:
        pass

    return JSONResponse({"status": "ok", "message": "تم إرسال رسالتك بنجاح"})


@app.get("/register", response_class=HTMLResponse, name="register")
async def register_page(request: Request, db: Session = Depends(get_db)):
    """Register page - MOE Jordan BTEC branded"""
    from app.security.csrf import issue_csrf_token, set_csrf_cookie

    user = get_current_user(request, db)
    error = request.query_params.get("error", "")
    error_messages = {
        "email_exists": "البريد الإلكتروني مسجل مسبقاً",
        "phone_exists": "رقم الهاتف مسجل مسبقاً",
        "password_mismatch": "كلمتا المرور غير متطابقتين",
        "password_weak": "كلمة المرور ضعيفة — 8+ أحرف مع حرف كبير وصغير ورقم",
        "missing_fields": "يرجى ملء جميع الحقول المطلوبة",
        "csrf": "انتهت صلاحية النموذج — أعد تحميل الصفحة",
    }
    token = issue_csrf_token()
    response = templates.TemplateResponse(
        "register.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv(
                "APP_TITLE", "منظومة تصحيح الواجبات بالذكاء الاصطناعي"
            ),
            "error": error_messages.get(error, ""),
            "csrf_token": token,
        },
    )
    set_csrf_cookie(response, token)
    return response


@app.post("/register", name="register_submit")
async def register_submit(request: Request, db: Session = Depends(get_db)):
    """Process registration form"""
    from fastapi.responses import RedirectResponse  # type: ignore

    from app.security.csrf import validate_csrf_request
    from app.security.password_policy import validate_password

    form = await request.form()
    try:
        validate_csrf_request(request, str(form.get("csrf_token", "")))
    except HTTPException:
        return RedirectResponse(url="/register?error=csrf", status_code=302)

    first_name = str(form.get("first_name", "")).strip()  # type: ignore
    last_name = str(form.get("last_name", "")).strip()  # type: ignore
    email = str(form.get("email", "")).strip()  # type: ignore
    country_code = str(form.get("country_code", "")).strip()  # type: ignore
    raw_phone = str(form.get("phone", "")).strip()  # type: ignore
    phone = (country_code + raw_phone) if raw_phone else ""
    job_title = str(form.get("role", "")).strip()  # type: ignore
    password = str(form.get("password", ""))  # type: ignore
    confirm_password = str(form.get("confirm_password", ""))  # type: ignore

    # Validation
    if not all([first_name, last_name, email, password]):
        return RedirectResponse(
            url="/register?error=missing_fields", status_code=302
        )

    if password != confirm_password:
        return RedirectResponse(
            url="/register?error=password_mismatch", status_code=302
        )

    ok_pw, _pw_errors = validate_password(password)
    if not ok_pw:
        return RedirectResponse(url="/register?error=password_weak", status_code=302)

    # Check if email already exists
    existing_email = db.query(models.User).filter(models.User.email == email).first()
    if existing_email:
        return RedirectResponse(
            url="/register?error=email_exists", status_code=302
        )

    # Check if phone already exists
    if phone:
        existing_phone = db.query(models.User).filter(models.User.phone == phone).first()
        if existing_phone:
            return RedirectResponse(
                url="/register?error=phone_exists", status_code=302
            )

    try:
        user = register_user(
            db, email, password, first_name, last_name, job_title, phone
        )
    except ValueError:
        return RedirectResponse(
            url="/register?error=email_exists", status_code=302
        )

    # Auto-login after registration
    session_id = create_session(int(user.id))  # type: ignore
    rname = f"{first_name} {last_name}".strip()
    log_activity(db, "register", "auth", f"تسجيل مستخدم جديد: {rname} ({email})", user_id=int(user.id), user_name=rname, user_email=email, ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="success")

    response = RedirectResponse(url="/dashboard", status_code=302)
    _set_session_cookie(response, session_id)
    return response


@app.get("/logout", response_class=HTMLResponse, name="logout")
async def logout_user(request: Request):
    """Logout user"""
    from fastapi.responses import RedirectResponse  # type: ignore

    session_id = request.cookies.get("session_id")
    if session_id:
        uid = get_current_user_id(request)
        delete_session(session_id)
        from app.database import SessionLocal  # type: ignore
        _db = SessionLocal()
        try:
            log_activity(_db, "logout", "auth", "تسجيل خروج", user_id=uid, user_name=_get_user_display(_db, uid), user_email=_get_user_email(_db, uid), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request))
        finally:
            _db.close()

    response = RedirectResponse(url="/login", status_code=302)
    response.delete_cookie("session_id")
    return response


@app.get(
    "/forgot-password", response_class=HTMLResponse, name="forgot_password"
)
async def forgot_password_page(request: Request):
    """Forgot password page"""
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    """Root - Redirect to login or dashboard"""
    from fastapi.responses import RedirectResponse  # type: ignore

    # Check if user is logged in
    user_id = get_current_user_id(request)

    if user_id:
        # User is logged in, redirect to dashboard
        return RedirectResponse(url="/dashboard", status_code=302)
    else:
        # User not logged in, show dashboard (home page)
        return RedirectResponse(url="/dashboard", status_code=302)


@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request, db: Session = Depends(get_db)):
    """Dashboard - Main home page (public)"""
    user_id = get_current_user_id(request)
    user = None
    sub_info = None
    if user_id:
        user = db.query(models.User).filter(models.User.id == user_id).first()
        sub_info = get_subscription_info(db, user_id)
    packages = (
        db.query(models.Package)
        .filter(models.Package.is_active.is_(True))
        .all()
    )
    from app.database import get_package_rows, get_student_package_rows  # type: ignore

    package_rows = get_package_rows(db)
    student_package_rows = get_student_package_rows(db)

    return templates.TemplateResponse(
        "home.html",
        {
            "request": request,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "user": user,
            "packages": packages,
            "package_rows": package_rows,
            "student_package_rows": student_package_rows,
            "subscription": sub_info,
        },
        headers={"Cache-Control": "no-store"},
    )


@app.post("/api/send-verification")
async def send_verification(request: Request, db: Session = Depends(get_db)):
    """Save verification request to DB for admin review"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse(
            {"status": "error", "message": "يجب تسجيل الدخول أولاً"}, status_code=401
        )

    try:
        data = await request.json()
        phone = data.get("phone", "").strip()
        transaction_id = data.get("transaction_id", "").strip()
        package_id = data.get("package_id")
        subjects = data.get("subjects", "")

        if not phone or not transaction_id or not package_id:
            return JSONResponse(
                {
                    "status": "error",
                    "message": "رقم الهاتف ورقم الحوالة والباقة مطلوبة",
                },
                status_code=400,
            )

        # Validate Jordanian phone number
        import re

        if not re.match(r"^07\d{8}$", phone):
            return JSONResponse(
                {"status": "error", "message": "رقم هاتف غير صحيح"},
                status_code=400,
            )

        # Check package exists
        package = (
            db.query(models.Package)
            .filter(models.Package.id == package_id)
            .first()
        )
        if not package:
            return JSONResponse(
                {"status": "error", "message": "الباقة غير موجودة"},
                status_code=404,
            )

        # Check if there's already a pending request for this user+package
        existing = (
            db.query(models.VerificationRequest)
            .filter(
                models.VerificationRequest.user_id == user_id,
                models.VerificationRequest.package_id == package_id,
                models.VerificationRequest.transaction_id == transaction_id,
                models.VerificationRequest.status
                == models.VerificationStatus.PENDING,
            )
            .first()
        )
        if existing is not None:  # type: ignore
            return JSONResponse(
                {
                    "status": "info",
                    "message": (
                        "لديك طلب معلق بالفعل لنفس الباقة ورقم الحوالة. "
                        "يرجى انتظار مراجعة الإدارة."
                    ),
                },
                status_code=200,
            )

        subjects_str = (
            ",".join(subjects) if isinstance(subjects, list) else subjects
        )
        verification = models.VerificationRequest(
            user_id=user_id,
            package_id=package_id,
            transaction_id=transaction_id,
            phone=phone,
            subjects=subjects_str,
            amount=float(package.price or 0),
            status=models.VerificationStatus.PENDING,
        )
        db.add(verification)
        db.commit()

        print(
            f"📋 New verification request: user={user_id}, "
            f"phone={phone}, tx={transaction_id}, package={package_id}"
        )

        return JSONResponse(
            {
                "status": "success",
                "message": (
                    "تم إرسال طلبك بنجاح! سيقوم المسؤول بمراجعة "
                    "الحوالة وإرسال رمز التفعيل."
                ),
            }
        )

    except Exception as e:
        print(f"Send Verification Error: {e}")
        return JSONResponse(
            {"status": "error", "message": "حدث خطأ أثناء إرسال الطلب"},
            status_code=500,
        )


@app.post("/api/subscribe")
async def subscribe_api(request: Request, db: Session = Depends(get_db)):
    """Verify activation code and activate subscription"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse(
            {
                "status": "error",
                "message": "يجب تسجيل الدخول أولاً",
            },
            status_code=401,
        )

    try:
        data = await request.json()
        activation_code = data.get("otp", "").strip()
        phone = data.get("phone", "").strip()

        if not activation_code or not phone:
            return JSONResponse(
                {
                    "status": "error",
                    "message": "رمز التفعيل ورقم الهاتف مطلوبان",
                },
                status_code=400,
            )

        # Find the verification request with matching activation code
        verification = (
            db.query(models.VerificationRequest)
            .filter(
                models.VerificationRequest.user_id == user_id,
                models.VerificationRequest.phone == phone,
                models.VerificationRequest.activation_code == activation_code,
                models.VerificationRequest.status
                == models.VerificationStatus.CODE_SENT,
            )
            .order_by(models.VerificationRequest.created_at.desc())
            .first()
        )

        if verification is None:  # type: ignore
            return JSONResponse(
                {
                    "status": "error",
                    "message": "رمز التفعيل غير صحيح. يرجى المحاولة مرة أخرى.",
                },
                status_code=400,
            )

        # Code matches - activate subscription
        package = (
            db.query(models.Package)
            .filter(models.Package.id == verification.package_id)
            .first()
        )
        if package is None:
            return JSONResponse(
                {"status": "error", "message": "الباقة غير موجودة"},
                status_code=404,
            )

        # Create active subscription
        new_sub = models.Subscription(
            user_id=user_id,
            package_id=package.id,
            status=models.SubscriptionStatus.ACTIVE,  # type: ignore
            transaction_id=verification.transaction_id,
            subjects=verification.subjects,
            assignments_limit=package.assignment_limit,
            start_date=datetime.utcnow(),
        )
        db.add(new_sub)

        # Mark verification as completed
        # type: ignore
        verification.status = models.VerificationStatus.VERIFIED
        db.commit()

        return JSONResponse(
            {"status": "success", "message": "تم تفعيل الاشتراك بنجاح! 🎉"}
        )

    except Exception as e:
        print(f"Subscription Error: {e}")
        return JSONResponse(
            {"status": "error", "message": "حدث خطأ أثناء معالجة الطلب"},
            status_code=500,
        )


@app.get("/api/subscription-status")
async def subscription_status(request: Request, db: Session = Depends(get_db)):
    """Get user's current subscription status"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse(
            {"status": "error", "message": "يجب تسجيل الدخول أولاً"},
            status_code=401,
        )

    info = get_subscription_info(db, user_id)
    return JSONResponse({"status": "success", **info})


# ═══════════════════════════════════════════════════════
# Admin Panel - Verification Request Management
# ═══════════════════════════════════════════════════════


@app.get("/admin", response_class=HTMLResponse)
async def admin_panel(request: Request, db: Session = Depends(get_db)):
    """Admin panel for managing verification requests"""
    user_id = get_current_user_id(request)
    if not user_id:
        return HTMLResponse("<script>window.location='/login'</script>")

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:  # type: ignore
        return HTMLResponse(
            "<h2 style='text-align:center;margin-top:100px;'> غير مصرح لك بالوصول</h2>",
            status_code=403,
        )

    sub_info = get_subscription_info(db, user_id)
    return templates.TemplateResponse(
        "admin.html",
        {"request": request, "user": user, "app_title": "لوحة الإدارة", "subscription": sub_info},
    )


@app.get("/api/admin/verifications")
async def get_verifications(request: Request, db: Session = Depends(get_db)):
    """Get all verification requests for admin"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user or user.role != models.UserRole.ADMIN:  # type: ignore
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    requests_list = (
        db.query(models.VerificationRequest)
        .order_by(models.VerificationRequest.created_at.desc())
        .all()
    )

    result = []
    for vr in requests_list:
        # Get user info
        req_user = db.query(models.User).filter(models.User.id == vr.user_id).first()
        pkg = (
            db.query(models.Package).filter(models.Package.id == vr.package_id).first()
        )

        # Calculate remaining limit if user has an active sub
        remaining_limit = 0
        if req_user:
            active_sub = (
                db.query(models.Subscription)
                .filter(
                    models.Subscription.user_id == req_user.id,
                    models.Subscription.status == models.SubscriptionStatus.ACTIVE,
                )
                .order_by(models.Subscription.created_at.desc())
                .first()
            )
            if active_sub is not None:
                assignments_limit = int(active_sub.assignments_limit or 0)
                assignments_used = int(active_sub.assignments_used or 0)
                remaining_limit = max(0, assignments_limit - assignments_used)

        result.append(
            {
                "id": vr.id,
                "user_id": vr.user_id,
                "user_name": (
                    f"{req_user.first_name or ''} {req_user.last_name or ''}".strip()
                    if req_user
                    else "غير معروف"
                ),
                "user_email": req_user.email if req_user else "",
                "phone": vr.phone,
                "transaction_id": vr.transaction_id,
                "package_id": vr.package_id,
                "package_name": pkg.name if pkg else "-",
                "subjects": vr.subjects or "",
                "remaining_limit": remaining_limit,
                "activation_code": vr.activation_code or "",
                "verified_teacher_name": vr.verified_teacher_name or "",
                "amount": (
                    vr.amount
                    if vr.amount is not None
                    else (float(pkg.price or 0) if pkg else "")
                ),
                "status": vr.status.value,
                "admin_note": vr.admin_note or "",
                "created_at": vr.created_at.isoformat() if vr.created_at else "",  # type: ignore
            }
        )

    return JSONResponse({"status": "success", "requests": result})


@app.get("/api/check-ai-balance")
async def check_ai_balance():
    """Legacy balance check (OpenRouter removed). Always returns no popup."""
    return JSONResponse({
        "balance": None,
        "total_credits": None,
        "total_usage": None,
        "show_popup": False,
        "whatsapp_sent": False,
        "error": None,
    })


@app.get("/api/admin/whatsapp-status")
async def admin_whatsapp_status(request: Request, db: Session = Depends(get_db)):
    """Return WhatsApp integration readiness for admin dashboard."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:  # type: ignore
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    from app.whatsapp_config import get_whatsapp_settings  # type: ignore
    from app.whatsapp_manager import is_whatsapp_service_running  # type: ignore

    wa_settings = get_whatsapp_settings()

    if not is_whatsapp_service_running():
        return JSONResponse({
            "status": "success",
            "configured": False,
            "connected": False,
            "service_running": False,
            "expected_phone": wa_settings["sender_phone"],
            "expected_phone_e164": wa_settings["sender_phone_e164"],
            "error": "خدمة الواتساب غير متاحة — جاري التشغيل التلقائي أو شغّل whatsapp_service يدوياً",
        })

    try:
        import httpx
        async with httpx.AsyncClient(timeout=5.0) as hc:
            resp = await hc.get(f"{WHATSAPP_SERVICE_URL}/status")
        data = resp.json()
        return JSONResponse({
            "status": "success",
            "configured": True,
            "service_running": True,
            "connected": data.get("ready", False),
            "hasQR": data.get("hasQR", False),
            "info": data.get("info"),
            "error": data.get("error"),
            "expected_phone": wa_settings["sender_phone"],
            "expected_phone_e164": wa_settings["sender_phone_e164"],
            "phone_match": data.get("phoneMatch", None),
        })
    except Exception:
        return JSONResponse({
            "status": "success",
            "configured": False,
            "service_running": False,
            "connected": False,
            "expected_phone": wa_settings["sender_phone"],
            "expected_phone_e164": wa_settings["sender_phone_e164"],
            "error": "خدمة الواتساب غير متاحة — تأكد من تشغيلها",
        })


@app.get("/api/admin/whatsapp-settings")
async def admin_whatsapp_settings_get(request: Request, db: Session = Depends(get_db)):
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    from app.whatsapp_config import get_whatsapp_settings  # type: ignore

    return JSONResponse({"status": "success", **get_whatsapp_settings()})


@app.post("/api/admin/whatsapp-settings")
async def admin_whatsapp_settings_save(request: Request, db: Session = Depends(get_db)):
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    from app.whatsapp_config import get_whatsapp_settings, set_whatsapp_sender_phone  # type: ignore

    try:
        body = await request.json()
    except Exception:
        body = {}
    phone = str(body.get("sender_phone", "")).strip()
    try:
        saved = set_whatsapp_sender_phone(phone)
    except ValueError as exc:
        return JSONResponse({"status": "error", "message": str(exc)}, status_code=400)

    log_activity(
        db,
        action="whatsapp_settings",
        category="admin",
        details=f"تحديث رقم واتساب النظام إلى {saved}",
        user_id=user.id,
        user_name=user.name or user.email or "",
        user_email=user.email or "",
        ip_address=_get_client_ip(request),
        user_agent=_get_user_agent(request),
        level="warning",
    )

    from app.whatsapp_manager import restart_whatsapp_service  # type: ignore

    restart_whatsapp_service()

    return JSONResponse({
        "status": "success",
        "message": f"تم حفظ رقم الواتساب: {saved}. إذا غيّرت الرقم، امسح QR من الهاتف الجديد.",
        **get_whatsapp_settings(),
    })


@app.get("/api/admin/whatsapp-qr")
async def admin_whatsapp_qr(request: Request, db: Session = Depends(get_db)):
    """Get WhatsApp QR code for scanning."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    try:
        import httpx
        async with httpx.AsyncClient(timeout=5.0) as hc:
            resp = await hc.get(f"{WHATSAPP_SERVICE_URL}/qr")
        data = resp.json()
        if data.get("qr"):
            return JSONResponse({"status": "success", "qr": data["qr"]})
        elif data.get("status") == "connected":
            return JSONResponse({"status": "success", "connected": True})
        else:
            return JSONResponse({"status": "waiting", "message": "لم يتم إنشاء رمز QR بعد"})
    except Exception:
        return JSONResponse({"status": "error", "message": "خدمة الواتساب غير متاحة"}, status_code=503)


@app.post("/api/admin/whatsapp-send-test")
async def admin_whatsapp_send_test(request: Request, db: Session = Depends(get_db)):
    """Send a test WhatsApp message."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    data = await request.json()
    phone = data.get("phone", "").strip()
    message = data.get("message", "رسالة تجريبية من منظومة التصحيح 🎓").strip()
    if not phone:
        return JSONResponse({"status": "error", "message": "رقم الهاتف مطلوب"}, status_code=400)
    if not message:
        return JSONResponse({"status": "error", "message": "نص الرسالة مطلوب"}, status_code=400)

    try:
        import httpx
        async with httpx.AsyncClient(timeout=15.0) as hc:
            resp = await hc.post(f"{WHATSAPP_SERVICE_URL}/send", json={"phone": phone, "message": message})
        result = resp.json()
        if result.get("success"):
            return JSONResponse({"status": "success", "message": "تم إرسال الرسالة بنجاح ✅"})
        return JSONResponse({"status": "error", "message": f"فشل الإرسال: {result.get('error', 'خطأ غير معروف')}"}, status_code=500)
    except Exception as e:
        return JSONResponse({"status": "error", "message": f"خدمة الواتساب غير متاحة: {e}"}, status_code=503)


@app.post("/api/admin/whatsapp-logout")
async def admin_whatsapp_logout(request: Request, db: Session = Depends(get_db)):
    """Disconnect WhatsApp session."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    try:
        import httpx
        async with httpx.AsyncClient(timeout=10.0) as hc:
            resp = await hc.post(f"{WHATSAPP_SERVICE_URL}/logout")
        return JSONResponse(resp.json())
    except Exception:
        return JSONResponse({"status": "error", "message": "خدمة الواتساب غير متاحة"}, status_code=503)


@app.post("/api/admin/set-code")
async def admin_set_code(request: Request, db: Session = Depends(get_db)):
    """Admin approves verification request, activates subscription, and notifies user on WhatsApp."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:  # type: ignore
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    try:
        data = await request.json()
        request_id = data.get("request_id")
        teacher_name = data.get("teacher_name", "").strip()

        if not request_id:
            return JSONResponse(
                {"status": "error", "message": "رقم الطلب مطلوب"},
                status_code=400,
            )

        vr = (
            db.query(models.VerificationRequest)
            .filter(models.VerificationRequest.id == request_id)
            .first()
        )

        if not vr:
            return JSONResponse(
                {"status": "error", "message": "الطلب غير موجود"}, status_code=404
            )

        if vr.status == models.VerificationStatus.VERIFIED:  # type: ignore
            return JSONResponse(
                {"status": "info", "message": "تمت الموافقة مسبقًا على هذا الطلب"},
                status_code=200,
            )

        package = (
            db.query(models.Package)
            .filter(models.Package.id == vr.package_id)
            .first()
        )
        if package is None:
            return JSONResponse(
                {"status": "error", "message": "الباقة غير موجودة"},
                status_code=404,
            )

        old_active = (
            db.query(models.Subscription)
            .filter(
                models.Subscription.user_id == vr.user_id,
                models.Subscription.status == models.SubscriptionStatus.ACTIVE,
            )
            .order_by(models.Subscription.created_at.desc())
            .first()
        )

        # Calculate carry-over: remaining balance from old subscription
        carry_over = 0
        if old_active is not None:
            old_limit = int(old_active.assignments_limit or 0)
            old_used = int(old_active.assignments_used or 0)
            carry_over = max(0, old_limit - old_used)
            old_active.status = models.SubscriptionStatus.CANCELLED  # type: ignore

        new_limit = int(package.assignment_limit or 0) + carry_over

        new_sub = models.Subscription(
            user_id=vr.user_id,
            package_id=package.id,
            status=models.SubscriptionStatus.ACTIVE,  # type: ignore
            transaction_id=vr.transaction_id,
            subjects=vr.subjects,
            assignments_limit=new_limit,
            start_date=datetime.utcnow(),
        )
        db.add(new_sub)

        # --- Per-subject balance update ---
        sub_subjects = [s.strip() for s in (vr.subjects or "").replace("|", ",").split(",") if s.strip()]
        if sub_subjects:
            pkg_limit = int(package.assignment_limit or 0)
            # Split the package limit evenly across selected subjects
            limit_each = pkg_limit // len(sub_subjects)
            remainder = pkg_limit % len(sub_subjects)
            for i, subj in enumerate(sub_subjects):
                extra = remainder if i == 0 else 0
                existing_balance = (
                    db.query(models.SubjectBalance)
                    .filter(
                        models.SubjectBalance.user_id == vr.user_id,
                        models.SubjectBalance.subject == subj,
                    )
                    .first()
                )
                if existing_balance is not None:
                    # Carry over remaining per-subject balance + add new limit
                    old_remaining = max(0, (existing_balance.assignments_limit or 0) - (existing_balance.assignments_used or 0))
                    existing_balance.assignments_limit = existing_balance.assignments_used + old_remaining + limit_each + extra  # type: ignore
                else:
                    db.add(models.SubjectBalance(
                        user_id=vr.user_id,
                        subject=subj,
                        assignments_limit=limit_each + extra,
                        assignments_used=0,
                    ))
        # --- End per-subject balance update ---

        if teacher_name:
            vr.verified_teacher_name = teacher_name
        vr.status = models.VerificationStatus.VERIFIED  # type: ignore

        db.commit()

        amount_value = float(vr.amount if vr.amount is not None else (package.price or 0))
        whatsapp_ok, whatsapp_status = await send_whatsapp_activation_message(
            vr.phone, amount_value
        )

        print(
            f"✅ Admin approved request #{request_id}: Name='{teacher_name}', WhatsApp='{whatsapp_status}'"
        )
        carry_info = f" (رصيد مُرحّل: {carry_over})" if carry_over > 0 else ""
        log_activity(db, "approve_subscription", "admin", f"موافقة على طلب اشتراك #{request_id} - المعلم: {teacher_name} - الرصيد: {new_limit}{carry_info}", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="success")

        if whatsapp_ok:
            message = "تمت الموافقة وتفعيل الاشتراك، وتم إرسال رسالة واتساب بنجاح"
        else:
            message = (
                "تمت الموافقة وتفعيل الاشتراك، لكن تعذر إرسال واتساب. "
                "تحقق من إعدادات WHATSAPP_PHONE_NUMBER_ID و WHATSAPP_ACCESS_TOKEN"
            )

        return JSONResponse(
            {"status": "success", "message": message}
        )

    except Exception as e:
        print(f"Admin Set Code Error: {e}")
        return JSONResponse({"status": "error", "message": "حدث خطأ"}, status_code=500)


@app.post("/api/admin/reject")
async def admin_reject(request: Request, db: Session = Depends(get_db)):
    """Admin rejects a verification request"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user or user.role != models.UserRole.ADMIN:  # type: ignore
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    try:
        data = await request.json()
        request_id = data.get("request_id")

        vr = (
            db.query(models.VerificationRequest)
            .filter(models.VerificationRequest.id == request_id)
            .first()
        )

        if vr is None:  # type: ignore
            return JSONResponse(
                {"status": "error", "message": "الطلب غير موجود"}, status_code=404
            )

        vr.status = models.VerificationStatus.REJECTED  # type: ignore
        db.commit()

        log_activity(db, "reject_subscription", "admin", f"رفض طلب اشتراك #{request_id}", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="warning")
        print(f"❌ Admin rejected request #{request_id}")

        return JSONResponse({"status": "success", "message": "تم رفض الطلب"})

    except Exception as e:
        print(f"Admin Reject Error: {e}")
        return JSONResponse({"status": "error", "message": "حدث خطأ"}, status_code=500)


@app.delete("/api/admin/verifications/{req_id}")
async def admin_delete_verification(req_id: int, request: Request, db: Session = Depends(get_db)):
    """Admin deletes a single verification request."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user or user.role != models.UserRole.ADMIN:  # type: ignore
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    vr = db.query(models.VerificationRequest).filter(models.VerificationRequest.id == req_id).first()
    if not vr:
        return JSONResponse({"status": "error", "message": "الطلب غير موجود"}, status_code=404)

    db.delete(vr)
    db.commit()
    log_activity(db, "delete_verification", "admin", f"حذف طلب تفعيل #{req_id}", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="warning")
    return JSONResponse({"status": "success", "message": "تم حذف الطلب"})


@app.delete("/api/admin/verifications")
async def admin_delete_all_verifications(request: Request, db: Session = Depends(get_db)):
    """Admin deletes all verification requests."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user or user.role != models.UserRole.ADMIN:  # type: ignore
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    count = db.query(models.VerificationRequest).count()
    db.query(models.VerificationRequest).delete()
    db.commit()
    log_activity(db, "delete_all_verifications", "admin", f"حذف جميع طلبات التفعيل ({count} طلب)", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="warning")
    return JSONResponse({"status": "success", "message": f"تم حذف {count} طلب"})


@app.get("/api/admin/reports")
async def admin_reports(request: Request, db: Session = Depends(get_db)):
    """Get aggregated subscription and usage data for reports dashboard"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:  # type: ignore
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    # Get all subscriptions
    subscriptions = db.query(models.Subscription).all()

    report_data = []

    for sub in subscriptions:
        sub_user = db.query(models.User).filter(models.User.id == sub.user_id).first()
        pkg = (
            db.query(models.Package).filter(models.Package.id == sub.package_id).first()
        )

        # Check if the user has a verification request to get amount paid and phone
        vr = (
            db.query(models.VerificationRequest)
            .filter(
                models.VerificationRequest.user_id == sub.user_id,
                models.VerificationRequest.package_id == sub.package_id,
            )
            .order_by(models.VerificationRequest.created_at.desc())
            .first()
        )

        amount_paid = vr.amount if vr and vr.amount else (pkg.price if pkg else 0.0)
        phone = vr.phone if vr else "غير متوفر"

        report_data.append(
            {
                "subscription_id": sub.id,
                "user_id": sub_user.id if sub_user else None,
                "user_name": (
                    f"{sub_user.first_name or ''} {sub_user.last_name or ''}".strip()
                    if sub_user
                    else "غير معروف"
                ),
                "phone": phone,
                "package_name": pkg.name if pkg else "غير معروف",
                "status": sub.status.value,
                "amount_paid": amount_paid,
                "assignments_used": sub.assignments_used or 0,
                "assignments_limit": sub.assignments_limit or 0,
            }
        )

    return JSONResponse({"status": "success", "reports": report_data})


@app.get("/api/admin/users")
async def get_admin_users(request: Request, db: Session = Depends(get_db)):
    """Get all users for User Management dashboard"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    users = db.query(models.User).filter(models.User.is_active.is_(True)).all()

    users_data = []
    for u in users:
        if u.role == models.UserRole.ADMIN:
            continue

        latest_sub = (
            db.query(models.Subscription)
            .filter(models.Subscription.user_id == u.id)
            .order_by(models.Subscription.created_at.desc())
            .first()
        )

        # Determine Phone - Since phone isn't on User model, get from last verification
        phone = ""
        vr = (
            db.query(models.VerificationRequest)
            .filter(models.VerificationRequest.user_id == u.id)
            .order_by(models.VerificationRequest.created_at.desc())
            .first()
        )
        if vr:
            phone = vr.phone

        remaining_limit = 0
        pkg_name = ""
        subject_details: list = []

        # Per-subject balances (new system)
        balances = (
            db.query(models.SubjectBalance)
            .filter(models.SubjectBalance.user_id == u.id)
            .all()
        )
        if balances:
            remaining_limit = sum(
                max(0, (b.assignments_limit or 0) - (b.assignments_used or 0))
                for b in balances
            )
            subject_details = [
                {
                    "subject": b.subject,
                    "remaining": max(0, (b.assignments_limit or 0) - (b.assignments_used or 0)),
                    "limit": b.assignments_limit or 0,
                    "used": b.assignments_used or 0,
                }
                for b in balances
            ]
        elif (
            latest_sub is not None
            and latest_sub.status == models.SubscriptionStatus.ACTIVE
        ):
            limit = int(latest_sub.assignments_limit or 0)
            used = int(latest_sub.assignments_used or 0)
            remaining_limit = max(0, limit - used)

        if (
            latest_sub is not None
            and latest_sub.status == models.SubscriptionStatus.ACTIVE
        ):
            pkg = (
                db.query(models.Package)
                .filter(models.Package.id == latest_sub.package_id)
                .first()
            )
            if pkg:
                pkg_name = pkg.name

        users_data.append(
            {
                "id": u.id,
                "name": f"{u.first_name or ''} {u.last_name or ''}".strip()
                or u.name
                or "غير معروف",
                "email": u.email,
                "phone": phone,
                "created_at": u.created_at.strftime("%Y-%m-%d"),
                "package_name": pkg_name,
                "remaining_limit": remaining_limit,
                "subject_details": subject_details,
            }
        )

    return JSONResponse(
        {"status": "success", "total_users": len(users_data), "users": users_data}
    )


@app.post("/api/admin/users")
async def add_admin_user(request: Request, db: Session = Depends(get_db)):
    """Manually add a new user from Admin panel"""
    from passlib.context import CryptContext  # type: ignore

    pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if not admin or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    data = await request.json()
    email = data.get("email", "").strip().lower()
    name = data.get("name", "").strip()
    password = data.get("password", "")

    if not email or not password:
        return JSONResponse(
            {"status": "error", "message": "البريد الإلكتروني ورقم المرور مطلوبان"},
            status_code=400,
        )

    # Check if user exists
    existing = db.query(models.User).filter(models.User.email == email).first()
    if existing:
        return JSONResponse(
            {"status": "error", "message": "البريد الإلكتروني مسجل مسبقاً"},
            status_code=400,
        )

    # Split name
    parts = name.split(" ", 1)
    first_name = parts[0]
    last_name = parts[1] if len(parts) > 1 else ""

    hash_pwd = pwd_context.hash(password)

    new_user = models.User(
        email=email,
        password_hash=hash_pwd,  # type: ignore
        first_name=first_name,
        last_name=last_name,
        name=name,
        login_method="local",
        role=models.UserRole.USER,
        is_active=True,  # type: ignore
    )

    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    log_activity(db, "admin_add_user", "admin", f"إضافة مستخدم جديد: {name} ({email})", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="info")

    # Optional: Save phone info into a verification record if needed or just skip it
    return JSONResponse({"status": "success", "message": "تمت إضافة المستخدم بنجاح!"})


@app.put("/api/admin/users/{edit_user_id}")
async def update_admin_user(
    edit_user_id: int, request: Request, db: Session = Depends(get_db)
):
    """Update existing user data from Admin panel"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    target_user = db.query(models.User).filter(models.User.id == edit_user_id).first()
    if target_user is None:
        return JSONResponse(
            {"status": "error", "message": "المستخدم غير موجود"}, status_code=404
        )

    data = await request.json()
    email = data.get("email", "").strip().lower()
    name = data.get("name", "").strip()
    password = data.get("password", "")

    if email:
        if email != target_user.email:
            existing = (
                db.query(models.User)
                .filter(models.User.email == email, models.User.id != target_user.id)
                .first()
            )
            if existing:
                return JSONResponse(
                    {"status": "error", "message": "الإيميل الجديد مسجل لمستخدم آخر"},
                    status_code=400,
                )
        target_user.email = email

    if name:
        target_user.name = name
        parts = name.split(" ", 1)
        target_user.first_name = parts[0]
        if len(parts) > 1:
            target_user.last_name = parts[1]

    if password:
        from passlib.context import CryptContext  # type: ignore

        pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
        target_user.password_hash = pwd_context.hash(password)

    db.commit()
    return JSONResponse(
        {"status": "success", "message": "تم تحديث بيانات التعديل بنجاح!"}
    )


@app.delete("/api/admin/users/{del_user_id}")
async def disable_admin_user(
    del_user_id: int, request: Request, db: Session = Depends(get_db)
):
    """Soft delete user - sets is_active to False"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    target_user = db.query(models.User).filter(models.User.id == del_user_id).first()
    if target_user is None:
        return JSONResponse(
            {"status": "error", "message": "المستخدم غير موجود"}, status_code=404
        )

    if target_user.role == models.UserRole.ADMIN:
        return JSONResponse(
            {"status": "error", "message": "لا يمكنك حذف الأدمن!"}, status_code=400
        )

    target_user.is_active = False  # type: ignore

    # Set active subscription to cancelled when deleted
    sub = (
        db.query(models.Subscription)
        .filter(
            models.Subscription.user_id == target_user.id,
            models.Subscription.status == models.SubscriptionStatus.ACTIVE,
        )
        .order_by(models.Subscription.created_at.desc())
        .first()
    )

    if sub is not None:
        _orm_set(sub, "status", models.SubscriptionStatus.CANCELLED)

    db.commit()
    return JSONResponse({"status": "success", "message": "تم حذف المستخدم بنجاح!"})


# ═══════════════════════════════════════════════════════
# Admin: Assignment Management
# ═══════════════════════════════════════════════════════

@app.get("/api/admin/assignments")
async def list_admin_assignments(request: Request, db: Session = Depends(get_db)):
    """List all assignments with reference solution status for admin."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    assignments = db.query(Assignment).order_by(Assignment.created_at.desc()).all()
    result = []
    for a in assignments:
        creator = db.query(models.User).filter(models.User.id == a.created_by).first()
        has_solution = bool(a.reference_solution_json or a.reference_solution_text)
        criteria_count = db.query(models.GradingCriteria).filter(
            models.GradingCriteria.assignment_id == a.id
        ).count()
        result.append({
            "id": a.id,
            "title": a.title,
            "status": str(a.status.value) if hasattr(a.status, 'value') else str(a.status),
            "has_solution": has_solution,
            "criteria_count": criteria_count,
            "creator": creator.name if creator else "غير معروف",
            "created_at": a.created_at.strftime("%Y-%m-%d %H:%M") if a.created_at else "",
        })
    return {"status": "success", "assignments": result, "total": len(result)}


@app.delete("/api/admin/assignments/{assignment_id}/solution")
async def clear_assignment_solution(
    assignment_id: int, request: Request, db: Session = Depends(get_db)
):
    """Clear reference solution, criteria, and cache for a specific assignment."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        return JSONResponse({"status": "error", "message": "المهمة غير موجودة"}, status_code=404)

    # Clear reference solution fields
    assignment.reference_solution_json = None  # type: ignore
    assignment.reference_solution_text = None  # type: ignore
    assignment.solution_hash = None  # type: ignore
    assignment.status = AssignmentStatus.DRAFT  # type: ignore

    # Clear grading criteria
    db.query(models.GradingCriteria).filter(
        models.GradingCriteria.assignment_id == assignment_id
    ).delete(synchronize_session=False)

    # Clear all grading cache (no assignment_id column, clear all)
    db.query(models.GradingCache).delete(synchronize_session=False)

    # Clear BTEC-specific models
    try:
        from app.btec_models import BTECReferenceSolution, BTECPreAssignmentValidation  # type: ignore
        db.query(BTECReferenceSolution).filter(
            BTECReferenceSolution.assignment_id == assignment_id
        ).delete(synchronize_session=False)
        db.query(BTECPreAssignmentValidation).filter(
            BTECPreAssignmentValidation.assignment_id == assignment_id
        ).delete(synchronize_session=False)
    except Exception:
        pass  # BTEC tables may not exist

    db.commit()
    return JSONResponse({
        "status": "success",
        "message": f"تم مسح دليل المهمة للمهمة: {assignment.title}"
    })


@app.delete("/api/admin/cache/solutions")
async def clear_solution_cache(request: Request, db: Session = Depends(get_db)):
    """Clear the SolutionCache table — forces regeneration of reference solutions."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    count = db.query(models.SolutionCache).count()
    db.query(models.SolutionCache).delete(synchronize_session=False)
    db.commit()
    return JSONResponse({
        "status": "success",
        "message": f"تم مسح كاش دليل المهمة بنجاح ({count} سجل)"
    })


@app.delete("/api/admin/cache/grading")
async def clear_grading_cache(request: Request, db: Session = Depends(get_db)):
    """Clear AI grading cache and invalidate submission replay (same-file fast reuse)."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    count = db.query(models.GradingCache).count()
    db.query(models.GradingCache).delete(synchronize_session=False)

    from app.submission_replay_cache import bump_replay_cache_generation

    replay_gen = bump_replay_cache_generation()
    db.commit()
    return JSONResponse({
        "status": "success",
        "message": (
            f"تم مسح كاش التصحيح ({count} سجل AI) "
            f"وتعطيل إعادة استخدام النتائج السابقة (جيل {replay_gen}). "
            "المرة القادمة سيُعاد التصحيح فعلياً حتى لنفس الملف."
        ),
    })


@app.delete("/api/admin/assignments/all")
async def clear_all_assignments(
    request: Request, db: Session = Depends(get_db)
):
    """Delete ALL assignments and all related data from the database."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    # Delete in correct order (children first, then parents)
    # 1. GradingResult (depends on submission + criteria)
    db.query(models.GradingResult).delete(synchronize_session=False)
    # 2. GradingSummary (depends on submission)
    db.query(models.GradingSummary).delete(synchronize_session=False)
    # 3. StudentReport (depends on submission)
    db.query(models.StudentReport).delete(synchronize_session=False)
    # 4. PlagiarismCheck (depends on submission + assignment)
    db.query(models.PlagiarismCheck).delete(synchronize_session=False)
    # 5. Submission (depends on assignment + batch)
    db.query(models.Submission).delete(synchronize_session=False)
    # 6. BatchGrading (depends on assignment)
    db.query(models.BatchGrading).delete(synchronize_session=False)
    # 7. GradingCriteria (depends on assignment)
    db.query(models.GradingCriteria).delete(synchronize_session=False)
    # 8. GradingCache
    db.query(models.GradingCache).delete(synchronize_session=False)
    # 9. BTEC models
    try:
        from app.btec_models import BTECReferenceSolution, BTECPreAssignmentValidation  # type: ignore
        db.query(BTECReferenceSolution).delete(synchronize_session=False)
        db.query(BTECPreAssignmentValidation).delete(synchronize_session=False)
    except Exception:
        pass
    # 10. Assignments themselves
    count = db.query(Assignment).delete(synchronize_session=False)

    db.commit()
    return JSONResponse({
        "status": "success",
        "message": f"تم حذف جميع المهام ({count} مهمة) وجميع البيانات المرتبطة بها"
    })


@app.get("/api/admin/submissions")
async def list_all_submissions(
    request: Request, db: Session = Depends(get_db)
):
    """List all student submissions with grading info for admin panel."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    submissions = (
        db.query(Submission)
        .order_by(Submission.created_at.desc())
        .all()
    )

    items = []
    for s in submissions:
        summary = (
            db.query(GradingSummary)
            .filter(GradingSummary.submission_id == s.id)
            .first()
        )
        report = (
            db.query(models.StudentReport)
            .filter(models.StudentReport.submission_id == s.id)
            .first()
        )
        assignment_title = s.assignment.title if s.assignment else "—"
        submitter = s.submitter
        submitter_name = "—"
        submitter_email = ""
        submitter_id = 0
        if submitter:
            submitter_id = submitter.id
            submitter_name = submitter.name or f"{submitter.first_name or ''} {submitter.last_name or ''}".strip() or submitter.email
            submitter_email = submitter.email or ""
        items.append({
            "id": s.id,
            "student_name": s.student_name,
            "assignment_title": assignment_title,
            "assignment_id": s.assignment_id,
            "status": s.status.value if s.status else "pending",
            "grade_level": summary.grade_level if summary else "—",
            "percentage": round(float(summary.percentage), 1) if summary and summary.percentage else 0,
            "ai_likelihood": summary.ai_likelihood if summary else 0,
            "plagiarism": round(float(summary.plagiarism_max_similarity), 1) if summary and summary.plagiarism_max_similarity else 0,
            "has_report": bool(report and report.report_file_path),
            "has_file": bool(s.submission_file_path),
            "created_at": s.created_at.strftime("%Y-%m-%d %H:%M") if s.created_at else "—",
            "submitter_id": submitter_id,
            "submitter_name": submitter_name,
            "submitter_email": submitter_email,
        })

    return JSONResponse({
        "status": "success",
        "total": len(items),
        "submissions": items
    })


@app.delete("/api/admin/submissions/all")
async def delete_all_submissions(
    request: Request, db: Session = Depends(get_db)
):
    """Delete ALL submissions and related grading data."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    db.query(models.GradingResult).delete(synchronize_session=False)
    db.query(models.GradingSummary).delete(synchronize_session=False)
    db.query(models.StudentReport).delete(synchronize_session=False)
    db.query(models.PlagiarismCheck).delete(synchronize_session=False)
    count = db.query(Submission).delete(synchronize_session=False)
    db.query(models.BatchGrading).delete(synchronize_session=False)
    db.query(models.GradingCache).delete(synchronize_session=False)
    db.commit()

    return JSONResponse({
        "status": "success",
        "message": f"تم حذف جميع واجبات الطلاب ({count} واجب) وجميع نتائج التصحيح"
    })


@app.delete("/api/admin/submissions/{submission_id}")
async def delete_single_submission(
    submission_id: int, request: Request, db: Session = Depends(get_db)
):
    """Delete a single submission and all related data."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    sub = db.query(Submission).filter(Submission.id == submission_id).first()
    if not sub:
        return JSONResponse({"status": "error", "message": "الواجب غير موجود"}, status_code=404)

    student_name = sub.student_name
    # Delete children first
    db.query(models.GradingResult).filter(models.GradingResult.submission_id == submission_id).delete(synchronize_session=False)
    db.query(models.GradingSummary).filter(models.GradingSummary.submission_id == submission_id).delete(synchronize_session=False)
    db.query(models.StudentReport).filter(models.StudentReport.submission_id == submission_id).delete(synchronize_session=False)
    db.query(models.PlagiarismCheck).filter(models.PlagiarismCheck.submission_id == submission_id).delete(synchronize_session=False)
    db.delete(sub)
    db.commit()

    return JSONResponse({
        "status": "success",
        "message": f"تم حذف واجب الطالب: {student_name}"
    })


# ═══════════════════════════════════════════════════════
# Activity Logs API
# ═══════════════════════════════════════════════════════


@app.get("/api/admin/logs")
async def get_activity_logs(request: Request, db: Session = Depends(get_db)):
    """Get activity logs for admin dashboard"""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user or user.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    from sqlalchemy import desc  # type: ignore

    # Get query parameters for filtering
    category = request.query_params.get("category", "")
    level = request.query_params.get("level", "")
    search = request.query_params.get("search", "")
    limit = min(int(request.query_params.get("limit", "200")), 1000)

    query = db.query(ActivityLog)

    if category:
        query = query.filter(ActivityLog.category == category)
    if level:
        query = query.filter(ActivityLog.level == level)
    if search:
        query = query.filter(ActivityLog.details.ilike(f"%{search}%"))

    logs = query.order_by(desc(ActivityLog.created_at)).limit(limit).all()

    # Category stats
    from sqlalchemy import func  # type: ignore
    stats = {r[0]: r[1] for r in db.query(ActivityLog.category, func.count(ActivityLog.id)).group_by(ActivityLog.category).all()}
    db_total = db.query(func.count(ActivityLog.id)).scalar() or 0

    logs_data = []
    for log in logs:
        logs_data.append({
            "id": log.id,
            "user_name": log.user_name or "",
            "user_email": log.user_email or "",
            "action": log.action,
            "category": log.category,
            "details": log.details or "",
            "ip_address": log.ip_address or "",
            "user_agent": log.user_agent or "",
            "level": log.level or "info",
            "created_at": log.created_at.strftime("%Y-%m-%d %H:%M:%S") if log.created_at else "",
        })

    return JSONResponse({
        "status": "success",
        "total": len(logs_data),
        "db_total": db_total,
        "stats": stats,
        "logs": logs_data,
    })


@app.delete("/api/admin/logs/all")
async def delete_all_activity_logs(request: Request, db: Session = Depends(get_db)):
    """Delete ALL activity logs from the database (admin only)."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"status": "error", "message": "غير مسجل"}, status_code=401)

    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        return JSONResponse({"status": "error", "message": "غير مصرح"}, status_code=403)

    from sqlalchemy import func  # type: ignore

    count = db.query(func.count(ActivityLog.id)).scalar() or 0
    db.query(ActivityLog).delete(synchronize_session=False)
    db.commit()

    log_activity(
        db,
        action="clear_logs",
        category="admin",
        details=f"مسح {count} سجل نشاط من قاعدة البيانات",
        user_id=admin.id,
        user_name=admin.name or f"{admin.first_name or ''} {admin.last_name or ''}".strip() or admin.email or "",
        user_email=admin.email or "",
        ip_address=_get_client_ip(request),
        user_agent=(request.headers.get("user-agent") or "")[:500],
        level="warning",
    )

    return JSONResponse({
        "status": "success",
        "message": f"تم مسح جميع سجلات النشاط ({count} سجل)",
        "deleted_count": count,
    })


@app.get("/smart-grader", response_class=HTMLResponse)
async def smart_grader_page(request: Request, db: Session = Depends(get_db)):
    """تصحيح الطلبة — عرض المهام الجاهزة للتصحيح المباشر"""
    from fastapi.responses import RedirectResponse  # type: ignore
    user = get_current_user(request, db)
    if not user:
        return RedirectResponse(url="/login", status_code=302)
    # Fetch assignments that have a reference solution (READY status) for current user
    from sqlalchemy import or_  # type: ignore

    _linked_ids = [
        row[0]
        for row in db.query(models.UserAssignmentLink.assignment_id)
        .filter(models.UserAssignmentLink.user_id == user.id)
        .all()
    ]
    if _linked_ids:
        ready_assignments = (
            db.query(Assignment)
            .filter(
                Assignment.status == AssignmentStatus.READY,
                or_(Assignment.created_by == user.id, Assignment.id.in_(_linked_ids)),
            )
            .order_by(Assignment.created_at.desc())
            .all()
        )
    else:
        ready_assignments = (
            db.query(Assignment)
            .filter(
                Assignment.status == AssignmentStatus.READY,
                Assignment.created_by == user.id,
            )
            .order_by(Assignment.created_at.desc())
            .all()
        )
    ready_assignments = _dedupe_assignments_by_content_hash(ready_assignments)
    assignments_data = []
    for a in ready_assignments:
        criteria_count = db.query(models.GradingCriteria).filter(
            models.GradingCriteria.assignment_id == a.id
        ).count()
        assignments_data.append({
            "id": a.id,
            "title": a.title,
            "unit_name": a.unit_name or "",
            "criteria_count": criteria_count,
            "created_at": a.created_at,
        })
    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "smart_grader.html",
        {
            "request": request,
            "user": user,
            "assignments": assignments_data,
            "subscription": sub_info,
        },
    )


@app.get("/create-assignment", response_class=HTMLResponse)
async def create_assignment_page(request: Request, db: Session = Depends(get_db)):
    """Create assignment page"""
    user = get_current_user(request, db)
    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "create_assignment.html",
        {
            "request": request,
            "user": user,
            "is_admin": bool(user and user.role == models.UserRole.ADMIN),
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "subscription": sub_info,
        },
    )


@app.get("/batch-grade/{assignment_id}", response_class=HTMLResponse)
async def batch_grade_page(
    request: Request, assignment_id: int, db: Session = Depends(get_db)
):
    """Batch grading page"""
    user = get_current_user(request, db)
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()

    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    # Extract dynamic criteria from database — strictly tied to THIS assignment.
    # Deduplicate by criteria_level (older grading runs may have inserted dupes)
    # and sort properly: P (Pass) → M (Merit) → D (Distinction), then by number,
    # while keeping any prefix grouping (A.*, B.*, C.* ...) together.
    db_criteria = db.query(models.GradingCriteria).filter(
        models.GradingCriteria.assignment_id == assignment_id
    ).all()

    _seen_levels = set()
    _unique_levels = []
    for c in db_criteria:
        lvl = (c.criteria_level or "").strip()
        if not lvl or lvl in _seen_levels:
            continue
        _seen_levels.add(lvl)
        _unique_levels.append(lvl)

    def _criteria_sort_key(level: str):
        import re as _re_local
        short = level.split(".")[-1] if "." in level else level
        type_order = {"P": 0, "M": 1, "D": 2}
        letter = short[:1].upper() if short else "Z"
        m = _re_local.search(r"\d+", short)
        num = int(m.group()) if m else 99
        prefix = level.split(".")[0] if "." in level else ""
        return (prefix, type_order.get(letter, 9), num)

    _unique_levels.sort(key=_criteria_sort_key)
    assignment_criteria = _unique_levels

    # Fallback if no criteria found
    if not assignment_criteria:
        assignment_criteria = ["P1", "P2", "M1", "D1"]

    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    is_admin = user_is_admin(user)

    return templates.TemplateResponse(
        "batch_grade.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "assignment": assignment,
            "assignment_criteria": assignment_criteria,
            "subscription": sub_info,
            "is_admin": is_admin,
            "intake_ignore_dirs": INTAKE_IGNORE_DIRS_LIST,
        },
        headers={"Cache-Control": "no-store"},
    )


@app.get("/batch-results/{batch_id}", response_class=HTMLResponse)
async def batch_results_page(
    request: Request,
    batch_id: int,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    """Batch results page"""
    import json as _json

    user = get_current_user(request, db)
    batch = db.query(BatchGrading).filter(BatchGrading.id == batch_id).first()

    if not batch and assignment_id:
        batch = (
            db.query(BatchGrading)
            .filter(BatchGrading.assignment_id == assignment_id)
            .order_by(BatchGrading.id.desc())
            .first()
        )
        if batch:
            return RedirectResponse(
                url=f"/batch-results/{batch.id}",
                status_code=302,
            )

    if not batch:
        return templates.TemplateResponse(
            "batch_not_found.html",
            {
                "request": request,
                "user": user,
                "batch_id": batch_id,
                "assignment_id": assignment_id or 1,
            },
            status_code=404,
        )

    submissions = db.query(Submission).filter(Submission.batch_id == batch_id).all()

    results_data = []
    for submission in submissions:
        summary = (
            db.query(GradingSummary)
            .filter(GradingSummary.submission_id == submission.id)
            .first()
        )
        report = (
            db.query(StudentReport)
            .filter(StudentReport.submission_id == submission.id)
            .first()
        )
        failure_error = None
        explainability = None
        pearson_criteria_rows = None
        pearson_engine_summary = None
        grade_display_metrics = None
        official_grade = None
        if submission.grading_snapshot_json:
            try:
                snap = _json.loads(str(submission.grading_snapshot_json))
                try:
                    from app.visual_evidence_registry import apply_game_criteria_pro_gate

                    apply_game_criteria_pro_gate(snap, grading_mode=snap.get("grading_mode"))
                except Exception:
                    pass
                from app.explainability_migration import extract_explainability_for_ui
                from app.official_grade import resolve_official_grade

                official = resolve_official_grade(
                    snap,
                    reapply_pipeline=True,
                    legacy_grade_level=summary.grade_level if summary else None,
                )
                if official.reapply_change_count > 0:
                    try:
                        from app.criteria_result_finalizer import sync_criteria_results_to_db

                        submission.grading_snapshot_json = _json.dumps(  # type: ignore
                            snap, ensure_ascii=False
                        )
                        sync_criteria_results_to_db(db, submission.id, snap)
                        db.commit()
                    except Exception as _persist_err:
                        print(
                            f"⚠️ [OFFICIAL-GRADE] persist batch_results "
                            f"submission {submission.id}: {_persist_err}"
                        )
                official_grade = official.to_dict()
                grade_display_metrics = official.grade_display_metrics
                explainability = extract_explainability_for_ui(snap)
                pearson_criteria_rows = None
                pearson_engine_summary = None
                try:
                    from app.grading_mode_policy import is_fast_grading_mode

                    if not is_fast_grading_mode(snap.get("grading_mode")):
                        from app.pro_btec_pearson import build_criteria_breakdown_for_ui

                        pearson_criteria_rows = build_criteria_breakdown_for_ui(snap)
                        pkg = snap.get("pearson_btec_pro") or {}
                        pearson_engine_summary = pkg.get("engine_runtime_summary")
                except Exception:
                    pearson_criteria_rows = None
                    pearson_engine_summary = None
            except Exception as _exp_err:
                print(f"⚠️ [EXPLAINABILITY] batch_results submission {submission.id}: {_exp_err}")
                explainability = None
                grade_display_metrics = None
        else:
            grade_display_metrics = None
        if (
            submission.status == SubmissionStatus.FAILED
            and submission.grading_snapshot_json
        ):
            try:
                snap = _json.loads(str(submission.grading_snapshot_json))
                failure_error = snap.get("error")
            except Exception:
                failure_error = None

        # VALIDATION: Log if summary is missing or has invalid data
        if not summary:
            print(
                f" [DISPLAY WARNING] No GradingSummary found for submission {submission.id} ({submission.student_name})"
            )
        elif summary.percentage == 0 and summary.total_score == 0:  # type: ignore
            print(
                f" [DISPLAY WARNING] GradingSummary for {submission.student_name} has all zeros!"
            )
            print(f"   - DB Summary ID: {summary.id}")
            print(f"   - total_score: {summary.total_score}")
            print(f"   - max_score: {summary.max_score}")
            print(f"   - percentage: {summary.percentage}")
            print(f"   - grade_level: '{summary.grade_level}'")
            print(f"   - ai_likelihood: {summary.ai_likelihood}")
        else:
            print(
                f" [DISPLAY OK] {submission.student_name}: {summary.percentage}% - {summary.grade_level}"
            )

        results_data.append(
            {
                "submission": submission,
                "summary": summary,
                "report": report,
                "failure_error": failure_error,
                "explainability": explainability,
                "grade_display_metrics": grade_display_metrics,
                "official_grade": official_grade,
                "pearson_criteria_rows": pearson_criteria_rows
                if submission.grading_snapshot_json
                else None,
                "pearson_engine_summary": pearson_engine_summary
                if submission.grading_snapshot_json
                else None,
            }
        )

    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None

    try:
        from app.report_generator import regenerate_batch_summary_pdf

        regenerate_batch_summary_pdf(db, batch_id)
    except Exception as _pdf_err:
        print(f"⚠️ [BATCH SUMMARY PDF] regenerate failed for batch {batch_id}: {_pdf_err}")

    return templates.TemplateResponse(
        "batch_results.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "batch": batch,
            "results": results_data,
            "assignment_id": batch.assignment_id,
            "subscription": sub_info,
        },
    )


@app.get("/model-answers", response_class=HTMLResponse)
async def model_answers_page(request: Request, db: Session = Depends(get_db)):
    """Model answers page listing all reference solutions - Requires subscription"""
    from fastapi.responses import RedirectResponse  # type: ignore

    user_id = get_current_user_id(request)
    if not user_id:
        return RedirectResponse(url="/login", status_code=302)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    is_admin = user and user.role == models.UserRole.ADMIN
    if not is_admin:
        sub = get_active_subscription(db, user_id)
        if not sub:
            packages = (
                db.query(models.Package)
                .filter(models.Package.is_active.is_(True))
                .all()
            )
            from app.database import get_package_rows, get_student_package_rows  # type: ignore
            sub_info = get_subscription_info(db, user_id)
            return templates.TemplateResponse(
                "home.html",
                {
                    "request": request,
                    "app_title": os.getenv(
                        "APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"
                    ),
                    "user": user,
                    "packages": packages,
                    "package_rows": get_package_rows(db),
                    "student_package_rows": get_student_package_rows(db),
                    "subscription": sub_info,
                    "error_message": "يجب الاشتراك للوصول إلى خدمة خطوات حل المهمة. يرجى الاشتراك في إحدى الباقات.",
                },
            )

    # Fetch all ready assignments that have a reference solution (legacy or BTEC)
    from sqlalchemy import or_  # type: ignore
    from app.btec_models import BTECReferenceSolution  # type: ignore

    _linked_ma_ids = [
        row[0]
        for row in db.query(models.UserAssignmentLink.assignment_id)
        .filter(models.UserAssignmentLink.user_id == user_id)
        .all()
    ]
    _ref_ok = (Assignment.reference_solution_json.isnot(None)) | (
        BTECReferenceSolution.id.isnot(None)
    )
    if _linked_ma_ids:
        assignments = (
            db.query(Assignment)
            .outerjoin(BTECReferenceSolution, Assignment.id == BTECReferenceSolution.assignment_id)
            .filter(
                or_(Assignment.created_by == user_id, Assignment.id.in_(_linked_ma_ids)),
                Assignment.status == AssignmentStatus.READY,
                _ref_ok,
            )
            .order_by(Assignment.created_at.desc())
            .all()
        )
    else:
        assignments = (
            db.query(Assignment)
            .outerjoin(BTECReferenceSolution, Assignment.id == BTECReferenceSolution.assignment_id)
            .filter(
                Assignment.created_by == user_id,
                Assignment.status == AssignmentStatus.READY,
                _ref_ok,
            )
            .order_by(Assignment.created_at.desc())
            .all()
        )
    assignments = _dedupe_assignments_by_content_hash(assignments)

    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "model_answers.html",
        {
            "request": request,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "user": user,
            "assignments": assignments,
            "subscription": sub_info,
        },
    )


@app.get("/grade", response_class=HTMLResponse)
async def grade_page(request: Request, db: Session = Depends(get_db)):
    """Single grading page"""
    user = get_current_user(request, db)
    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "grade.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "subscription": sub_info,
        },
    )


@app.get("/results/{submission_id}", response_class=HTMLResponse)
async def results_page(
    request: Request, submission_id: int, db: Session = Depends(get_db)
):
    """Single submission results page"""
    submission = db.query(Submission).filter(Submission.id == submission_id).first()

    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    summary = (
        db.query(GradingSummary)
        .filter(GradingSummary.submission_id == submission.id)
        .first()
    )
    results = (
        db.query(GradingResult)
        .filter(GradingResult.submission_id == submission.id)
        .all()
    )

    # Parse JSON fields for template rendering
    verdict_by_level = {}
    visual_by_level = {}
    awardable_by_level = {}
    requires_pro_by_level = {}
    btec_institutional_award = None
    if getattr(submission, "grading_snapshot_json", None):
        try:
            _vs = json.loads(str(submission.grading_snapshot_json))
            try:
                from app.visual_evidence_registry import apply_game_criteria_pro_gate

                apply_game_criteria_pro_gate(_vs, grading_mode=_vs.get("grading_mode"))
            except Exception:
                pass
            btec_institutional_award = _vs.get("btec_institutional_award")
            for cr in _vs.get("criteria_results") or []:
                if isinstance(cr, dict) and cr.get("criteria_level"):
                    lvl = str(cr["criteria_level"])
                    verdict_by_level[lvl] = cr.get("verdict_status")
                    if cr.get("requires_pro"):
                        requires_pro_by_level[lvl] = True
                    awardable_by_level[lvl] = {
                        "awardable": cr.get("awardable"),
                        "award_block_reason": cr.get("award_block_reason"),
                        "award_block_reason_ar": cr.get("award_block_reason_ar"),
                    }
                    reg = cr.get("evidence_registry") or {}
                    vis = cr.get("visual_evidence") or reg.get("visual_evidence")
                    if isinstance(vis, dict):
                        visual_by_level[lvl] = vis
        except (json.JSONDecodeError, TypeError):
            verdict_by_level = {}
            visual_by_level = {}
            awardable_by_level = {}

    from app.btec_criteria_governance import (
        ensure_clean_grading_result_feedback,
        teacher_facing_feedback,
    )

    snapshot_by_level: Dict[str, Dict[str, Any]] = {}
    _snap_dirty = False
    official_grade = None
    grade_display_metrics = None
    evidence_map: List[Dict[str, Any]] = []
    evidence_map_summary: Dict[str, Any] = {}
    if getattr(submission, "grading_snapshot_json", None):
        try:
            from app.btec_criteria_governance import ensure_clean_grading_result_feedback
            from app.official_grade import resolve_official_grade

            _snap_fix = json.loads(str(submission.grading_snapshot_json))
            _snap_before = json.dumps(
                _snap_fix.get("criteria_results") or [], ensure_ascii=False
            )
            official = resolve_official_grade(
                _snap_fix,
                reapply_pipeline=True,
                legacy_grade_level=summary.grade_level if summary else None,
            )
            ensure_clean_grading_result_feedback(_snap_fix)
            _snap_after = json.dumps(
                _snap_fix.get("criteria_results") or [], ensure_ascii=False
            )
            _snap_dirty = (
                _snap_before != _snap_after or official.reapply_change_count > 0
            )
            official_grade = official.to_dict()
            grade_display_metrics = official.grade_display_metrics
            try:
                from app.evidence_map import build_evidence_map, build_evidence_map_summary

                evidence_map = build_evidence_map(_snap_fix)
                evidence_map_summary = build_evidence_map_summary(evidence_map)
            except Exception as _emap_err:
                print(f"⚠️ [EVIDENCE-MAP] submission {submission.id}: {_emap_err}")
            if _snap_dirty:
                from app.criteria_result_finalizer import sync_criteria_results_to_db

                submission.grading_snapshot_json = json.dumps(  # type: ignore
                    _snap_fix, ensure_ascii=False
                )
                sync_criteria_results_to_db(db, submission.id, _snap_fix)
                db.commit()
            for cr in _snap_fix.get("criteria_results") or []:
                if isinstance(cr, dict) and cr.get("criteria_level"):
                    snapshot_by_level[str(cr["criteria_level"])] = cr
        except Exception:
            snapshot_by_level = {}

    parsed_results = []
    for r in results:
        criteria_level = ""
        criteria_name = ""
        if r.criteria:
            criteria_level = r.criteria.criteria_level or ""
            criteria_name = r.criteria.criteria_name or ""
        snap_cr = snapshot_by_level.get(criteria_level) or {}
        parsed = {
            "achieved": bool(snap_cr.get("achieved")) if snap_cr else r.achieved,
            "score": int(snap_cr.get("score") if snap_cr else r.score or 0),
            "max_score": r.max_score,
            "feedback": teacher_facing_feedback(
                snap_cr.get("feedback") or r.feedback or ""
            ),
            "criteria_level": criteria_level,
            "criteria_name": criteria_name,
            "missing_points": [],
            "verdict_status": snap_cr.get("verdict_status")
            or verdict_by_level.get(criteria_level),
            "requires_pro": requires_pro_by_level.get(criteria_level, False),
            "visual_evidence": visual_by_level.get(criteria_level),
        }
        _aw = awardable_by_level.get(criteria_level) or {}
        if _aw.get("awardable") is not None:
            parsed["awardable"] = _aw.get("awardable")
            parsed["award_block_reason"] = _aw.get("award_block_reason")
            parsed["award_block_reason_ar"] = _aw.get("award_block_reason_ar")
        elif parsed["achieved"]:
            parsed["awardable"] = True
        else:
            parsed["awardable"] = False
        if r.missing_points:
            try:
                pts = json.loads(str(r.missing_points)) if isinstance(r.missing_points, str) else r.missing_points
                parsed["missing_points"] = pts if isinstance(pts, list) else []
            except (json.JSONDecodeError, TypeError):
                parsed["missing_points"] = []
        parsed_results.append(parsed)

    # Sort criteria P → M → D
    def _crit_sort_key(item):
        lv = item.get('criteria_level', '')
        short = lv.split('.')[-1] if '.' in lv else lv
        _order = {'P': 0, 'M': 1, 'D': 2}
        letter = short[0].upper() if short else 'Z'
        try:
            num = int(short[1:]) if len(short) > 1 else 0
        except ValueError:
            num = 99
        return (_order.get(letter, 9), num)

    parsed_results.sort(key=_crit_sort_key)

    # Parse JSON string fields in summary for template rendering
    if summary:
        for field in ("strengths", "improvements"):
            raw = getattr(summary, field, None)
            if raw and isinstance(raw, str):
                try:
                    parsed_list = json.loads(raw)
                    if isinstance(parsed_list, list):
                        setattr(summary, field, parsed_list)
                    else:
                        setattr(summary, field, [])
                except (json.JSONDecodeError, TypeError):
                    setattr(summary, field, [])

    user = get_current_user(request, db)
    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None

    l5_playtest_status = None
    runtime_db_sync = None
    explainability = None
    if getattr(submission, "grading_snapshot_json", None):
        try:
            _snap = json.loads(str(submission.grading_snapshot_json))
            l5_playtest_status = (_snap.get("l5_human_playtest") or {}).get("status")
            runtime_db_sync = _snap.get("runtime_adjudication_db_sync")
            from app.explainability_migration import extract_explainability_for_ui

            explainability = extract_explainability_for_ui(_snap)
        except (json.JSONDecodeError, TypeError):
            pass

    if grade_display_metrics is None and getattr(submission, "grading_snapshot_json", None):
        try:
            from app.official_grade import resolve_official_grade

            _snap_gdm = json.loads(str(submission.grading_snapshot_json))
            official = resolve_official_grade(
                _snap_gdm,
                reapply_pipeline=False,
                legacy_grade_level=summary.grade_level if summary else None,
            )
            official_grade = official.to_dict()
            grade_display_metrics = official.grade_display_metrics
        except (json.JSONDecodeError, TypeError):
            grade_display_metrics = None

    return templates.TemplateResponse(
        "results.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "submission": submission,
            "summary": summary,
            "results": parsed_results,
            "percentage": summary.percentage if summary else 0,
            "grade_display_metrics": grade_display_metrics,
            "official_grade": official_grade,
            "evidence_map": evidence_map,
            "evidence_map_summary": evidence_map_summary,
            "subscription": sub_info,
            "has_governance_replay": bool(getattr(submission, "grading_snapshot_json", None)),
            "has_runtime_replay": _submission_has_runtime_replay(submission),
            "has_timeline_replay": bool(getattr(submission, "grading_snapshot_json", None)),
            "l5_playtest_status": l5_playtest_status,
            "runtime_db_sync": runtime_db_sync,
            "explainability": explainability,
            "btec_institutional_award": btec_institutional_award,
        },
    )


def _submission_has_runtime_replay(submission) -> bool:
    """True when grading snapshot or artifact inventory includes L4 runtime evidence."""
    grading_snapshot = None
    if getattr(submission, "grading_snapshot_json", None):
        try:
            grading_snapshot = json.loads(str(submission.grading_snapshot_json))
        except (json.JSONDecodeError, TypeError):
            grading_snapshot = None
    from app.runtime_replay_viewer import build_runtime_replay

    replay = build_runtime_replay(
        grading_snapshot,
        student_name=getattr(submission, "student_name", "") or "",
        batch_id=getattr(submission, "batch_id", None),
        submission_id=getattr(submission, "id", None),
    )
    return bool(replay.get("available"))


@app.get("/authority-replay/{submission_id}", response_class=HTMLResponse)
async def authority_replay_page(
    request: Request, submission_id: int, db: Session = Depends(get_db)
):
    """Authority Replay Viewer — temporal provenance chain (read-only)."""
    from app.authority_replay import build_authority_replay

    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    grading_snapshot = None
    if getattr(submission, "grading_snapshot_json", None):
        try:
            grading_snapshot = json.loads(str(submission.grading_snapshot_json))
        except (json.JSONDecodeError, TypeError):
            grading_snapshot = None

    replay = build_authority_replay(grading_snapshot)

    drift = None
    if grading_snapshot and grading_snapshot.get("governance_drift"):
        drift = grading_snapshot.get("governance_drift")
    else:
        try:
            from app.governance_drift_monitor import analyze_submission_governance_drift

            drift = analyze_submission_governance_drift(grading_snapshot)
        except Exception:
            drift = None

    user = get_current_user(request, db)
    user_id = get_current_user_id(request)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "authority_replay.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "submission": submission,
            "replay": replay,
            "drift": drift,
            "subscription": sub_info,
        },
    )


@app.post("/api/admin/explainability-backfill/preview")
async def explainability_backfill_preview_api(
    request: Request,
    db: Session = Depends(get_db),
):
    """
    Backfill impact preview — submissions / explainability missing / integrity risk.
    Body JSON: { "submission_id": N } or { "batch_id": N }, optional force.
    """
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="غير مسجل")
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="غير مصرح")

    try:
        body = await request.json()
    except Exception:
        body = {}

    force = bool(body.get("force"))
    submission_id = body.get("submission_id")
    batch_id = body.get("batch_id")

    from app.explainability_migration import preview_backfill_batch, preview_submission_backfill

    if submission_id is not None:
        sub = db.query(Submission).filter(Submission.id == int(submission_id)).first()
        if not sub:
            raise HTTPException(status_code=404, detail="Submission not found")
        row = preview_submission_backfill(sub, force=force)
        return {
            "ok": True,
            "preview": {
                "submissions": 1,
                "explainability_missing": 1 if row.get("explainability_missing") else 0,
                "would_apply": 1 if row.get("would_apply") else 0,
                "would_skip": 1 if row.get("would_skip") else 0,
                "integrity_risk": row.get("integrity_risk") or "medium",
                "rows": [row],
            },
        }

    if batch_id is not None:
        preview = preview_backfill_batch(db, int(batch_id), force=force)
        return {"ok": True, "preview": preview}

    raise HTTPException(status_code=400, detail="Provide submission_id or batch_id")


@app.post("/api/admin/explainability-backfill")
async def explainability_backfill_api(
    request: Request,
    db: Session = Depends(get_db),
):
    """
    Non-destructive explainability migration — adds diagnostics without regrading.
    Body JSON: { "submission_id": N } or { "batch_id": N }, optional dry_run, force.
    """
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="غير مسجل")
    admin = db.query(models.User).filter(models.User.id == user_id).first()
    if admin is None or admin.role != models.UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="غير مصرح")

    try:
        body = await request.json()
    except Exception:
        body = {}

    dry_run = bool(body.get("dry_run"))
    force = bool(body.get("force"))
    submission_id = body.get("submission_id")
    batch_id = body.get("batch_id")
    generated_by = str(getattr(admin, "email", None) or f"admin:{user_id}")

    from app.explainability_migration import (
        backfill_batch_submissions,
        backfill_submission_record,
    )

    if submission_id is not None:
        sub = db.query(Submission).filter(Submission.id == int(submission_id)).first()
        if not sub:
            raise HTTPException(status_code=404, detail="Submission not found")
        report = backfill_submission_record(
            sub,
            db=db,
            dry_run=dry_run,
            force=force,
            rerun_runtime=bool(body.get("rerun_runtime")),
            generated_by=generated_by,
        )
        if report.get("applied") and not dry_run:
            db.commit()
        return {"ok": True, "report": report}

    if batch_id is not None:
        summary = backfill_batch_submissions(
            db,
            int(batch_id),
            dry_run=dry_run,
            force=force,
            generated_by=generated_by,
        )
        if summary.get("applied") and not dry_run:
            db.commit()
        return {"ok": True, "summary": summary}

    raise HTTPException(status_code=400, detail="Provide submission_id or batch_id")


@app.get("/api/governance-drift/{submission_id}")
async def governance_drift_submission_api(submission_id: int, db: Session = Depends(get_db)):
    """Per-submission semantic drift vs GOVERNANCE_FREEZE_v1."""
    from app.governance_drift_monitor import analyze_submission_governance_drift

    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    grading_snapshot = None
    if getattr(submission, "grading_snapshot_json", None):
        try:
            grading_snapshot = json.loads(str(submission.grading_snapshot_json))
        except (json.JSONDecodeError, TypeError):
            grading_snapshot = None
    if not grading_snapshot:
        return {
            "freeze_id": "GOVERNANCE_FREEZE_v1",
            "status": "no_snapshot",
            "summary_ar": "لا grading_snapshot — أعد التصحيح لتفعيل drift monitor.",
        }
    return analyze_submission_governance_drift(grading_snapshot)


@app.get("/api/governance-drift/batch/{batch_id}")
async def governance_drift_batch_api(batch_id: int, db: Session = Depends(get_db)):
    """Cohort governance metrics for human pilot instrumentation."""
    from app.governance_drift_monitor import analyze_cohort_governance_metrics

    subs = (
        db.query(Submission)
        .filter(Submission.batch_id == batch_id)
        .all()
    )
    snapshots: List[Dict[str, Any]] = []
    for sub in subs:
        if not getattr(sub, "grading_snapshot_json", None):
            continue
        try:
            snapshots.append(json.loads(str(sub.grading_snapshot_json)))
        except (json.JSONDecodeError, TypeError):
            continue
    return analyze_cohort_governance_metrics(
        snapshots, cohort_id=f"batch_{batch_id}"
    )


@app.post("/api/governance-workshop/incident")
async def governance_workshop_incident_api(request: Request):
    """Classify manual pilot/workshop governance incident (observatory log)."""
    from app.governance_failure_taxonomy import classify_workshop_incident

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    incident = classify_workshop_incident(
        incident_type=str(data.get("incident_type") or "reviewer_confusion"),
        notes_ar=str(data.get("notes_ar") or ""),
        submission_id=data.get("submission_id"),
        reviewer_confused_l3=bool(data.get("reviewer_confused_l3")),
        trust_eroded=bool(data.get("trust_eroded")),
    )
    # Append to workshop log file (best-effort persistence)
    try:
        log_dir = Path("app/calibration/human_cohort_workshop")
        log_dir.mkdir(parents=True, exist_ok=True)
        log_path = log_dir / "incidents.jsonl"
        record = {**incident, "logged_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")}
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
        incident["log_path"] = str(log_path)
    except OSError:
        pass
    return incident


@app.post("/api/governance-mitigation/outcome")
async def governance_mitigation_outcome_api(request: Request):
    """Record mitigation effectiveness outcome (workshop / verifier)."""
    from app.governance_mitigation_memory import record_mitigation_outcome

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    result = record_mitigation_outcome(
        mitigation_id=str(data.get("mitigation_id") or ""),
        outcome=str(data.get("outcome") or "unknown"),
        notes_ar=str(data.get("notes_ar") or ""),
        recurrence_detected=bool(data.get("recurrence_detected")),
        reviewer_id=data.get("reviewer_id"),
    )
    if result.get("error"):
        raise HTTPException(status_code=404, detail=result.get("error"))
    return result


@app.get("/api/governance-mitigation/summary")
async def governance_mitigation_summary_api(
    batch_id: Optional[int] = None,
    failure_mode_id: Optional[str] = None,
):
    """Mitigation effectiveness learning summary."""
    from app.governance_mitigation_memory import analyze_mitigation_effectiveness

    cohort_id = f"batch_{batch_id}" if batch_id is not None else None
    return analyze_mitigation_effectiveness(
        cohort_id=cohort_id,
        batch_id=batch_id,
        failure_mode_id=failure_mode_id,
    )


@app.get("/governance-pilot", response_class=HTMLResponse)
async def governance_pilot_start_page(request: Request, db: Session = Depends(get_db)):
    """Pilot pool analysis + links to run observatory on a controlled batch."""
    from app.governance_pilot_observatory import scan_pilot_pool

    subs = (
        db.query(Submission)
        .filter(Submission.grading_snapshot_json.isnot(None))
        .order_by(Submission.id.desc())
        .all()
    )
    pool = scan_pilot_pool(subs)
    return templates.TemplateResponse(
        "governance_pilot_start.html",
        {"request": request, "pool": pool},
    )


@app.get("/api/governance-pilot/pool")
async def governance_pilot_pool_api(db: Session = Depends(get_db)):
    """JSON pilot pool mix analysis for cohort planning."""
    from app.governance_pilot_observatory import scan_pilot_pool

    subs = (
        db.query(Submission)
        .filter(Submission.grading_snapshot_json.isnot(None))
        .order_by(Submission.id.desc())
        .all()
    )
    return scan_pilot_pool(subs)


@app.get("/governance-pilot/batch/{batch_id}", response_class=HTMLResponse)
async def governance_pilot_observatory_page(
    batch_id: int,
    request: Request,
    db: Session = Depends(get_db),
):
    """Human governance observation worksheet UI — not grading."""
    from app.epistemic_leakage_lexicon import LEXICON_ENTRIES
    from app.facilitator_epistemic_worksheet import EPISTEMIC_QUESTIONS
    from app.governance_drift_monitor import analyze_submission_governance_drift
    from app.governance_pilot_observatory import (
        WORKSHEET_VERSION,
        build_worksheet_draft,
        load_observations,
        synthesize_cohort_governance_report,
    )
    from app.phase2_institutional_observation import get_phase2_cohort_state

    subs = db.query(Submission).filter(Submission.batch_id == batch_id).all()
    snapshots: List[Dict[str, Any]] = []
    items: List[Dict[str, Any]] = []
    observed_ids = {o.get("submission_id") for o in load_observations(batch_id=batch_id)}
    observation_count = len(observed_ids)

    for sub in subs:
        snap: Optional[Dict[str, Any]] = None
        if getattr(sub, "grading_snapshot_json", None):
            try:
                snap = json.loads(str(sub.grading_snapshot_json))
                snapshots.append(snap)
            except (json.JSONDecodeError, TypeError):
                snap = None
        drift = analyze_submission_governance_drift(snap) if snap else None
        draft = build_worksheet_draft(
            submission_id=sub.id,
            student_name=getattr(sub, "student_name", "") or "",
            batch_id=batch_id,
            snapshot=snap,
            drift=drift,
        )
        items.append({
            "submission_id": sub.id,
            "student_name": getattr(sub, "student_name", "") or "",
            "drift_status": (drift or {}).get("status"),
            "section_a": draft.get("section_a_runtime_evidence_state") or {},
            "section_f": draft.get("section_f_epistemic_trace") or {},
            "suggested_events": draft.get("section_c_governance_events") or [],
            "already_observed": sub.id in observed_ids,
        })

    def _workshop_sort_key(item: Dict[str, Any]) -> tuple:
        """Unobserved L3/contradictory cases first — balanced Phase 2 workshop."""
        observed = 1 if item.get("already_observed") else 0
        section_a = item.get("section_a") or {}
        level = int(section_a.get("runtime_level_numeric") or 0)
        contra = len(section_a.get("contradiction_flags") or [])
        exe = 1 if section_a.get("executable_detected") else 0
        return (observed, -level, -contra, -exe, item.get("submission_id") or 0)

    items.sort(key=_workshop_sort_key)

    synthesis = None
    if snapshots:
        from app.canonical_stability_metrics import (
            attach_to_cohort_synthesis,
            compute_canonical_stability_metrics,
        )
        from app.canonical_stability_trajectory import build_governance_trajectory_report

        synthesis = synthesize_cohort_governance_report(
            batch_id=batch_id,
            snapshots=snapshots,
        )
        trajectory = build_governance_trajectory_report(
            db, batch_id=batch_id, record_snapshot=False
        )
        stability = trajectory.get("current") or compute_canonical_stability_metrics(
            db, batch_id=batch_id
        )
        synthesis = attach_to_cohort_synthesis(synthesis, stability)
        synthesis["governance_trajectory"] = {
            "transitions": trajectory.get("stability_transitions") or [],
            "interpretation_ar": trajectory.get("trajectory_interpretation_ar") or [],
            "freeze_epoch": trajectory.get("freeze_epoch"),
            "history_points": trajectory.get("history_points", 0),
        }
        from app.governance_epoch_narrative import build_epoch_narrative
        from app.canonical_stability_trajectory import ACTIVE_FREEZE_EPOCH
        from app.governance_artifact_chain import build_governance_artifact_chain

        epoch_narrative = build_epoch_narrative(
            db, epoch_id=ACTIVE_FREEZE_EPOCH, assignment_id=None
        )
        artifact_chain = build_governance_artifact_chain(
            db, epoch_id=ACTIVE_FREEZE_EPOCH, assignment_id=None
        )
        synthesis["epoch_narrative"] = {
            "epoch_id": epoch_narrative.get("epoch_id"),
            "freeze_id": epoch_narrative.get("freeze_id"),
            "narrative_ar": epoch_narrative.get("narrative_ar") or [],
            "replay_trust": epoch_narrative.get("replay_trust_state"),
        }
        synthesis["artifact_chain_summary"] = {
            "node_count": artifact_chain.get("node_count"),
            "who_approved": (artifact_chain.get("provenance_answers") or {}).get("who_approved"),
            "design_principle_ar": artifact_chain.get("design_principle_ar"),
        }

    return templates.TemplateResponse(
        "governance_pilot_observatory.html",
        {
            "request": request,
            "batch_id": batch_id,
            "worksheet_version": WORKSHEET_VERSION,
            "submissions": items,
            "synthesis": synthesis,
            "canonical_stability": synthesis.get("canonical_stability_metrics") if synthesis else None,
            "governance_trajectory": synthesis.get("governance_trajectory") if synthesis else None,
            "epoch_narrative": synthesis.get("epoch_narrative") if synthesis else None,
            "epistemic_questions": EPISTEMIC_QUESTIONS,
            "epistemic_evidence": (
                synthesis.get("epistemic_behavioural_evidence") if synthesis else None
            ),
            "leakage_lexicon": LEXICON_ENTRIES,
            "phase2_state": get_phase2_cohort_state(batch_id, observation_count=observation_count),
            "observation_count": observation_count,
            "epistemic_trace_invariant_ar": (
                "هذه الواجهة تسجّل انتقالات الحالة المعرفية. لا تمنح سلطة."
            ),
            "epistemic_trace_invariant_en": (
                "This interface records epistemic state transitions. "
                "It does not assign authority."
            ),
            "vocabulary_hint_invariant_en": (
                "Human confirmation required. "
                "This hint does not imply authority escalation."
            ),
            "vocabulary_hint_invariant_ar": (
                "التأكيد البشري مطلوب. "
                "هذا التلميح لا يعني تصعيدًا للسلطة."
            ),
            "escalation_notice_invariant_en": (
                "The system may notice epistemic escalation. "
                "It may not silently conclude it."
            ),
            "escalation_notice_invariant_ar": (
                "النظام قد يلاحظ تصعيدًا معرفيًا. "
                "لا يجوز أن يستنتجه بصمت."
            ),
            "observability_invariant_en": (
                "Observability may illuminate authority formation. "
                "It may not silently inherit authority formation."
            ),
            "observability_invariant_ar": (
                "الرصد قد يُضيء تشكّل السلطة. "
                "لا يجوز أن يرث تشكّل السلطة بصمت."
            ),
        },
    )


@app.post("/api/governance-pilot/possible-vocabulary-escalation-hint")
async def governance_pilot_vocabulary_hint_api(request: Request):
    """Advisory phrase candidates — limited sources, never auto-populates lexicon field."""
    from app.possible_vocabulary_escalation_hint import (
        compute_possible_vocabulary_escalation_hint,
        validate_hint_request,
    )

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    sources, err = validate_hint_request(body)
    if err:
        raise HTTPException(status_code=400, detail=err)
    return compute_possible_vocabulary_escalation_hint(**sources)


@app.post("/api/governance-pilot/observer-role-drift-check")
async def governance_pilot_observer_drift_check_api(request: Request):
    """Record observer posture review — not grading, not compliance."""
    from app.observer_role_drift_check import append_drift_check, normalize_drift_responses

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    question_set = body.get("question_set") or "baseline"
    try:
        record = normalize_drift_responses(
            question_set=question_set,
            responses=body.get("responses") or [],
            submission_id=body.get("submission_id"),
            batch_id=body.get("batch_id"),
            session_label=body.get("session_label") or "",
            facilitator_notes_ar=body.get("facilitator_notes_ar") or "",
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    check_id = append_drift_check(record)
    return {"ok": True, "check_id": check_id, "assigns_authority": False}


@app.post("/api/governance-pilot/post-session-constitutional-review")
async def governance_pilot_post_session_review_api(request: Request):
    """Record post-session constitutional review — not grading, not feature planning."""
    from app.post_session_constitutional_review import (
        append_post_session_review,
        normalize_review,
    )

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    cohort_id = body.get("cohort_id") or ""
    session_label = body.get("session_label") or ""
    if not cohort_id or not session_label:
        raise HTTPException(status_code=400, detail="cohort_id and session_label required")
    record = normalize_review(
        cohort_id=cohort_id,
        session_label=session_label,
        facilitator_notes_ar=body.get("facilitator_notes_ar") or "",
        answers=body.get("answers") or [],
        organizing_cognition_detected=body.get("organizing_cognition_detected"),
        submission_ids=body.get("submission_ids") or [],
    )
    review_id = append_post_session_review(record)
    return {"ok": True, "review_id": review_id, "assigns_authority": False}


@app.post("/api/governance-pilot/runtime-observation-ledger")
async def governance_pilot_runtime_observation_ledger_api(request: Request):
    """Append runtime observation capture — observability only, not legitimacy."""
    from app.runtime_observation_ledger import append_runtime_observation, normalize_observation

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    observation_mode = body.get("observation_mode") or ""
    if not observation_mode:
        raise HTTPException(status_code=400, detail="observation_mode required")
    try:
        record = normalize_observation(
            observation_mode=observation_mode,
            artifact_name=body.get("artifact_name") or "",
            artifact_type=body.get("artifact_type") or "",
            submission_id=body.get("submission_id"),
            batch_id=body.get("batch_id"),
            session_label=body.get("session_label") or "",
            telemetry_graph=body.get("telemetry_graph"),
            provenance_chain=body.get("provenance_chain"),
            sandbox_result=body.get("sandbox_result"),
            phenomenology=body.get("phenomenology"),
            replay_phenomenology=body.get("replay_phenomenology"),
            telemetry_replay_wiring=body.get("telemetry_replay_wiring"),
            runtime_epistemic_governance=body.get("runtime_epistemic_governance"),
            facilitator_notes_ar=body.get("facilitator_notes_ar") or "",
            narrative_summary=body.get("narrative_summary") or "",
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    observation_id = append_runtime_observation(record)
    return {
        "ok": True,
        "observation_id": observation_id,
        "assigns_authority": False,
        "assigns_legitimacy": False,
    }


@app.post("/api/governance-pilot/telemetry-replay-capture/wire")
async def governance_pilot_telemetry_replay_wire_api(request: Request):
    """Build telemetry + replay wiring bundle — preview only, no legitimacy."""
    from app.telemetry_replay_capture import wire_from_sandbox_analyses, wire_telemetry_replay_capture

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    events = body.get("events")
    analyses = body.get("analyses")
    replay_descriptors = body.get("replay_phenomenology") or body.get("replay_descriptors")
    try:
        if analyses:
            bundle = wire_from_sandbox_analyses(
                analyses,
                contract_id=body.get("contract_id") or "",
                observation_mode=body.get("observation_mode") or "controlled_observational",
                replay_descriptors=replay_descriptors,
            )
        elif events:
            bundle = wire_telemetry_replay_capture(
                events,
                artifact_name=body.get("artifact_name") or "",
                observation_mode=body.get("observation_mode") or "controlled_observational",
                contract_id=body.get("contract_id") or "",
                replay_descriptors=replay_descriptors,
                prior_provenance_segment_id=body.get("prior_provenance_segment_id") or "",
            )
        else:
            raise HTTPException(status_code=400, detail="events or analyses required")
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return {
        "ok": True,
        "wiring_id": bundle.get("wiring_id"),
        "bundle": bundle,
        "assigns_authority": False,
        "assigns_legitimacy": False,
        "epistemic_interpretation_blocked": True,
    }


@app.post("/api/governance-pilot/runtime-epistemic-governance/build")
async def governance_pilot_runtime_epistemic_governance_api(request: Request):
    """Build runtime epistemic governance bundle — three layers, no auto-merge."""
    from app.runtime_epistemic_governance import build_runtime_epistemic_governance_bundle
    from app.telemetry_replay_capture import wire_from_sandbox_analyses, wire_telemetry_replay_capture

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    wiring = body.get("telemetry_replay_wiring") or body.get("wiring_bundle")
    if not wiring:
        events = body.get("events")
        analyses = body.get("analyses")
        replay_descriptors = body.get("replay_phenomenology") or body.get("replay_descriptors")
        try:
            if analyses:
                wiring = wire_from_sandbox_analyses(
                    analyses,
                    contract_id=body.get("contract_id") or "",
                    observation_mode=body.get("observation_mode") or "controlled_observational",
                    replay_descriptors=replay_descriptors,
                )
            elif events:
                wiring = wire_telemetry_replay_capture(
                    events,
                    artifact_name=body.get("artifact_name") or "",
                    observation_mode=body.get("observation_mode") or "controlled_observational",
                    contract_id=body.get("contract_id") or "",
                    replay_descriptors=replay_descriptors,
                )
            else:
                raise HTTPException(status_code=400, detail="wiring_bundle, events, or analyses required")
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc))
    try:
        bundle = build_runtime_epistemic_governance_bundle(
            wiring,
            provenance_chain=body.get("provenance_chain"),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return {
        "ok": True,
        "governance_id": bundle.get("governance_id"),
        "bundle": bundle,
        "assigns_authority": False,
        "assigns_legitimacy": False,
        "assigns_comprehension": False,
        "layers_merged": False,
    }


@app.get("/api/governance-pilot/redundancy-audit")
async def governance_pilot_redundancy_audit_api():
    """Redundancy audit — unresolved uncertainty framing, review only."""
    from app.redundancy_audit import run_redundancy_audit

    return {"ok": True, **run_redundancy_audit()}


@app.get("/api/governance-pilot/hardening-gate-review")
async def governance_pilot_hardening_gate_review_api():
    """Hardening gate review step 4 — requires less belief to trust restraint?"""
    from app.constitutional_maintenance import run_hardening_gate_review

    return {"ok": True, **run_hardening_gate_review()}


@app.post("/api/governance-pilot/runtime-observation-session")
async def governance_pilot_runtime_observation_session_api(request: Request):
    """Run runtime observation session — not final legitimacy; ledger not appended."""
    from app.runtime_observation_session import run_runtime_observation_session

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    artifacts = body.get("artifacts") or body.get("artifact_paths") or []
    if not artifacts:
        raise HTTPException(status_code=400, detail="artifacts required")
    try:
        session = run_runtime_observation_session(
            artifacts,
            submission_id=body.get("submission_id"),
            batch_id=body.get("batch_id"),
            session_label=body.get("session_label") or "",
            enable_smoke_test=body.get("enable_smoke_test", True),
            facilitator_notes_ar=body.get("facilitator_notes_ar") or "",
            save_session=body.get("save_session", True),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return {
        "ok": True,
        "session_id": session.get("session_id"),
        "mode": "runtime_observation_session",
        "session": session,
        "assigns_legitimacy": False,
        "institutional_ledger_appended": False,
    }


@app.get("/api/governance-pilot/constitutional-maintenance/status")
async def governance_pilot_constitutional_maintenance_status_api():
    """Constitutional maintenance posture — restraint over self-legitimation."""
    from app.constitutional_maintenance import maintenance_status

    return {"ok": True, **maintenance_status()}


@app.get("/api/governance-pilot/architectural-hardening/status")
async def governance_pilot_architectural_hardening_status_api():
    """Hardening phase status — clarity under restraint, Step 6 blocked."""
    from app.architectural_hardening import (
        assess_architecture_overload,
        load_constitutional_core,
        step_6_readiness,
    )

    return {
        "ok": True,
        "phase_id": "ARCHITECTURAL_HARDENING_v1",
        "mode": "clarity_under_restraint",
        "constitutional_core": load_constitutional_core(),
        "overload_assessment": assess_architecture_overload(),
        "step_6_readiness": step_6_readiness(),
        "assigns_legitimacy": False,
    }


@app.get("/api/governance-pilot/worksheet/{submission_id}")
async def governance_pilot_worksheet_api(
    submission_id: int,
    db: Session = Depends(get_db),
):
    """Prefill governance observation worksheet Section A from snapshot."""
    from app.governance_drift_monitor import analyze_submission_governance_drift
    from app.governance_pilot_observatory import build_worksheet_draft

    sub = db.query(Submission).filter(Submission.id == submission_id).first()
    if not sub:
        raise HTTPException(status_code=404, detail="Submission not found")
    snap = None
    if getattr(sub, "grading_snapshot_json", None):
        try:
            snap = json.loads(str(sub.grading_snapshot_json))
        except (json.JSONDecodeError, TypeError):
            snap = None
    drift = analyze_submission_governance_drift(snap) if snap else None
    return build_worksheet_draft(
        submission_id=submission_id,
        student_name=getattr(sub, "student_name", "") or "",
        batch_id=getattr(sub, "batch_id", None),
        snapshot=snap,
        drift=drift,
    )


@app.post("/api/governance-pilot/observation")
async def governance_pilot_observation_api(request: Request):
    """Save human governance observation worksheet (Sections B–E)."""
    from app.facilitator_epistemic_worksheet import normalize_section_e
    from app.governance_pilot_observatory import save_observation

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    if not data.get("submission_id"):
        raise HTTPException(status_code=400, detail="submission_id required")
    if "section_e_epistemic_behaviour" in data:
        data["section_e_epistemic_behaviour"] = normalize_section_e(
            data.get("section_e_epistemic_behaviour")
        )
    if data.get("section_f_epistemic_trace") is not None:
        from app.epistemic_trace_capture import normalize_epistemic_trace

        data["section_f_epistemic_trace"] = normalize_epistemic_trace(
            data.get("section_f_epistemic_trace"),
            replay_consulted_at=data.get("replay_consulted_at"),
        )
    result = save_observation(data)
    if not result.get("ok"):
        raise HTTPException(
            status_code=403,
            detail=result.get("message_ar") or result.get("code") or "observation rejected",
        )
    return result


@app.get("/api/governance-pilot/epistemic-evidence/batch/{batch_id}")
async def governance_pilot_epistemic_evidence_api(batch_id: int):
    """Qualitative epistemic behavioural evidence — no score, no gate."""
    from app.facilitator_epistemic_worksheet import synthesize_epistemic_behavioural_evidence
    from app.governance_pilot_observatory import load_observations
    from app.phase2_institutional_observation import mark_epistemic_synthesis_generated

    observations = load_observations(batch_id=batch_id)
    synthesis = synthesize_epistemic_behavioural_evidence(observations, batch_id=batch_id)
    mark_result = mark_epistemic_synthesis_generated(
        batch_id, observation_count=len(observations)
    )
    synthesis["phase2_synthesis_meta"] = mark_result
    if not mark_result.get("ok") and mark_result.get("code") == "no_human_evidence":
        synthesis["warning_ar"] = mark_result.get("message_ar")
    return synthesis


@app.post("/api/governance-pilot/phase2/cohort/{batch_id}/activate")
async def governance_pilot_phase2_activate_api(batch_id: int, request: Request):
    """Register real cohort for Phase 2 live institutional observation."""
    from app.phase2_institutional_observation import activate_phase2_cohort

    try:
        data = await request.json()
    except Exception:
        data = {}
    return activate_phase2_cohort(
        batch_id=batch_id,
        facilitator=str(data.get("facilitator") or ""),
        target_submissions=int(data.get("target_submissions") or 25),
        notes_ar=str(data.get("notes_ar") or ""),
    )


@app.post("/api/governance-pilot/phase2/cohort/{batch_id}/complete")
async def governance_pilot_phase2_complete_api(batch_id: int, request: Request):
    """End workshop and start mandatory cooling period."""
    from app.governance_pilot_observatory import load_observations
    from app.phase2_institutional_observation import complete_workshop_and_start_cooling

    try:
        data = await request.json()
    except Exception:
        data = {}
    observations = load_observations(batch_id=batch_id)
    return complete_workshop_and_start_cooling(
        batch_id=batch_id,
        cooling_days=int(data.get("cooling_days") or 5),
        observation_count=len(observations),
        force=bool(data.get("force")),
    )


@app.get("/api/governance-pilot/phase2/cohort/{batch_id}")
async def governance_pilot_phase2_state_api(batch_id: int):
    """Phase 2 cohort state + exit checklist preview."""
    from app.governance_pilot_observatory import load_observations, synthesize_cohort_governance_report
    from app.facilitator_epistemic_worksheet import synthesize_epistemic_behavioural_evidence
    from app.phase2_institutional_observation import build_phase2_exit_checklist, get_phase2_cohort_state

    observations = load_observations(batch_id=batch_id)
    epistemic = synthesize_epistemic_behavioural_evidence(observations, batch_id=batch_id)
    state = get_phase2_cohort_state(batch_id, observation_count=len(observations))
    synthesis = None
    if observations:
        from app.database import SessionLocal
        from app.models import Submission

        db = SessionLocal()
        try:
            subs = db.query(Submission).filter(Submission.batch_id == batch_id).all()
            snapshots = []
            for sub in subs:
                if getattr(sub, "grading_snapshot_json", None):
                    try:
                        snapshots.append(json.loads(str(sub.grading_snapshot_json)))
                    except (json.JSONDecodeError, TypeError):
                        pass
            if snapshots:
                synthesis = synthesize_cohort_governance_report(
                    batch_id=batch_id, snapshots=snapshots
                )
        finally:
            db.close()
    exit_checklist = build_phase2_exit_checklist(
        synthesis=synthesis or {},
        epistemic=epistemic,
        batch_id=batch_id,
    )
    return {"cohort": state, "exit_checklist": exit_checklist, "epistemic_preview": epistemic}


@app.get("/governance-pilot/batch/{batch_id}/ritual-reading", response_class=HTMLResponse)
async def governance_pilot_ritual_reading_page(
    batch_id: int,
    request: Request,
):
    """Slow ritual reading — after cooling, before semantic memory."""
    from app.phase2_institutional_observation import (
        RITUAL_READING_LENSES,
        get_phase2_cohort_state,
        load_ritual_reading,
    )
    from app.governance_pilot_observatory import load_observations

    obs_count = len(load_observations(batch_id=batch_id))
    phase2 = get_phase2_cohort_state(batch_id, observation_count=obs_count)
    ritual = load_ritual_reading(batch_id)
    return templates.TemplateResponse(
        "phase2_ritual_reading.html",
        {
            "request": request,
            "batch_id": batch_id,
            "phase2_state": phase2,
            "ritual_reading": ritual,
            "lenses": RITUAL_READING_LENSES,
        },
    )


@app.get("/api/governance-pilot/phase2/ritual-reading/{batch_id}")
async def governance_pilot_ritual_reading_get_api(batch_id: int):
    from app.phase2_institutional_observation import load_ritual_reading, RITUAL_READING_LENSES

    return {"ritual_reading": load_ritual_reading(batch_id), "lenses": RITUAL_READING_LENSES}


@app.post("/api/governance-pilot/phase2/ritual-reading/{batch_id}")
async def governance_pilot_ritual_reading_save_api(batch_id: int, request: Request):
    from app.phase2_institutional_observation import advance_to_ritual_reading, save_ritual_reading

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    if data.get("advance"):
        gate = advance_to_ritual_reading(batch_id)
        if not gate.get("ok"):
            raise HTTPException(status_code=403, detail=gate.get("message_ar") or gate.get("code"))
    result = save_ritual_reading(batch_id, data)
    if not result.get("ok"):
        raise HTTPException(status_code=400, detail="save failed")
    return result


@app.post("/api/governance/semantic-memory/record")
async def governance_semantic_memory_record_api(request: Request):
    """
    Record deliberative semantic memory snapshot after workshop.
    Facilitator-initiated — not automatic gate.
    """
    from app.facilitator_epistemic_worksheet import synthesize_epistemic_behavioural_evidence
    from app.governance_pilot_observatory import load_observations
    from app.institutional_semantic_memory import record_semantic_memory_snapshot

    from app.phase2_institutional_observation import (
        assert_semantic_memory_allowed,
        mark_semantic_memory_recorded,
    )

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    batch_id = data.get("batch_id")
    if not batch_id:
        raise HTTPException(status_code=400, detail="batch_id required")
    allowed, gate = assert_semantic_memory_allowed(int(batch_id))
    if not allowed:
        raise HTTPException(
            status_code=403,
            detail=gate.get("message_ar") or gate.get("code") or "semantic memory locked",
        )
    observations = load_observations(batch_id=int(batch_id))
    synthesis = synthesize_epistemic_behavioural_evidence(observations, batch_id=int(batch_id))
    result = record_semantic_memory_snapshot(
        batch_id=int(batch_id),
        epistemic_synthesis=synthesis,
        freeze_epoch_id=str(data.get("freeze_epoch_id") or "epoch_1"),
        facilitator_note_ar=str(data.get("facilitator_note_ar") or ""),
        workshop_context_ar=str(data.get("workshop_context_ar") or ""),
    )
    mark_semantic_memory_recorded(int(batch_id))
    return result


@app.get("/api/governance/semantic-memory/trajectory")
async def governance_semantic_memory_trajectory_api(
    freeze_epoch_id: Optional[str] = None,
):
    """Cross-epoch semantic adaptation traces — deliberation only."""
    from app.institutional_semantic_memory import build_semantic_memory_trajectory

    return build_semantic_memory_trajectory(freeze_epoch_id=freeze_epoch_id)


@app.get("/api/governance/epistemic-leakage-lexicon")
async def governance_epistemic_leakage_lexicon_api():
    """Institutional semantic drift vocabulary — reference only, not a gate."""
    from app.epistemic_leakage_lexicon import build_lexicon_report

    return build_lexicon_report()


@app.get("/api/governance-pilot/synthesis/batch/{batch_id}")
async def governance_pilot_synthesis_batch_api(
    batch_id: int,
    db: Session = Depends(get_db),
):
    """Institutional governance stability report — not AI accuracy."""
    from app.governance_pilot_observatory import synthesize_cohort_governance_report

    subs = db.query(Submission).filter(Submission.batch_id == batch_id).all()
    snapshots: List[Dict[str, Any]] = []
    meta: List[Dict[str, Any]] = []
    for sub in subs:
        if not getattr(sub, "grading_snapshot_json", None):
            continue
        try:
            snap = json.loads(str(sub.grading_snapshot_json))
            snapshots.append(snap)
            meta.append({
                "submission_id": sub.id,
                "student_name": getattr(sub, "student_name", "") or "",
            })
        except (json.JSONDecodeError, TypeError):
            continue
    report = synthesize_cohort_governance_report(
        batch_id=batch_id,
        snapshots=snapshots,
        submission_meta=meta,
    )
    from app.canonical_stability_metrics import (
        attach_to_cohort_synthesis,
        compute_canonical_stability_metrics,
    )

    stability = compute_canonical_stability_metrics(db, batch_id=batch_id)
    return attach_to_cohort_synthesis(report, stability)


@app.get("/api/governance/canonical-stability/batch/{batch_id}")
async def canonical_stability_batch_api(batch_id: int, db: Session = Depends(get_db)):
    """Institutional reproducibility health metrics for one batch."""
    from app.canonical_stability_metrics import compute_canonical_stability_metrics

    return compute_canonical_stability_metrics(db, batch_id=batch_id)


@app.get("/api/governance/canonical-stability/assignment/{assignment_id}")
async def canonical_stability_assignment_api(
    assignment_id: int, db: Session = Depends(get_db)
):
    """Institutional reproducibility health metrics across an assignment."""
    from app.canonical_stability_metrics import compute_canonical_stability_metrics

    return compute_canonical_stability_metrics(db, assignment_id=assignment_id)


@app.get("/api/governance/canonical-stability/trajectory/assignment/{assignment_id}")
async def canonical_stability_trajectory_assignment_api(
    assignment_id: int,
    record: bool = False,
    db: Session = Depends(get_db),
):
    """Governance trajectory — transitions + freeze-epoch-relative stability."""
    from app.canonical_stability_trajectory import build_governance_trajectory_report

    return build_governance_trajectory_report(
        db, assignment_id=assignment_id, record_snapshot=record
    )


@app.get("/api/governance/canonical-stability/trajectory/batch/{batch_id}")
async def canonical_stability_trajectory_batch_api(
    batch_id: int,
    record: bool = False,
    db: Session = Depends(get_db),
):
    from app.models import BatchGrading
    from app.canonical_stability_trajectory import build_governance_trajectory_report

    batch = db.query(BatchGrading).filter(BatchGrading.id == batch_id).first()
    if not batch:
        raise HTTPException(status_code=404, detail="Batch not found")
    return build_governance_trajectory_report(
        db,
        assignment_id=batch.assignment_id,
        batch_id=batch_id,
        record_snapshot=record,
    )


@app.get("/api/governance/epoch/{epoch_id}/narrative")
async def governance_epoch_narrative_api(
    epoch_id: str,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    """Institutional governance memory narrative for one freeze epoch."""
    from app.governance_epoch_narrative import build_epoch_narrative

    return build_epoch_narrative(db, epoch_id=epoch_id, assignment_id=assignment_id)


@app.get("/api/governance/epoch/{epoch_id}/rfc-review")
async def governance_epoch_rfc_review_api(
    epoch_id: str,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    """Pre-RFC review package — justify epoch transition institutionally."""
    from app.canonical_stability_trajectory import ACTIVE_FREEZE_EPOCH
    from app.governance_epoch_narrative import build_epoch_review_rfc_package

    return build_epoch_review_rfc_package(
        db,
        target_epoch_id=epoch_id,
        current_epoch_id=ACTIVE_FREEZE_EPOCH,
        assignment_id=assignment_id,
    )


@app.get("/governance-epoch-workshop", response_class=HTMLResponse)
async def governance_epoch_workshop_page(
    request: Request,
    target_epoch: str = "epoch_2",
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    """Epoch Workshop Review — institutional gate (not technical review)."""
    from app.canonical_stability_trajectory import ACTIVE_FREEZE_EPOCH
    from app.governance_epoch_workshop import build_epoch_workshop_review

    review = build_epoch_workshop_review(
        db,
        current_epoch_id=ACTIVE_FREEZE_EPOCH,
        target_epoch_id=target_epoch,
        assignment_id=assignment_id,
    )
    return templates.TemplateResponse(
        "governance_epoch_workshop.html",
        {"request": request, "review": review},
    )


@app.get("/api/governance/epoch/workshop-review")
async def governance_epoch_workshop_review_api(
    target_epoch: str = "epoch_2",
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    from app.canonical_stability_trajectory import ACTIVE_FREEZE_EPOCH
    from app.governance_epoch_workshop import build_epoch_workshop_review

    return build_epoch_workshop_review(
        db,
        current_epoch_id=ACTIVE_FREEZE_EPOCH,
        target_epoch_id=target_epoch,
        assignment_id=assignment_id,
    )


@app.post("/api/governance/epoch/workshop-review")
async def governance_epoch_workshop_save_api(request: Request):
    from app.governance_epoch_workshop import save_epoch_workshop_review

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    if not data.get("transition_verdict"):
        raise HTTPException(status_code=400, detail="transition_verdict required")
    if not data.get("institution_affirmation"):
        raise HTTPException(
            status_code=400,
            detail="institution_affirmation required — signed institutional artifact",
        )
    try:
        return save_epoch_workshop_review(data)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/governance/epoch/{epoch_id}/mitigation-ledger")
async def epoch_mitigation_ledger_api(
    epoch_id: str,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    from app.governance_epoch_mitigation_ledger import build_mitigation_ledger_report

    return build_mitigation_ledger_report(db, epoch_id=epoch_id, assignment_id=assignment_id)


@app.post("/api/governance/epoch/mitigation-ledger")
async def epoch_mitigation_ledger_append_api(request: Request):
    from app.governance_epoch_mitigation_ledger import append_ledger_entry

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")
    for field in ("epoch_id", "issue", "mitigation", "result"):
        if not data.get(field):
            raise HTTPException(status_code=400, detail=f"{field} required")
    return append_ledger_entry(
        epoch_id=data["epoch_id"],
        issue=data["issue"],
        mitigation=data["mitigation"],
        result=data["result"],
        failure_mode_id=data.get("failure_mode_id"),
        result_status=data.get("result_status", "pending"),
        recorded_by=data.get("recorded_by", ""),
        result_evidence_ar=data.get("result_evidence_ar", ""),
        lineage_refs=data.get("lineage_refs"),
        artifact_refs=data.get("artifact_refs"),
    )


@app.get("/api/governance/artifact/{artifact_id}")
async def institutional_artifact_api(artifact_id: str):
    from app.institutional_artifact import load_artifact, verify_artifact_integrity

    art = load_artifact(artifact_id)
    if not art:
        raise HTTPException(status_code=404, detail="Artifact not found")
    return {"artifact": art, "integrity": verify_artifact_integrity(art)}


@app.get("/api/governance/artifact-chain/epoch/{epoch_id}")
async def governance_artifact_chain_epoch_api(
    epoch_id: str,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    from app.governance_artifact_chain import build_governance_artifact_chain

    return build_governance_artifact_chain(
        db, epoch_id=epoch_id, assignment_id=assignment_id
    )


@app.get("/api/governance/artifact-chain/artifact/{artifact_id}")
async def governance_artifact_chain_anchor_api(
    artifact_id: str,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    from app.governance_artifact_chain import build_governance_artifact_chain
    from app.institutional_artifact import load_artifact

    art = load_artifact(artifact_id)
    if not art:
        raise HTTPException(status_code=404, detail="Artifact not found")
    epoch_id = (art.get("governance_context") or {}).get("freeze_epoch_id") or "epoch_1"
    chain = build_governance_artifact_chain(
        db,
        epoch_id=epoch_id,
        assignment_id=assignment_id,
        anchor_artifact_id=artifact_id,
    )
    chain["anchored_on"] = artifact_id
    return chain


@app.get("/governance/l4-decision", response_class=HTMLResponse)
async def governance_l4_decision_page(
    request: Request,
    batch_id: Optional[int] = None,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    """L4 institutional decision point — epoch review before sandbox / FREEZE v2."""
    from app.l4_institutional_decision import build_l4_decision_package
    from app.phase2_institutional_observation import is_governance_expansion_locked

    if batch_id is not None and is_governance_expansion_locked(batch_id):
        return templates.TemplateResponse(
            "phase2_governance_locked.html",
            {
                "request": request,
                "batch_id": batch_id,
                "message_ar": (
                    "L4 / sandbox / constitutional evolution مقفول — "
                    "أكمل Phase 2: workshop → synthesis → cooling → ritual reading → semantic memory → epoch deliberation."
                ),
            },
        )

    package = build_l4_decision_package(
        db, batch_id=batch_id, assignment_id=assignment_id
    )
    return templates.TemplateResponse(
        "governance_l4_decision.html",
        {
            "request": request,
            "package": package,
            "batch_id": batch_id,
            "assignment_id": assignment_id,
        },
    )


@app.get("/api/governance/l4-decision/package")
async def governance_l4_decision_package_api(
    batch_id: Optional[int] = None,
    assignment_id: Optional[int] = None,
    db: Session = Depends(get_db),
):
    from app.l4_institutional_decision import build_l4_decision_package

    return build_l4_decision_package(
        db, batch_id=batch_id, assignment_id=assignment_id
    )


@app.get("/settings", response_class=HTMLResponse)
async def settings_page(request: Request, db: Session = Depends(get_db)):
    """Settings page"""
    from fastapi.responses import RedirectResponse  # type: ignore
    user_id = get_current_user_id(request)
    if not user_id:
        return RedirectResponse(url="/login", status_code=302)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if user is None or user.role != models.UserRole.ADMIN:
        return RedirectResponse(url="/dashboard", status_code=302)
    sub_info = get_subscription_info(db, user_id) if user_id else None
    return templates.TemplateResponse(
        "settings.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "subscription": sub_info,
        },
    )


# API Endpoints


@app.post("/api/upload-textbook")
async def upload_textbook(
    request: Request,
    file: UploadFile = File(...),
    title: str = Form(...),
    db: Session = Depends(get_db),
):
    """Upload a textbook PDF — hardened with full validation + clear errors."""
    try:
        # 1. Validate inputs
        if not file or not file.filename:
            raise HTTPException(status_code=400, detail="لم يتم استلام أي ملف. الرجاء اختيار ملف PDF صالح.")
        ext = Path(file.filename).suffix.lower()
        if ext not in (".pdf",):
            raise HTTPException(status_code=400,
                detail=f"نوع الملف غير مدعوم ({ext or 'بدون امتداد'}). الرجاء رفع ملف PDF فقط.")
        if not title or not title.strip():
            raise HTTPException(status_code=400, detail="عنوان الكتاب مطلوب.")

        # 2. Ensure target directory exists
        try:
            TEXTBOOKS_DIR.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            raise HTTPException(status_code=500,
                detail=f"فشل في تجهيز مجلد الرفع: {e}.")

        # 3. Sanitize filename WITHOUT regex (avoid character-class quirks)
        _UNSAFE_FILENAME_CHARS = set('<>:"/\\|?*') | {chr(_i) for _i in range(32)}
        _raw_name = Path(file.filename).name  # strip path components
        safe_name = ''.join(
            ('_' if _ch in _UNSAFE_FILENAME_CHARS else _ch) for _ch in _raw_name
        ).strip()
        while '__' in safe_name:
            safe_name = safe_name.replace('__', '_')
        if not safe_name or safe_name in ('.', '..'):
            import time as _time_mod
            safe_name = f"book_{int(_time_mod.time())}.pdf"
        # Unique prefix avoids overwriting / DB dedup quirks when same filename is reused
        import uuid as _uuid_mod
        unique_name = f"{_uuid_mod.uuid4().hex[:10]}_{safe_name}"
        file_path = TEXTBOOKS_DIR / unique_name

        # 4. Save with size guard (async UploadFile reads — avoids empty/partial uploads on some setups)
        try:
            try:
                await file.seek(0)
            except Exception:
                pass
            with open(file_path, "wb") as buffer:
                MAX_SIZE = 200 * 1024 * 1024
                bytes_written = 0
                while True:
                    chunk = await file.read(1024 * 1024)
                    if not chunk:
                        break
                    buffer.write(chunk)
                    bytes_written += len(chunk)
                    if bytes_written > MAX_SIZE:
                        buffer.close()
                        try: file_path.unlink(missing_ok=True)
                        except Exception: pass
                        raise HTTPException(status_code=413,
                            detail="حجم الملف يتجاوز الحد الأقصى المسموح (200 ميغابايت).")
            if bytes_written == 0:
                try: file_path.unlink(missing_ok=True)
                except Exception: pass
                raise HTTPException(status_code=400, detail="الملف المرفوع فارغ.")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"فشل في حفظ الملف على القرص: {e}")

        # 5. Validate it's actually a PDF
        try:
            with open(file_path, "rb") as fp:
                head = fp.read(4096).lstrip(b"\xef\xbb\xbf \t\r\n")
            if len(head) < 5 or not head.startswith(b"%PDF"):
                try: file_path.unlink(missing_ok=True)
                except Exception: pass
                raise HTTPException(status_code=400,
                    detail="الملف المرفوع ليس PDF صالحاً (لا يبدأ بمحتوى PDF معروف).")
        except HTTPException:
            raise
        except Exception:
            pass

        # 6. Resolve uploading user (avoid FK violation)
        try:
            uid = get_current_user_id(request)
        except Exception:
            uid = None
        if not uid:
            uid = 1
        try:
            _u_exists = db.query(models.User.id).filter(models.User.id == uid).first()
        except Exception:
            _u_exists = None
        if not _u_exists:
            try:
                _any_user = db.query(models.User.id).order_by(models.User.id.asc()).first()
                if _any_user:
                    uid = int(_any_user[0])
            except Exception:
                uid = 1

        # 7. DEDUP
        try:
            existing_textbook = db.query(Textbook).filter(Textbook.file_path == str(file_path)).first()
            if existing_textbook:
                return {"success": True, "id": existing_textbook.id, "title": existing_textbook.title}
        except Exception as e:
            print(f"DEDUP lookup failed (non-fatal): {e}")

        # 8. Create record
        try:
            textbook = Textbook(
                title=title.strip(),
                file_url=f"/uploads/textbooks/{unique_name}",
                file_path=str(file_path),
                uploaded_by=uid,
            )
            db.add(textbook); db.commit(); db.refresh(textbook)
        except Exception as e:
            db.rollback()
            raise HTTPException(status_code=500,
                detail=f"فشل في حفظ سجل الكتاب في قاعدة البيانات: {type(e).__name__}: {e}")

        return {"success": True, "id": textbook.id, "title": title.strip()}

    except HTTPException:
        raise
    except Exception as e:
        _traceback_print_safe()
        raise HTTPException(status_code=500,
            detail=f"خطأ غير متوقع أثناء رفع الكتاب: {type(e).__name__}: {str(e) or '(بدون رسالة)'}") from e


@app.get("/api/get-units")
async def get_units(level: str | None = None, grade: str | None = None):
    """Get list of all BTEC units from specification, optionally filtered by level or grade"""
    try:
        units_file = Path("uploads/specification/btec_units.json")
        if not units_file.exists():
            return {
                "success": True,
                "units_list": [],
                "total": 0,
                "grade_levels": [],
                "missing_spec": True,
                "detail": "ملف المواصفات غير موجود: uploads/specification/btec_units.json",
            }

        with open(units_file, encoding="utf-8") as f:
            data = json.load(f)

        units_list = data.get("units_list", [])
        if level:
            units_list = [
                u for u in units_list
                if u.get("qualification_level", "L3") == level
            ]
        if grade:
            units_list = [
                u for u in units_list
                if u.get("grade_level", "") == grade
            ]

        # Collect unique grade levels for dropdown population
        all_units = data.get("units_list", [])
        grade_levels = sorted(set(
            u["grade_level"] for u in all_units if u.get("grade_level")
        ))

        return {
            "success": True,
            "units_list": units_list,
            "total": len(units_list),
            "grade_levels": grade_levels,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e


@app.get("/api/get-unit/{unit_number}")
async def get_unit_details(unit_number: str):
    """Get detailed information for a specific unit"""
    try:
        units_file = Path("uploads/specification/btec_units.json")
        if not units_file.exists():
            raise HTTPException(status_code=404, detail="Units file not found")

        with open(units_file, encoding="utf-8") as f:
            data = json.load(f)

        units = data.get("units", {})
        key = _resolve_btec_unit_key(unit_number, units)
        if key:
            return {"success": True, "unit": units[key]}

        raise HTTPException(status_code=404, detail="Unit not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e


def _dedupe_assignments_by_content_hash(assignments_list: list) -> list:
    """Keep one card per content_hash (input order preserved; use newest-first query)."""
    seen: set[str] = set()
    out = []
    for a in assignments_list:
        key = (getattr(a, "content_hash", None) or "").strip() or f"__id_{a.id}__"
        if key in seen:
            continue
        seen.add(key)
        out.append(a)
    return out


def _ensure_user_assignment_link(db: Session, user_id: int, assignment_id: int) -> None:
    """Let this user list/open a canonical shared assignment without duplicating rows."""
    a = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not a or a.created_by == user_id:
        return
    exists = (
        db.query(models.UserAssignmentLink)
        .filter(
            models.UserAssignmentLink.user_id == user_id,
            models.UserAssignmentLink.assignment_id == assignment_id,
        )
        .first()
    )
    if exists:
        return
    db.add(
        models.UserAssignmentLink(
            user_id=user_id,
            assignment_id=assignment_id,
        )
    )


@app.get("/api/available-assignments")
async def available_assignments(
    request: Request,
    unit_number: str,
    subject: str | None = None,
    db: Session = Depends(get_db),
):
    """List admin-created READY assignments the user can request for a unit."""
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="يجب تسجيل الدخول أولاً")

    user = db.query(models.User).filter(models.User.id == user_id).first()
    is_admin = user and user.role == models.UserRole.ADMIN
    if is_admin:
        raise HTTPException(status_code=403, detail="هذا المسار مخصص لطلب الدليل من حساب المستخدم.")

    unit_key = (unit_number or "").strip()
    if not unit_key:
        return JSONResponse(
            {"success": False, "detail": "يجب تحديد الوحدة."},
            status_code=400,
        )

    from app.assignment_access import (  # type: ignore
        assignment_request_item,
        check_assignment_creation_access,
        list_ready_assignments_for_unit,
    )

    access_error, _resolved_subject = check_assignment_creation_access(
        db,
        user_id,
        subject=subject,
        unit_number=unit_key,
        is_admin=False,
        for_request=True,
    )
    if access_error is not None:
        return access_error

    assignments = list_ready_assignments_for_unit(db, unit_number=unit_key)
    return {
        "success": True,
        "assignments": [assignment_request_item(a) for a in assignments],
    }


@app.post("/api/request-assignment")
async def request_assignment(
    request: Request,
    db: Session = Depends(get_db),
):
    """Link the user to an admin-created assignment guide (no upload, no AI)."""
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="يجب تسجيل الدخول أولاً")

    user = db.query(models.User).filter(models.User.id == user_id).first()
    is_admin = user and user.role == models.UserRole.ADMIN
    if is_admin:
        raise HTTPException(status_code=403, detail="هذا المسار مخصص لطلب الدليل من حساب المستخدم.")

    try:
        payload = await request.json()
    except Exception:
        payload = {}

    assignment_id = payload.get("assignment_id")
    subject = payload.get("subject")
    unit_number = payload.get("unit_number")

    if not assignment_id:
        return JSONResponse(
            {"success": False, "detail": "يجب اختيار المهمة أولاً."},
            status_code=400,
        )

    from app.assignment_access import (  # type: ignore
        assignment_not_available_response,
        check_assignment_creation_access,
    )

    assignment = db.query(Assignment).filter(Assignment.id == int(assignment_id)).first()
    if not assignment:
        return JSONResponse(
            {"success": False, "detail": "المهمة غير موجودة."},
            status_code=404,
        )

    if assignment.status != AssignmentStatus.READY or not assignment.reference_solution_json:
        return assignment_not_available_response()

    unit_key = (unit_number or assignment.unit_number or "").strip()
    if unit_key and (assignment.unit_number or "").strip() != unit_key:
        return JSONResponse(
            {
                "success": False,
                "detail": "المهمة المختارة لا تتطابق مع الوحدة المحددة.",
            },
            status_code=400,
        )

    access_error, _resolved_subject = check_assignment_creation_access(
        db,
        user_id,
        subject=subject or getattr(assignment, "subject", None),
        unit_number=(assignment.unit_number or unit_key or None),
        is_admin=False,
        for_request=True,
    )
    if access_error is not None:
        return access_error

    if int(getattr(assignment, "created_by", 0) or 0) != user_id:
        _ensure_user_assignment_link(db, user_id, assignment.id)
    db.commit()

    return {
        "success": True,
        "id": assignment.id,
        "title": assignment.title,
        "status": "ready",
        "cached": True,
        "requested": True,
        "reference_solution": strip_excluded_guide_sections(
            json.loads(str(assignment.reference_solution_json))
        ),
    }


@app.post("/api/create-assignment")
async def create_assignment(
    request: Request,
    assignment_file: UploadFile = File(...),
    title: str = Form(...),
    unit_number: str = Form(None),  # NEW - Unit selection
    subject: str = Form(None),      # e.g. "برمجة" — matches SubjectBalance
    description: str = Form(None),
    textbook_id: int = Form(None),
    page_from: int = Form(None),
    page_to: int = Form(None),
    db: Session = Depends(get_db),
):
    """Create assignment (admin) or request an existing one (users)."""
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="يجب تسجيل الدخول أولاً")

    user = db.query(models.User).filter(models.User.id == user_id).first()
    is_admin = user and user.role == models.UserRole.ADMIN

    from app.assignment_access import (  # type: ignore
        assignment_not_available_response,
        check_assignment_creation_access,
        find_ready_assignment_for_request,
    )

    assignment_raw = await assignment_file.read()
    if not assignment_raw:
        raise HTTPException(status_code=400, detail="ملف المهمة فارغ")

    # ── Regular users: request an admin-created assignment only ─────────────
    if not is_admin:
        unit_key = (unit_number or "").strip()
        if not unit_key:
            return JSONResponse(
                {
                    "success": False,
                    "detail": "يجب اختيار المادة/الوحدة من اشتراكك لطلب دليل المهمة.",
                },
                status_code=400,
            )

        access_error, _resolved_subject = check_assignment_creation_access(
            db,
            user_id,
            subject=subject,
            unit_number=unit_key,
            is_admin=False,
            for_request=True,
        )
        if access_error is not None:
            return access_error

        cached_hit, cache_source = find_ready_assignment_for_request(
            db,
            unit_number=unit_key,
            assignment_raw=assignment_raw,
            assignments_dir=ASSIGNMENTS_DIR,
        )
        if cached_hit is None:
            return assignment_not_available_response()

        print(
            f"✅ [REQUEST:{cache_source}] Linking user {user_id} "
            f"to assignment #{cached_hit.id}"
        )
        if int(getattr(cached_hit, "created_by", 0) or 0) != user_id:
            _ensure_user_assignment_link(db, user_id, cached_hit.id)
        db.commit()
        return {
            "success": True,
            "id": cached_hit.id,
            "title": cached_hit.title,
            "status": "ready",
            "cached": True,
            "requested": True,
            "reference_solution": strip_excluded_guide_sections(
                json.loads(str(cached_hit.reference_solution_json))
            ),
        }

    # ── Admin: full create flow (textbook + AI on cache miss) ───────────────
    if textbook_id is None or page_from is None or page_to is None:
        raise HTTPException(
            status_code=400,
            detail="يجب رفع الكتاب وتحديد نطاق الصفحات لإنشاء مهمة جديدة.",
        )

    try:
        # Load unit data if unit_number provided
        unit_data: dict | None = None
        if unit_number:
            try:
                units_file = Path("uploads/specification/btec_units.json")
                if units_file.exists():
                    with open(units_file, encoding="utf-8") as f:
                        units_json = json.load(f)
                        all_u = units_json.get("units", {})
                        unit_key = _resolve_btec_unit_key(unit_number, all_u)
                        if unit_key:
                            unit_dict = all_u[unit_key]
                            unit_data = unit_dict
                            unit_number = unit_key
                            print(
                                f" Loaded Unit {unit_key}: {unit_dict.get('unit_name', '')}"
                            )
                        else:
                            print(f" Unit {unit_number} not found in specification")
            except Exception as e:
                print(f" Could not load unit data: {e}")

        from app.assignment_access import (  # type: ignore
            check_assignment_creation_access,
            find_ready_assignment_in_db,
        )

        access_error, resolved_subject = check_assignment_creation_access(
            db,
            user_id,
            subject=subject,
            unit_number=unit_number,
            is_admin=True,
        )
        if access_error is not None:
            return access_error

        # Get textbook
        textbook = db.query(Textbook).filter(Textbook.id == textbook_id).first()
        if not textbook:
            raise HTTPException(status_code=404, detail="Textbook not found")

        _unit_key_part = (unit_number or "").strip() if unit_number else ""
        _hash_payload = (
            f"{textbook_id}|{_unit_key_part}|{page_from}|"
            f"{page_to}|{hashlib.sha256(assignment_raw).hexdigest()}"
        )
        content_hash = hashlib.sha256(_hash_payload.encode("utf-8")).hexdigest()

        cached_hit, cache_source = find_ready_assignment_in_db(
            db,
            user_id=user_id,
            content_hash=content_hash,
            unit_number=unit_number,
            assignment_raw=assignment_raw,
            assignments_dir=ASSIGNMENTS_DIR,
        )
        if cached_hit is not None:
            print(
                f"✅ [DB CACHE:{cache_source}] Returning assignment "
                f"#{cached_hit.id} for user {user_id}"
            )
            if int(getattr(cached_hit, "created_by", 0) or 0) != user_id:
                _ensure_user_assignment_link(db, user_id, cached_hit.id)
            db.commit()
            return {
                "success": True,
                "id": cached_hit.id,
                "title": cached_hit.title,
                "status": "ready",
                "cached": True,
                "reference_solution": strip_excluded_guide_sections(
                    json.loads(str(cached_hit.reference_solution_json))
                ),
            }

        # ── DEDUP: Clean up incomplete (DRAFT/ANALYZING) assignments with
        #    the same content hash from this user to avoid orphan duplicates ───
        stale_drafts = (
            db.query(Assignment)
            .filter(
                Assignment.content_hash == content_hash,
                Assignment.created_by == user_id,
                Assignment.status.in_([
                    AssignmentStatus.DRAFT,
                    AssignmentStatus.ANALYZING,
                ]),
            )
            .all()
        )
        for stale in stale_drafts:
            print(
                f"🗑️ [DEDUP] Removing stale {stale.status.name} "
                f"assignment #{stale.id} (same content hash)"
            )
            db.delete(stale)
        if stale_drafts:
            db.commit()

        # ── PERSISTENT CACHE: check SolutionCache (survives assignment deletion)
        solution_cache_hit = (
            db.query(models.SolutionCache)
            .filter(models.SolutionCache.content_hash == content_hash)
            .first()
        )
        if solution_cache_hit:
            print(
                f"✅ [SOLUTION CACHE] Found persistent cached solution "
                f"for content hash {content_hash[:16]}..."
            )

        # Save assignment file (basename only — avoids Errno 22 from fake browser paths)
        safe_assignment_name = _safe_upload_basename(
            assignment_file.filename, default="assignment_upload.bin"
        )
        assignment_file_path = ASSIGNMENTS_DIR / safe_assignment_name
        with open(assignment_file_path, "wb") as buffer:
            buffer.write(assignment_raw)

        # Extract assignment text (absolute path — stable on Windows)
        assignment_file_path = assignment_file_path.resolve()
        assignment_text = extract_text_from_file(str(assignment_file_path))

        # Extract required criteria from assignment and filter unit criteria (NEW)
        filtered_criteria = None
        if unit_data and assignment_text:
            from app.criteria_extractor import (  # type: ignore
                extract_required_criteria_from_text,
                filter_unit_criteria,
            )

            # Extract criteria mentioned in the assignment
            required_criteria = extract_required_criteria_from_text(assignment_text)

            if required_criteria:
                print(
                    f" Required Criteria found in assignment: {', '.join(required_criteria)}"
                )

                # Filter unit criteria to include only required ones
                all_criteria = unit_data.get("criteria", [])
                filtered_criteria = filter_unit_criteria(
                    all_criteria, required_criteria
                )

                print(
                    f" Filtered to {len(filtered_criteria)} criteria (out of {len(all_criteria)} total)"
                )
            else:
                print(
                    " No specific criteria found in assignment, using all unit criteria"
                )
                filtered_criteria = unit_data.get("criteria", [])

        # Create assignment record
        assignment = Assignment(
            title=title,
            description=description,
            textbook_id=textbook_id,
            page_from=page_from,
            page_to=page_to,
            unit_number=unit_number if unit_data else None,
            unit_name=unit_data.get("unit_name") if unit_data else None,
            subject=resolved_subject or (subject or "").strip() or None,
            unit_criteria_json=(
                json.dumps(filtered_criteria, ensure_ascii=False)
                if filtered_criteria
                else None
            ),  # UPDATED
            assignment_file_url=f"/uploads/assignments/{safe_assignment_name}",
            assignment_text=assignment_text,
            content_hash=content_hash,
            status=AssignmentStatus.ANALYZING,
            created_by=user_id,
        )

        db.add(assignment)
        db.commit()
        db.refresh(assignment)

        # ── PERSISTENT CACHE HIT: reuse from SolutionCache table
        if solution_cache_hit:
            print("✅ [SOLUTION CACHE HIT] Restoring solution from persistent cache")
            _cached_guide = strip_excluded_guide_sections(
                json.loads(str(solution_cache_hit.solution_json))
            )
            _cached_guide_json = json.dumps(_cached_guide, ensure_ascii=False)
            assignment.reference_solution_json = _cached_guide_json  # type: ignore
            assignment.reference_solution_text = _cached_guide_json  # type: ignore
            assignment.solution_hash = solution_cache_hit.solution_hash
            _orm_set(assignment, "is_locked", True)
            _orm_set(assignment, "status", AssignmentStatus.READY)

            cached_criteria = json.loads(solution_cache_hit.criteria_json)
            for c in cached_criteria:
                db.add(models.GradingCriteria(
                    assignment_id=assignment.id,
                    criteria_level=c["criteria_level"],
                    criteria_name=c["criteria_name"],
                    criteria_description=c["criteria_description"],
                    key_points=json.dumps(c.get("key_points", []), ensure_ascii=False),
                    weight=c.get("weight", 25),
                ))

            db.commit()
            return {
                "success": True,
                "id": assignment.id,
                "title": title,
                "status": "ready",
                "cached": True,
                "reference_solution": strip_excluded_guide_sections(
                    json.loads(str(assignment.reference_solution_json))
                ),
            }

        # Process textbook and generate reference solution (async)
        try:
            textbook_pdf = _resolve_uploaded_path(str(textbook.file_path))
            if not textbook_pdf.is_file():
                raise HTTPException(
                    status_code=404,
                    detail=f"ملف الكتاب غير موجود على القرص: {textbook_pdf.name}",
                )
            page_from_i = max(1, page_from)
            page_to_i = max(page_from_i, page_to)

            analysis, reference_solution, grading_criteria = (
                await process_textbook_and_assignment(
                    str(textbook_pdf),
                    page_from_i,
                    page_to_i,
                    assignment_text,
                    known_criteria_list=filtered_criteria,  # type: ignore # NEW - Pass filtered official specs
                    assignment_file_path=str(assignment_file_path.resolve()),  # NEW - Pass file path for image analysis
                )
            )

            # Print the reference solution for verification
            print("\n" + "=" * 80)
            print(" دليل المهمة المُنشأ / Generated Reference Solution:")
            print("=" * 80)
            print(json.dumps(reference_solution, ensure_ascii=False, indent=2))
            print("=" * 80 + "\n")

            # Update assignment with reference solution (sections 6/7/9 excluded)
            reference_solution = strip_excluded_guide_sections(reference_solution)
            assignment.reference_solution_text = json.dumps(reference_solution, ensure_ascii=False)  # type: ignore
            assignment.reference_solution_json = json.dumps(reference_solution, ensure_ascii=False)  # type: ignore

            # GOLDEN HASH LOCK: Calculate SHA256 and lock the solution
            solution_hash = hashlib.sha256(
                assignment.reference_solution_json.encode("utf-8") if assignment.reference_solution_json else b""  # type: ignore
            ).hexdigest()
            assignment.solution_hash = solution_hash  # type: ignore
            assignment.is_locked = True  # type: ignore # Lock permanently - no regeneration allowed
            print(f" Golden Hash Lock: {solution_hash[:16]}...")

            assignment.status = AssignmentStatus.READY  # type: ignore

            # Create grading criteria records
            for criteria in grading_criteria:
                db_criteria = models.GradingCriteria(
                    assignment_id=assignment.id,
                    criteria_level=str(criteria["criteria_level"]),  # Fix KeyError: arbitrary BTEC criteria
                    criteria_name=criteria["criteria_name"],
                    criteria_description=criteria["criteria_description"],
                    key_points=json.dumps(criteria["key_points"], ensure_ascii=False),
                    weight=criteria["weight"],
                )
                db.add(db_criteria)

            # ── Save to persistent SolutionCache so it survives assignment deletion
            existing_sc = (
                db.query(models.SolutionCache)
                .filter(models.SolutionCache.content_hash == content_hash)
                .first()
            )
            if not existing_sc:
                criteria_for_cache = [
                    {
                        "criteria_level": str(c["criteria_level"]),
                        "criteria_name": c["criteria_name"],
                        "criteria_description": c["criteria_description"],
                        "key_points": c["key_points"],
                        "weight": c["weight"],
                    }
                    for c in grading_criteria
                ]
                db.add(models.SolutionCache(
                    content_hash=content_hash,
                    solution_json=assignment.reference_solution_json,
                    solution_hash=solution_hash,
                    criteria_json=json.dumps(criteria_for_cache, ensure_ascii=False),
                ))
                print(f"💾 [SOLUTION CACHE] Saved to persistent cache for hash {content_hash[:16]}...")

            db.commit()

            return {
                "success": True,
                "id": assignment.id,
                "title": title,
                "status": "ready",
                "reference_solution": reference_solution,
            }

        except Exception as e:
            _traceback_print_safe()
            _stderr_line_utf8(f"[create-assignment] {type(e).__name__}: {e}")

            try:
                assignment.status = AssignmentStatus.FAILED  # type: ignore
                db.commit()
            except Exception:
                db.rollback()
            raise HTTPException(
                status_code=500,
                detail=f"فشل في إنشاء دليل المهمة ({type(e).__name__}): {e}",
            ) from e

    except OSError as e:
        raise HTTPException(
            status_code=500,
            detail=f"خطأ في الملفات أو مسار PDF ({type(e).__name__}): {e}",
        ) from e
    except HTTPException:
        raise
    except Exception as e:
        _traceback_print_safe()
        raise HTTPException(
            status_code=500,
            detail=f"خطأ أثناء إنشاء المهمة ({type(e).__name__}): {e}",
        ) from e


@app.post("/api/generate-solution/{assignment_id}")
async def generate_solution_for_assignment(
    assignment_id: int,
    request: Request,
    db: Session = Depends(get_db),
):
    """
    Generate (or retrieve cached) reference solution for an existing assignment.
    - If a solution already exists in DB → return it immediately (no AI call).
    - If another assignment with the same content exists → copy its solution.
    - Otherwise → call AI, generate, save, return.
    """
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="يجب تسجيل الدخول أولاً")

    user = db.query(models.User).filter(models.User.id == user_id).first()
    is_admin = user and user.role == models.UserRole.ADMIN

    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    # ── CACHE HIT 1: solution already stored for this assignment ──────────────
    if assignment.reference_solution_json:
        print(f"✅ [CACHE] Returning stored solution for assignment #{assignment_id}")
        return {
            "success": True,
            "cached": True,
            "source": "this_assignment",
            "reference_solution": strip_excluded_guide_sections(
                json.loads(str(assignment.reference_solution_json))
            ),
        }

    # ── CACHE HIT 2: same content exists in another assignment ────────────────
    if assignment.content_hash:
        sibling = (
            db.query(Assignment)
            .filter(
                Assignment.content_hash == assignment.content_hash,
                Assignment.reference_solution_json.isnot(None),
                Assignment.status == AssignmentStatus.READY,
                Assignment.id != assignment_id,
            )
            .first()
        )
        if sibling:
            print(f"✅ [CACHE] Copying solution from sibling assignment #{sibling.id}")
            _sib_guide = strip_excluded_guide_sections(
                json.loads(str(sibling.reference_solution_json))
            )
            _sib_json = json.dumps(_sib_guide, ensure_ascii=False)
            assignment.reference_solution_json = _sib_json  # type: ignore
            assignment.reference_solution_text = _sib_json  # type: ignore
            assignment.solution_hash = sibling.solution_hash
            _orm_set(assignment, "is_locked", True)
            _orm_set(assignment, "status", AssignmentStatus.READY)
            db.commit()
            return {
                "success": True,
                "cached": True,
                "source": f"assignment_{sibling.id}",
                "reference_solution": strip_excluded_guide_sections(
                    json.loads(str(assignment.reference_solution_json))
                ),
            }

    # ── CACHE MISS: generate with AI (admin only) ───────────────────────────
    if not is_admin:
        raise HTTPException(
            status_code=403,
            detail=(
                "المهمة غير جاهزة بعد. يرجى التواصل مع الإدارة "
                "لإنشاء دليل المهمة أولاً."
            ),
        )

    textbook = db.query(Textbook).filter(Textbook.id == assignment.textbook_id).first()
    if not textbook:
        raise HTTPException(status_code=404, detail="Textbook not found")

    print(f"🤖 [AI] Generating new solution for assignment #{assignment_id}...")
    _orm_set(assignment, "status", AssignmentStatus.ANALYZING)
    db.commit()

    try:
        # Load filtered criteria if available
        filtered_criteria = None
        if assignment.unit_criteria_json:
            filtered_criteria = json.loads(str(assignment.unit_criteria_json))

        analysis, reference_solution, grading_criteria = (
            await process_textbook_and_assignment(
                textbook.file_path,  # type: ignore
                assignment.page_from or 1,
                assignment.page_to or 1,
                assignment.assignment_text or "",
                known_criteria_list=filtered_criteria,
                assignment_file_path=str(assignment.assignment_file_url).replace("/uploads/assignments/", "uploads/assignments/") if assignment.assignment_file_url else None,
            )
        )

        import hashlib as _hl
        reference_solution = strip_excluded_guide_sections(reference_solution)
        _orm_set(
            assignment,
            "reference_solution_text",
            json.dumps(reference_solution, ensure_ascii=False),
        )
        _orm_set(
            assignment,
            "reference_solution_json",
            json.dumps(reference_solution, ensure_ascii=False),
        )
        _ref_json = _orm_str(assignment.reference_solution_json)
        _orm_set(
            assignment,
            "solution_hash",
            _hl.sha256(_ref_json.encode("utf-8")).hexdigest(),
        )
        _orm_set(assignment, "is_locked", True)
        _orm_set(assignment, "status", AssignmentStatus.READY)

        # Save grading criteria (remove old ones first)
        db.query(models.GradingCriteria).filter(
            models.GradingCriteria.assignment_id == assignment_id
        ).delete()
        for criteria in grading_criteria:
            db.add(models.GradingCriteria(
                assignment_id=assignment.id,
                criteria_level=str(criteria["criteria_level"]),
                criteria_name=criteria["criteria_name"],
                criteria_description=criteria["criteria_description"],
                key_points=json.dumps(criteria["key_points"], ensure_ascii=False),
                weight=criteria["weight"],
            ))

        db.commit()
        log_activity(db, "create_assignment", "grading", f"إنشاء مهمة جديدة: {assignment.title} (ID: {assignment.id})", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="success")
        return {
            "success": True,
            "cached": False,
            "source": "ai_generated",
            "reference_solution": reference_solution,
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        _orm_set(assignment, "status", AssignmentStatus.FAILED)
        db.commit()
        log_activity(db, "create_assignment_error", "error", f"فشل إنشاء مهمة: {assignment.title} - {str(e)[:300]}", user_id=user_id, user_name=_get_user_display(db, user_id), user_email=_get_user_email(db, user_id), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request), level="error")
        raise HTTPException(status_code=500, detail=f"فشل في إنشاء دليل المهمة: {str(e)}")


@app.get("/api/download-reference-solution/{assignment_id}")
async def download_reference_solution(
    assignment_id: int, request: Request, db: Session = Depends(get_db)
):
    """Download the reference solution for an assignment as Word file"""
    try:
        uid = get_current_user_id(request)
        log_activity(db, "download_reference_solution", "export", f"تحميل دليل حل النموذجي - مهمة #{assignment_id}", user_id=uid, user_name=_get_user_display(db, uid), user_email=_get_user_email(db, uid), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request))
        return await _download_reference_solution_impl(assignment_id, db)
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating Word: {str(e)}")


def _normalize_slim_guide_sections(reference_data: dict) -> dict:
    """Map slim_v3 section 1 (and legacy keys) to sec1."""
    apply_slim_v3_section_titles(reference_data)
    sec1 = (
        reference_data.get("section_1_criteria_guide")
        or reference_data.get("section_3_criteria_guide")
        or reference_data.get("section_2_teacher_reference")
        or {}
    )
    if isinstance(sec1, dict):
        sec1 = {**sec1, "title": SLIM_V3_SECTION_TITLES[0]}
    return {"sec1": sec1, "sec2": {}, "sec3": {}}


def _structured_teacher_guide_has_body(reference_data: dict) -> bool:
    """True when slim guide JSON has real structured content."""
    norm = _normalize_slim_guide_sections(reference_data)
    s1 = norm["sec1"]
    degraded = (s1.get("_degraded_note") or "").strip()
    if degraded and not s1.get("criteria_guide"):
        return False
    if s1.get("criteria_guide"):
        return True
    legacy_extract = reference_data.get("section_1_criteria_extraction") or {}
    if legacy_extract.get("learning_outcomes"):
        return True
    s8 = reference_data.get("section_8_common_errors") or {}
    if s8.get("errors"):
        return True
    s10 = reference_data.get("section_10_student_model_answer") or {}
    if s10.get("sections"):
        return True
    s2 = reference_data.get("section_2_mission_interpretation") or {}
    if s2.get("skills_assessed"):
        return True
    obj = (s2.get("objective") or "").strip()
    if obj and (not degraded or degraded not in obj[:400]):
        return True
    return False


async def _download_reference_solution_impl(assignment_id: int, db: Session):
    from fastapi.responses import Response  # type: ignore
    from docx import Document  # type: ignore
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.enum.text import WD_BREAK  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    import io
    import os

    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    if not assignment.reference_solution_json:  # type: ignore
        raise HTTPException(status_code=404, detail="Reference solution not found")

    # Parse reference solution (strip legacy sections 6/7/9 if present in cache/DB)
    _raw_guide = json.loads(str(assignment.reference_solution_json))  # type: ignore
    reference_data = strip_excluded_guide_sections(_raw_guide)
    _slim_json = json.dumps(reference_data, ensure_ascii=False)
    _orig_json = json.dumps(_raw_guide, ensure_ascii=False)
    if _slim_json != _orig_json:
        assignment.reference_solution_json = _slim_json  # type: ignore
        assignment.reference_solution_text = _slim_json  # type: ignore
        db.commit()

    if _structured_teacher_guide_has_body(reference_data):
        reference_data.pop("markdown_guide", None)
        reference_data.pop("teacher_guide", None)

    _is_slim_guide = str(reference_data.get("guide_version", "")).startswith("slim_v")

    # Create Word document from pre-built RTL template
    import os as _os
    _rtl_template = _os.path.join(_os.path.dirname(__file__), "app", "templates", "rtl_template.docx")
    doc = Document(_rtl_template)

    # ── Gamma-inspired palette (modern presentation) ──
    PRIMARY = RGBColor(30, 27, 75)       # deep indigo
    ACCENT = RGBColor(6, 182, 212)     # cyan
    BODY_TEXT = RGBColor(30, 41, 59)
    WHITE = RGBColor(255, 255, 255)
    NAVY = RGBColor(30, 27, 75)
    NAVY_MID = RGBColor(79, 70, 229)
    SLATE = RGBColor(100, 116, 139)
    TEAL = RGBColor(13, 148, 136)
    GREEN = RGBColor(16, 185, 129)
    BLUE = RGBColor(59, 130, 246)
    PURPLE = RGBColor(139, 92, 246)
    INDIGO = RGBColor(99, 102, 241)
    RED = RGBColor(239, 68, 68)
    GOLD = RGBColor(245, 158, 11)
    SURFACE = RGBColor(248, 250, 252)

    # ── Global document settings (override template defaults) ──
    style = doc.styles['Normal']
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        _margin = Cm(1.5) if _is_slim_guide else Cm(2)
        _side = Cm(1.5) if _is_slim_guide else Cm(2.5)
        section.top_margin = _margin
        section.bottom_margin = _margin
        section.left_margin = _side
        section.right_margin = _side
        # Add footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("BTEC Teacher Guide  |  AI Grader Platform")
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(148, 163, 184)
        fr.font.name = "Calibri"
        fp2 = footer.add_paragraph()
        fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr2 = fp2.add_run("دليل المعلم — BTEC AI Grader")
        fr2.font.size = Pt(8)
        fr2.font.color.rgb = RGBColor(199, 210, 254)
        fr2.font.name = "Calibri"

    # ── Document-level RTL/bidi settings ──
    doc_settings = doc.settings.element
    bidi_doc = OxmlElement('w:bidi')
    bidi_doc.set(qn('w:val'), '1')
    doc_settings.append(bidi_doc)

    # ── Helper: set table visual RTL ──
    def set_table_bidi(tbl):
        tblPr = tbl._tbl.tblPr
        bidiVisual = OxmlElement('w:bidiVisual')
        bidiVisual.set(qn('w:val'), '1')
        tblPr.append(bidiVisual)

    # ── Helper functions ──
    def clean_text(text):
        """Remove markdown asterisks from text"""
        if not isinstance(text, str):
            return str(text)
        return text.replace('*', '')

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        # Remove existing bidi to avoid duplicates
        for existing in pPr.findall(qn('w:bidi')):
            pPr.remove(existing)
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        # REMOVE any jc - with bidi, Word default alignment IS right
        # Setting jc='right' in bidi context actually means LEFT!
        for existing in pPr.findall(qn('w:jc')):
            pPr.remove(existing)

    def set_run_cs_font(run, font_name='Calibri'):
        """Set Complex Script font for Arabic/RTL text on a run"""
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:cs'), font_name)
        # Also mark the run as RTL
        for existing in rPr.findall(qn('w:rtl')):
            rPr.remove(existing)
        rtl_el = OxmlElement('w:rtl')
        rtl_el.set(qn('w:val'), '1')
        rPr.append(rtl_el)

    def set_cell_shading(cell, hex_color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_borders(cell, color="CCCCCC", sz="4", sides=None):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in (sides or ('top', 'left', 'bottom', 'right')):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def clear_cell_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'FFFFFF')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def set_cell_vertical_alignment(cell, align="center"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), align)
        tcPr.append(va)

    def set_cell_margin(cell, top=0, bottom=0, start=0, end=0):
        """Set internal cell padding (margins) in twips (1pt = 20twips)"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for name, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
            mar = OxmlElement(f'w:{name}')
            mar.set(qn('w:w'), str(val))
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        tcPr.append(tcMar)

    section_counter = [0]

    def add_gamma_section_header(
        section_num: int,
        title: str,
        subtitle: str = "",
        accent_hex: str = "6366F1",
        light_bg: str = "EEF2FF",
    ):
        """Gamma-style section opener: numbered badge + title card."""
        section_counter[0] += 1
        num_str = f"{section_num:02d}"
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        set_table_bidi(tbl)
        badge_w = 1134
        force_table_full_width(tbl, [PAGE_WIDTH_TWIPS - badge_w, badge_w])
        content_cell = tbl.cell(0, 0)
        set_cell_shading(content_cell, light_bg)
        clear_cell_borders(content_cell)
        set_cell_borders(content_cell, color=accent_hex, sz="6", sides=("bottom",))
        set_cell_margin(content_cell, top=200, bottom=200, start=260, end=180)
        cp = content_cell.paragraphs[0]
        set_rtl(cp)
        tr = cp.add_run(clean_text(title))
        tr.bold = True
        tr.font.size = Pt(17)
        tr.font.color.rgb = PRIMARY
        tr.font.name = "Calibri"
        set_run_cs_font(tr)
        if subtitle:
            cp2 = content_cell.add_paragraph()
            set_rtl(cp2)
            sr = cp2.add_run(clean_text(subtitle))
            sr.font.size = Pt(10)
            sr.font.color.rgb = SLATE
            sr.font.name = "Calibri"
            set_run_cs_font(sr)
        badge_cell = tbl.cell(0, 1)
        set_cell_shading(badge_cell, accent_hex)
        clear_cell_borders(badge_cell)
        set_cell_vertical_alignment(badge_cell, "center")
        set_cell_margin(badge_cell, top=140, bottom=140, start=80, end=80)
        bp = badge_cell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bp.add_run(num_str)
        br.bold = True
        br.font.size = Pt(24)
        br.font.color.rgb = WHITE
        br.font.name = "Calibri"
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        return tbl

    def add_styled_heading(text, icon="", color_hex=None):
        section_counter[0] += 1
        accent = color_hex or "6366F1"
        return add_gamma_section_header(section_counter[0], text, accent_hex=accent)

    def add_rtl_paragraph(text, bold=False, color=None, size=None, font_name='Calibri'):
        p = doc.add_paragraph()
        set_rtl(p)
        run = p.add_run(clean_text(text))
        run.font.name = font_name
        run.font.color.rgb = color or BODY_TEXT
        if bold:
            run.bold = True
        run.font.size = size or Pt(11)
        set_run_cs_font(run, font_name)
        p.paragraph_format.space_after = Pt(3 if _is_slim_guide else 5)
        p.paragraph_format.line_spacing = 1.35 if _is_slim_guide else 1.6
        return p

    def add_rtl_bullet(text, indent_level=0, marker_color=None):
        p = doc.add_paragraph()
        set_rtl(p)
        bullet_chars = ["-", "-", "-"]
        bullet = bullet_chars[min(indent_level, 2)]
        indent_px = 18 * indent_level
        # Use right_indent for RTL layout (not left_indent)
        p.paragraph_format.right_indent = Pt(indent_px + 14)
        # Use ind start/end via XML for proper RTL hanging indent
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:start'), str(int((indent_px + 14) * 20)))  # twips
        ind.set(qn('w:hanging'), str(14 * 20))  # twips
        p.paragraph_format.space_after = Pt(2 if _is_slim_guide else 4)
        p.paragraph_format.line_spacing = 1.35 if _is_slim_guide else 1.5
        b_run = p.add_run(f"{bullet}  ")
        b_run.font.size = Pt(11)
        b_run.font.color.rgb = marker_color or TEAL
        b_run.font.name = 'Calibri'
        b_run.bold = True
        set_run_cs_font(b_run)
        t_run = p.add_run(clean_text(text))
        t_run.font.size = Pt(11)
        t_run.font.name = 'Calibri'
        t_run.font.color.rgb = BODY_TEXT
        set_run_cs_font(t_run)
        return p

    def add_sub_heading(text, color=None):
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_before = Pt(6 if _is_slim_guide else 12)
        p.paragraph_format.space_after = Pt(3 if _is_slim_guide else 5)
        p.paragraph_format.right_indent = Pt(6)
        if not _is_slim_guide:
            # Right border accent for RTL sub-heading (full layout only)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            right_bdr = OxmlElement('w:right')
            c_hex = color or NAVY_MID
            c_str = f"{c_hex[0]:02X}{c_hex[1]:02X}{c_hex[2]:02X}" if isinstance(c_hex, RGBColor) else '1E3A8A'
            right_bdr.set(qn('w:val'), 'single')
            right_bdr.set(qn('w:sz'), '18')
            right_bdr.set(qn('w:space'), '8')
            right_bdr.set(qn('w:color'), c_str)
            pBdr.append(right_bdr)
            pPr.append(pBdr)
        run = p.add_run(clean_text(text))
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = color or NAVY_MID
        run.font.name = 'Calibri'
        set_run_cs_font(run)
        return p

    # Total usable width: 21cm - 2.5cm - 2.5cm = 16cm = 9072 twips
    PAGE_WIDTH_TWIPS = 9072

    def force_table_full_width(tbl, col_widths_twips):
        """Force a table to span the full page width with fixed layout.
        col_widths_twips: list of column widths in twips, must sum to PAGE_WIDTH_TWIPS."""
        tblPr = tbl._tbl.tblPr
        # Remove any jc
        for jc in tblPr.findall(qn('w:jc')):
            tblPr.remove(jc)
        # Table width
        for old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(PAGE_WIDTH_TWIPS))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        # Fixed layout
        for old in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        # Zero indent
        for old in tblPr.findall(qn('w:tblInd')):
            tblPr.remove(old)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '0')
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        # Set column widths
        for i, w in enumerate(col_widths_twips):
            tbl.columns[i].width = Emu(w * 635)
            cell = tbl.cell(0, i)
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    def add_info_box(text, bg_color="F0F7FF", border_color="3B82F6", text_color=None):
        """Styled callout box with right accent (RTL)"""
        ACCENT_W = 227  # ~0.4cm
        CONTENT_W = PAGE_WIDTH_TWIPS - ACCENT_W

        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        set_table_bidi(tbl)
        force_table_full_width(tbl, [CONTENT_W, ACCENT_W])

        # Content (first column in RTL)
        cc = tbl.cell(0, 0)
        set_cell_shading(cc, bg_color)
        clear_cell_borders(cc)
        set_cell_margin(
            cc,
            top=60 if _is_slim_guide else 120,
            bottom=60 if _is_slim_guide else 120,
            start=120 if _is_slim_guide else 200,
            end=100 if _is_slim_guide else 160,
        )
        p = cc.paragraphs[0]
        set_rtl(p)
        run = p.add_run(clean_text(text))
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = text_color or BODY_TEXT
        set_run_cs_font(run)

        # Right accent (second column in RTL)
        ac = tbl.cell(0, 1)
        set_cell_shading(ac, border_color)
        clear_cell_borders(ac)
        ac.paragraphs[0].add_run(" ")
        ac.paragraphs[0].paragraph_format.space_before = Pt(0)
        ac.paragraphs[0].paragraph_format.space_after = Pt(0)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return tbl

    def add_separator():
        """Visible thin horizontal line separator between sections"""
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        # Add bottom border to paragraph for a visible line
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'E2E8F0')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_card_table(data_rows, headers, header_bg="0F1E4B", header_text_color=WHITE,
                       alt_bg_1="FFFFFF", alt_bg_2="F8FAFC", border_color="E2E8F0"):
        """Professional table with header and alternating rows"""
        num_cols = len(headers)
        tbl = doc.add_table(rows=len(data_rows) + 1, cols=num_cols)
        tbl.autofit = False
        # Set table RTL
        tblPr = tbl._tbl.tblPr
        bidiVisual = OxmlElement('w:bidiVisual')
        bidiVisual.set(qn('w:val'), '1')
        tblPr.append(bidiVisual)
        # Full width with equal columns
        col_w = PAGE_WIDTH_TWIPS // num_cols
        force_table_full_width(tbl, [col_w] * num_cols)
        # Header
        for ci, hdr_text in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            set_cell_shading(cell, header_bg)
            set_cell_borders(cell, color=header_bg, sz="4")
            set_cell_margin(cell, top=80, bottom=80, start=120, end=120)
            p = cell.paragraphs[0]
            set_rtl(p)
            r = p.add_run(clean_text(hdr_text))
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = header_text_color
            r.font.name = 'Calibri'
            set_run_cs_font(r)
        # Data rows
        for ri, row_data in enumerate(data_rows):
            bg = alt_bg_1 if ri % 2 == 0 else alt_bg_2
            for ci, cell_text in enumerate(row_data):
                cell = tbl.rows[ri + 1].cells[ci]
                set_cell_shading(cell, bg)
                set_cell_borders(cell, color=border_color, sz="2")
                set_cell_margin(cell, top=60, bottom=60, start=120, end=120)
                p = cell.paragraphs[0]
                set_rtl(p)
                r = p.add_run(clean_text(str(cell_text)))
                r.font.size = Pt(10)
                r.font.name = 'Calibri'
                r.font.color.rgb = BODY_TEXT
                set_run_cs_font(r)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return tbl

    # ════════════════════════════════════════════════════════════════
    #  COVER — compact for slim_v3 (matches PDF); full Gamma for legacy
    # ════════════════════════════════════════════════════════════════
    from datetime import datetime

    if _is_slim_guide:
        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ctr = cover_title.add_run("دليل المعلم للتقييم والتصحيح")
        ctr.bold = True
        ctr.font.size = Pt(20)
        ctr.font.color.rgb = PRIMARY
        ctr.font.name = "Calibri"
        set_run_cs_font(ctr)
        cover_sub = doc.add_paragraph()
        cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        csr = cover_sub.add_run("Teacher Assessment Guide")
        csr.font.size = Pt(10)
        csr.font.color.rgb = SLATE
        csr.font.name = "Calibri"
        cover_assign = doc.add_paragraph()
        cover_assign.alignment = WD_ALIGN_PARAGRAPH.CENTER
        car = cover_assign.add_run(f"المهمة: {assignment.title or ''}")
        car.bold = True
        car.font.size = Pt(13)
        car.font.color.rgb = NAVY_MID
        car.font.name = "Calibri"
        set_run_cs_font(car)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        foot_cover = doc.add_paragraph()
        foot_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fcr = foot_cover.add_run("سري — للاستخدام المعلّم فقط")
        fcr.font.size = Pt(8)
        fcr.font.color.rgb = SLATE
        fcr.italic = True
        set_run_cs_font(fcr)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
    else:
        logo_path = os.path.join(os.path.dirname(__file__), "app", "static", "images", "logo.png")

        hero_tbl = doc.add_table(rows=2, cols=1)
        hero_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hero_main = hero_tbl.cell(0, 0)
        set_cell_shading(hero_main, "312E81")
        clear_cell_borders(hero_main)
        set_cell_margin(hero_main, top=520, bottom=400, start=280, end=280)
        hp = hero_main.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(logo_path):
            lr = hp.add_run()
            lr.add_picture(logo_path, width=Inches(1.35))
            hp.add_run("\n")
        ht = hp.add_run("دليل المعلم للتقييم والتصحيح")
        ht.bold = True
        ht.font.size = Pt(28)
        ht.font.color.rgb = WHITE
        ht.font.name = "Calibri"
        set_run_cs_font(ht)
        he = hero_main.add_paragraph()
        he.alignment = WD_ALIGN_PARAGRAPH.CENTER
        her = he.add_run("BTEC Teacher Assessment Guide")
        her.font.size = Pt(11)
        her.font.color.rgb = RGBColor(199, 210, 254)
        her.font.name = "Calibri"

        hero_accent = hero_tbl.cell(1, 0)
        set_cell_shading(hero_accent, "06B6D4")
        clear_cell_borders(hero_accent)
        set_cell_margin(hero_accent, top=20, bottom=20, start=0, end=0)
        hero_accent.paragraphs[0].add_run(" ")

        doc.add_paragraph().paragraph_format.space_after = Pt(18)

        assign_tbl = doc.add_table(rows=1, cols=1)
        assign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        assign_cell = assign_tbl.cell(0, 0)
        set_cell_shading(assign_cell, "FFFFFF")
        set_cell_borders(assign_cell, color="6366F1", sz="16")
        set_cell_margin(assign_cell, top=280, bottom=280, start=320, end=320)
        ap = assign_cell.paragraphs[0]
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        al = ap.add_run("المهمة")
        al.font.size = Pt(10)
        al.font.color.rgb = SLATE
        al.font.name = "Calibri"
        set_run_cs_font(al)
        ap2 = assign_cell.add_paragraph()
        ap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = ap2.add_run(assignment.title or "")
        ar.bold = True
        ar.font.size = Pt(22)
        ar.font.color.rgb = NAVY_MID
        ar.font.name = "Calibri"
        set_run_cs_font(ar)

        doc.add_paragraph().paragraph_format.space_after = Pt(14)

        toc_tbl = doc.add_table(rows=1, cols=3)
        toc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_bidi(toc_tbl)
        tw = PAGE_WIDTH_TWIPS // 3
        force_table_full_width(toc_tbl, [tw, tw, PAGE_WIDTH_TWIPS - 2 * tw])
        toc_items = [
            ("01", "تفسير المعايير", "EEF2FF", "6366F1"),
            ("02", "المتطلبات النظرية", "ECFEFF", "0891B2"),
            ("03", "المتطلبات العملية", "ECFDF5", "10B981"),
        ]
        for ci, (num, label, bg, accent) in enumerate(toc_items):
            tc = toc_tbl.cell(0, ci)
            set_cell_shading(tc, bg)
            clear_cell_borders(tc)
            set_cell_borders(tc, color=accent, sz="4")
            set_cell_margin(tc, top=120, bottom=120, start=80, end=80)
            tp = tc.paragraphs[0]
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tn = tp.add_run(f"{num}\n")
            tn.bold = True
            tn.font.size = Pt(14)
            tn.font.color.rgb = RGBColor(
                int(accent[0:2], 16), int(accent[2:4], 16), int(accent[4:6], 16)
            )
            tl = tp.add_run(label)
            tl.font.size = Pt(9)
            tl.font.color.rgb = BODY_TEXT
            tl.font.name = "Calibri"
            set_run_cs_font(tl)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        meta_tbl = doc.add_table(rows=1, cols=2)
        meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_bidi(meta_tbl)
        mw = PAGE_WIDTH_TWIPS // 2
        force_table_full_width(meta_tbl, [mw, mw])
        for mi, (mlabel, mval) in enumerate(
            [
                ("التاريخ", datetime.now().strftime("%Y-%m-%d")),
                ("المنصة", "BTEC AI Grader"),
            ]
        ):
            mc = meta_tbl.cell(0, mi)
            set_cell_shading(mc, "F1F5F9")
            clear_cell_borders(mc)
            set_cell_margin(mc, top=80, bottom=80, start=100, end=100)
            mp = mc.paragraphs[0]
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mr1 = mp.add_run(f"{mlabel}\n")
            mr1.font.size = Pt(8)
            mr1.font.color.rgb = SLATE
            mr2 = mp.add_run(mval)
            mr2.bold = True
            mr2.font.size = Pt(10)
            mr2.font.color.rgb = PRIMARY
            mr2.font.name = "Calibri"

        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        foot_cover = doc.add_paragraph()
        foot_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fcr = foot_cover.add_run("سري — للاستخدام المعلّم فقط")
        fcr.font.size = Pt(9)
        fcr.font.color.rgb = SLATE
        fcr.italic = True
        set_run_cs_font(fcr)

        br_p = doc.add_paragraph()
        br_p.add_run().add_break(WD_BREAK.PAGE)

    guide_text = reference_data.get("markdown_guide") or reference_data.get("teacher_guide")
    use_markdown_primary = bool(
        guide_text and not _structured_teacher_guide_has_body(reference_data)
    )

    if use_markdown_primary:
        degraded_note = _normalize_slim_guide_sections(reference_data)["sec1"].get(
            "_degraded_note"
        )
        if degraded_note:
            add_info_box(
                clean_text(str(degraded_note)),
                bg_color="FEF3C7",
                border_color="D97706",
                text_color=RGBColor(146, 64, 14),
            )
        add_gamma_section_header(
            1, "دليل المعلم التفصيلي", "نسخة نصية كاملة", "6366F1", "EEF2FF"
        )
        for paragraph in str(guide_text).split("\n"):
            text = paragraph.strip()
            if not text:
                doc.add_paragraph()
                continue
            if text.startswith("### "):
                add_sub_heading(text.lstrip("#").strip(), color=NAVY_MID)
            elif text.startswith("## "):
                add_styled_heading(text.lstrip("#").strip(), color_hex="7C3AED")
            elif text.startswith("# "):
                add_styled_heading(text.lstrip("#").strip(), color_hex="1E3A8A")
            elif text.startswith("- ") or text.startswith("* "):
                add_rtl_bullet(text.lstrip("-* ").strip())
            else:
                add_rtl_paragraph(text)
    else:
        _render_structured_teacher_guide_sections(
            reference_data,
            add_gamma_section_header=add_gamma_section_header,
            add_styled_heading=add_styled_heading,
            add_sub_heading=add_sub_heading,
            add_rtl_paragraph=add_rtl_paragraph,
            add_rtl_bullet=add_rtl_bullet,
            add_info_box=add_info_box,
            add_card_table=add_card_table,
            add_separator=add_separator,
            clean_text=clean_text,
            set_rtl=set_rtl,
            set_run_cs_font=set_run_cs_font,
            set_table_bidi=set_table_bidi,
            force_table_full_width=force_table_full_width,
            set_cell_shading=set_cell_shading,
            clear_cell_borders=clear_cell_borders,
            set_cell_borders=set_cell_borders,
            set_cell_margin=set_cell_margin,
            set_cell_vertical_alignment=set_cell_vertical_alignment,
            doc=doc,
            PRIMARY=PRIMARY,
            NAVY_MID=NAVY_MID,
            SLATE=SLATE,
            TEAL=TEAL,
            GREEN=GREEN,
            BLUE=BLUE,
            PURPLE=PURPLE,
            INDIGO=INDIGO,
            RED=RED,
            BODY_TEXT=BODY_TEXT,
            PAGE_WIDTH_TWIPS=PAGE_WIDTH_TWIPS,
            RGBColor=RGBColor,
            Pt=Pt,
            WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
            WD_BREAK=WD_BREAK,
        )

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # URL encode filename for Arabic support (RFC 5987)
    from urllib.parse import quote

    encoded_filename = quote(f"{assignment.title}_Reference_Solution.docx")

    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
        "Access-Control-Expose-Headers": "Content-Disposition",
    }

    return Response(
        content=file_stream.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


def _render_structured_teacher_guide_sections(reference_data, **ctx):
    """Render sections 1–10 and legacy criterion keys into an open Document."""
    doc = ctx["doc"]
    add_styled_heading = ctx["add_styled_heading"]
    add_sub_heading = ctx["add_sub_heading"]
    add_rtl_paragraph = ctx["add_rtl_paragraph"]
    add_rtl_bullet = ctx["add_rtl_bullet"]
    add_info_box = ctx["add_info_box"]
    add_card_table = ctx["add_card_table"]
    add_separator = ctx["add_separator"]
    clean_text = ctx["clean_text"]
    set_rtl = ctx["set_rtl"]
    set_run_cs_font = ctx["set_run_cs_font"]
    set_table_bidi = ctx["set_table_bidi"]
    force_table_full_width = ctx["force_table_full_width"]
    set_cell_shading = ctx["set_cell_shading"]
    clear_cell_borders = ctx["clear_cell_borders"]
    set_cell_borders = ctx["set_cell_borders"]
    set_cell_margin = ctx["set_cell_margin"]
    set_cell_vertical_alignment = ctx["set_cell_vertical_alignment"]
    NAVY_MID = ctx["NAVY_MID"]
    SLATE = ctx["SLATE"]
    TEAL = ctx["TEAL"]
    GREEN = ctx["GREEN"]
    BLUE = ctx["BLUE"]
    PURPLE = ctx["PURPLE"]
    INDIGO = ctx["INDIGO"]
    RED = ctx["RED"]
    BODY_TEXT = ctx["BODY_TEXT"]
    PAGE_WIDTH_TWIPS = ctx["PAGE_WIDTH_TWIPS"]
    RGBColor = ctx["RGBColor"]
    Pt = ctx["Pt"]
    WD_ALIGN_PARAGRAPH = ctx["WD_ALIGN_PARAGRAPH"]
    WD_BREAK = ctx["WD_BREAK"]

    _cleaned_guide = strip_excluded_guide_sections(dict(reference_data))
    reference_data.clear()
    reference_data.update(_cleaned_guide)

    _gv = str(reference_data.get("guide_version", ""))
    _slim_guide = _gv.startswith("slim_v")

    # Legacy sections 1–2 (pre–slim guides only)
    sec1 = reference_data.get("section_1_criteria_extraction", {})
    if sec1.get("learning_outcomes"):
        add_styled_heading(sec1.get("title", "القسم الأول: استخراج المعايير"), color_hex="1E3A8A")
        for lo in sec1.get("learning_outcomes", []):
            add_sub_heading(lo.get("outcome_name", ""), color=NAVY_MID)
            crits = lo.get("criteria", [])
            if crits:
                data_rows = [[c.get('original_text', ''), c.get('code', '')] for c in crits]
                add_card_table(data_rows, ["نص المعيار", "الرمز"],
                               header_bg="1E3A8A", alt_bg_2="EEF2FF", border_color="C7D2FE")
        add_separator()

    sec2 = reference_data.get("section_2_mission_interpretation", {})
    if sec2.get("objective") or sec2.get("skills_assessed"):
        add_styled_heading(sec2.get("title", "القسم الثاني: تفسير المهمة للمعلم"), color_hex="7C3AED")
        if sec2.get("objective"):
            add_info_box(sec2.get("objective", ""), bg_color="F5F3FF", border_color="8B5CF6")
        if sec2.get("skills_assessed"):
            add_sub_heading("المهارات المقيّمة:", color=PURPLE)
            for skill in sec2.get("skills_assessed", []):
                add_rtl_bullet(skill, marker_color=PURPLE)
        add_separator()

    add_gamma = ctx.get("add_gamma_section_header") or add_styled_heading
    norm = _normalize_slim_guide_sections(reference_data)
    g1 = norm["sec1"]

    if g1 and (g1.get("criteria_guide") or g1.get("mission_intro")):
        if _slim_guide:
            sec_p = doc.add_paragraph()
            set_rtl(sec_p)
            sec_p.paragraph_format.space_before = Pt(4)
            sec_p.paragraph_format.space_after = Pt(6)
            sr = sec_p.add_run(clean_text(SLIM_V3_SECTION_TITLES[0]))
            sr.bold = True
            sr.font.size = Pt(14)
            sr.font.color.rgb = NAVY_MID
            sr.font.name = "Calibri"
            set_run_cs_font(sr)
        else:
            add_gamma(
                1,
                SLIM_V3_SECTION_TITLES[0],
                "المعايير، الأدلة، وما يبحث عنه المقيّم",
                "6366F1",
                "EEF2FF",
            )
        if g1.get("mission_intro"):
            add_info_box(g1.get("mission_intro", ""), bg_color="EEF2FF", border_color="6366F1")
        for guide in g1.get("criteria_guide", []):
            if _slim_guide:
                add_sub_heading(
                    clean_text(f"{guide.get('code', '')}  —  {guide.get('command_verb', '')}"),
                    color=NAVY_MID,
                )
            else:
                cr_tbl = doc.add_table(rows=1, cols=2)
                cr_tbl.autofit = False
                set_table_bidi(cr_tbl)
                force_table_full_width(cr_tbl, [PAGE_WIDTH_TWIPS - 227, 227])
                cr_cell = cr_tbl.cell(0, 0)
                set_cell_shading(cr_cell, "F8FAFC")
                clear_cell_borders(cr_cell)
                set_cell_borders(cr_cell, color="C7D2FE", sz="4", sides=('bottom',))
                cp = cr_cell.paragraphs[0]
                set_rtl(cp)
                # Right accent bar (col 1 = left side in RTL)
                cr_ac = cr_tbl.cell(0, 1)
                set_cell_shading(cr_ac, "6366F1")
                clear_cell_borders(cr_ac)
                cr_ac.paragraphs[0].add_run(" ")
                cr_ac.paragraphs[0].paragraph_format.space_before = Pt(0)
                cr_ac.paragraphs[0].paragraph_format.space_after = Pt(0)
                cr = cp.add_run(clean_text(f"{guide.get('code', '')}  —  {guide.get('command_verb', '')}"))
                cr.bold = True
                cr.font.size = Pt(13)
                cr.font.color.rgb = NAVY_MID
                cr.font.name = 'Calibri'
                set_run_cs_font(cr)
                cp.paragraph_format.space_before = Pt(8)
                cp.paragraph_format.space_after = Pt(8)

            if guide.get("what_student_must_do"):
                add_sub_heading("ما يجب على الطالب فعله:", color=GREEN)
                add_rtl_paragraph(guide.get("what_student_must_do", ""))

            if guide.get("required_evidence"):
                add_sub_heading("الأدلة المطلوبة:", color=INDIGO)
                for ev in guide.get("required_evidence", []):
                    add_rtl_bullet(ev, marker_color=INDIGO)

            if guide.get("assessor_look_for"):
                add_sub_heading("ما يبحث عنه المقيّم:", color=RED)
                for item in guide.get("assessor_look_for", []):
                    add_rtl_bullet(item, marker_color=RED)

            errs = guide.get("common_errors") or []
            if errs:
                add_sub_heading("أخطاء شائعة:", color=RED)
                for err in errs:
                    if isinstance(err, dict):
                        line = f"{err.get('error', '')} — تجنّب: {err.get('how_to_avoid', '')}"
                        add_rtl_bullet(line.strip(" —"), marker_color=RED)
                    else:
                        add_rtl_bullet(str(err), marker_color=RED)

            doc.add_paragraph().paragraph_format.space_after = Pt(2 if _slim_guide else 4)
        if not _slim_guide:
            add_separator()

    # Never render excluded legacy sections (2–3 theoretical/practical, 6/7/9).
    for _drop_key in (
        "section_2_theoretical",
        "section_3_theoretical",
        "section_3_practical",
        "section_4_theoretical",
        "section_4_practical",
        "section_5_practical",
    ):
        reference_data.pop(_drop_key, None)

    # Never render excluded legacy sections (6/7/9) even if stale JSON slipped through.
    for _legacy_key in (
        "section_5_evidence",
        "section_6_evidence",
        "section_6_grade_levels",
        "section_7_grade_levels",
        "section_7_model_answer_framework",
        "section_9_marking_checklist",
        "section_9_common_errors",
    ):
        reference_data.pop(_legacy_key, None)

    # ════════════════════════════════════════════════════════════
    #  SECTION 8: Common Errors
    # ════════════════════════════════════════════════════════════
    sec8 = reference_data.get("section_8_common_errors", {})
    if not _slim_guide and sec8.get("errors"):
        add_styled_heading(sec8.get("title", "القسم الثامن: أخطاء الطلاب الشائعة"), color_hex="DC2626")
        errors = sec8.get("errors", [])
        if errors:
            data_rows = [[e.get('error', ''), e.get('how_to_avoid', '')] for e in errors]
            add_card_table(data_rows, ["الخطأ الشائع", "كيفية تجنبه"],
                           header_bg="DC2626", alt_bg_2="FEF2F2", border_color="FECACA")
        add_separator()

    # ════════════════════════════════════════════════════════════
    #  SECTION 10: Student Model Answer Framework
    # ════════════════════════════════════════════════════════════
    sec10 = reference_data.get("section_10_student_model_answer", {})
    if not _slim_guide and (
        sec10.get("sections")
        or sec10.get("introduction")
    ):
        doc.add_page_break()
        add_styled_heading(sec10.get("title", "القسم العاشر: إطار الإجابة النموذجية للطالب"), color_hex="A855F7")
        if sec10.get("introduction"):
            add_info_box(sec10.get("introduction", ""), bg_color="FAF5FF", border_color="A855F7")
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        for idx, section in enumerate(sec10.get("sections", []), 1):
            # Section card with accent (RTL)
            s_tbl = doc.add_table(rows=1, cols=2)
            s_tbl.autofit = False
            set_table_bidi(s_tbl)
            s_accent_w = 227  # ~0.4cm
            force_table_full_width(s_tbl, [PAGE_WIDTH_TWIPS - s_accent_w, s_accent_w])
            # Content (col 0 = right side in RTL)
            s_cell = s_tbl.cell(0, 0)
            set_cell_shading(s_cell, "FAF5FF")
            clear_cell_borders(s_cell)
            set_cell_margin(s_cell, top=120, bottom=120, start=200, end=120)
            sp = s_cell.paragraphs[0]
            set_rtl(sp)
            sr = sp.add_run(clean_text(f"القسم {idx}  —  المعيار {section.get('criterion_code', '')}"))
            sr.bold = True
            sr.font.size = Pt(13)
            sr.font.color.rgb = PURPLE
            sr.font.name = 'Calibri'
            set_run_cs_font(sr)
            # Right accent bar (col 1 = left side in RTL)
            s_ac = s_tbl.cell(0, 1)
            set_cell_shading(s_ac, "A855F7")
            clear_cell_borders(s_ac)
            s_ac.paragraphs[0].add_run(" ")
            s_ac.paragraphs[0].paragraph_format.space_before = Pt(0)
            s_ac.paragraphs[0].paragraph_format.space_after = Pt(0)

            if section.get("section_title"):
                add_rtl_paragraph(section.get("section_title", ""), bold=True, size=Pt(12))

            if section.get("what_to_include"):
                add_sub_heading("ما يجب تضمينه:", color=GREEN)
                for item in section.get("what_to_include", []):
                    add_rtl_bullet(item, marker_color=GREEN)

            if section.get("required_evidence"):
                add_sub_heading("الأدلة المطلوبة:", color=INDIGO)
                for item in section.get("required_evidence", []):
                    add_rtl_bullet(item, marker_color=INDIGO)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ════════════════════════════════════════════════════════════
    #  FALLBACK: Old section_7_model_answer_framework format
    # ════════════════════════════════════════════════════════════
    old_sec7 = reference_data.get("section_7_model_answer_framework", {})
    if old_sec7 and not sec10:
        add_styled_heading(old_sec7.get("title", "إطار الإجابة النموذجية"), color_hex="A855F7")
        if old_sec7.get("report_structure"):
            add_sub_heading("هيكل التقرير:", color=NAVY_MID)
            for item in old_sec7.get("report_structure", []):
                add_rtl_bullet(item)
        if old_sec7.get("evidence_types"):
            add_sub_heading("أنواع الأدلة:", color=INDIGO)
            for item in old_sec7.get("evidence_types", []):
                add_rtl_bullet(item, marker_color=INDIGO)
        if old_sec7.get("general_examples"):
            add_sub_heading("أمثلة:", color=GREEN)
            for item in old_sec7.get("general_examples", []):
                add_rtl_bullet(item, marker_color=GREEN)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # FALLBACK: Old per-criterion format (P1, M1, D1, etc.)
    # ══════════════════════════════════════════════════
    skip_keys = {"metadata", "guide_version", "overall_structure", "word_count_guidance",
                 "markdown_guide",
                 "section_1_criteria_extraction", "section_1_criteria_guide",
                 "section_2_mission_interpretation", "section_2_theoretical",
                 "section_2_teacher_reference",
                 "section_3_criteria_guide", "section_3_theoretical", "section_3_practical",
                 "section_4_theoretical", "section_4_practical",
                 "section_5_practical", "section_5_evidence",
                 "section_6_evidence", "section_6_grade_levels",
                 "section_7_grade_levels", "section_7_model_answer_framework",
                 "section_8_common_errors", "section_9_marking_checklist",
                 "section_10_student_model_answer", "teacher_guide"}

    import re as _re

    def criteria_sort_key(key):
        match = _re.search(r"([PMD])(\d+)", key)
        if match:
            level_order = {"P": 0, "M": 1, "D": 2}
            return (level_order.get(match.group(1), 9), int(match.group(2)))
        return (9, 0)

    criterion_keys = sorted(
        [k for k in reference_data.keys() if k not in skip_keys and isinstance(reference_data[k], dict)],
        key=criteria_sort_key,
    )

    if criterion_keys:
        add_styled_heading("تفاصيل المعايير", color_hex="1E3A8A")
        for level_key in criterion_keys:
            criterion = reference_data[level_key]
            if not isinstance(criterion, dict):
                continue
            add_sub_heading(f"المعيار {level_key}", color=NAVY_MID)
            if criterion.get("command_verb"):
                add_info_box(f"الفعل المطلوب: {criterion['command_verb']}", bg_color="F5F3FF", border_color="A855F7")
            if criterion.get("answer"):
                add_sub_heading("الإجابة النموذجية:", color=NAVY_MID)
                for para in criterion["answer"].split("\n"):
                    if para.strip():
                        add_rtl_paragraph(para.strip())
            if criterion.get("evidence_requirements"):
                add_sub_heading("متطلبات الأدلة:", color=INDIGO)
                for req in criterion["evidence_requirements"]:
                    add_rtl_bullet(f"[ ] {req}", marker_color=INDIGO)
            if criterion.get("key_points"):
                add_sub_heading("النقاط الرئيسية:", color=GREEN)
                for point in criterion["key_points"]:
                    add_rtl_bullet(point)
            doc.add_paragraph()


def _resolve_arabic_pdf_font() -> str:
    """Pick a TTF that can render Arabic for ReportLab PDF export."""
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore

    candidates = [
        (os.path.join(os.path.dirname(__file__), "fonts", "NotoSansArabic-Regular.ttf"), "Arabic"),
        (r"C:\Windows\Fonts\tahoma.ttf", "Tahoma"),
        (r"C:\Windows\Fonts\arial.ttf", "Arial"),
        (r"C:\Windows\Fonts\trado.ttf", "TraditionalArabic"),
    ]
    for path, name in candidates:
        if not os.path.exists(path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            return name
        except Exception:
            continue
    return "Helvetica"


def _append_slim_v3_guide_pdf_elements(
    reference_data: dict,
    *,
    elements: list,
    _ar,
    heading_style,
    sub_heading_style,
    normal_style,
    bullet_style,
    colors,
    Spacer,
    HRFlowable,
    inch,
    Paragraph,
    ParagraphStyle,
) -> None:
    """Render slim_v3 section_1_criteria_guide into ReportLab flowables."""
    norm = _normalize_slim_guide_sections(reference_data)
    g1 = norm.get("sec1") or {}
    if not (g1.get("criteria_guide") or g1.get("mission_intro")):
        return

    elements.append(Paragraph(_ar(SLIM_V3_SECTION_TITLES[0]), heading_style))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=8))

    mission_intro = (g1.get("mission_intro") or "").strip()
    if mission_intro:
        elements.append(Paragraph(_ar(mission_intro), ParagraphStyle(
            "MissionIntro", parent=normal_style,
            backColor=colors.HexColor("#eef2ff"),
            borderColor=colors.HexColor("#6366f1"),
            borderWidth=1,
            borderPadding=8,
            spaceAfter=12,
        )))

    for guide in g1.get("criteria_guide", []):
        if not isinstance(guide, dict):
            continue
        code = str(guide.get("code") or "").strip()
        verb = str(guide.get("command_verb") or "").strip()
        header = f"{code} — {verb}" if code and verb else (code or verb or "معيار")
        elements.append(Paragraph(_ar(header), sub_heading_style))

        must_do = (guide.get("what_student_must_do") or "").strip()
        if must_do:
            elements.append(Paragraph(_ar("ما يجب على الطالب فعله:"), ParagraphStyle(
                "MustDoLabel", parent=normal_style, textColor=colors.HexColor("#15803d"),
            )))
            elements.append(Paragraph(_ar(must_do), normal_style))

        for label, key, color in (
            ("الأدلة المطلوبة:", "required_evidence", "#4338ca"),
            ("ما يبحث عنه المقيّم:", "assessor_look_for", "#b91c1c"),
        ):
            items = guide.get(key) or []
            if not items:
                continue
            elements.append(Paragraph(_ar(label), ParagraphStyle(
                "ListLabel", parent=normal_style, textColor=colors.HexColor(color),
            )))
            for item in items:
                elements.append(Paragraph(_ar(f"• {item}"), bullet_style))

        errs = guide.get("common_errors") or []
        if errs:
            elements.append(Paragraph(_ar("أخطاء شائعة:"), ParagraphStyle(
                "ErrLabel", parent=normal_style, textColor=colors.HexColor("#b91c1c"),
            )))
            for err in errs:
                if isinstance(err, dict):
                    line = f"{err.get('error', '')} — تجنّب: {err.get('how_to_avoid', '')}".strip(" —")
                else:
                    line = str(err)
                if line:
                    elements.append(Paragraph(_ar(f"• {line}"), bullet_style))

        elements.append(Spacer(1, 0.2 * inch))


@app.get("/api/download-reference-solution-pdf/{assignment_id}")
async def download_reference_solution_pdf(
    assignment_id: int, db: Session = Depends(get_db)
):
    """Download the reference solution / teacher guide for an assignment as PDF with RTL right-aligned content"""
    import io as _io
    from fastapi.responses import Response  # type: ignore
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
    from reportlab.lib.units import inch, cm  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable  # type: ignore
    from reportlab.lib import colors  # type: ignore
    from reportlab.lib.enums import TA_RIGHT, TA_CENTER  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
    from arabic_reshaper import reshape  # type: ignore
    from bidi.algorithm import get_display  # type: ignore

    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")
    if not assignment.reference_solution_json:  # type: ignore
        raise HTTPException(status_code=404, detail="Reference solution not found")

    reference_data = strip_excluded_guide_sections(
        json.loads(str(assignment.reference_solution_json))  # type: ignore
    )
    if _structured_teacher_guide_has_body(reference_data):
        reference_data.pop("markdown_guide", None)
        reference_data.pop("teacher_guide", None)

    # Register font (Arabic-capable TTF required — bundled Noto or Windows fallback)
    _arabic_font = _resolve_arabic_pdf_font()

    def _ar(text: str) -> str:
        """Reshape and apply BiDi to Arabic text for PDF rendering."""
        if not text:
            return ""
        try:
            reshaped = reshape(text)
            result = get_display(reshaped)
            return result if isinstance(result, str) else result.decode("utf-8")
        except Exception:
            return text

    buf = _io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=1.5 * cm, leftMargin=1.5 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "GuideTitle", parent=styles["Heading1"],
        fontName=_arabic_font, fontSize=22,
        textColor=colors.HexColor("#0f1e4b"),
        alignment=TA_CENTER, spaceAfter=16, spaceBefore=8,
        backColor=colors.HexColor("#eff6ff"),
        borderColor=colors.HexColor("#3b82f6"), borderWidth=2, borderPadding=10,
    )
    heading_style = ParagraphStyle(
        "GuideHeading", parent=styles["Heading2"],
        fontName=_arabic_font, fontSize=15,
        textColor=colors.HexColor("#7c3aed"),
        alignment=TA_RIGHT, spaceAfter=10, spaceBefore=16,
    )
    sub_heading_style = ParagraphStyle(
        "GuideSubHeading", parent=styles["Heading3"],
        fontName=_arabic_font, fontSize=13,
        textColor=colors.HexColor("#1e3a8a"),
        alignment=TA_RIGHT, spaceAfter=8, spaceBefore=12,
    )
    normal_style = ParagraphStyle(
        "GuideNormal", parent=styles["Normal"],
        fontName=_arabic_font, fontSize=11, leading=18,
        alignment=TA_RIGHT,
    )
    bullet_style = ParagraphStyle(
        "GuideBullet", parent=styles["Normal"],
        fontName=_arabic_font, fontSize=11, leading=16,
        alignment=TA_RIGHT, leftIndent=20,
    )

    elements: list = []

    # Cover title
    elements.append(Paragraph(_ar("دليل المعلم للتقييم والتصحيح"), title_style))
    elements.append(Paragraph(_ar("Teacher Assessment Guide"), ParagraphStyle(
        "CoverSub", parent=styles["Normal"], fontName=_arabic_font, fontSize=12,
        alignment=TA_CENTER, textColor=colors.HexColor("#6b7280"), spaceAfter=6,
    )))
    elements.append(Paragraph(_ar(f"المهمة: {assignment.title}"), ParagraphStyle(
        "AssignTitle", parent=styles["Normal"], fontName=_arabic_font, fontSize=13,
        alignment=TA_CENTER, textColor=colors.HexColor("#1e3a8a"), spaceAfter=20,
        fontWeight="bold",
    )))
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#3b82f6"), spaceAfter=20))

    def _add_criteria_section(data: dict):
        import re as _re_sort

        def _crit_sort(key):
            m = _re_sort.search(r"([PMD])(\d+)", str(key))
            if m:
                return ({'P': 0, 'M': 1, 'D': 2}.get(m.group(1), 9), int(m.group(2)))
            return (9, 0)

        sorted_items = sorted(data.items(), key=lambda x: _crit_sort(x[0]))
        for level_key, criterion in sorted_items:
            if not isinstance(criterion, dict):
                continue
            elements.append(Paragraph(_ar(f"المعيار: {level_key}"), sub_heading_style))
            if criterion.get("command_verb"):
                elements.append(Paragraph(_ar(f"الفعل الأمري: {criterion['command_verb']}"), normal_style))
            if criterion.get("answer"):
                elements.append(Paragraph(_ar("نموذج الإجابة:"), ParagraphStyle(
                    "AnswerLabel", parent=normal_style,
                    textColor=colors.HexColor("#1e3a8a"),
                )))
                for para in str(criterion["answer"]).split("\n"):
                    if para.strip():
                        elements.append(Paragraph(_ar(para.strip()), normal_style))
            if criterion.get("key_points"):
                elements.append(Paragraph(_ar("النقاط الرئيسية:"), normal_style))
                for pt in criterion["key_points"]:
                    elements.append(Paragraph(_ar(f"• {pt}"), bullet_style))
            if criterion.get("evidence_requirements"):
                elements.append(Paragraph(_ar("متطلبات الدليل:"), normal_style))
                for req in criterion["evidence_requirements"]:
                    elements.append(Paragraph(_ar(f"• {req}"), bullet_style))
            elements.append(Spacer(1, 0.2 * inch))

    # Sections
    sections_map = {
        "section_1_pass_criteria": "معايير النجاح (Pass)",
        "section_2_merit_criteria": "معايير الجدارة (Merit)",
        "section_3_distinction_criteria": "معايير التميز (Distinction)",
    }
    for sec_key, sec_label in sections_map.items():
        sec = reference_data.get(sec_key, {})
        if sec:
            elements.append(Paragraph(_ar(f"📋 {sec_label}"), heading_style))
            elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=8))
            _add_criteria_section(sec)

    if _structured_teacher_guide_has_body(reference_data):
        _append_slim_v3_guide_pdf_elements(
            reference_data,
            elements=elements,
            _ar=_ar,
            heading_style=heading_style,
            sub_heading_style=sub_heading_style,
            normal_style=normal_style,
            bullet_style=bullet_style,
            colors=colors,
            Spacer=Spacer,
            HRFlowable=HRFlowable,
            inch=inch,
            Paragraph=Paragraph,
            ParagraphStyle=ParagraphStyle,
        )

    # Teacher guide / markdown guide (legacy fallback only)
    guide_text = None if _structured_teacher_guide_has_body(reference_data) else (
        reference_data.get("markdown_guide") or reference_data.get("teacher_guide")
    )
    if guide_text:
        elements.append(Paragraph(_ar("📝 دليل المعلم التفصيلي"), heading_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=8))
        for line in str(guide_text).split("\n"):
            line = line.strip().lstrip("#").lstrip("-").strip()
            if line:
                elements.append(Paragraph(_ar(line), normal_style))
            else:
                elements.append(Spacer(1, 0.1 * inch))

    doc_pdf.build(elements)
    buf.seek(0)

    from urllib.parse import quote
    encoded_name = quote(f"{assignment.title}_دليل_الحل.pdf")
    return Response(
        content=buf.read(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}",
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


@app.post("/api/preflight-evidence/{assignment_id}")
async def preflight_evidence_scan(
    request: Request,
    assignment_id: int,
    file: Optional[UploadFile] = File(None),
    paths_json: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    """
    Fast path-only evidence scan before full PRO grading (no AI / vision / runtime).
    Accepts either a ZIP/RAR upload or a JSON list of relative paths (folder picks).
    """
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="يجب تسجيل الدخول أولاً")

    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="الواجب غير موجود")

    from app.preflight_evidence_scan import (  # type: ignore
        scan_relative_paths,
        scan_upload_archive,
    )

    try:
        if paths_json:
            import json as _json

            raw = _json.loads(paths_json)
            if isinstance(raw, dict):
                paths = raw.get("paths") or []
            elif isinstance(raw, list):
                paths = raw
            else:
                paths = []
            result = scan_relative_paths([str(p) for p in paths if p])
            return {"success": True, "preflight": result}

        if file and getattr(file, "filename", None):
            ext = Path(file.filename).suffix.lower()
            if ext not in (".zip", ".rar"):
                raise HTTPException(
                    status_code=400,
                    detail="الفحص السريع يدعم ملف ZIP أو RAR فقط — أو أرسل قائمة مسارات للمجلد.",
                )
            result = await scan_upload_archive(file)
            return {"success": True, "preflight": result}

        raise HTTPException(
            status_code=400,
            detail="أرسل ملف أرشيف أو paths_json لقائمة المسارات.",
        )
    except HTTPException:
        raise
    except ValueError as exc:
        code = str(exc)
        if code == "not_a_zip":
            raise HTTPException(status_code=400, detail="الملف ليس أرشيف ZIP صالحاً.")
        if code == "unsupported_archive":
            raise HTTPException(status_code=400, detail="صيغة الأرشيف غير مدعومة للفحص السريع.")
        raise HTTPException(status_code=400, detail=str(exc))
    except Exception as exc:
        print(f"⚠️ [PREFLIGHT] assignment={assignment_id} error: {exc}")
        raise HTTPException(status_code=500, detail="فشل الفحص السريع للأدلة.")


@app.post("/api/batch-grade/{assignment_id}")
async def batch_grade(
    request: Request,
    background_tasks: BackgroundTasks,
    assignment_id: int,
    files: List[UploadFile] = File(...),
    batch_name: str = Form(...),
    selected_criteria: str = Form("P1,P2,M1,D1"),
    folder_mode: str = Form("false"),
    folder_map: str = Form(""),
    single_student_archive: str = Form("false"),
    force_regrade: str = Form("false"),
    grading_mode: str = Form("deep"),
    db: Session = Depends(get_db),
):
    """Grade multiple students in batch"""
    _force_regrade = force_regrade == "true"
    print(
        f"📥 [BATCH-REQ] assignment={assignment_id} "
        f"parts={len(files)} folder_mode={folder_mode} "
        f"single_student_archive={single_student_archive} "
        f"force_regrade={_force_regrade} grading_mode={grading_mode}"
    )
    # --- Subscription check ---
    user_id = get_current_user_id(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="يجب تسجيل الدخول أولاً")

    # Admin bypass - no subscription needed (needed before lock for PRO vs BASIC)
    user = db.query(models.User).filter(models.User.id == user_id).first()
    is_admin = user and user.role == models.UserRole.ADMIN
    admin_grading_mode = (
        normalize_grading_mode_choice(grading_mode) if is_admin else None
    )
    _request_grading_mode = admin_grading_mode
    if _request_grading_mode is None:
        _sub_early = get_active_subscription(db, user_id)
        if _sub_early:
            _pkg_early = (
                db.query(models.Package)
                .filter(models.Package.id == _sub_early.package_id)
                .first()
            )
            _request_grading_mode = resolve_grading_policy(
                _pkg_early.name if _pkg_early else ""
            ).get("grading_mode", "deep")
        else:
            _request_grading_mode = "deep"

    _is_pro_request = _request_grading_mode == "deep"
    _is_pro_single_upload = _is_pro_request and (
        single_student_archive == "true"
        or (folder_mode != "true" and len(files) == 1)
    )

    from app.batch_grade_worker import (
        _ASSIGNMENT_JOB_GEN,
        assignment_batch_is_locked,
        pro_clear_assignment_lock_for_new_upload,
        release_assignment_batch_lock,
        request_batch_cancel,
        should_supersede_assignment_lock,
    )

    _upload_job_generation: int | None = None

    if _is_pro_request and (_is_pro_single_upload or _force_regrade or len(files) == 1):
        pro_clear_assignment_lock_for_new_upload(
            batch_progress,
            assignment_id,
            db,
            reason=(
                "pro_single_upload"
                if _is_pro_single_upload
                else ("pro_force_regrade" if _force_regrade else "pro_single_file")
            ),
        )
        _upload_job_generation = _ASSIGNMENT_JOB_GEN.get(assignment_id)
    elif _request_grading_mode == "fast":
        # BASIC: keep strict assignment lock (unchanged behaviour).
        if assignment_batch_is_locked(batch_progress, assignment_id, db):
            _active_batch = batch_progress.get(assignment_id)
            _total = int((_active_batch or {}).get("total") or 0)
            _lock_msg = (
                f"يوجد تصحيح BASIC جارٍ لهذا الواجب ({_total} طالب). "
                "انتظر حتى ينتهي ثم حاول مجدداً."
            )
            raise HTTPException(status_code=409, detail=_lock_msg)
    elif assignment_batch_is_locked(batch_progress, assignment_id, db):
        _active_batch = batch_progress.get(assignment_id)
        if should_supersede_assignment_lock(
            _active_batch,
            single_student_archive=single_student_archive == "true",
            force_regrade=_force_regrade,
        ):
            if _active_batch:
                request_batch_cancel(batch_progress, assignment_id)
            release_assignment_batch_lock(batch_progress, assignment_id)
            print(
                f"🔁 [BATCH-LOCK] superseded prior in-flight job for assignment={assignment_id}"
            )
        else:
            _total = int((_active_batch or {}).get("total") or 0)
            _lock_msg = (
                f"يوجد تصحيح PRO جارٍ لهذا الواجب ({_total} طالب). "
                "انتظر حتى ينتهي ثم حاول مجدداً."
            )
            raise HTTPException(status_code=409, detail=_lock_msg)

    sub: models.Subscription | None = None
    subject_bal: models.SubjectBalance | None = None  # used for per-subject deduction

    # Function-scope variable so quota enforcement after extraction can read it
    remaining_quota = None  # type: ignore

    if not is_admin:
        # Fetch assignment early to get its subject for per-subject balance check
        _asgn_early = db.query(Assignment).filter(Assignment.id == assignment_id).first()
        asgn_subject = (_asgn_early.subject or "").strip() if _asgn_early else ""

        if asgn_subject:
            subject_bal = (
                db.query(models.SubjectBalance)
                .filter(
                    models.SubjectBalance.user_id == user_id,
                    models.SubjectBalance.subject == asgn_subject,
                )
                .first()
            )

        if subject_bal is not None:
            # Per-subject enforcement
            remaining = max(0, (subject_bal.assignments_limit or 0) - (subject_bal.assignments_used or 0))
            remaining_quota = remaining
            if remaining <= 0:
                return JSONResponse(
                    {
                        "success": False,
                        "subscription_required": True,
                        "detail": f"لقد استنفدت رصيدك لمادة {asgn_subject}. يرجى تجديد الاشتراك.",
                    },
                    status_code=403,
                )
        else:
            # Fall back to global subscription (no per-subject balance exists)
            sub = get_active_subscription(db, user_id)
            if sub is None:
                return JSONResponse(
                    {
                        "success": False,
                        "subscription_required": True,
                        "detail": "يجب الاشتراك أولاً لتصحيح الواجبات",
                    },
                    status_code=403,
                )
            remaining = (sub.assignments_limit or 0) - (sub.assignments_used or 0)
            remaining_quota = remaining
            if remaining <= 0:
                return JSONResponse(
                    {
                        "success": False,
                        "subscription_required": True,
                        "detail": "لقد استنفدت عدد الواجبات المتاح. يرجى تجديد الاشتراك.",
                    },
                    status_code=403,
                )
    # --- End subscription check ---
    try:
        # Validate file count
        # ── Student-count limit (NOT file-count limit) ──
        # In folder mode the user uploads ONE folder per student, but the
        # folder may contain 50+ files (images, Word, Packet Tracer, ...).
        # We must count STUDENTS, not raw files.
        MAX_STUDENTS_PER_BATCH = 30
        if folder_mode == "true" and folder_map:
            try:
                _fm_check = json.loads(folder_map)
                _student_count_estimate = len(_fm_check) if isinstance(_fm_check, dict) else 0
            except Exception:
                _student_count_estimate = 1
        else:
            # Regular mode: each file is one student.
            # Archives (.zip/.rar) count as 1 here — actual student count is
            # validated later after extraction.
            _student_count_estimate = len(files)
        if _student_count_estimate > MAX_STUDENTS_PER_BATCH:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"لا يمكن تصحيح أكثر من {MAX_STUDENTS_PER_BATCH} طالب في الدفعة الواحدة. "
                    f"العدد المُحدَّد: {_student_count_estimate} طالب. "
                    f"الرجاء تقليل العدد أو رفعه على دفعات أصغر."
                ),
            )

        from app.project_intelligence.submission_intake import (  # type: ignore
            INTAKE_IGNORE_DIR_NAMES,
            analyze_multipart_upload_manifest,
            get_ingestion_limits,
            intake_rejection_message,
            path_matches_intake_ignore,
        )

        _ing_limits = get_ingestion_limits()
        # Folder uploads: skip reading every file twice for manifest (can be thousands of parts).
        if folder_mode == "true" and folder_map:
            _est_bytes = 0
            for _uf in files:
                try:
                    _sp = getattr(_uf, "file", None)
                    if _sp is not None:
                        _pos = _sp.tell()
                        _sp.seek(0, 2)
                        _est_bytes += int(_sp.tell())
                        _sp.seek(_pos)
                except Exception:
                    pass
            batch_intake_manifest = {
                "upload_diagnostics": {
                    "multipart_parts": len(files),
                    "total_bytes_uploaded": _est_bytes,
                    "folder_mode": True,
                },
                "submission_noise_flags": [],
            }
        else:
            batch_intake_manifest = await analyze_multipart_upload_manifest(files)
        if len(files) > _ing_limits["max_multipart_files"]:
            return JSONResponse(
                {
                    "success": False,
                    "intake_rejected": True,
                    "detail": intake_rejection_message(
                        "too_many_multipart_parts",
                        limit=_ing_limits["max_multipart_files"],
                        actual=len(files),
                    ),
                    "upload_diagnostics": batch_intake_manifest.get("upload_diagnostics"),
                    "submission_noise_flags": batch_intake_manifest.get("submission_noise_flags"),
                },
                status_code=413,
            )
        _ing_total_bytes = int(batch_intake_manifest["upload_diagnostics"].get("total_bytes_uploaded") or 0)
        if _ing_total_bytes > _ing_limits["max_upload_bytes_total"]:
            return JSONResponse(
                {
                    "success": False,
                    "intake_rejected": True,
                    "detail": intake_rejection_message(
                        "upload_too_large",
                        limit=_ing_limits["max_upload_bytes_total"],
                        actual=_ing_total_bytes,
                    ),
                    "upload_diagnostics": batch_intake_manifest.get("upload_diagnostics"),
                    "submission_noise_flags": batch_intake_manifest.get("submission_noise_flags"),
                },
                status_code=413,
            )

        # Get assignment
        assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
        if not assignment or assignment.status != AssignmentStatus.READY:  # type: ignore
            raise HTTPException(status_code=404, detail="Assignment not ready")

        # Parse reference solution
        reference_solution = json.loads(str(assignment.reference_solution_json))  # type: ignore

        # Get grading criteria
        criteria = (
            db.query(models.GradingCriteria)
            .filter(models.GradingCriteria.assignment_id == assignment_id)
            .all()
        )

        def _criteria_sort_key(item):
            lv = item.get('criteria_level', '') if isinstance(item, dict) else (item.criteria_level or '')
            short = lv.split('.')[-1] if '.' in lv else lv
            _order = {'P': 0, 'M': 1, 'D': 2}
            letter = short[0].upper() if short else 'Z'
            try:
                num = int(short[1:]) if len(short) > 1 else 0
            except ValueError:
                num = 99
            return (_order.get(letter, 9), num)

        grading_criteria = []
        existing_levels = set()
        for c in criteria:
            grading_criteria.append(
                {
                    "criteria_level": c.criteria_level,  # Fix: It's a string now, not an Enum
                    "criteria_name": c.criteria_name,
                    "criteria_description": c.criteria_description,
                    "key_points": json.loads(str(c.key_points)),  # type: ignore
                }
            )
            existing_levels.add(c.criteria_level)

        grading_criteria.sort(key=_criteria_sort_key)

        # Fallback: only add default criteria if NO criteria found
        if not existing_levels:
            default_criteria = {
                "P1": {
                    "name": "الوصف (Pass 1)",
                    "description": "وصف المفاهيم والمبادئ الأساسية",
                },
                "P2": {
                    "name": "الشرح (Pass 2)",
                    "description": "شرح كيفية عمل المفاهيم والعمليات",
                },
                "M1": {
                    "name": "التحليل (Merit 1)",
                    "description": "تحليل العلاقات والروابط بين المفاهيم",
                },
                "D1": {
                    "name": "التقييم (Distinction 1)",
                    "description": "تقييم نقدي وتوصيات مدعمة بالأدلة",
                },
            }

            for level, info in default_criteria.items():
                if level not in existing_levels:
                    grading_criteria.append(
                        {
                            "criteria_level": level,
                            "criteria_name": info["name"],
                            "criteria_description": info["description"],
                            "key_points": [],  # Will rely on reference solution
                        }
                    )

        # Create batch record
        batch = BatchGrading(
            assignment_id=assignment_id,
            batch_name=batch_name,
            total_students=len(files),
            status=BatchStatus.PROCESSING,
            created_by=user_id,
        )

        db.add(batch)
        db.commit()
        db.refresh(batch)

        # Save uploaded files and prepare student info
        # Check for existing graded submissions for this assignment
        from app.archive_extraction_utils import hash_submission_file as _hash_submission_file
        from app.grading_snapshot_governance import submission_replay_eligible

        existing_submissions = (
            db.query(Submission)
            .filter(
                Submission.assignment_id == assignment_id,
                Submission.status == SubmissionStatus.COMPLETED,
            )
            .order_by(Submission.id.desc())
            .all()
        )
        # Map by student_name — keep newest submission per name; hash from file bytes
        existing_by_name = {}
        _skip_submission_replay = _force_regrade or _is_pro_request
        if _skip_submission_replay:
            print(
                "🔄 [PRO/FORCE] Skipping submission replay cache — full re-grade "
                f"(force_regrade={_force_regrade}, pro={_is_pro_request})"
            )
        for s in existing_submissions:
            if _skip_submission_replay or s.student_name in existing_by_name:
                continue
            summary = (
                db.query(GradingSummary)
                .filter(GradingSummary.submission_id == s.id)
                .first()
            )
            if not submission_replay_eligible(s, summary):
                continue
            source_hash = _hash_submission_file(
                str(s.submission_file_path or ""),
                "",
            )
            existing_by_name[s.student_name] = {
                "submission": s,
                "source_hash": source_hash,
            }

        student_files = []
        cached_results = []

        # Helper: extract files from ZIP/RAR archive
        import zipfile

        # Folders/paths to skip (build artifacts, IDE files, OS junk) — shared with intake layer
        _SKIP_DIRS = set(INTAKE_IGNORE_DIR_NAMES)
        # Document extensions that are actual assignment files
        _DOC_EXTENSIONS = ('.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt')
        # Code file extensions — covers all BTEC IT (Jordan) units
        _CODE_EXTENSIONS = (
            '.py', '.java', '.cs', '.cpp', '.c', '.h', '.hpp',
            '.js', '.ts', '.jsx', '.tsx', '.html', '.css', '.scss',
            '.rb', '.go', '.rs', '.php', '.swift', '.kt', '.scala', '.r',
            '.sql', '.sh', '.bat', '.ps1',
            '.xml', '.yaml', '.yml', '.ini', '.cfg', '.conf', '.log',
            '.csv', '.tsv', '.toml',
            '.dart', '.gradle', '.kts',                       # Flutter / Android
            '.gml', '.gd', '.uc',                             # Game Dev
            '.ipynb', '.rmd',                                 # Data Science
            '.vue', '.svelte', '.razor', '.cshtml',           # Web Frameworks
        )
        # Image extensions
        _IMAGE_EXTENSIONS = (
            '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff',
            '.svg', '.ico', '.heic', '.heif', '.psd', '.ai',
        )
        # Networking simulation files (Cisco Packet Tracer, etc.)
        _NETWORKING_EXTENSIONS = ('.pkt', '.pka', '.pcap', '.pcapng', '.cap')
        # Game/Multimedia project files (Unity, Unreal, GameMaker, etc.)
        _GAME_PROJECT_EXTENSIONS = (
            '.unity', '.prefab', '.scene', '.uasset', '.umap',
            '.gms2', '.yyp', '.gmproj', '.fla', '.swf', '.gma',
            '.rbxl', '.rbxmx', '.sb3',
        )
        # Runnable build artifacts — extracted for inventory, NOT auto-executed
        from app.artifact_inventory import EXECUTABLE_ARTIFACT_EXTENSIONS  # type: ignore

        _EXECUTABLE_ARTIFACT_EXTENSIONS = tuple(EXECUTABLE_ARTIFACT_EXTENSIONS)
        # Database files
        _DATABASE_EXTENSIONS = ('.db', '.sqlite', '.sqlite3', '.mdb', '.accdb')
        # Audio / Video — for evidence
        _MEDIA_EXTENSIONS = (
            '.mp3', '.wav', '.ogg', '.m4a', '.flac',
            '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.webm', '.m4v',
        )
        _VIDEO_INTAKE_EXTENSIONS = (
            '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.webm', '.m4v', '.flv',
        )
        # All supported extensions for student submissions
        _ALL_EXTENSIONS = (
            _DOC_EXTENSIONS + ('.txt', '.md', '.json', '.odt', '.ods', '.odp')
            + _CODE_EXTENSIONS + _IMAGE_EXTENSIONS
            + _NETWORKING_EXTENSIONS + _GAME_PROJECT_EXTENSIONS
            + _DATABASE_EXTENSIONS + _MEDIA_EXTENSIONS
            + _EXECUTABLE_ARTIFACT_EXTENSIONS
        )
        _INTAKE_EXTENSIONS = _ALL_EXTENSIONS

        def _is_in_skip_dir(filepath: str) -> bool:
            """Check if file is inside a build/system directory."""
            parts = Path(filepath).parts
            return any(p.lower() in _SKIP_DIRS for p in parts)

        def _merge_intake_videos_from_display(
            all_arc: list,
            folder_prefix: str | None,
            display_files: list | None,
        ) -> None:
            """Add video paths from archive listing (skipped on disk in BASIC) for keyframe vision."""
            if not display_files:
                return
            seen = {a.replace("\\", "/").lower() for a in all_arc}
            prefix = f"{folder_prefix}/" if folder_prefix else None
            for disp in display_files:
                norm = disp.replace("\\", "/")
                if not any(norm.lower().endswith(ext) for ext in _VIDEO_INTAKE_EXTENSIONS):
                    continue
                if prefix and not norm.startswith(prefix):
                    continue
                if prefix is None and "/" in norm:
                    continue
                if norm.lower() not in seen:
                    seen.add(norm.lower())
                    all_arc.append(norm)

        def _pick_best_file(file_list: list) -> tuple | None:
            """From a list of rows (student_name, path[, archive_rel]) for ONE student folder,
            pick the best assignment file (doc → game .exe → source code).
            Returns (student_name, path) or None."""
            from app.archive_extraction_utils import pick_best_submission_file

            return pick_best_submission_file(
                file_list,
                doc_extensions=_DOC_EXTENSIONS,
            )

        def _collect_archive_entries(archive_obj, archive_type: str,
                                     all_display_files=None) -> list:
            """Collect valid file entries from archive.
            Returns [(archive_path, decoded_name)].
            If all_display_files list is provided, appends ALL visible
            filenames (no extension filter) for UI display."""
            entries = []
            for info in archive_obj.infolist():
                if info.is_dir():
                    continue
                raw_path = info.filename
                # Handle Arabic filenames encoding (ZIP cp437 issue)
                if archive_type == 'zip':
                    try:
                        raw_bytes = raw_path.encode('cp437')
                        decoded_path = raw_bytes.decode('utf-8')
                    except (UnicodeDecodeError, UnicodeEncodeError):
                        decoded_path = raw_path
                else:
                    decoded_path = raw_path
                from app.archive_extraction_utils import collapse_redundant_archive_path

                decoded_path = collapse_redundant_archive_path(decoded_path)

                fname = Path(decoded_path).name
                # Skip hidden/system files
                if fname.startswith('.') or fname.startswith('__'):
                    continue
                # Skip build/system directories
                if _is_in_skip_dir(decoded_path):
                    continue
                # Add ALL visible files for UI display
                if all_display_files is not None:
                    all_display_files.append(decoded_path)
                # Only supported file types for extraction
                fname_lower = fname.lower()
                if fname_lower.endswith(".exe"):
                    from app.archive_extraction_utils import is_primary_game_executable

                    if not is_primary_game_executable(decoded_path):
                        continue
                if not any(fname_lower.endswith(e) for e in _ALL_EXTENSIONS):
                    continue
                try:
                    from app.grading_mode_policy import should_skip_archive_extract_to_disk

                    if should_skip_archive_extract_to_disk(
                        decoded_path, _pipeline_grading_mode
                    ):
                        continue
                except Exception:
                    pass
                entries.append((info, decoded_path))
            return entries

        # Shared dict for archive display info (populated by extract_archive)
        _archive_display_data = {'all_files': [], 'student_map': {}}
        _archive_pipeline_errors: list[str] = []

        # Track which students have code / executable artifacts in their archive folders
        _student_has_code_files = {}  # student_name -> bool
        _student_has_executable_artifacts: dict[str, bool] = {}
        _PROGRAMMING_EXTS = {
            '.py', '.java', '.cs', '.cpp', '.c', '.js', '.ts',
            '.html', '.jsx', '.tsx', '.rb', '.go', '.php',
            '.gml', '.gd',  # GameMaker / Godot scripts
            '.lua',         # Common in game curricula
        }
        _NESTED_ARCHIVE_EXTENSIONS = ('.zip', '.rar')

        def _collect_nested_archives(scan_root: Path) -> list[tuple[str, str]]:
            """Find ZIP/RAR members inside an already-extracted tree (bundle-of-archives)."""
            from app.archive_extraction_utils import collapse_redundant_archive_path

            out: list[tuple[str, str]] = []
            if not scan_root.exists():
                return out
            for fpath in scan_root.rglob('*'):
                if not fpath.is_file():
                    continue
                if fpath.suffix.lower() not in _NESTED_ARCHIVE_EXTENSIONS:
                    continue
                rel = collapse_redundant_archive_path(
                    fpath.relative_to(scan_root).as_posix()
                )
                fname = Path(rel).name
                if fname.startswith('.') or fname.startswith('__'):
                    continue
                if _is_in_skip_dir(rel):
                    continue
                out.append((rel, str(fpath)))
            return out

        def _apply_archive_display(all_display_files: list[str]) -> None:
            display_names: list[str] = []
            student_map: dict[str, list[str]] = {}
            for fpath in all_display_files:
                fname = Path(fpath).name
                parts = Path(fpath).parts
                if len(parts) > 1:
                    student = parts[0]
                else:
                    student = Path(fpath).stem
                display_names.append(fname)
                student_map.setdefault(student, []).append(fname)
            _archive_display_data['all_files'] = display_names
            _archive_display_data['student_map'] = student_map

        def _merge_archive_display(other: dict) -> None:
            for fname in other.get('all_files') or []:
                if fname not in _archive_display_data['all_files']:
                    _archive_display_data['all_files'].append(fname)
            for student, fnames in (other.get('student_map') or {}).items():
                bucket = _archive_display_data['student_map'].setdefault(student, [])
                for fname in fnames:
                    if fname not in bucket:
                        bucket.append(fname)

        _pipeline_grading_mode = admin_grading_mode
        if _pipeline_grading_mode is None and sub is not None:
            try:
                _pkg_row = (
                    db.query(models.Package)
                    .filter(models.Package.id == sub.package_id)
                    .first()
                )
                _pipeline_grading_mode = resolve_grading_policy(
                    _pkg_row.name if _pkg_row else None
                ).get("grading_mode", "deep")
            except Exception:
                _pipeline_grading_mode = "deep"
        if _pipeline_grading_mode is None:
            _pipeline_grading_mode = "deep"
        if _pipeline_grading_mode == "fast":
            print("⚡ [BASIC] Selective archive extract — Word/code only (skip exe/images/media)")

        def extract_archive(archive_path: str, batch_id: int, *, _depth: int = 0) -> list:
            """Extract files from ZIP/RAR and return list of
            (student_name, file_path, submission_paths[], intake_relative_paths[]) tuples.

            Smart extraction logic:
            - If archive contains folders, each top-level folder = one student
            - If archive contains flat files, each file = one student
            - Skips build artifacts (bin/, obj/, .vs/, Debug/, etc.)
            - Picks the best document file per student folder (.docx/.pdf preferred)
            """
            ext = Path(archive_path).suffix.lower()
            # Short extract root — avoids WinError 206 with deep Arabic folder names
            extract_dir = STUDENTS_DIR / f"bx{batch_id}"
            if _depth > 0:
                import hashlib

                _apath = str(Path(archive_path).resolve())
                _digest = hashlib.sha1(_apath.encode("utf-8", errors="replace")).hexdigest()[:10]
                extract_dir = extract_dir / "_nested" / _digest
            extract_dir.mkdir(parents=True, exist_ok=True)

            # وضع «طالب واحد → ZIP/RAR»: مجلدان (مثل قبل/بعد التعديل) ≠ طالبان
            _is_single_student_archive = single_student_archive == "true"

            def _progress_student_total(raw_groups: int) -> int:
                if _is_single_student_archive:
                    return 1
                return max(raw_groups, 1)

            def _emit_archive_progress(
                *,
                student: str = "",
                frac: float | None = None,
                phase_label: str = "",
                nested_idx: int | None = None,
                nested_total: int | None = None,
                listed: int = 0,
                sub_phase: str = "",
            ) -> None:
                from app.archive_extraction_utils import archive_ui_percent

                _pi = batch_progress.get(assignment_id) or {}
                _pi["current_phase"] = "extracting_archive"
                if student:
                    _pi["current_student"] = student
                if phase_label:
                    _pi["phase_label"] = phase_label
                if nested_total is not None and nested_total > 0:
                    _pi["total"] = _progress_student_total(nested_total)
                    if nested_idx is not None:
                        frac = (nested_idx + 1) / nested_total
                if sub_phase == "listing":
                    _pi["percent"] = max(
                        int(_pi.get("percent") or 0),
                        archive_ui_percent(phase="listing", listed=listed),
                    )
                elif sub_phase == "manifest":
                    _pi["percent"] = max(
                        int(_pi.get("percent") or 0),
                        archive_ui_percent(phase="manifest"),
                    )
                elif frac is not None:
                    _pi["percent"] = max(
                        int(_pi.get("percent") or 0),
                        archive_ui_percent(phase="extract", frac=frac),
                    )
                batch_progress[assignment_id] = _pi

            def _archive_out_path(decoded_path: str) -> Path:
                from app.archive_extraction_utils import safe_archive_out_path

                out, _ = safe_archive_out_path(extract_dir, decoded_path)
                return out

            def _write_archive_member(data: bytes, decoded_path: str) -> tuple[str, str]:
                from app.archive_extraction_utils import safe_archive_out_path, write_bytes_safe

                out, rel_key = safe_archive_out_path(extract_dir, decoded_path)
                write_bytes_safe(out, data)
                return rel_key, str(out)

            entries = []  # [(info, decoded_path), ...]
            all_display_files = []  # ALL visible files for UI

            if ext == '.zip':
                from app.archive_extraction_utils import selective_extract_zip

                def _emit_zip_manifest(student_groups: int, files_to_extract: int) -> None:
                    groups = max(student_groups, 1)
                    folder_note = ""
                    if _is_single_student_archive and groups > 1:
                        folder_note = f" ({groups} مجلدات داخل أرشيف طالب واحد)"
                    _emit_archive_progress(
                        phase_label=(
                            f"جاري استخراج {files_to_extract} ملفاً"
                            f"{folder_note}..."
                        ),
                        nested_total=groups,
                    )

                def _emit_zip_progress(done: int, total: int, name: str) -> None:
                    frac = (done + 1) / max(total, 1)
                    _emit_archive_progress(student=name, frac=frac)

                extracted_files: list = []
                _zip_bytes = os.path.getsize(archive_path)
                with zipfile.ZipFile(archive_path, 'r') as zf:
                    _entry_count = sum(1 for i in zf.infolist() if not i.is_dir())
                from app.archive_extraction_utils import (
                    archive_should_use_selective_extract,
                    max_archive_extract_files,
                )

                _use_selective = archive_should_use_selective_extract(
                    _entry_count, _zip_bytes, _pipeline_grading_mode
                )
                if _use_selective:
                    sel_files, sel_display = selective_extract_zip(
                        archive_path,
                        extract_dir,
                        skip_dir_names=frozenset(_SKIP_DIRS),
                        gradable_extensions=_ALL_EXTENSIONS,
                        on_progress=_emit_zip_progress,
                        on_manifest=_emit_zip_manifest,
                        grading_mode=_pipeline_grading_mode,
                        max_extract_files=max_archive_extract_files(
                            _pipeline_grading_mode, archive_bytes=_zip_bytes
                        ),
                    )
                    extracted_files = sel_files
                    if sel_display:
                        all_display_files.extend(sel_display)
                    print(
                        f"📦 [ZIP-SEL] {Path(archive_path).name}: "
                        f"{_entry_count} entries → {len(extracted_files)} gradable file(s)"
                    )
                else:
                    with zipfile.ZipFile(archive_path, 'r') as zf:
                        entries = _collect_archive_entries(
                            zf, 'zip',
                            all_display_files=all_display_files)
                        _zip_total = len(entries)
                        for _zi, (info, decoded_path) in enumerate(entries):
                            if _zi == 0 or (_zi + 1) % 3 == 0 or _zi + 1 == _zip_total:
                                _emit_archive_progress(
                                    student=Path(decoded_path).name,
                                    frac=(_zi + 1) / max(_zip_total, 1),
                                )
                            with zf.open(info) as src:
                                rel_key, disk = _write_archive_member(src.read(), decoded_path)
                            extracted_files.append((rel_key, disk))

            elif ext == '.rar':
                extracted_files = []
                try:
                    from app.archive_extraction_utils import (
                        bulk_extract_rar,
                        collapse_redundant_archive_path,
                        find_unrar_tool,
                        safe_archive_out_path,
                        selective_extract_rar,
                        write_bytes_safe,
                    )

                    def _register_archive_file(decoded_path: str) -> bool:
                        fname = Path(decoded_path).name
                        if fname.startswith('.') or fname.startswith('__'):
                            return False
                        if _is_in_skip_dir(decoded_path):
                            return False
                        if all_display_files is not None:
                            all_display_files.append(decoded_path)
                        fname_lower = fname.lower()
                        if fname_lower.endswith(".exe"):
                            from app.archive_extraction_utils import is_primary_game_executable

                            if not is_primary_game_executable(decoded_path):
                                return False
                        return any(fname_lower.endswith(e) for e in _ALL_EXTENSIONS)

                    def _collect_from_raw(raw_root: Path) -> list:
                        out = []
                        for fpath in raw_root.rglob('*'):
                            if not fpath.is_file():
                                continue
                            rel = fpath.relative_to(raw_root).as_posix()
                            decoded_path = collapse_redundant_archive_path(rel)
                            if not _register_archive_file(decoded_path):
                                continue
                            if len(str(fpath)) > 240:
                                out_path, rel_key = safe_archive_out_path(extract_dir, decoded_path)
                                if out_path.resolve() != fpath.resolve():
                                    write_bytes_safe(out_path, fpath.read_bytes())
                                    try:
                                        fpath.unlink()
                                    except OSError:
                                        pass
                                    out.append((rel_key, str(out_path)))
                                else:
                                    out.append((decoded_path, str(fpath)))
                            else:
                                out.append((decoded_path, str(fpath)))
                        return out

                    def _emit_rar_list_progress(listed: int) -> None:
                        _emit_archive_progress(
                            phase_label="جاري قراءة قائمة ملفات RAR...",
                            listed=listed,
                            sub_phase="listing",
                        )

                    def _emit_rar_manifest(student_groups: int, files_to_extract: int) -> None:
                        groups = max(student_groups, 1)
                        folder_note = ""
                        if _is_single_student_archive and groups > 1:
                            folder_note = f" ({groups} مجلدات داخل أرشيف طالب واحد)"
                        _emit_archive_progress(
                            phase_label=(
                                f"جاري استخراج {files_to_extract} ملفاً"
                                f"{folder_note}..."
                            ),
                            nested_total=groups,
                            sub_phase="manifest",
                        )

                    def _emit_rar_progress(done: int, total: int, name: str) -> None:
                        frac = (done + 1) / max(total, 1)
                        _emit_archive_progress(
                            student=name,
                            frac=frac,
                            sub_phase="extract",
                        )

                    _rar_bytes = os.path.getsize(archive_path)
                    _bulk_ok = False
                    from app.archive_extraction_utils import (
                        archive_should_use_selective_extract,
                        max_archive_extract_files,
                    )

                    # Do not run a full `unrar lb` pass here — selective_extract_rar lists once.
                    _use_selective = archive_should_use_selective_extract(
                        0, _rar_bytes, _pipeline_grading_mode
                    )

                    if _use_selective:
                        sel_files, sel_display = selective_extract_rar(
                            archive_path,
                            extract_dir,
                            skip_dir_names=frozenset(_SKIP_DIRS),
                            gradable_extensions=_ALL_EXTENSIONS,
                            on_progress=_emit_rar_progress,
                            on_list_progress=_emit_rar_list_progress,
                            on_manifest=_emit_rar_manifest,
                            grading_mode=_pipeline_grading_mode,
                            max_extract_files=max_archive_extract_files(
                                _pipeline_grading_mode, archive_bytes=_rar_bytes
                            ),
                        )
                        extracted_files = sel_files
                        if all_display_files is not None:
                            all_display_files.extend(sel_display)
                        _bulk_ok = bool(extracted_files)
                        print(
                            f"📦 [RAR-SEL] {Path(archive_path).name} "
                            f"({_rar_bytes / (1024 * 1024):.0f} MB, "
                            f"{len(sel_display)} visible path(s)) "
                            f"→ {len(extracted_files)} gradable file(s)"
                        )
                    else:
                        raw_dir = extract_dir / "_raw"
                        if find_unrar_tool():
                            try:
                                bulk_extract_rar(archive_path, raw_dir)
                                extracted_files = _collect_from_raw(raw_dir)
                                if extracted_files:
                                    _bulk_ok = True
                                    print(f"📦 [RAR] bulk extract → {len(extracted_files)} file(s)")
                            except Exception as _bulk_exc:
                                print(f"⚠️ [RAR] bulk extract failed: {_bulk_exc}")

                    if not _bulk_ok:
                        sel_files, sel_display = selective_extract_rar(
                            archive_path,
                            extract_dir,
                            skip_dir_names=frozenset(_SKIP_DIRS),
                            gradable_extensions=_ALL_EXTENSIONS,
                            on_progress=_emit_rar_progress,
                            on_list_progress=_emit_rar_list_progress,
                            on_manifest=_emit_rar_manifest,
                            grading_mode=_pipeline_grading_mode,
                        )
                        extracted_files = sel_files
                        if all_display_files is not None and sel_display:
                            if not all_display_files:
                                all_display_files.extend(sel_display)
                        if not extracted_files:
                            raise RuntimeError(
                                "لم يُستخرج أي ملف مدعوم من RAR — احذف مجلدات Library و.godot "
                                "ثم اضغط ZIP"
                            )
                except ImportError:
                    msg = (
                        "تعذّر فك ملف RAR: مكتبة rarfile غير مثبتة على الخادم. "
                        "استخدم ZIP أو ثبّت rarfile + WinRAR UnRAR."
                    )
                    print(f"⚠️ {msg}")
                    _archive_pipeline_errors.append(msg)
                    return []
                except Exception as _rar_exc:
                    msg = f"تعذّر فك ملف RAR: {_rar_exc}"
                    print(f"⚠️ {msg}")
                    _archive_pipeline_errors.append(msg)
                    return []
            else:
                return []

            if not extracted_files:
                scan_root = extract_dir / "_raw" if (extract_dir / "_raw").exists() else extract_dir
                nested_archives = _collect_nested_archives(scan_root)
                if nested_archives:
                    from app.archive_extraction_utils import (
                        strip_archive_wrapper_prefix,
                        unwrap_single_wrapper_folder,
                    )

                    nested_archives, _wrap = unwrap_single_wrapper_folder(nested_archives)
                    if _wrap:
                        for rel, disk in nested_archives:
                            if all_display_files is not None:
                                all_display_files.append(
                                    strip_archive_wrapper_prefix(rel, _wrap)
                                )
                    elif all_display_files is not None:
                        for rel, _disk in nested_archives:
                            all_display_files.append(rel)

                if nested_archives and _depth < 2:
                    combined: list = []
                    bundle_display: dict = {'all_files': [], 'student_map': {}}
                    print(
                        f"📦 [ARCHIVE-BUNDLE] depth={_depth} "
                        f"nested={len(nested_archives)} in {Path(archive_path).name}"
                    )
                    for _ni, (rel, disk) in enumerate(nested_archives):
                        _emit_archive_progress(
                            student=Path(rel).name,
                            phase_label=(
                                f"جاري فك أرشيف {_ni + 1} من {len(nested_archives)}..."
                            ),
                            nested_idx=_ni,
                            nested_total=len(nested_archives),
                        )
                        inner = extract_archive(disk, batch_id, _depth=_depth + 1)
                        if inner:
                            combined.extend(inner)
                            bundle_display['all_files'].extend(
                                _archive_display_data.get('all_files') or []
                            )
                            for student, fnames in (
                                _archive_display_data.get('student_map') or {}
                            ).items():
                                bucket = bundle_display['student_map'].setdefault(student, [])
                                for fname in fnames:
                                    if fname not in bucket:
                                        bucket.append(fname)
                    if combined:
                        _archive_display_data['all_files'] = list(bundle_display['all_files'])
                        _archive_display_data['student_map'] = {
                            k: list(v) for k, v in bundle_display['student_map'].items()
                        }
                        print(
                            f"📦 [ARCHIVE-BUNDLE] → {len(combined)} student(s) "
                            f"from {len(nested_archives)} archive(s)"
                        )
                        from app.archive_extraction_utils import consolidate_archive_student_results

                        return consolidate_archive_student_results(combined)

                _apply_archive_display(all_display_files)
                return []

            from app.archive_extraction_utils import (
                consolidate_archive_student_results,
                merge_single_student_category_bundle,
                strip_archive_wrapper_prefix,
                unwrap_single_wrapper_folder,
            )

            from app.archive_extraction_utils import materialize_nested_zip_game_executables

            nested_runtime = materialize_nested_zip_game_executables(
                extract_dir, grading_mode=_pipeline_grading_mode
            )
            if nested_runtime:
                extracted_files.extend(nested_runtime)
                if all_display_files is not None:
                    for rel_key, _disk in nested_runtime:
                        all_display_files.append(rel_key)

            extracted_files, _wrapper = unwrap_single_wrapper_folder(extracted_files)
            extracted_files = merge_single_student_category_bundle(extracted_files)
            if _wrapper and all_display_files is not None:
                all_display_files[:] = [
                    strip_archive_wrapper_prefix(p, _wrapper) for p in all_display_files
                ]

            # Bundle: outer archive contains only nested ZIP/RAR files (common teacher upload)
            nested_only = [
                (p, d) for p, d in extracted_files
                if Path(p).suffix.lower() in _NESTED_ARCHIVE_EXTENSIONS
            ]
            if nested_only and len(nested_only) == len(extracted_files) and _depth < 2:
                combined = []
                bundle_display = {'all_files': [], 'student_map': {}}
                print(
                    f"📦 [ARCHIVE-BUNDLE] depth={_depth} "
                    f"nested={len(nested_only)} in {Path(archive_path).name}"
                )
                for _ni, (_rel, disk) in enumerate(nested_only):
                    _emit_archive_progress(
                        student=Path(_rel).name,
                        phase_label=(
                            f"جاري فك أرشيف {_ni + 1} من {len(nested_only)}..."
                        ),
                        nested_idx=_ni,
                        nested_total=len(nested_only),
                    )
                    inner = extract_archive(disk, batch_id, _depth=_depth + 1)
                    if inner:
                        combined.extend(inner)
                        bundle_display['all_files'].extend(
                            _archive_display_data.get('all_files') or []
                        )
                        for student, fnames in (
                            _archive_display_data.get('student_map') or {}
                        ).items():
                            bucket = bundle_display['student_map'].setdefault(student, [])
                            for fname in fnames:
                                if fname not in bucket:
                                    bucket.append(fname)
                if combined:
                    _archive_display_data['all_files'] = list(bundle_display['all_files'])
                    _archive_display_data['student_map'] = {
                        k: list(v) for k, v in bundle_display['student_map'].items()
                    }
                    print(
                        f"📦 [ARCHIVE-BUNDLE] → {len(combined)} student(s) "
                        f"from {len(nested_only)} archive(s)"
                    )
                    from app.archive_extraction_utils import consolidate_archive_student_results

                    return consolidate_archive_student_results(combined)

            # ── Build display data for archive UI ──
            def _build_archive_display():
                _apply_archive_display(all_display_files)
                print(
                    f"📁 [ARCHIVE-UI] {len(_archive_display_data['all_files'])} file(s) "
                    f"for display, mapped to {len(_archive_display_data['student_map'])} student group(s)"
                )

            # Determine structure: flat files vs folder-based
            # Get top-level entries (first path component after root)
            top_level_dirs = set()
            root_files = []
            for arc_path, disk_path in extracted_files:
                parts = Path(arc_path).parts
                if len(parts) == 1:
                    # File at root level
                    root_files.append((arc_path, disk_path))
                else:
                    # File inside a folder
                    top_level_dirs.add(parts[0])

            # Case 1: All files at root → each file = one student
            if not top_level_dirs:
                print(f"📁 [ARCHIVE] Flat structure: {len(root_files)} student file(s)")
                result = []
                seen_names = set()
                for arc_path, disk_path in root_files:
                    student_name = Path(arc_path).stem
                    if student_name not in seen_names:
                        seen_names.add(student_name)
                        arc_norm = arc_path.replace("\\", "/")
                        intake = [arc_norm]
                        _merge_intake_videos_from_display(intake, None, all_display_files)
                        result.append((student_name, disk_path, [disk_path], intake))
                _build_archive_display()
                return consolidate_archive_student_results(result)

            # Case 2: Files in folders → each top-level folder = one student
            # Group files by their top-level folder
            folder_files = {}  # {folder_name: [(student_name, disk_path, archive_rel), ...]}
            for arc_path, disk_path in extracted_files:
                parts = Path(arc_path).parts
                arc_norm = arc_path.replace("\\", "/")
                if len(parts) == 1:
                    # Root file alongside folders — treat as its own student
                    student_name = Path(arc_path).stem
                    folder_files.setdefault('__root__', []).append((student_name, disk_path, arc_norm))
                else:
                    folder_name = parts[0]
                    folder_files.setdefault(folder_name, []).append((folder_name, disk_path, arc_norm))

            result = []
            seen_names = set()
            for folder_name, files_in_folder in folder_files.items():
                if folder_name == '__root__':
                    for row in files_in_folder:
                        student_name = row[0]
                        disk_path = row[1]
                        if student_name not in seen_names:
                            seen_names.add(student_name)
                            ar = row[2] if len(row) > 2 else Path(student_name).name
                            result.append((student_name, disk_path, [disk_path], [ar]))
                else:
                    # Pick the best doc/pdf from this student's folder
                    best = _pick_best_file(files_in_folder)
                    if best:
                        student_name = folder_name if folder_name != "__root__" else best[0]
                        if student_name.startswith("_student_bundle"):
                            from app.archive_extraction_utils import resolve_bundle_student_name

                            all_arc_preview = [
                                str(row[2]) if len(row) > 2 else Path(str(row[1])).name
                                for row in files_in_folder
                            ]
                            resolved = resolve_bundle_student_name(
                                str(best[1]),
                                all_arc_preview,
                                all_paths=[str(row[1]) for row in files_in_folder],
                            )
                            if resolved:
                                student_name = resolved
                            else:
                                student_name = best[0]
                        if student_name not in seen_names:
                            seen_names.add(student_name)
                            all_paths = sorted({str(row[1]) for row in files_in_folder}, key=lambda x: x.lower())
                            all_arc = sorted(
                                {
                                    str(row[2]) if len(row) > 2 else Path(str(row[1])).name
                                    for row in files_in_folder
                                },
                                key=lambda x: x.lower(),
                            )
                            _merge_intake_videos_from_display(
                                all_arc,
                                folder_name if folder_name != "__root__" else None,
                                all_display_files,
                            )
                            result.append((student_name, best[1], all_paths, all_arc))
                            # Check if any code files exist in this folder
                            has_code = any(
                                any(str(row[1]).lower().endswith(e) for e in _PROGRAMMING_EXTS)
                                for row in files_in_folder
                            )
                            _student_has_code_files[student_name] = has_code
                            has_exe = any(
                                any(str(row[1]).lower().endswith(e) for e in _EXECUTABLE_ARTIFACT_EXTENSIONS)
                                for row in files_in_folder
                            )
                            _student_has_executable_artifacts[student_name] = has_exe
                            if has_code:
                                print(f"   💻 [CODE] Student '{student_name}' has code files in folder")

            print(f"📁 [ARCHIVE] Folder structure: {len(top_level_dirs)} folder(s) → {len(result)} student(s)")
            from app.archive_extraction_utils import (
                consolidate_archive_student_results,
                merge_likely_single_student_bundle,
            )

            result = merge_likely_single_student_bundle(
                result,
                top_level_folder_names=set(top_level_dirs) | (
                    {"__root__"} if "__root__" in folder_files else set()
                ),
                archive_name=Path(archive_path).name,
            )
            _build_archive_display()
            return consolidate_archive_student_results(result)

        # Accumulate archive display files across all uploaded archives
        accumulated_archive_files = []
        accumulated_archive_map = {}

        async def _run_batch_pipeline() -> None:
            """Spool uploads, prepare submissions, then grade (runs after HTTP response)."""
            import asyncio

            staged_uploads: list = []
            try:
                _staging_root = STUDENTS_DIR / f"batch_{batch.id}_upload"
                _staging_root.mkdir(parents=True, exist_ok=True)
                print(f"📦 [BATCH-SPOOL] start batch {batch.id} parts={len(files)}")
                _spool_skipped = 0
                _spool_total = max(len(files), 1)
                _SPOOL_CHUNK = 4 * 1024 * 1024
                from app.project_intelligence.submission_intake import _multipart_part_size

                def _spool_phase_percent(file_index: int, file_frac: float) -> int:
                    """Map spool progress into 1–12% (prepare band before archive extract)."""
                    overall = (file_index + min(1.0, max(0.0, file_frac))) / _spool_total
                    return max(1, min(12, round(overall * 12)))

                for _si, _uf in enumerate(files):
                    _rel = (_uf.filename or "unknown").replace("\\", "/")
                    if path_matches_intake_ignore(_rel):
                        await _uf.read()
                        _spool_skipped += 1
                        continue
                    _dest = _staging_root / _rel
                    _dest.parent.mkdir(parents=True, exist_ok=True)
                    try:
                        _file_bytes = await _multipart_part_size(_uf)
                    except Exception:
                        _file_bytes = 0
                    _pi = batch_progress.get(assignment_id) or {}
                    _pi.update({
                        "status": "preparing",
                        "current_phase": "preparing",
                        "current_student": Path(_rel).name,
                        "percent": _spool_phase_percent(_si, 0.0),
                        "phase_label": (
                            f"جاري حفظ الملف على الخادم ({_file_bytes / (1024 * 1024):.0f} MB)..."
                            if _file_bytes > 0
                            else "جاري حفظ الملف على الخادم..."
                        ),
                    })
                    batch_progress[assignment_id] = _pi
                    _written = 0
                    with open(_dest, "wb") as _buf:
                        while True:
                            _chunk = await _uf.read(_SPOOL_CHUNK)
                            if not _chunk:
                                break
                            await asyncio.to_thread(_buf.write, _chunk)
                            _written += len(_chunk)
                            if _file_bytes > 0:
                                _file_frac = min(1.0, _written / _file_bytes)
                            else:
                                _file_frac = min(1.0, _written / max(_written + _SPOOL_CHUNK, 1))
                            _pi = batch_progress.get(assignment_id) or {}
                            _pi.update({
                                "status": "preparing",
                                "current_phase": "preparing",
                                "current_student": Path(_rel).name,
                                "percent": _spool_phase_percent(_si, _file_frac),
                                "phase_label": (
                                    f"جاري حفظ الملف: {Path(_rel).name} "
                                    f"({_written / (1024 * 1024):.0f}/"
                                    f"{max(1, _file_bytes // (1024 * 1024))} MB)"
                                    if _file_bytes > 0
                                    else f"جاري حفظ الملف: {Path(_rel).name}"
                                ),
                            })
                            batch_progress[assignment_id] = _pi
                    staged_uploads.append({"filename": _rel, "path": str(_dest)})
                    _pi = batch_progress.get(assignment_id) or {}
                    _pi.update({
                        "status": "preparing",
                        "current_phase": "preparing",
                        "current_student": Path(_rel).name,
                        "percent": _spool_phase_percent(_si + 1, 0.0),
                        "phase_label": f"اكتمل حفظ {Path(_rel).name}",
                    })
                    batch_progress[assignment_id] = _pi
                    print(
                        f"📦 [BATCH-SPOOL] saved {_rel} ({_written / (1024 * 1024):.1f} MB)"
                    )
                print(
                    f"📦 [BATCH-SPOOL] done batch {batch.id}: "
                    f"{len(staged_uploads)} kept, {_spool_skipped} auto-ignored"
                )
                from app.batch_checkpoint import save_batch_checkpoint

                save_batch_checkpoint(
                    int(batch.id),
                    {
                        "stage": "staged",
                        "assignment_id": assignment_id,
                        "batch_id": int(batch.id),
                        "batch_name": batch_name,
                        "user_id": user_id,
                        "subject_bal_id": subject_bal.id if subject_bal is not None else None,
                        "sub_id": sub.id if sub is not None else None,
                        "grading_mode": _pipeline_grading_mode,
                        "single_student_archive": single_student_archive == "true",
                        "folder_mode": folder_mode == "true",
                        "selected_criteria": selected_criteria,
                        "batch_intake_manifest": batch_intake_manifest,
                        "skip_grading_cache": _force_regrade,
                        "staged_uploads": staged_uploads,
                    },
                )
                if _spool_skipped:
                    _pi = batch_progress.get(assignment_id) or {}
                    _pi["intake_auto_filtered"] = {
                        "skipped_count": _spool_skipped,
                        "kept_count": len(staged_uploads),
                    }
                    batch_progress[assignment_id] = _pi

                # ── Folder Mode: each folder = one student ──
                _folder_mode_active = folder_mode == "true" and bool(folder_map)
                if _folder_mode_active:
                    try:
                        fm = json.loads(folder_map)
                    except Exception:
                        fm = {}

                    # Map staged uploads (already on disk)
                    _saved_paths = {}  # original_filename -> disk_path
                    for _su in staged_uploads:
                        _saved_paths[_su["filename"]] = _su["path"]

                    # Group by folder (student)
                    folder_student_count = 0
                    for student_folder, rel_paths in fm.items():
                        # Collect saved disk paths for this student's files
                        student_folder_files = []
                        for rp in rel_paths:
                            # Match by filename since browser sends webkitRelativePath as filename
                            for orig_name, disk_path in _saved_paths.items():
                                # orig_name may be like "RootFolder/StudentFolder/file.ext"
                                if orig_name == rp or orig_name.replace('\\', '/') == rp.replace('\\', '/'):
                                    student_folder_files.append((student_folder, disk_path))
                                    break
    
                        if not student_folder_files:
                            continue
    
                        folder_student_count += 1
                        # Pick best file for grading
                        best = _pick_best_file(student_folder_files)
                        if best:
                            sname = best[0]
                            spath = best[1]
    
                            # Check for code files
                            has_code = any(
                                any(p.lower().endswith(e) for e in _PROGRAMMING_EXTS)
                                for _, p in student_folder_files
                            )
                            _student_has_code_files[sname] = has_code
    
                            # Try to extract real student name from document content
                            doc_name = await asyncio.to_thread(
                                extract_student_name_from_file, spath
                            )
                            if doc_name:
                                print(f"📛 [NAME] Folder '{sname}' → document name: '{doc_name}'")
                                sname = doc_name

                            new_source_hash = await asyncio.to_thread(
                                _hash_submission_file, spath, ""
                            )

                            if sname in existing_by_name:
                                existing_data = existing_by_name[sname]
                                existing_sub = existing_data["submission"]
                                old_hash = existing_data["source_hash"]
                                if old_hash and new_source_hash and old_hash == new_source_hash:
                                    existing_summary = (
                                        db.query(GradingSummary)
                                        .filter(GradingSummary.submission_id == existing_sub.id)
                                        .first()
                                    )
                                    if existing_summary:
                                        print(
                                            f"✅ [CACHE] Student '{sname}' — identical file, "
                                            f"reusing grade {existing_summary.grade_level}"
                                        )
                                        cached_results.append({
                                            "student_name": sname,
                                            "success": True,
                                            "cached": True,
                                            "submission_id": existing_sub.id,
                                        })
                                        continue
                                else:
                                    print(f"🔄 [RESUBMIT] Student '{sname}' has MODIFIED content - will re-grade")
    
                            submission_paths = sorted({str(p) for (_, p) in student_folder_files}, key=lambda x: x.lower())
                            rel_for_intake = sorted({rp.replace("\\", "/") for rp in rel_paths}, key=lambda x: x.lower())
                            student_files.append({
                                "name": sname,
                                "path": spath,
                                "email": "",
                                "student_id": "",
                                "has_code_files": has_code,
                                "submission_paths": submission_paths,
                                "intake_relative_paths": rel_for_intake,
                            })
    
                    # ── Quota enforcement for folder mode ──
                    if remaining_quota is not None and folder_student_count > remaining_quota:
                        try:
                            batch.status = BatchStatus.FAILED  # type: ignore
                            db.commit()
                        except Exception:
                            db.rollback()
                        batch_progress[assignment_id] = {
                            **(batch_progress.get(assignment_id) or {}),
                            "finished": True,
                            "failed": True,
                            "error": (
                                f"العدد ({folder_student_count}) يتجاوز الرصيد ({remaining_quota})"
                            ),
                            "final_response": {
                                "success": False,
                                "subscription_required": True,
                                "detail": (
                                    f"العدد الإجمالي ({folder_student_count}) يتجاوز رصيدك المتبقي "
                                    f"({remaining_quota}). الرجاء تقليل عدد الطلاب أو تجديد الاشتراك."
                                ),
                            },
                        }
                        return
    
                    # Update batch total with folder count
                    _orm_set(batch, "total_students", folder_student_count)
                    db.commit()
                    print(f"📁 [FOLDER-MODE] {folder_student_count} student folder(s) processed")
    
                # ── Regular Mode (no folder grouping) ──
                if not _folder_mode_active:
                    for _su in staged_uploads:
                        file_path = Path(_su["path"])
                        file_ext = file_path.suffix.lower()
                        _orig_name = _su["filename"]

                        # If it's an archive, extract and process each file inside
                        if file_ext in ('.zip', '.rar'):
                            _archive_display_data['all_files'] = []
                            _archive_display_data['student_map'] = {}
                            _pi = batch_progress.get(assignment_id) or {}
                            _pi.update({
                                "status": "preparing",
                                "current_phase": "extracting_archive",
                                "current_student": Path(_orig_name).name,
                                "percent": max(int(_pi.get("percent") or 0), 13),
                            })
                            batch_progress[assignment_id] = _pi
                            print(
                                f"📦 [ARCHIVE] extracting {_orig_name} "
                                f"(batch {batch.id}) — may take several minutes"
                            )
                            try:
                                db.commit()
                            except Exception:
                                db.rollback()
                            from app.archive_extraction_utils import (
                                archive_extract_timeout_seconds,
                            )

                            _ARCHIVE_EXTRACT_TIMEOUT = archive_extract_timeout_seconds(
                                _pipeline_grading_mode
                            )
                            try:
                                archive_files = await asyncio.wait_for(
                                    asyncio.to_thread(
                                        extract_archive, str(file_path), batch.id
                                    ),
                                    timeout=_ARCHIVE_EXTRACT_TIMEOUT,
                                )
                            except asyncio.TimeoutError:
                                _timeout_msg = (
                                    f"انتهت مهلة فك الأرشيف ({_ARCHIVE_EXTRACT_TIMEOUT // 60} دقيقة) "
                                    f"لـ «{Path(_orig_name).name}». "
                                    "احذف مجلدات Library وTemp/.godot/_build ثم اضغط ZIP، أو جرّب BASIC."
                                )
                                print(f"❌ [ARCHIVE-TIMEOUT] {_timeout_msg}")
                                _archive_pipeline_errors.append(_timeout_msg)
                                continue
                            if single_student_archive == "true" and len(archive_files) > 1:
                                from app.archive_extraction_utils import (
                                    force_single_student_archive_result,
                                )

                                archive_files = force_single_student_archive_result(
                                    archive_files,
                                    archive_name=_orig_name,
                                )
                                merged_name = str(archive_files[0][0]) if archive_files else ""
                                merged_files = list(_archive_display_data.get("all_files") or [])
                                _archive_display_data["student_map"] = (
                                    {merged_name: merged_files} if merged_name and merged_files else {}
                                )
                                print(
                                    f"📦 [ARCHIVE-SINGLE] forced merge → 1 student "
                                    f"({len(merged_files)} file(s) in bundle)"
                                )
                            accumulated_archive_files.extend(
                                _archive_display_data.get('all_files', []))
                            accumulated_archive_map.update(
                                _archive_display_data.get('student_map', {}))
                            if not archive_files:
                                print(f"⚠️ Archive {_orig_name} is empty or contains no supported files")
                                _display_count = len(_archive_display_data.get('all_files') or [])
                                if file_ext == '.rar' and not _archive_pipeline_errors:
                                    if _display_count > 0:
                                        _archive_pipeline_errors.append(
                                            f"ملف RAR «{Path(_orig_name).name}» يحتوي {_display_count} ملفاً "
                                            "لكن لا Word/PDF/كود/صور مدعومة للتصحيح. "
                                            "احذف مجلدات Library وTemp/.godot ثم اضغط ZIP."
                                        )
                                    else:
                                        _archive_pipeline_errors.append(
                                            f"ملف RAR «{Path(_orig_name).name}» فارغ أو لا يحتوي ملفات مدعومة "
                                            "(Word/PDF/كود/صور). يُفضّل ZIP بعد حذف Library وTemp."
                                        )
                                elif file_ext == '.zip' and not _archive_pipeline_errors:
                                    _archive_pipeline_errors.append(
                                        f"ملف ZIP «{Path(_orig_name).name}» فارغ أو لا يحتوي ملفات مدعومة."
                                    )
                                continue

                            if remaining_quota is not None:
                                projected_total = batch.total_students - 1 + len(archive_files)
                                if projected_total > remaining_quota:
                                    try:
                                        batch.status = BatchStatus.FAILED  # type: ignore
                                        db.commit()
                                    except Exception:
                                        db.rollback()
                                    batch_progress[assignment_id] = {
                                        **(batch_progress.get(assignment_id) or {}),
                                        "finished": True,
                                        "failed": True,
                                        "error": (
                                            f"العدد الإجمالي ({projected_total}) يتجاوز الرصيد ({remaining_quota})"
                                        ),
                                    }
                                    return

                            batch.total_students = batch.total_students - 1 + len(archive_files)  # type: ignore
                            db.commit()

                            for _ai, arc_entry in enumerate(archive_files):
                                if len(arc_entry) >= 4:
                                    student_name, extracted_path, submission_paths, intake_rel_paths = arc_entry  # type: ignore[misc]
                                    intake_rel = list(intake_rel_paths or ())
                                elif len(arc_entry) == 3:
                                    student_name, extracted_path, submission_paths = arc_entry  # type: ignore[misc]
                                    intake_rel = []
                                else:
                                    student_name, extracted_path = arc_entry  # type: ignore[misc]
                                    submission_paths = [str(extracted_path)]
                                    intake_rel = []
                                _pi = batch_progress.get(assignment_id) or {}
                                _pi.update({
                                    "current_phase": "extracting_archive",
                                    "current_student": Path(str(extracted_path)).name,
                                    "total": len(archive_files),
                                    "phase_label": (
                                        f"جاري تجهيز الطالب {_ai + 1} من {len(archive_files)}..."
                                    ),
                                    "percent": min(
                                        25,
                                        13 + round(((_ai + 1) / max(len(archive_files), 1)) * 12),
                                    ),
                                })
                                batch_progress[assignment_id] = _pi
                                doc_name = await asyncio.to_thread(
                                    extract_student_name_from_file, extracted_path
                                )
                                if doc_name:
                                    student_name = doc_name
                                new_source_hash = await asyncio.to_thread(
                                    _hash_submission_file, extracted_path, ""
                                )
                                if student_name in existing_by_name:
                                    existing_data = existing_by_name[student_name]
                                    if (
                                        existing_data["source_hash"]
                                        and new_source_hash
                                        and existing_data["source_hash"] == new_source_hash
                                    ):
                                        existing_summary = (
                                            db.query(GradingSummary)
                                            .filter(
                                                GradingSummary.submission_id
                                                == existing_data["submission"].id
                                            )
                                            .first()
                                        )
                                        if existing_summary:
                                            print(
                                                f"✅ [CACHE] Student '{student_name}' — identical file, "
                                                f"reusing grade {existing_summary.grade_level}"
                                            )
                                            cached_results.append({
                                                "student_name": student_name,
                                                "success": True,
                                                "cached": True,
                                                "submission_id": existing_data["submission"].id,
                                            })
                                            continue
                                student_files.append({
                                    "name": student_name,
                                    "path": extracted_path,
                                    "email": "",
                                    "student_id": "",
                                    "has_code_files": _student_has_code_files.get(student_name, False),
                                    "has_executable_artifacts": _student_has_executable_artifacts.get(student_name, False),
                                    "submission_paths": list(submission_paths),
                                    "intake_relative_paths": intake_rel,
                                    "source_archive_path": str(file_path),
                                })
                            continue

                        student_name = Path(_orig_name).stem
                        doc_name = await asyncio.to_thread(
                            extract_student_name_from_file, str(file_path)
                        )
                        if doc_name:
                            student_name = doc_name
                        new_source_hash = await asyncio.to_thread(
                            _hash_submission_file, str(file_path), ""
                        )
                        if student_name in existing_by_name:
                            existing_data = existing_by_name[student_name]
                            if (
                                existing_data["source_hash"]
                                and new_source_hash
                                and existing_data["source_hash"] == new_source_hash
                            ):
                                existing_summary = (
                                    db.query(GradingSummary)
                                    .filter(
                                        GradingSummary.submission_id
                                        == existing_data["submission"].id
                                    )
                                    .first()
                                )
                                if existing_summary:
                                    print(
                                        f"✅ [CACHE] Student '{student_name}' — identical file, "
                                        f"reusing grade {existing_summary.grade_level}"
                                    )
                                    cached_results.append({
                                        "student_name": student_name,
                                        "success": True,
                                        "cached": True,
                                        "submission_id": existing_data["submission"].id,
                                    })
                                    continue
                        single_file_is_code = any(str(file_path).lower().endswith(e) for e in _PROGRAMMING_EXTS)
                        _single_rel = _orig_name.replace("\\", "/")
                        student_files.append({
                            "name": student_name,
                            "path": str(file_path),
                            "email": "",
                            "student_id": "",
                            "has_code_files": single_file_is_code,
                            "submission_paths": [str(file_path)],
                            "intake_relative_paths": [_single_rel] if _single_rel else [],
                        })

                # Grade only NEW students (not cached)
                selected_criteria_list = selected_criteria.split(",")
    
                # --- Post-processing subscription check (actual student count) ---
                asgn_subject = ""
                actual_student_count = len(student_files) + len(cached_results)
                if actual_student_count == 0:
                    _empty_detail = (
                        _archive_pipeline_errors[0]
                        if _archive_pipeline_errors
                        else "لم يُستخرج أي ملف طالب من التسليم. تحقق من ZIP/RAR أو ارفع مجلد/Word."
                    )
                    try:
                        batch.status = BatchStatus.FAILED  # type: ignore
                        batch.processed_students = 0  # type: ignore
                        batch.failure_message = _empty_detail  # type: ignore
                        db.commit()
                    except Exception:
                        db.rollback()
                    batch_progress[assignment_id] = {
                        **(batch_progress.get(assignment_id) or {}),
                        "finished": True,
                        "failed": True,
                        "error": _empty_detail,
                        "final_response": {
                            "success": False,
                            "detail": _empty_detail,
                            "batch_id": batch.id,
                            "processed": 0,
                        },
                    }
                    print(f"❌ [BATCH] No students extracted batch={batch.id}: {_empty_detail}")
                    return
                if subject_bal is not None:
                    remaining = max(0, (subject_bal.assignments_limit or 0) - (subject_bal.assignments_used or 0))
                    if actual_student_count > remaining:
                        batch_progress[assignment_id] = {
                            **(batch_progress.get(assignment_id) or {}),
                            "finished": True,
                            "failed": True,
                            "final_response": {
                                "success": False,
                                "subscription_required": True,
                                "detail": f"لديك {remaining} واجب متبقي لمادة {asgn_subject}، لكنك حاولت تحميل {actual_student_count} طالب.",
                            },
                        }
                        return
                elif sub is not None:
                    remaining = (sub.assignments_limit or 0) - (sub.assignments_used or 0)
                    if actual_student_count > remaining:
                        batch_progress[assignment_id] = {
                            **(batch_progress.get(assignment_id) or {}),
                            "finished": True,
                            "failed": True,
                            "final_response": {
                                "success": False,
                                "subscription_required": True,
                                "detail": f"لديك {remaining} واجب متبقي فقط، لكنك حاولت تحميل {actual_student_count} طالب.",
                            },
                        }
                        return
                # --- End post-processing subscription check ---
    
                # Initialize real-time progress tracker
                total_to_grade = len(student_files) + len(cached_results)
                all_names = [r.get("student_name", "") for r in cached_results] + [s["name"] for s in student_files]

                # Keep DB counter aligned with what we actually grade
                try:
                    batch.total_students = int(total_to_grade)  # type: ignore
                    db.commit()
                except Exception:
                    db.rollback()

                if is_admin:
                    _effective_grading_mode = admin_grading_mode or "deep"
                else:
                    _grading_policy = resolve_grading_policy(None)
                    if sub is not None:
                        try:
                            _pkg = (
                                db.query(models.Package)
                                .filter(models.Package.id == sub.package_id)
                                .first()
                            )
                            _grading_policy = resolve_grading_policy(
                                _pkg.name if _pkg else None
                            )
                        except Exception:
                            _grading_policy = resolve_grading_policy(None)
                    _effective_grading_mode = str(
                        _grading_policy.get("grading_mode") or "deep"
                    )

                _archive_ready = bool(accumulated_archive_files)
                _init_pct = round(
                    (len(cached_results) / max(total_to_grade, 1)) * 100
                )
                if _archive_ready:
                    _init_phase = "queued"
                    _init_pct = max(_init_pct, 8)
                    _init_label = "في قائمة انتظار التصحيح على الخادم..."
                else:
                    _init_phase = "starting"
                    _init_pct = max(_init_pct, 1)
                    _init_label = "بدء التصحيح..."
                batch_progress[assignment_id] = {
                    "completed": len(cached_results),
                    "total": total_to_grade,
                    "batch_id": batch.id,
                    "current_student": all_names[0] if len(all_names) == 1 else "",
                    "current_phase": _init_phase,
                    "student_progress": 0.08 if _archive_ready else 0.0,
                    "percent": _init_pct,
                    "phase_label": _init_label,
                    "start_time": time.time(),
                    "student_times": [],
                    "completed_students": [r.get("student_name", "") for r in cached_results],
                    "all_student_names": all_names,
                    "archive_all_files": accumulated_archive_files,
                    "archive_student_map": accumulated_archive_map,
                    "upload_intake": batch_intake_manifest,
                    "grading_mode": _effective_grading_mode,
                    "grading_mode_label": grading_mode_display_label(
                        _effective_grading_mode
                    ),
                }

                from app.batch_checkpoint import save_batch_checkpoint

                save_batch_checkpoint(
                    int(batch.id),
                    {
                        "stage": "grading",
                        "assignment_id": assignment_id,
                        "batch_id": int(batch.id),
                        "batch_name": batch_name,
                        "user_id": user_id,
                        "subject_bal_id": subject_bal.id if subject_bal is not None else None,
                        "sub_id": sub.id if sub is not None else None,
                        "grading_mode": _effective_grading_mode,
                        "single_student_archive": single_student_archive == "true",
                        "folder_mode": folder_mode == "true",
                        "selected_criteria_list": selected_criteria_list,
                        "batch_intake_manifest": batch_intake_manifest,
                        "skip_grading_cache": _force_regrade,
                        "student_files": student_files,
                        "cached_results": cached_results,
                        "accumulated_archive_files": accumulated_archive_files,
                        "accumulated_archive_map": accumulated_archive_map,
                    },
                )
    
                # Run grading in background so the HTTP connection closes before AI finishes
                # (prevents browser "Failed to fetch" on long batches).
                from app.batch_grade_worker import schedule_batch_grading_job  # type: ignore

                schedule_batch_grading_job(
                    batch_progress=batch_progress,
                    assignment_id=assignment_id,
                    batch_id=batch.id,
                    batch_name=batch_name,
                    student_files=student_files,
                    cached_results=cached_results,
                    grading_criteria=grading_criteria,
                    selected_criteria_list=selected_criteria_list,
                    reference_solution=reference_solution,
                    batch_intake_manifest=batch_intake_manifest,
                    user_id=user_id,
                    subject_bal_id=subject_bal.id if subject_bal is not None else None,
                    sub_id=sub.id if sub is not None else None,
                    grading_mode=_effective_grading_mode,
                    skip_grading_cache=_force_regrade,
                    job_generation=_upload_job_generation,
                )
    
            except Exception as _bg_exc:
                import traceback as _tb
                print("❌ [BATCH-BG]", _bg_exc)
                print(_tb.format_exc())
                _pi = batch_progress.get(assignment_id) or {}
                _pi.update({"finished": True, "failed": True, "error": str(_bg_exc)})
                batch_progress[assignment_id] = _pi
                try:
                    _bf = db.query(BatchGrading).filter(BatchGrading.id == batch.id).first()
                    if _bf:
                        _bf.status = BatchStatus.FAILED  # type: ignore
                        db.commit()
                except Exception:
                    pass

        batch_progress[assignment_id] = {
            "completed": 0,
            "total": 0,
            "batch_id": batch.id,
            "current_student": "جاري استلام الملفات...",
            "current_phase": "preparing",
            "student_progress": 0.0,
            "percent": 1,
            "phase_label": "جاري استلام الملفات من المتصفح...",
            "start_time": time.time(),
            "student_times": [],
            "completed_students": [],
            "all_student_names": [],
            "archive_all_files": [],
            "archive_student_map": {},
            "upload_intake": batch_intake_manifest,
            "status": "preparing",
        }
        # Return HTTP response immediately; spool + grading run after response (files still valid).
        background_tasks.add_task(_run_batch_pipeline)
        print(f"✅ [BATCH-REQ] queued background job batch={batch.id} assignment={assignment_id}")
        return {
            "success": True,
            "async": True,
            "batch_id": batch.id,
            "total_students": 0,
            "processed": 0,
            "upload_intake": batch_intake_manifest,
        }



    except Exception as e:
        import traceback

        error_details = traceback.format_exc()
        # Guarantee a meaningful detail even if str(e) is empty
        _exc_msg = str(e).strip() or repr(e) or type(e).__name__
        _detail = f"فشل التصحيح: {type(e).__name__}: {_exc_msg}"
        print(f"خطأ في التصحيح الجماعي: {_exc_msg}")
        print(f"تفاصيل الخطأ:\n{error_details}")
        try:
            log_activity(db, "batch_grade_error", "error",
                         f"خطأ في التصحيح الجماعي: {_exc_msg[:500]}",
                         user_id=user_id, user_name=_get_user_display(db, user_id),
                         user_email=_get_user_email(db, user_id),
                         ip_address=_get_client_ip(request),
                         user_agent=_get_user_agent(request), level="error")
        except Exception:
            pass

        try:
            if 'batch' in dir() and batch is not None:  # type: ignore[possibly-undefined]
                batch.status = BatchStatus.FAILED  # type: ignore
                db.commit()
        except Exception:
            pass
        raise HTTPException(status_code=500, detail=_detail) from e


@app.get("/api/assignments")
async def list_assignments(db: Session = Depends(get_db)):
    """List all assignments"""
    assignments = db.query(Assignment).order_by(Assignment.created_at.desc()).all()

    return [
        {
            "id": a.id,
            "title": a.title,
            "status": a.status,
            "created_at": a.created_at.isoformat(),
        }
        for a in assignments
    ]


@app.get("/plagiarism-highlight/{submission_id}", response_class=HTMLResponse)
async def plagiarism_highlight_page(
    submission_id: int, request: Request, db: Session = Depends(get_db)
):
    """Interactive HTML report — color-coded phrases copied from classmates."""
    from app.plagiarism_highlight_report import (
        build_comparisons_from_db,
        generate_plagiarism_highlight_html,
    )

    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    if not submission.submission_text:
        raise HTTPException(status_code=404, detail="لا يوجد نص للمقارنة")

    summary = (
        db.query(GradingSummary)
        .filter(GradingSummary.submission_id == submission_id)
        .first()
    )
    batch = (
        db.query(BatchGrading).filter(BatchGrading.id == submission.batch_id).first()
        if submission.batch_id
        else None
    )
    comparisons = build_comparisons_from_db(db, submission)
    html_body = generate_plagiarism_highlight_html(
        student_name=submission.student_name,
        submission_id=submission_id,
        batch_name=str(getattr(batch, "batch_name", "") or ""),
        batch_id=int(submission.batch_id) if submission.batch_id else None,
        student_text=str(submission.submission_text),
        comparisons=comparisons,
        max_similarity=float(summary.plagiarism_max_similarity or 0) if summary else 0.0,
        suspicious_count=int(summary.plagiarism_suspicious_count or 0) if summary else 0,
    )
    return HTMLResponse(content=html_body)


@app.get("/api/download-plagiarism-report/{submission_id}")
async def download_plagiarism_report(
    submission_id: int, request: Request, db: Session = Depends(get_db)
):
    """Download plagiarism highlight report as HTML file."""
    from fastapi.responses import Response  # type: ignore
    from urllib.parse import quote
    from app.plagiarism_highlight_report import (
        build_comparisons_from_db,
        generate_plagiarism_highlight_html,
    )

    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    if not submission.submission_text:
        raise HTTPException(status_code=404, detail="لا يوجد نص للمقارنة")

    summary = (
        db.query(GradingSummary)
        .filter(GradingSummary.submission_id == submission_id)
        .first()
    )
    batch = (
        db.query(BatchGrading).filter(BatchGrading.id == submission.batch_id).first()
        if submission.batch_id
        else None
    )
    uid = get_current_user_id(request)
    log_activity(
        db,
        "download_plagiarism_report",
        "export",
        f"تحميل تقرير التشابه - طالب #{submission_id}",
        user_id=uid,
        user_name=_get_user_display(db, uid),
        user_email=_get_user_email(db, uid),
        ip_address=_get_client_ip(request),
        user_agent=_get_user_agent(request),
    )
    comparisons = build_comparisons_from_db(db, submission)
    html_body = generate_plagiarism_highlight_html(
        student_name=submission.student_name,
        submission_id=submission_id,
        batch_name=str(getattr(batch, "batch_name", "") or ""),
        batch_id=int(submission.batch_id) if submission.batch_id else None,
        student_text=str(submission.submission_text),
        comparisons=comparisons,
        max_similarity=float(summary.plagiarism_max_similarity or 0) if summary else 0.0,
        suspicious_count=int(summary.plagiarism_suspicious_count or 0) if summary else 0,
    )
    safe_name = re.sub(r"[^\w\s\-]", "", _orm_str(submission.student_name) or "student")[:40]
    filename = f"plagiarism_{safe_name}_{submission_id}.html"
    return Response(
        content=html_body.encode("utf-8"),
        media_type="text/html; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"
        },
    )


@app.get("/api/download-report/{submission_id}")
async def download_report(submission_id: int, request: Request, db: Session = Depends(get_db)):
    """Download individual student report as PDF"""
    report = (
        db.query(StudentReport)
        .filter(StudentReport.submission_id == submission_id)
        .first()
    )

    if not report or not os.path.exists(str(report.report_file_path)):  # type: ignore
        raise HTTPException(status_code=404, detail="Report not found")

    uid = get_current_user_id(request)
    log_activity(db, "download_report_pdf", "export", f"تحميل تقرير PDF - طالب #{submission_id}", user_id=uid, user_name=_get_user_display(db, uid), user_email=_get_user_email(db, uid), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request))

    return FileResponse(
        str(report.report_file_path),  # type: ignore
        media_type="application/pdf",
        filename=f"report_{submission_id}.pdf",
    )


@app.get("/api/download-report-word/{submission_id}")
async def download_report_word(submission_id: int, request: Request, db: Session = Depends(get_db)):
    """Download individual student report as Word document"""
    uid = get_current_user_id(request)
    log_activity(db, "download_report_word", "export", f"تحميل تقرير Word - طالب #{submission_id}", user_id=uid, user_name=_get_user_display(db, uid), user_email=_get_user_email(db, uid), ip_address=_get_client_ip(request), user_agent=_get_user_agent(request))
    from fastapi.responses import Response  # type: ignore
    from docx import Document  # type: ignore
    from docx.shared import RGBColor, Cm, Pt, Inches  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    import io
    import json
    import os as _os
    from urllib.parse import quote
    from datetime import datetime as _dt
    import re as _re

    # Get submission and grading summary
    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    # Governance export gate (GOVERNANCE_RESPONSE_PROTOCOLS_v1)
    _gov_snapshot = None
    if getattr(submission, "grading_snapshot_json", None):
        try:
            _gov_snapshot = json.loads(str(submission.grading_snapshot_json))
        except (json.JSONDecodeError, TypeError):
            _gov_snapshot = None
    if _gov_snapshot:
        _gov_drift = _gov_snapshot.get("governance_drift") or {}
        if not _gov_drift and _gov_snapshot:
            try:
                from app.governance_drift_monitor import analyze_submission_governance_drift
                _gov_drift = analyze_submission_governance_drift(_gov_snapshot)
            except Exception:
                _gov_drift = {}
        _drift_export = (_gov_drift.get("governance_responses") or {}).get("export_policy") or {}
        _guardrails = (_gov_snapshot.get("criterion_authority_guardrails") or {})
        from app.criterion_authority_guardrails import merge_export_policy_with_guardrails
        _export_policy = merge_export_policy_with_guardrails(_drift_export, _guardrails)
        if _export_policy.get("gate") == "block_until_review":
            raise HTTPException(
                status_code=403,
                detail=_export_policy.get("message_ar") or "تصدير التقرير موقوف — مراجعة governance مطلوبة.",
            )

    summary = (
        db.query(GradingSummary)
        .filter(GradingSummary.submission_id == submission_id)
        .first()
    )
    grading_results = (
        db.query(GradingResult)
        .filter(GradingResult.submission_id == submission_id)
        .all()
    )

    _rtl_template = _os.path.join(_os.path.dirname(__file__), "app", "templates", "rtl_template.docx")
    doc = Document(_rtl_template)

    # Document-level RTL: colons, digits, and % render correctly next to Arabic.
    _root_settings = doc.settings.element
    if _root_settings.find(qn("w:bidi")) is None:
        _bd0 = OxmlElement("w:bidi")
        _bd0.set(qn("w:val"), "1")
        _root_settings.append(_bd0)

    # ── Color Palette ──
    PRIMARY = RGBColor(30, 58, 138)  # #1e3a8a
    BODY_TEXT = RGBColor(31, 41, 55)  # #1f2937
    WHITE = RGBColor(255, 255, 255)
    SLATE = RGBColor(107, 114, 128)  # #6b7280
    GREEN = RGBColor(16, 185, 129)   # #10b981
    BLUE = RGBColor(59, 130, 246)    # #3b82f6
    PURPLE = RGBColor(124, 58, 237)  # #7c3aed
    RED = RGBColor(239, 68, 68)      # #ef4444
    GOLD = RGBColor(234, 179, 8)     # #eab308

    style = doc.styles['Normal']
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.right_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("BTEC Student Report  |  AI Grader Platform")
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(148, 163, 184)
        fr.font.name = 'Calibri'

    _LRE = "\u202a"  # LEFT-TO-RIGHT EMBEDDING
    _PDF = "\u202c"  # POP DIRECTIONAL FORMATTING

    def _ltr_embed(text):
        """Numbers, %, /, Latin tokens inside Arabic paragraphs (LRE…PDF)."""
        if text is None:
            return ""
        return f"{_LRE}{text}{_PDF}"

    def _report_text(text) -> str:
        from app.report_feedback_formatter import clean_report_text
        return clean_report_text(str(text or ""))

    def _set_run_cs(run, font_name='Calibri'):
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:cs'), font_name)
        for _ex in rPr.findall(qn('w:rtl')):
            rPr.remove(_ex)
        _rtl = OxmlElement('w:rtl')
        _rtl.set(qn('w:val'), '1')
        rPr.append(_rtl)

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        for existing in pPr.findall(qn('w:bidi')):
            pPr.remove(existing)
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        for existing in pPr.findall(qn('w:jc')):
            pPr.remove(existing)

    def set_cell_shading(cell, hex_color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def clear_cell_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'FFFFFF')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def set_cell_borders(cell, color="CCCCCC", sz="4", sides=None):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in (sides or ('top', 'left', 'bottom', 'right')):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def set_cell_margin(cell, top=0, bottom=0, start=0, end=0):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for name, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
            mar = OxmlElement(f'w:{name}')
            mar.set(qn('w:w'), str(val))
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        tcPr.append(tcMar)

    def set_cell_vertical_alignment(cell, align="center"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), align)
        tcPr.append(va)

    def set_table_bidi(_tbl):
        # Intentionally empty: w:bidiVisual mirrors column order and places labels on the
        # visual left / values on the right. With col0=value, col1=label, plain LTR column
        # order gives label on the right as expected for Arabic.

        return

    def add_heading(text, level=1, color=PRIMARY, bg_color=None):
        p = doc.add_paragraph()
        set_rtl(p)
        if level == 1:
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            # In RTL paragraphs MS Word interprets WD_ALIGN_PARAGRAPH.RIGHT
            # (<w:jc val="right"/>) as the LOGICAL right side, which then
            # renders on the visual LEFT in Arabic — the opposite of what
            # we want. Use the bidi-aware "start" value which always lands
            # on the reading-start side: visual RIGHT in RTL, visual LEFT
            # in LTR. Since these headings always sit in RTL paragraphs
            # (bidi was set by set_rtl above), this lands on the visual
            # RIGHT in Arabic where Arabic headings belong.
            _pPr = p._p.get_or_add_pPr()
            for _ex in _pPr.findall(qn('w:jc')):
                _pPr.remove(_ex)
            _jc = OxmlElement('w:jc')
            _jc.set(qn('w:val'), 'start')
            _pPr.append(_jc)

        if bg_color:
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color)
            pPr.append(shd)

            pBdr = OxmlElement('w:pBdr')
            for bdr_name in ('top', 'left', 'bottom', 'right'):
                bdr = OxmlElement(f'w:{bdr_name}')
                bdr.set(qn('w:val'), 'single')
                bdr.set(qn('w:sz'), '12')
                bdr.set(qn('w:space'), '4')
                bdr.set(qn('w:color'), '3B82F6')
                pBdr.append(bdr)
            pPr.append(pBdr)

        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(22) if level == 1 else Pt(16)
        r.font.color.rgb = color
        r.font.name = 'Calibri'
        _set_run_cs(r)
        return p

    def add_bullet(text):
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"• {text}")
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r.font.color.rgb = BODY_TEXT
        _set_run_cs(r)
        return p

    add_heading(" تقرير تصحيح واجب BTEC", level=1, bg_color="EFF6FF")

    sub_p = doc.add_paragraph()
    set_rtl(sub_p)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(16)
    sub_r = sub_p.add_run("BTEC Assignment Grading Report")
    sub_r.font.size = Pt(12)
    sub_r.font.color.rgb = SLATE
    sub_r.font.name = 'Calibri'
    _set_run_cs(sub_r)

    grading_snapshot: dict[str, Any] | None = None
    if getattr(submission, "grading_snapshot_json", None):
        try:
            _parsed_snapshot = json.loads(str(submission.grading_snapshot_json))
            if isinstance(_parsed_snapshot, dict):
                grading_snapshot = _parsed_snapshot
                from app.official_grade import resolve_official_grade

                official_w = resolve_official_grade(
                    grading_snapshot,
                    reapply_pipeline=True,
                )
                if official_w.reapply_change_count > 0:
                    from app.criteria_result_finalizer import sync_criteria_results_to_db

                    submission.grading_snapshot_json = json.dumps(  # type: ignore
                        grading_snapshot, ensure_ascii=False
                    )
                    sync_criteria_results_to_db(db, submission.id, grading_snapshot)
                    db.commit()
        except Exception:
            grading_snapshot = None

    # ──────────────────────────────────────────────────────────────────
    #  FALLBACK: If grading_snapshot_json is missing (legacy submissions
    #  graded before the column was added), reconstruct a snapshot-like
    #  dict from the database so the Word report matches the PDF format
    #  instead of falling back to the degraded legacy renderer.
    # ──────────────────────────────────────────────────────────────────
    if not grading_snapshot:
        # 1) Try GradingCache lookup (best fidelity if available)
        try:
            from app.models import GradingCache as _GradingCache  # type: ignore
            from app.batch_grader import generate_content_fingerprint as _gen_fp  # type: ignore
            if submission.submission_text:
                _fp = _gen_fp(str(submission.submission_text))
                _fp_id = _fp.get("fingerprint_id") if isinstance(_fp, dict) else None
                if _fp_id:
                    _candidates = (
                        db.query(_GradingCache)
                        .filter(_GradingCache.result_json.like(f"%{_fp_id}%"))
                        .order_by(_GradingCache.created_at.desc())
                        .limit(5)
                        .all()
                    )
                    for _cand in _candidates:
                        try:
                            _cand_json = json.loads(_cand.result_json)
                            if isinstance(_cand_json, dict):
                                _cf = _cand_json.get("content_fingerprint", {}) or {}
                                if (_cf.get("fingerprint_id") == _fp_id or
                                        _cand_json.get("student_name") == submission.student_name):
                                    grading_snapshot = _cand_json
                                    break
                        except Exception:
                            continue
        except Exception:
            pass

    if not grading_snapshot:
        # 2) Reconstruct snapshot from DB tables
        _crit_results_list = []
        for _r in grading_results:
            _crit = _r.criteria
            if not _crit:
                continue

            _key_pts = []
            try:
                _kp_raw = _crit.key_points if _crit.key_points else "[]"
                _kp_parsed = json.loads(_kp_raw) if isinstance(_kp_raw, str) else _kp_raw
                if isinstance(_kp_parsed, list):
                    _key_pts = [str(x).strip() for x in _kp_parsed if str(x).strip()]
            except Exception:
                _key_pts = []

            _is_ok = bool(_r.achieved)

            _decision_matrix = []
            if _key_pts:
                for _kp in _key_pts:
                    _decision_matrix.append({
                        "requirement": _kp,
                        "met": _is_ok,
                        "evidence": ("متحقق بناءً على تقييم النظام." if _is_ok else "-"),
                    })
            else:
                _req_text = (_crit.criteria_description or _crit.criteria_name or _crit.criteria_level or "").strip()
                _decision_matrix.append({
                    "requirement": _req_text or (_crit.criteria_level or ""),
                    "met": _is_ok,
                    "evidence": ("متحقق بناءً على تقييم النظام." if _is_ok else "-"),
                })

            _nlr = []
            if _r.next_level_requirements:
                try:
                    _nlr = json.loads(_r.next_level_requirements)
                except Exception:
                    _nlr = []

            _missing = []
            if _r.missing_points:
                try:
                    _missing = json.loads(_r.missing_points)
                except Exception:
                    _missing = []
            if not isinstance(_missing, list):
                _missing = []

            _crit_results_list.append({
                "criteria_level": _crit.criteria_level or "",
                "achieved": _is_ok,
                "feedback": _r.feedback or "",
                "explanation": "",
                "decision_matrix": _decision_matrix,
                "next_level_requirements": _nlr,
                "missing_points": _missing,
                "rule_validation": {},
            })

        grading_snapshot = {
            "criteria_results": _crit_results_list,
            "student_name": _orm_str(submission.student_name),
            "student_email": _orm_str(submission.student_email),
        }

        if summary:
            grading_snapshot["total_score"] = _orm_int(summary.total_score, 0)
            grading_snapshot["max_score"] = _orm_int(summary.max_score, 100)
            grading_snapshot["percentage"] = _orm_float(summary.percentage, 0.0)
            grading_snapshot["grade_level"] = _orm_str(summary.grade_level) or "-"
            grading_snapshot["ai_likelihood"] = _orm_int(summary.ai_likelihood, 0)

            _ai_score = _orm_int(summary.ai_likelihood, 0)
            if _ai_score <= 20:
                _ai_label = "محتوى بشري"
                _ai_icon = "🟢"
            elif _ai_score <= 40:
                _ai_label = "محتوى بشري مع مساعدة بسيطة"
                _ai_icon = "🟡"
            elif _ai_score <= 60:
                _ai_label = "محتوى هجين"
                _ai_icon = "🟠"
            elif _ai_score <= 80:
                _ai_label = "محتوى مولد آلياً بشكل كبير"
                _ai_icon = "🔴"
            else:
                _ai_label = "محتوى مولد آلياً بالكامل"
                _ai_icon = "⛔"

            grading_snapshot["ai_detection_info"] = {
                "score": _ai_score,
                "risk_classification": {"icon": _ai_icon, "label_ar": _ai_label},
                "indicators_detected": [],
            }

            grading_snapshot["plagiarism_info"] = {
                "max_similarity": _orm_float(summary.plagiarism_max_similarity, 0.0),
                "matches": [],
            }

            try:
                grading_snapshot["strengths"] = json.loads(summary.strengths) if summary.strengths else []
            except Exception:
                grading_snapshot["strengths"] = []
            try:
                grading_snapshot["improvements"] = json.loads(summary.improvements) if summary.improvements else []
            except Exception:
                grading_snapshot["improvements"] = []

            grading_snapshot["overall_feedback"] = summary.overall_feedback or ""

        # Persist this reconstructed snapshot so future Word exports skip the rebuild
        try:
            submission.grading_snapshot_json = json.dumps(grading_snapshot, ensure_ascii=False)  # type: ignore
            db.commit()
        except Exception:
            db.rollback()

    def _fill_student_info_table(info_data):
        info_tbl = doc.add_table(rows=len(info_data), cols=2)
        info_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
        info_tbl.autofit = False
        set_table_bidi(info_tbl)
        info_tbl.columns[0].width = Cm(10)
        info_tbl.columns[1].width = Cm(5)
        for ri, (label, value) in enumerate(info_data):
            label_cell = info_tbl.cell(ri, 1)
            set_cell_shading(label_cell, "F3F4F6")
            set_cell_borders(label_cell, color="E5E7EB", sz="4")
            set_cell_margin(label_cell, top=100, bottom=100, start=120, end=120)
            set_cell_vertical_alignment(label_cell, "center")
            lp = label_cell.paragraphs[0]
            set_rtl(lp)
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(11)
            lr.font.color.rgb = BODY_TEXT
            lr.font.name = 'Calibri'
            _set_run_cs(lr)

            val_cell = info_tbl.cell(ri, 0)
            set_cell_shading(val_cell, "FFFFFF")
            set_cell_borders(val_cell, color="E5E7EB", sz="4")
            set_cell_margin(val_cell, top=100, bottom=100, start=120, end=120)
            set_cell_vertical_alignment(val_cell, "center")
            vp = val_cell.paragraphs[0]
            set_rtl(vp)
            vr = vp.add_run(value)
            vr.font.size = Pt(11)
            vr.font.color.rgb = SLATE
            vr.font.name = 'Calibri'
            _set_run_cs(vr)

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def _criteria_sort_from_dict(crit):
        level = crit.get("criteria_level", "")
        short = level.split(".")[-1] if "." in level else level
        type_order = {"P": 0, "M": 1, "D": 2}
        letter = short[0].upper() if short else "Z"
        num = short[1:] if len(short) > 1 else "0"
        try:
            num_val = int(num)
        except ValueError:
            num_val = 99
        return (type_order.get(letter, 9), num_val)

    if grading_snapshot:
        # ── Body from same payload as PDF (grading_snapshot) ──────────
        gs = grading_snapshot
        info_data_snap = [
            ("اسم الطالب:", submission.student_name or "—"),
            ("تاريخ التصحيح:", _ltr_embed(_dt.now().strftime('%Y-%m-%d %H:%M'))),
        ]
        fp_snap = gs.get("content_fingerprint") or {}
        if fp_snap:
            info_data_snap.append(("عدد الكلمات:", _ltr_embed(str(fp_snap.get("word_count", "-")))))
        _fill_student_info_table(info_data_snap)

        add_heading(" الملخص التنفيذي", level=2, color=PURPLE)
        from app.official_grade import resolve_official_grade

        official_snap = resolve_official_grade(gs, reapply_pipeline=False)
        gdm = official_snap.grade_display_metrics or gs.get("grade_display_metrics") or {}
        erg = gdm.get("expected_runtime_grade") or {}
        grade_level_s = official_snap.grade_label
        inst_s = gs.get("institutional_resolution") or {}
        inst_label_s = inst_s.get("display_grade_ar") if inst_s else None
        if inst_label_s and str(gdm.get("final_btec_grade", "U")).upper() == "U":
            grade_level_s = f"{grade_level_s} — {inst_label_s}"
        expected_grade_s = erg.get("expected_btec_grade_label") if erg else None
        percentage_s = float(gdm.get("criteria_completion_pct", gs.get("percentage", 0)))
        total_score_s = gs.get("total_score", 0)
        max_score_s = gs.get("max_score", 100)
        highest_crit_s = gdm.get("highest_criterion_achieved") or "—"
        exec_mode_s = gdm.get("execution_mode") or "PRO"
        from app.rule_bundle import format_rule_bundle_label, provenance_from_payload

        rule_bundle_s = format_rule_bundle_label(provenance_from_payload(gs))
        ai_info_s = gs.get("ai_detection_info") or {}
        try:
            ai_score_s = int(ai_info_s.get("score", gs.get("ai_likelihood", 0)))
        except (TypeError, ValueError):
            ai_score_s = 0
        ai_risk_s = ai_info_s.get("risk_classification") or {}
        ai_icon_s = ai_risk_s.get("icon", "❓")
        plag_info_s = gs.get("plagiarism_info") or {}
        try:
            plag_max_s = float(plag_info_s.get("max_similarity", 0))
        except (TypeError, ValueError):
            plag_max_s = 0.0

        summary_data_s = [
            ("التقدير المعتمد:", _ltr_embed(grade_level_s)),
        ]
        if expected_grade_s and erg.get("expected_btec_grade") != gdm.get("final_btec_grade"):
            summary_data_s.append(
                ("التقدير المتوقع (عند إكمال التشغيل والتأكد من اللعبة):", _ltr_embed(str(expected_grade_s)))
            )
            summary_data_s.append(
                ("ملاحظة التقدير المتوقع:", erg.get("disclaimer_ar") or "")
            )
        summary_data_s.extend([
            ("أعلى معيار متحقق:", _ltr_embed(str(highest_crit_s))),
            ("نسبة المعايير (تحليلي):", _ltr_embed(f"{percentage_s:.1f}%")),
            ("الدرجة الكلية:", _ltr_embed(f"{total_score_s} / {max_score_s}")),
            ("نسبة الذكاء الاصطناعي (إرشادي):", f"{ai_icon_s} {_ltr_embed(f'{ai_score_s}%')}"),
            ("نسبة الانتحال:", _ltr_embed(f"{plag_max_s:.1f}%")),
            ("وضع التصحيح:", _ltr_embed(str(exec_mode_s))),
            ("Rule Bundle:", _ltr_embed(rule_bundle_s)),
        ])

        sum_tbl_s = doc.add_table(rows=len(summary_data_s), cols=2)
        sum_tbl_s.alignment = WD_TABLE_ALIGNMENT.RIGHT
        sum_tbl_s.autofit = False
        set_table_bidi(sum_tbl_s)
        sum_tbl_s.columns[0].width = Cm(10)
        sum_tbl_s.columns[1].width = Cm(5)
        for ri, (label, value) in enumerate(summary_data_s):
            label_cell = sum_tbl_s.cell(ri, 1)
            set_cell_shading(label_cell, "EFF6FF")
            set_cell_borders(label_cell, color="60A5FA", sz="12")
            set_cell_margin(label_cell, top=120, bottom=120, start=140, end=140)
            set_cell_vertical_alignment(label_cell, "center")
            lp = label_cell.paragraphs[0]
            set_rtl(lp)
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(12)
            lr.font.color.rgb = BODY_TEXT
            lr.font.name = 'Calibri'
            _set_run_cs(lr)

            val_cell = sum_tbl_s.cell(ri, 0)
            set_cell_shading(val_cell, "3B82F6")
            set_cell_borders(val_cell, color="60A5FA", sz="12")
            set_cell_margin(val_cell, top=120, bottom=120, start=140, end=140)
            set_cell_vertical_alignment(val_cell, "center")
            vp = val_cell.paragraphs[0]
            set_rtl(vp)
            vr = vp.add_run(value)
            vr.bold = True
            vr.font.size = Pt(12)
            vr.font.color.rgb = WHITE
            vr.font.name = 'Calibri'
            _set_run_cs(vr)

        doc.add_paragraph().paragraph_format.space_after = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        add_heading("🔗 تحليل الانتحال (Plagiarism Analysis)", level=2, color=PURPLE)
        if plag_max_s <= 10:
            plag_text_s = f"✅ مقبول: تشابه طبيعي ({_ltr_embed(f'{plag_max_s:.1f}%')})"
            plag_color_s = GREEN
        elif plag_max_s <= 25:
            plag_text_s = f"🟡 تشابه ملحوظ - يحتاج مراقبة ({_ltr_embed(f'{plag_max_s:.1f}%')})"
            plag_color_s = GOLD
        elif plag_max_s <= 50:
            plag_text_s = f"🟠 انتحال مشتبه به - يحتاج تحقيق ({_ltr_embed(f'{plag_max_s:.1f}%')})"
            plag_color_s = RGBColor(249, 115, 22)
        elif plag_max_s <= 75:
            plag_text_s = f"🔴 انتحال واضح ({_ltr_embed(f'{plag_max_s:.1f}%')})"
            plag_color_s = RED
        else:
            plag_text_s = f"⛔ نسخ شبه كامل ({_ltr_embed(f'{plag_max_s:.1f}%')})"
            plag_color_s = RGBColor(127, 29, 29)
        p_plag_s = doc.add_paragraph()
        set_rtl(p_plag_s)
        r_plag_s = p_plag_s.add_run(plag_text_s)
        r_plag_s.font.color.rgb = plag_color_s
        r_plag_s.font.size = Pt(12)
        r_plag_s.bold = True
        r_plag_s.font.name = 'Calibri'
        _set_run_cs(r_plag_s)

        matches_s = plag_info_s.get("matches") or []
        if matches_s:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            hdr_m = doc.add_paragraph()
            set_rtl(hdr_m)
            r_hdr_m = hdr_m.add_run("أعلى حالات التشابه مع طلاب آخرين:")
            r_hdr_m.font.size = Pt(12)
            r_hdr_m.font.name = 'Calibri'
            r_hdr_m.font.color.rgb = BODY_TEXT
            _set_run_cs(r_hdr_m)
            mtbl = doc.add_table(rows=1, cols=3)
            mtbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
            mtbl.autofit = False
            set_table_bidi(mtbl)
            mtbl.columns[0].width = Cm(2.5)
            mtbl.columns[1].width = Cm(2.5)
            mtbl.columns[2].width = Cm(8.0)
            mh = ["الحالة", "النسبة", "الطالب المقارن"]
            for ci, hd in enumerate(mh):
                cell = mtbl.cell(0, ci)
                set_cell_shading(cell, "F3F4F6")
                set_cell_borders(cell, color="D1D5DB", sz="8")
                p = cell.paragraphs[0]
                set_rtl(p)
                r = p.add_run(hd)
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = 'Calibri'
                _set_run_cs(r)
            for mi, m in enumerate(matches_s):
                row_cells = mtbl.add_row().cells
                susp = "مشبوه" if m.get("is_suspicious") else "عادي"
                pct = _ltr_embed(f"{m.get('percentage', 0)}%")
                stu = str(m.get("student", ""))
                for ci, val in enumerate([susp, pct, stu]):
                    cell = row_cells[ci]
                    set_cell_borders(cell, color="D1D5DB", sz="8")
                    set_cell_margin(cell, top=60, bottom=60, start=80, end=80)
                    p = cell.paragraphs[0]
                    set_rtl(p)
                    r = p.add_run(val)
                    r.font.size = Pt(11)
                    r.font.name = 'Calibri'
                    _set_run_cs(r)
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        add_heading(" تحليل الذكاء الاصطناعي (AI Content Analysis)", level=2, color=PURPLE)
        if ai_score_s <= 20:
            ai_color_s = GREEN
        elif ai_score_s <= 40:
            ai_color_s = GOLD
        elif ai_score_s <= 60:
            ai_color_s = RGBColor(249, 115, 22)
        elif ai_score_s <= 80:
            ai_color_s = RED
        else:
            ai_color_s = RGBColor(127, 29, 29)
        ai_label_snap = ai_risk_s.get("label_ar", "غير محدد")
        p_ai_s = doc.add_paragraph()
        set_rtl(p_ai_s)
        r_ai_s = p_ai_s.add_run(f"{ai_icon_s} النسبة: {_ltr_embed(f'{ai_score_s}%')} — {ai_label_snap}")
        r_ai_s.font.color.rgb = ai_color_s
        r_ai_s.font.size = Pt(12)
        r_ai_s.bold = True
        r_ai_s.font.name = 'Calibri'
        _set_run_cs(r_ai_s)

        indicators_s = ai_info_s.get("indicators_detected") or []
        if indicators_s:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            p_ind = doc.add_paragraph()
            set_rtl(p_ind)
            r_ind = p_ind.add_run("المؤشرات المكتشفة:")
            r_ind.bold = True
            r_ind.font.size = Pt(12)
            r_ind.font.color.rgb = BLUE
            r_ind.font.name = 'Calibri'
            _set_run_cs(r_ind)
            for ind in indicators_s[:10]:
                add_bullet(str(ind))
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        _crs_raw = gs.get("criteria_results") or []
        crs_snap: list[Any] = _crs_raw if isinstance(_crs_raw, list) else []
        achieved_lvls = [c.get("criteria_level", "") for c in crs_snap if c.get("achieved")]
        if achieved_lvls:
            def _lvl_key_snap(lv):
                lv_str = lv.split(".")[-1].strip() if "." in lv else lv.strip()
                t = lv_str[:1].upper()
                m = _re.search(r'\d+', lv_str)
                n = int(m.group()) if m else 999
                return ({"P": 0, "M": 1, "D": 2}.get(t, 3), n)

            achieved_lvls.sort(key=_lvl_key_snap)
            cur_lv = achieved_lvls[-1]
            add_heading(" أعلى معيار متحقق (ليس الدرجة BTEC النهائية)", level=2, color=PURPLE)
            p_lv = doc.add_paragraph()
            set_rtl(p_lv)
            r1 = p_lv.add_run("أنت حالياً في مستوى: ")
            r1.font.size = Pt(13)
            r1.font.color.rgb = BLUE
            r1.font.name = 'Calibri'
            _set_run_cs(r1)
            r2 = p_lv.add_run(cur_lv)
            r2.bold = True
            r2.font.size = Pt(13)
            r2.font.color.rgb = BLUE
            r2.font.name = 'Calibri'
            _set_run_cs(r2)
            doc.add_paragraph().paragraph_format.space_after = Pt(10)

        add_heading(" تفاصيل المعايير", level=2, color=PURPLE)
        for criteria in sorted(crs_snap, key=_criteria_sort_from_dict):
            level_c = criteria.get("criteria_level", "")
            is_ok = bool(criteria.get("achieved", False))
            human_review = (criteria.get("achievement_authority") or "") == "HUMAN_REVIEW_REQUIRED"
            if human_review:
                card_ac = "FEF3C7"
                card_bd = "F59E0B"
                st_icon = "⏸"
                st_txt = "مراجعة بشرية مطلوبة (Human Review Required)"
            elif is_ok:
                card_ac = "D1FAE5"
                card_bd = "10B981"
                st_icon = "✅"
                st_txt = "متحقق (Achieved)"
            else:
                card_ac = "FEE2E2"
                card_bd = "EF4444"
                st_icon = "❌"
                st_txt = "غير متحقق (Not Achieved)"

            card_tbl_c = doc.add_table(rows=1, cols=2)
            card_tbl_c.alignment = WD_TABLE_ALIGNMENT.RIGHT
            card_tbl_c.autofit = False
            set_table_bidi(card_tbl_c)
            card_tbl_c.columns[0].width = Cm(7.5)
            card_tbl_c.columns[1].width = Cm(7.5)
            for ci in range(2):
                cell = card_tbl_c.cell(0, ci)
                set_cell_shading(cell, card_ac)
                set_cell_borders(cell, color=card_bd, sz="24")
                set_cell_margin(cell, top=100, bottom=100, start=120, end=120)
                set_cell_vertical_alignment(cell, "center")

            c0s = card_tbl_c.cell(0, 0)
            p0s = c0s.paragraphs[0]
            set_rtl(p0s)
            r0s = p0s.add_run(st_txt)
            r0s.font.size = Pt(14)
            r0s.font.color.rgb = PRIMARY
            r0s.font.name = 'Calibri'
            _set_run_cs(r0s)

            c1s = card_tbl_c.cell(0, 1)
            p1s = c1s.paragraphs[0]
            set_rtl(p1s)
            r1s = p1s.add_run(f"{st_icon} المعيار {level_c}")
            r1s.font.size = Pt(14)
            r1s.font.color.rgb = PRIMARY
            r1s.font.name = 'Calibri'
            _set_run_cs(r1s)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)

            expl = str(criteria.get("explanation", "") or "")
            if expl:
                ex_p = doc.add_paragraph()
                set_rtl(ex_p)
                ex_l = ex_p.add_run("الشرح المفصل: ")
                ex_l.bold = True
                ex_l.font.size = Pt(13)
                ex_l.font.color.rgb = BLUE
                ex_l.font.name = 'Calibri'
                _set_run_cs(ex_l)
                ex_b = doc.add_paragraph()
                set_rtl(ex_b)
                ex_r = ex_b.add_run(_report_text(expl))
                ex_r.font.size = Pt(12)
                ex_r.font.color.rgb = BODY_TEXT
                ex_r.font.name = 'Calibri'
                _set_run_cs(ex_r)
                ex_b.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

            fb_raw = str(criteria.get("feedback", "") or "")
            if fb_raw or criteria.get("runtime_observation_note_ar"):
                from app.report_feedback_formatter import format_criterion_feedback_for_report
                fb = format_criterion_feedback_for_report(
                    fb_raw,
                    runtime_note_ar=criteria.get("runtime_observation_note_ar"),
                )
                fb_p = doc.add_paragraph()
                set_rtl(fb_p)
                fb_l = fb_p.add_run("الملاحظات: ")
                fb_l.bold = True
                fb_l.font.size = Pt(13)
                fb_l.font.color.rgb = BLUE
                fb_l.font.name = 'Calibri'
                _set_run_cs(fb_l)
                fb_b = doc.add_paragraph()
                set_rtl(fb_b)
                fb_ru = fb_b.add_run(fb)
                fb_ru.font.size = Pt(12)
                fb_ru.font.color.rgb = BODY_TEXT
                fb_ru.font.name = 'Calibri'
                _set_run_cs(fb_ru)
                fb_b.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

            dm = criteria.get("decision_matrix") or []
            if isinstance(dm, list) and dm:
                dm_p = doc.add_paragraph()
                set_rtl(dm_p)
                dm_l = dm_p.add_run("متطلبات المعيار: ")
                dm_l.bold = True
                dm_l.font.size = Pt(13)
                dm_l.font.color.rgb = BLUE
                dm_l.font.name = 'Calibri'
                _set_run_cs(dm_l)

                dm_tbl = doc.add_table(rows=1, cols=3)
                dm_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
                dm_tbl.autofit = False
                set_table_bidi(dm_tbl)
                dm_tbl.columns[0].width = Cm(7.0)
                dm_tbl.columns[1].width = Cm(3.0)
                dm_tbl.columns[2].width = Cm(5.0)

                headers = ["الدليل (Evidence)", "الحالة (Status)", "المتطلب (Requirement)"]
                for ci, hd in enumerate(headers):
                    cell = dm_tbl.cell(0, ci)
                    set_cell_shading(cell, "6366F1")
                    set_cell_borders(cell, color="D1D5DB", sz="12")
                    set_cell_margin(cell, top=80, bottom=80, start=80, end=80)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_rtl(p)
                    r = p.add_run(hd)
                    r.bold = True
                    r.font.size = Pt(11)
                    r.font.color.rgb = WHITE
                    r.font.name = 'Calibri'
                    _set_run_cs(r)

                for row_data in dm:
                    if not isinstance(row_data, dict):
                        continue
                    row = dm_tbl.add_row()
                    met = bool(row_data.get("met", False))
                    ev = _report_text(row_data.get("evidence", "-"))
                    if not ev or ev == "Not found":
                        ev = "-"
                    req = _report_text(row_data.get("requirement", ""))
                    status_symbol = "✅" if met else "❌"
                    status_str = "متحقق" if met else "غير متحقق"
                    cell_data = [ev, f"{status_symbol} {status_str}", req]
                    for ci, cd in enumerate(cell_data):
                        cell = row.cells[ci]
                        set_cell_shading(cell, "FAFAFA" if ci == 1 else "FFFFFF")
                        set_cell_borders(cell, color="D1D5DB", sz="12")
                        set_cell_margin(cell, top=80, bottom=80, start=80, end=80)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        set_rtl(p)
                        r = p.add_run(_report_text(cd) if ci != 1 else cd)
                        r.font.size = Pt(11)
                        r.font.color.rgb = BODY_TEXT
                        r.font.name = 'Calibri'
                        _set_run_cs(r)

                doc.add_paragraph().paragraph_format.space_after = Pt(10)

            rv = criteria.get("rule_validation") or {}
            if isinstance(rv, dict) and rv:
                mk = rv.get("matched_keywords") or []
                mp = rv.get("matched_patterns") or []
                if not isinstance(mk, list):
                    mk = []
                if not isinstance(mp, list):
                    mp = []
                if mk or mp:
                    rv_p = doc.add_paragraph()
                    set_rtl(rv_p)
                    rv_l = rv_p.add_run("التحقق من الكلمات المفتاحية: ")
                    rv_l.bold = True
                    rv_l.font.size = Pt(13)
                    rv_l.font.color.rgb = BLUE
                    rv_l.font.name = 'Calibri'
                    _set_run_cs(rv_l)
                    found_terms = [str(t) for t in mk + mp]
                    top_t = found_terms[:5]
                    terms_line = ", ".join(top_t)
                    if len(found_terms) > 5:
                        terms_line += "..."
                    rv_b = doc.add_paragraph()
                    set_rtl(rv_b)
                    rv_r = rv_b.add_run(f"تم العثور على مصطلحات: {terms_line}")
                    rv_r.font.size = Pt(11)
                    rv_r.font.name = 'Calibri'
                    _set_run_cs(rv_r)
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)

            nl = criteria.get("next_level_requirements")
            if nl and not is_ok:
                nl_p = doc.add_paragraph()
                set_rtl(nl_p)
                nl_l = nl_p.add_run("للوصول للمعيار التالي: ")
                nl_l.bold = True
                nl_l.font.size = Pt(13)
                nl_l.font.color.rgb = BLUE
                nl_l.font.name = 'Calibri'
                _set_run_cs(nl_l)
                if isinstance(nl, list):
                    for req_text in nl:
                        add_bullet(str(req_text))
                else:
                    r_p = doc.add_paragraph()
                    set_rtl(r_p)
                    r_run = r_p.add_run(str(nl))
                    r_run.font.size = Pt(11)
                    r_run.font.name = 'Calibri'
                    _set_run_cs(r_run)
            doc.add_paragraph().paragraph_format.space_after = Pt(10)

        strengths_s = gs.get("strengths") or []
        improvements_s = gs.get("improvements") or []
        if strengths_s:
            add_heading("🟢 نقاط قوة الطالب", level=2, color=PURPLE)
            for i, st in enumerate(strengths_s, 1):
                p = doc.add_paragraph()
                set_rtl(p)
                r = p.add_run(f"{i}. {st}")
                r.font.size = Pt(12)
                r.font.color.rgb = BODY_TEXT
                r.font.name = 'Calibri'
                _set_run_cs(r)
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
        if improvements_s:
            add_heading("💡 التحسينات المطلوبة", level=2, color=PURPLE)
            mid_s = max(1, len(improvements_s) // 2)
            critical_s = improvements_s[:mid_s]
            suggested_s = improvements_s[mid_s:]
            if critical_s:
                p_c = doc.add_paragraph()
                set_rtl(p_c)
                r_c = p_c.add_run("تحسينات حرجة (يجب إكمالها) 🔴")
                r_c.bold = True
                r_c.font.size = Pt(13)
                r_c.font.color.rgb = BLUE
                r_c.font.name = 'Calibri'
                _set_run_cs(r_c)
                for imp in critical_s:
                    add_bullet(str(imp))
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            if suggested_s:
                p_s = doc.add_paragraph()
                set_rtl(p_s)
                r_s = p_s.add_run("تحسينات مقترحة (للحصول على درجة أعلى) 🟡")
                r_s.bold = True
                r_s.font.size = Pt(13)
                r_s.font.color.rgb = BLUE
                r_s.font.name = 'Calibri'
                _set_run_cs(r_s)
                for imp in suggested_s:
                    add_bullet(str(imp))
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

        ofb = gs.get("overall_feedback", "")
        if ofb:
            add_heading("التقييم العام", level=2, color=PURPLE)
            fb_o = doc.add_paragraph()
            set_rtl(fb_o)
            fb_or = fb_o.add_run(_report_text(ofb))
            fb_or.font.size = Pt(12)
            fb_or.font.color.rgb = BODY_TEXT
            fb_or.font.name = 'Calibri'
            _set_run_cs(fb_or)
            fb_o.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    else:
        # ════════════════════════════════════════════════════════════
        #  STUDENT INFO TABLE (legacy: no grading_snapshot)
        # ════════════════════════════════════════════════════════════
        info_tbl = doc.add_table(rows=2, cols=2)
        info_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
        info_tbl.autofit = False
        set_table_bidi(info_tbl)
        info_tbl.columns[0].width = Cm(10)
        info_tbl.columns[1].width = Cm(5)

        info_data = [
            ("اسم الطالب:", submission.student_name or "—"),
            ("تاريخ التصحيح:", _ltr_embed(_dt.now().strftime('%Y-%m-%d %H:%M'))),
        ]

        for ri, (label, value) in enumerate(info_data):
            label_cell = info_tbl.cell(ri, 1)
            set_cell_shading(label_cell, "F3F4F6")
            set_cell_borders(label_cell, color="E5E7EB", sz="4")
            set_cell_margin(label_cell, top=100, bottom=100, start=120, end=120)
            set_cell_vertical_alignment(label_cell, "center")
            lp = label_cell.paragraphs[0]
            set_rtl(lp)
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(11)
            lr.font.color.rgb = BODY_TEXT
            lr.font.name = 'Calibri'
            _set_run_cs(lr)

            val_cell = info_tbl.cell(ri, 0)
            set_cell_shading(val_cell, "FFFFFF")
            set_cell_borders(val_cell, color="E5E7EB", sz="4")
            set_cell_margin(val_cell, top=100, bottom=100, start=120, end=120)
            set_cell_vertical_alignment(val_cell, "center")
            vp = val_cell.paragraphs[0]
            set_rtl(vp)
            vr = vp.add_run(value)
            vr.font.size = Pt(11)
            vr.font.color.rgb = SLATE
            vr.font.name = 'Calibri'
            _set_run_cs(vr)

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # ════════════════════════════════════════════════════════════
        #  EXECUTIVE SUMMARY
        # ════════════════════════════════════════════════════════════
        if summary:
            add_heading(" الملخص التنفيذي", level=2, color=PURPLE)

            grade_level = summary.grade_level or "-"
            percentage = summary.percentage or 0
            total_score = summary.total_score or 0
            max_score = summary.max_score or 100
            ai_score = summary.ai_likelihood or 0
            plag_max = summary.plagiarism_max_similarity or 0
            highest_crit = "—"
            exec_mode = "PRO"
            rule_bundle = "—"
            try:
                if getattr(submission, "grading_snapshot_json", None):
                    _gdm_snap = json.loads(str(submission.grading_snapshot_json))
                    _gdm = _gdm_snap.get("grade_display_metrics") or {}
                    if _gdm:
                        grade_level = _gdm.get("final_btec_grade_label") or grade_level
                        percentage = _gdm.get("criteria_completion_pct", percentage)
                        highest_crit = _gdm.get("highest_criterion_achieved") or "—"
                        exec_mode = _gdm.get("execution_mode") or exec_mode
                    from app.rule_bundle import format_rule_bundle_label, provenance_from_payload

                    _prov = provenance_from_payload(_gdm_snap)
                    rule_bundle = format_rule_bundle_label(_prov)
            except Exception:
                pass

            # Determine AI risk icon
            ai_icon = "❓"
            if ai_score <= 20: ai_icon = "🟢"
            elif ai_score <= 40: ai_icon = "🟡"
            elif ai_score <= 60: ai_icon = "🟠"
            elif ai_score <= 80: ai_icon = "🔴"
            else: ai_icon = "⛔"

            summary_data = [
                ("الدرجة BTEC النهائية:", _ltr_embed(str(grade_level))),
                ("أعلى معيار متحقق:", _ltr_embed(str(highest_crit))),
                ("إكمال المعايير:", _ltr_embed(f"{percentage:.1f}%")),
                ("الدرجة الكلية:", _ltr_embed(f"{total_score} / {max_score}")),
                ("نسبة الذكاء الاصطناعي (إرشادي):", f"{ai_icon} {_ltr_embed(f'{ai_score}%')}"),
                ("نسبة الانتحال:", _ltr_embed(f"{plag_max:.1f}%")),
                ("وضع التصحيح:", _ltr_embed(str(exec_mode))),
                ("Rule Bundle:", _ltr_embed(rule_bundle)),
            ]

            sum_tbl = doc.add_table(rows=len(summary_data), cols=2)
            sum_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
            sum_tbl.autofit = False
            set_table_bidi(sum_tbl)
            sum_tbl.columns[0].width = Cm(10)
            sum_tbl.columns[1].width = Cm(5)

            for ri, (label, value) in enumerate(summary_data):
                label_cell = sum_tbl.cell(ri, 1)
                set_cell_shading(label_cell, "EFF6FF")
                set_cell_borders(label_cell, color="60A5FA", sz="12")
                set_cell_margin(label_cell, top=120, bottom=120, start=140, end=140)
                set_cell_vertical_alignment(label_cell, "center")
                lp = label_cell.paragraphs[0]
                set_rtl(lp)
                lr = lp.add_run(label)
                lr.bold = True
                lr.font.size = Pt(12)
                lr.font.color.rgb = BODY_TEXT
                lr.font.name = 'Calibri'
                _set_run_cs(lr)

                val_cell = sum_tbl.cell(ri, 0)
                set_cell_shading(val_cell, "3B82F6")
                set_cell_borders(val_cell, color="60A5FA", sz="12")
                set_cell_margin(val_cell, top=120, bottom=120, start=140, end=140)
                set_cell_vertical_alignment(val_cell, "center")
                vp = val_cell.paragraphs[0]
                set_rtl(vp)
                vr = vp.add_run(value)
                vr.bold = True
                vr.font.size = Pt(12)
                vr.font.color.rgb = WHITE
                vr.font.name = 'Calibri'
                _set_run_cs(vr)

            doc.add_paragraph().paragraph_format.space_after = Pt(10)
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

            # ════════════════════════════════════════════════════════════
            #  PLAGIARISM ANALYSIS
            # ════════════════════════════════════════════════════════════
            if plag_max >= 0:
                add_heading("🔗 تحليل الانتحال (Plagiarism Analysis)", level=2, color=PURPLE)

                if plag_max <= 10:
                    plag_text = f"✅ مقبول: تشابه طبيعي ({_ltr_embed(f'{plag_max:.1f}%')})"
                    plag_color = GREEN
                elif plag_max <= 25:
                    plag_text = f"🟡 تشابه ملحوظ - يحتاج مراقبة ({_ltr_embed(f'{plag_max:.1f}%')})"
                    plag_color = GOLD
                elif plag_max <= 50:
                    plag_text = f"🟠 انتحال مشتبه به - يحتاج تحقيق ({_ltr_embed(f'{plag_max:.1f}%')})"
                    plag_color = RGBColor(249, 115, 22)
                elif plag_max <= 75:
                    plag_text = f"🔴 انتحال واضح ({_ltr_embed(f'{plag_max:.1f}%')})"
                    plag_color = RED
                else:
                    plag_text = f"⛔ نسخ شبه كامل ({_ltr_embed(f'{plag_max:.1f}%')})"
                    plag_color = RGBColor(127, 29, 29)

                p_plag = doc.add_paragraph()
                set_rtl(p_plag)
                r_plag = p_plag.add_run(plag_text)
                r_plag.font.color.rgb = plag_color
                r_plag.font.size = Pt(12)
                r_plag.bold = True
                r_plag.font.name = 'Calibri'
                _set_run_cs(r_plag)
                doc.add_paragraph().paragraph_format.space_after = Pt(10)

            # ════════════════════════════════════════════════════════════
            #  AI DETECTION ANALYSIS
            # ════════════════════════════════════════════════════════════
            if ai_score >= 0:
                add_heading(" تحليل الذكاء الاصطناعي (AI Content Analysis)", level=2, color=PURPLE)

                ai_label_val = "غير محدد"
                if ai_score <= 20:
                    ai_color = GREEN
                    ai_label_val = "محتوى بشري"
                elif ai_score <= 40:
                    ai_color = GOLD
                    ai_label_val = "محتوى بشري مع مساعدة بسيطة"
                elif ai_score <= 60:
                    ai_color = RGBColor(249, 115, 22)
                    ai_label_val = "محتوى هجين"
                elif ai_score <= 80:
                    ai_color = RED
                    ai_label_val = "محتوى مولد آلياً بشكل كبير"
                else:
                    ai_color = RGBColor(127, 29, 29)
                    ai_label_val = "محتوى مولد آلياً بالكامل"

                p_ai = doc.add_paragraph()
                set_rtl(p_ai)
                r_ai = p_ai.add_run(f"{ai_icon} النسبة: {_ltr_embed(f'{ai_score}%')} — {ai_label_val}")
                r_ai.font.color.rgb = ai_color
                r_ai.font.size = Pt(12)
                r_ai.bold = True
                r_ai.font.name = 'Calibri'
                _set_run_cs(r_ai)
                doc.add_paragraph().paragraph_format.space_after = Pt(10)

            # ════════════════════════════════════════════════════════════
            #  CURRENT LEVEL ACHIEVEMENT
            # ════════════════════════════════════════════════════════════
            achieved_levels = [r.criteria.criteria_level for r in grading_results if r.achieved and r.criteria]
            if achieved_levels:
                # Sort to get highest
                def _word_sort_key_lvl(lv):
                    lv_str = lv.split(".")[-1].strip() if "." in lv else lv.strip()
                    t = lv_str[:1].upper()
                    m = _re.search(r'\d+', lv_str)
                    n = int(m.group()) if m else 999
                    return ({"P": 0, "M": 1, "D": 2}.get(t, 3), n)
            
                achieved_levels.sort(key=_word_sort_key_lvl)
                current_level = achieved_levels[-1]
            
                add_heading(" أعلى معيار متحقق (ليس الدرجة BTEC النهائية)", level=2, color=PURPLE)
                p_lvl = doc.add_paragraph()
                set_rtl(p_lvl)
                r_lvl1 = p_lvl.add_run("أعلى مستوى معيار تم تحقيقه: ")
                r_lvl1.font.size = Pt(13)
                r_lvl1.font.color.rgb = BLUE
                r_lvl1.font.name = 'Calibri'
                _set_run_cs(r_lvl1)
            
                r_lvl2 = p_lvl.add_run(current_level)
                r_lvl2.bold = True
                r_lvl2.font.size = Pt(13)
                r_lvl2.font.color.rgb = BLUE
                r_lvl2.font.name = 'Calibri'
                _set_run_cs(r_lvl2)
                doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # ════════════════════════════════════════════════════════════
        #  CRITERIA RESULTS — Detailed breakdown
        # ════════════════════════════════════════════════════════════
        add_heading(" تفاصيل المعايير", level=2, color=PURPLE)

        def _word_sort_key(r):
            lv = ""
            if r.criteria:
                lv = (r.criteria.criteria_level or "").split(".")[-1].strip()
            t = lv[:1].upper()
            m = _re.search(r'\d+', lv)
            n = int(m.group()) if m else 999
            return ({"P": 0, "M": 1, "D": 2}.get(t, 3), n)

        grading_results.sort(key=_word_sort_key)

        for result in grading_results:
            criteria_level = ""
            if result.criteria:
                criteria_level = result.criteria.criteria_level or ""

            is_achieved = bool(result.achieved)
            if is_achieved:
                card_accent = "D1FAE5"  # green bg
                card_border = "10B981"
                _st_icon = "✅"
                status_text_plain = "متحقق (Achieved)"
            else:
                card_accent = "FEE2E2"  # red bg
                card_border = "EF4444"
                _st_icon = "❌"
                status_text_plain = "غير متحقق (Not Achieved)"

            # Criterion header table
            card_tbl = doc.add_table(rows=1, cols=2)
            card_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
            card_tbl.autofit = False
            set_table_bidi(card_tbl)
            card_tbl.columns[0].width = Cm(7.5)
            card_tbl.columns[1].width = Cm(7.5)

            for ci in range(2):
                cell = card_tbl.cell(0, ci)
                set_cell_shading(cell, card_accent)
                set_cell_borders(cell, color=card_border, sz="24")
                set_cell_margin(cell, top=100, bottom=100, start=120, end=120)
                set_cell_vertical_alignment(cell, "center")

            # Status cell (left side in RTL -> col 0)
            c0 = card_tbl.cell(0, 0)
            p0 = c0.paragraphs[0]
            set_rtl(p0)
            r0 = p0.add_run(status_text_plain)
            r0.font.size = Pt(14)
            r0.font.color.rgb = PRIMARY
            r0.font.name = 'Calibri'
            _set_run_cs(r0)

            # Title cell (right side in RTL -> col 1)
            c1 = card_tbl.cell(0, 1)
            p1 = c1.paragraphs[0]
            set_rtl(p1)
            r1 = p1.add_run(f"{_st_icon} المعيار {criteria_level}")
            r1.font.size = Pt(14)
            r1.font.color.rgb = PRIMARY
            r1.font.name = 'Calibri'
            _set_run_cs(r1)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # Feedback
            if result.feedback:
                from app.report_feedback_formatter import format_criterion_feedback_for_report
                fb_formatted = format_criterion_feedback_for_report(str(result.feedback or ""))
                fb_p = doc.add_paragraph()
                set_rtl(fb_p)
                fb_label = fb_p.add_run("الملاحظات: ")
                fb_label.bold = True
                fb_label.font.size = Pt(13)
                fb_label.font.color.rgb = BLUE
                fb_label.font.name = 'Calibri'
                _set_run_cs(fb_label)
            
                fb_text = doc.add_paragraph()
                set_rtl(fb_text)
                fb_text_run = fb_text.add_run(fb_formatted)
                fb_text_run.font.size = Pt(12)
                fb_text_run.font.color.rgb = BODY_TEXT
                fb_text_run.font.name = 'Calibri'
                _set_run_cs(fb_text_run)
                fb_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # Decision Matrix Reconstruction
            missing_pts_data = []
            if result.missing_points:
                try:
                    missing_pts_data = json.loads(result.missing_points)
                except Exception:
                    pass
        
            # We construct a mock decision matrix row since 'evidence' is not explicitly saved 
            # in the DB unless we fetch it from Cache. We will use the reasoning if it exists.
            decision_matrix = []
            if is_achieved:
                decision_matrix.append({"requirement": criteria_level, "met": True, "evidence": "متحقق بناءً على تقييم النظام."})
            else:
                reasoning = missing_pts_data[0] if isinstance(missing_pts_data, list) and missing_pts_data else "غير متحقق."
                decision_matrix.append({"requirement": criteria_level, "met": False, "evidence": "-", "reasoning": reasoning})

            if decision_matrix:
                dm_p = doc.add_paragraph()
                set_rtl(dm_p)
                dm_label = dm_p.add_run("متطلبات المعيار: ")
                dm_label.bold = True
                dm_label.font.size = Pt(13)
                dm_label.font.color.rgb = BLUE
                dm_label.font.name = 'Calibri'
                _set_run_cs(dm_label)

                dm_tbl = doc.add_table(rows=1, cols=3)
                dm_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
                dm_tbl.autofit = False
                set_table_bidi(dm_tbl)
                dm_tbl.columns[0].width = Cm(7.0)  # Evidence (Left in RTL)
                dm_tbl.columns[1].width = Cm(3.0)  # Status
                dm_tbl.columns[2].width = Cm(5.0)  # Requirement (Right in RTL)

                headers = ["الدليل (Evidence)", "الحالة (Status)", "المتطلب (Requirement)"]
                for ci, hd in enumerate(headers):
                    cell = dm_tbl.cell(0, ci)
                    set_cell_shading(cell, "6366F1")
                    set_cell_borders(cell, color="D1D5DB", sz="12")
                    set_cell_margin(cell, top=80, bottom=80, start=80, end=80)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_rtl(p)
                    r = p.add_run(hd)
                    r.bold = True
                    r.font.size = Pt(11)
                    r.font.color.rgb = WHITE
                    r.font.name = 'Calibri'
                    _set_run_cs(r)

                for row_data in decision_matrix:
                    row = dm_tbl.add_row()
                    met = row_data.get("met", False)
                    ev = _report_text(row_data.get("evidence", "-"))
                    if not ev or ev == "Not found":
                        ev = "-"
                    req = _report_text(row_data.get("requirement", ""))

                    status_symbol = "✅" if met else "❌"
                    status_str = "متحقق" if met else "غير متحقق"

                    cell_data = [ev, f"{status_symbol} {status_str}", req]
                    for ci, cd in enumerate(cell_data):
                        cell = row.cells[ci]
                        set_cell_shading(cell, "FAFAFA" if ci == 1 else "FFFFFF")
                        set_cell_borders(cell, color="D1D5DB", sz="12")
                        set_cell_margin(cell, top=80, bottom=80, start=80, end=80)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        set_rtl(p)
                        r = p.add_run(_report_text(cd) if ci != 1 else cd)
                        r.font.size = Pt(11)
                        r.font.color.rgb = BODY_TEXT
                        r.font.name = 'Calibri'
                        _set_run_cs(r)

                doc.add_paragraph().paragraph_format.space_after = Pt(10)

            # Next Level Requirements
            if not is_achieved and result.next_level_requirements:
                try:
                    nl_data = json.loads(result.next_level_requirements)
                except Exception:
                    nl_data = []
            
                if nl_data:
                    nl_p = doc.add_paragraph()
                    set_rtl(nl_p)
                    nl_label = nl_p.add_run("للوصول للمعيار التالي: ")
                    nl_label.bold = True
                    nl_label.font.size = Pt(13)
                    nl_label.font.color.rgb = BLUE
                    nl_label.font.name = 'Calibri'
                    _set_run_cs(nl_label)
                
                    if isinstance(nl_data, list):
                        for req_text in nl_data:
                            add_bullet(str(req_text))
                    else:
                        r_p = doc.add_paragraph()
                        set_rtl(r_p)
                        r_run = r_p.add_run(str(nl_data))
                        r_run.font.size = Pt(11)
                        r_run.font.name = 'Calibri'
                        _set_run_cs(r_run)

            doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # ════════════════════════════════════════════════════════════
        #  STRENGTHS & IMPROVEMENTS
        # ════════════════════════════════════════════════════════════
        if summary:
            strengths = []
            if summary.strengths:
                try: strengths = json.loads(summary.strengths)
                except Exception: pass
            
            improvements = []
            if summary.improvements:
                try: improvements = json.loads(summary.improvements)
                except Exception: pass

            if strengths:
                add_heading("🟢 نقاط قوة الطالب", level=2, color=PURPLE)
                for i, st in enumerate(strengths, 1):
                    p = doc.add_paragraph()
                    set_rtl(p)
                    r = p.add_run(f"{i}. {st}")
                    r.font.size = Pt(12)
                    r.font.color.rgb = BODY_TEXT
                    r.font.name = 'Calibri'
                    _set_run_cs(r)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                
            if improvements:
                add_heading("💡 التحسينات المطلوبة", level=2, color=PURPLE)
                mid = max(1, len(improvements) // 2)
                critical = improvements[:mid]
                suggested = improvements[mid:]
            
                if critical:
                    p_c = doc.add_paragraph()
                    set_rtl(p_c)
                    r_c = p_c.add_run("تحسينات حرجة (يجب إكمالها) 🔴")
                    r_c.bold = True
                    r_c.font.size = Pt(13)
                    r_c.font.color.rgb = BLUE
                    r_c.font.name = 'Calibri'
                    _set_run_cs(r_c)
                
                    for i, imp in enumerate(critical, 1):
                        add_bullet(str(imp))
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
                if suggested:
                    p_s = doc.add_paragraph()
                    set_rtl(p_s)
                    r_s = p_s.add_run("تحسينات مقترحة (للحصول على درجة أعلى) 🟡")
                    r_s.bold = True
                    r_s.font.size = Pt(13)
                    r_s.font.color.rgb = BLUE
                    r_s.font.name = 'Calibri'
                    _set_run_cs(r_s)
                
                    for i, imp in enumerate(suggested, 1):
                        add_bullet(str(imp))
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # Overall Feedback
            if summary.overall_feedback:
                add_heading("التقييم العام", level=2, color=PURPLE)
                fb_text = doc.add_paragraph()
                set_rtl(fb_text)
                fb_run = fb_text.add_run(_report_text(summary.overall_feedback))
                fb_run.font.size = Pt(12)
                fb_run.font.color.rgb = BODY_TEXT
                fb_run.font.name = 'Calibri'
                _set_run_cs(fb_run)
                fb_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    ascii_filename = f"report_{submission_id}.docx"
    encoded_filename = quote(f"تقرير_{submission.student_name}.docx")

    headers = {
        "Content-Disposition": (
            f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{encoded_filename}"
        ),
        "Access-Control-Expose-Headers": "Content-Disposition",
    }

    return Response(
        content=file_stream.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


# ==================== Settings API ====================


@app.post("/api/test-provider")
async def test_provider_endpoint(request: Request):
    """Test Gemini or Ollama only."""
    try:
        data = await request.json()
        provider_name = (data.get("provider") or "").strip().lower()
        api_key = (data.get("api_key") or "").strip()

        if provider_name not in ("gemini", "ollama"):
            return JSONResponse(
                {
                    "status": "failed",
                    "message": " المدعوم فقط: gemini أو ollama",
                }
            )
        if provider_name == "gemini" and not api_key:
            return JSONResponse(
                {"status": "failed", "message": " أدخل مفتاح Gemini (GEMINI_API_KEY)"}
            )

        env_key = f"{provider_name.upper()}_API_KEY"
        old_value = os.getenv(env_key)
        if provider_name == "gemini":
            os.environ[env_key] = api_key

        try:
            provider = AIProvider(provider_name)
            response = provider.chat_completion(
                messages=[{"role": "user", "content": "قل 'مرحباً' فقط"}],
                temperature=0.0,
                seed=42,
                max_tokens=50,
            )

            return JSONResponse(
                {
                    "status": "success",
                    "message": f" {provider_name} يعمل بشكل صحيح! الرد: {response}",
                }
            )

        except Exception as e:
            return JSONResponse(
                {
                    "status": "failed",
                    "message": f" فشل الاتصال بـ {provider_name}: {str(e)}",
                }
            )
        finally:
            if provider_name == "gemini":
                if old_value:
                    os.environ[env_key] = old_value
                elif env_key in os.environ:
                    os.environ.pop(env_key, None)

    except Exception as e:
        return JSONResponse({"status": "failed", "message": f" خطأ: {str(e)}"})


@app.post("/api/save-settings")
async def save_settings_endpoint(request: Request):
    """Save API settings"""
    try:
        settings: dict = await request.json()  # type: ignore

        # Save to .env file
        env_path = Path(".env")
        env_lines = []

        if env_path.exists():
            with open(env_path, "r") as f:
                env_lines = f.readlines()

        # Update or add settings
        settings_map = {
            "grading_mode": "GRADING_MODE",
            "primary_provider": "AI_PROVIDER",
            "gemini_api_key": "GEMINI_API_KEY",
            "gemini_model": "GEMINI_MODEL",
            "ollama_base_url": "OLLAMA_BASE_URL",
            "ollama_model": "OLLAMA_MODEL",
            "ollama_vision_model": "OLLAMA_VISION_MODEL",
        }

        for key, env_var in settings_map.items():
            val = settings.get(key) if isinstance(settings, dict) else None
            if val:
                # Update environment
                os.environ[env_var] = str(val)

                # Update .env file
                found = False
                for i, line in enumerate(env_lines):
                    if line.startswith(f"{env_var}="):
                        env_lines[i] = f"{env_var}={val}\n"
                        found = True
                        break

                if not found:
                    env_lines.append(f"{env_var}={val}\n")

        # Write back to .env
        with open(env_path, "w") as f:
            f.writelines(env_lines)

        # Reset global provider to pick up new settings
        reset_global_provider()

        return JSONResponse({"success": True, "message": " تم حفظ الإعدادات بنجاح!"})

    except Exception as e:
        return JSONResponse(
            {"success": False, "message": f" فشل حفظ الإعدادات: {str(e)}"}
        )


@app.get("/api/get-settings")
async def get_settings_endpoint():
    """Get current settings"""
    try:
        settings = {
            "grading_mode": os.getenv("GRADING_MODE", "single"),
            "primary_provider": os.getenv("AI_PROVIDER", "gemini"),
            "gemini_api_key": os.getenv("GEMINI_API_KEY", ""),
            "gemini_model": os.getenv("GEMINI_MODEL", "gemini-2.5-pro"),
            "ollama_base_url": os.getenv(
                "OLLAMA_BASE_URL", "http://localhost:11434/v1"
            ),
            "ollama_model": os.getenv("OLLAMA_MODEL", "deepseek-coder"),
            "ollama_vision_model": os.getenv("OLLAMA_VISION_MODEL", ""),
        }
        return JSONResponse(settings)
    except Exception as e:
        return JSONResponse({"error": str(e)})


    """
مسارات ميزة "إنشاء سجلات التقييم" — إضافة هذا المحتوى إلى نهاية main.py

انقل هذا المحتوى إلى نهاية main.py في مشروعك السابق (قبل أي كتلة
if __name__ == "__main__":).

يتطلب هذا الكود:
  - app/template_filler.py (موجود في نفس مجلد المشروع)
  - app/templates/graded_students.html (موجود في app/templates/)
  - قوالب docx في uploads/templates/
"""

# ─── 1) صفحة عرض كل الطلاب المُصحَّحين مجمَّعين بالواجب ───
@app.get("/api/pearson-templates-status")
async def pearson_templates_status_api(
    request: Request,
    level: str = "DEFAULT",
    db: Session = Depends(get_db),
):
    """Return readiness of Pearson BTEC Word templates (Assessment Record, IV, Evidence)."""
    user = get_current_user(request, db)
    if not user:
        return JSONResponse(status_code=401, content={"error": "يجب تسجيل الدخول أولاً"})
    from app.pearson_templates import pearson_templates_status  # type: ignore
    return JSONResponse(pearson_templates_status(level or "DEFAULT"))


@app.get("/graded-students", response_class=HTMLResponse)
async def graded_students_page(request: Request, db: Session = Depends(get_db)):
    """Page showing all batches with their graded students, grouped by assignment/subject."""
    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login", status_code=302)

    user_id = user.id
    is_admin = user.role and user.role.value == "admin"

    # Get batches: admin sees all, regular user sees own
    query = db.query(BatchGrading)
    if not is_admin:
        query = query.filter(BatchGrading.created_by == user_id)
    batches = query.order_by(BatchGrading.created_at.desc()).all()

    # Group batches by assignment (subject/مادة)
    from collections import OrderedDict
    assignments_map: OrderedDict = OrderedDict()

    for batch in batches:
        assignment = db.query(Assignment).filter(Assignment.id == batch.assignment_id).first()
        assignment_id = batch.assignment_id

        if assignment_id not in assignments_map:
            from app.template_filler import build_unit_title
            from app.pearson_templates import resolve_for_assignment  # type: ignore
            unit_info = build_unit_title(assignment.unit_number or "", assignment.unit_name or "") if assignment else ""
            _pearson = resolve_for_assignment(
                assignment.title if assignment else "",
                assignment.unit_name if assignment else "",
                unit_info,
            )
            assignments_map[assignment_id] = {
                "assignment": assignment,
                "assignment_id": assignment_id,
                "assignment_title": assignment.title if assignment else f"مادة #{assignment_id}",
                "unit_info": unit_info,
                "btec_level": _pearson["level_hint"],
                "pearson_status": _pearson["status"],
                "batches": [],
            }

        subs = (
            db.query(Submission)
            .filter(Submission.batch_id == batch.id)
            .all()
        )
        students = []
        for sub in subs:
            summary = (
                db.query(GradingSummary)
                .filter(GradingSummary.submission_id == sub.id)
                .first()
            )
            # Check if assessment record file exists for this submission
            import re as _re
            _safe = _re.sub(r'[^\w\s\u0600-\u06FF-]', '', _orm_str(sub.student_name)).strip()
            _rec_path = os.path.join(
                "uploads", "reports", "assessment_records",
                f"assessment_record_sub{sub.id}_{_safe}.docx"
            )
            has_record = os.path.exists(_rec_path)
            from app.evidence_templates import student_has_evidence_files  # type: ignore
            has_evidence = student_has_evidence_files(sub.id, sub.student_name or "")
            students.append({
                "submission": sub,
                "summary": summary,
                "has_record": has_record,
                "has_evidence": has_evidence,
            })

        assignments_map[assignment_id]["batches"].append({
            "batch": batch,
            "students": students,
        })

    sub_info = get_subscription_info(db, user_id)
    from app.pearson_templates import pearson_templates_status  # type: ignore
    pearson_global = pearson_templates_status("DEFAULT")
    return templates.TemplateResponse(
        "graded_students.html",
        {
            "request": request,
            "user": user,
            "app_title": os.getenv("APP_TITLE", "أداة تصحيح واجبات الذكاء الاصطناعي"),
            "assignments_map": list(assignments_map.values()),
            "subscription": sub_info,
            "pearson_global": pearson_global,
        },
    )



# ─── 2) إنشاء سجل التقييم (Assessment Record) لطلاب دفعة ───
@app.post("/api/fill-assessment-records/{batch_id}")
async def fill_assessment_records_api(
    request: Request,
    batch_id: int,
    assessor_name: str = Form(""),
    assessment_date: str = Form(""),
    submission_date: str = Form(""),
    deadline_date: str = Form(""),
    student_ids_json: str = Form("{}"),
    selected_submission_ids_json: str = Form("[]"),
    db: Session = Depends(get_db),
):
    """Generate filled BTEC Assessment Record documents for all students in a batch."""
    try:
        user = get_current_user(request, db)
        if not user:
            return JSONResponse(status_code=401, content={"error": "يجب تسجيل الدخول أولاً"})

        from app.template_filler import fill_batch_assessment_records
        from app.pearson_templates import resolve_assessment_record_template  # type: ignore
        import json as json_mod

        template_path = resolve_assessment_record_template()
        if not template_path:
            return JSONResponse(
                status_code=404,
                content={"error": "قالب Assessment Record غير موجود. ضع assessment_record_template.docx في uploads/templates/"}
            )

        if not assessor_name.strip():
            return JSONResponse(status_code=400, content={"error": "يرجى إدخال اسم المُقيّم"})

        if not assessment_date.strip():
            assessment_date = datetime.now().strftime("%d/%m/%Y")

        # Parse student IDs JSON
        try:
            student_ids = json_mod.loads(student_ids_json) if student_ids_json.strip() else {}
        except (json_mod.JSONDecodeError, ValueError):
            student_ids = {}

        # Parse selected submission IDs (optional filter)
        try:
            selected_submission_ids = json_mod.loads(selected_submission_ids_json) if selected_submission_ids_json.strip() else []
            selected_submission_ids = [int(x) for x in selected_submission_ids] if selected_submission_ids else []
        except (json_mod.JSONDecodeError, ValueError):
            selected_submission_ids = []

        generated_files = await fill_batch_assessment_records(
            batch_id=batch_id,
            template_path=template_path,
            assessor_name=assessor_name.strip(),
            assessment_date=assessment_date.strip(),
            submission_date=submission_date.strip(),
            deadline_date=deadline_date.strip(),
            student_ids=student_ids if student_ids else None,
            selected_submission_ids=selected_submission_ids if selected_submission_ids else None,
            db_session=db,
        )

        if not generated_files:
            return JSONResponse(
                status_code=404,
                content={"error": "لم يتم العثور على نتائج تصحيح في هذه الدفعة"}
            )

        total = len(generated_files)

        # Create ZIP if multiple files
        if total > 1:
            import zipfile
            zip_path = f"uploads/reports/assessment_records/batch_{batch_id}_records.zip"
            os.makedirs(os.path.dirname(zip_path), exist_ok=True)
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for fp in generated_files:
                    zf.write(fp, os.path.basename(fp))
            return JSONResponse(content={
                "success": True,
                "message": f"تم تجهيز {total} سجل تقييم (سجل واحد لكل طالب)",
                "zip_url": f"/{zip_path}",
                "files": [os.path.basename(f) for f in generated_files],
                "count": total,
            })
        else:
            return JSONResponse(content={
                "success": True,
                "message": "تم تجهيز سجل التقييم (سجل واحد لكل طالب)",
                "file_url": f"/{generated_files[0]}",
                "files": [os.path.basename(generated_files[0])],
                "count": 1,
            })

    except Exception as e:
        print(f"❌ Assessment record generation failed: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": f"فشل في إنشاء سجلات التقييم: {str(e)}"}
        )



# ─── 3) تنزيل سجل التقييم بصيغة Word ───
@app.get("/api/download-assessment-record/{filename}")
async def download_assessment_record(request: Request, filename: str, db: Session = Depends(get_db)):
    """Download a generated assessment record file."""
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    # Sanitize filename to prevent path traversal
    safe_filename = os.path.basename(filename)
    file_path = os.path.join("uploads", "reports", "assessment_records", safe_filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="الملف غير موجود")

    return FileResponse(
        path=file_path,
        filename=safe_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ═══════════════════════════════════════════════════════════════════════════
# IV OF ASSIGNMENT BRIEF AUTO-FILL
# ═══════════════════════════════════════════════════════════════════════════


# ─── 4) إنشاء IV Assignment Brief (لكل واجب) ───
@app.post("/api/fill-iv-assignment-brief")
async def fill_iv_assignment_brief_api(
    request: Request,
    assignment_id: int = Form(...),
    assessor_name: str = Form(""),
    iv_name: str = Form(""),
    assessment_date: str = Form(""),
    db: Session = Depends(get_db),
):
    """Generate a filled Internal Verification of Assignment Brief document."""
    try:
        user = get_current_user(request, db)
        if not user:
            return JSONResponse(status_code=401, content={"error": "يجب تسجيل الدخول أولاً"})

        if not assessor_name.strip():
            return JSONResponse(status_code=400, content={"error": "يرجى إدخال اسم المُقيّم"})
        if not iv_name.strip():
            return JSONResponse(status_code=400, content={"error": "يرجى إدخال اسم المُدقق الداخلي"})

        if not assessment_date.strip():
            assessment_date = datetime.now().strftime("%d/%m/%Y")

        from app.models import Assignment, GradingCriteria
        from app.template_filler import fill_iv_assignment_brief
        from app.pearson_templates import resolve_iv_brief_template  # type: ignore
        import asyncio

        # Load assignment info
        assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
        if not assignment:
            return JSONResponse(status_code=404, content={"error": "الواجب غير موجود"})

        assignment_title = _orm_str(assignment.title)
        from app.template_filler import build_unit_title
        unit_title = build_unit_title(assignment.unit_number or "", assignment.unit_name or "")

        # Determine program level from unit_number (L2_x → 2, else → 3)
        unit_num_upper = (assignment.unit_number or "").upper()
        if unit_num_upper.startswith("L2"):
            level = 2
        elif unit_num_upper.startswith("L3"):
            level = 3
        else:
            level = 3
        program_title = f"شهادات Pearson BTEC International من المستوى {level} في تكنولوجيا المعلومات"

        # Build criteria text grouped by learning objective section
        # Priority: unit_criteria_json (from assignment spec) → GradingCriteria table
        import json as _json
        letter_to_arabic = {
            'A': 'أ', 'B': 'ب', 'C': 'ج', 'D': 'د',
            'E': 'هـ', 'F': 'و', 'G': 'ز',
        }
        groups: dict = {}

        if assignment.unit_criteria_json:
            try:
                spec_criteria = _json.loads(assignment.unit_criteria_json)
                for item in spec_criteria:
                    code = (item.get("code") or "").strip()
                    aim_ref = (item.get("learning_aim_ref") or (code[0] if code else "?")).upper()
                    if code:
                        groups.setdefault(aim_ref, []).append(code)
            except Exception:
                pass

        # Fallback to GradingCriteria table if spec json was empty/invalid
        if not groups:
            criteria_rows = (
                db.query(GradingCriteria)
                .filter(GradingCriteria.assignment_id == assignment_id)
                .order_by(GradingCriteria.criteria_level)
                .all()
            )
            for cr in criteria_rows:
                lvl = (cr.criteria_level or "").strip()
                letter = lvl[0].upper() if lvl else "?"
                groups.setdefault(letter, []).append(lvl)

        criteria_lines = []
        for letter in sorted(groups.keys()):
            ar_letter = letter_to_arabic.get(letter, letter)
            items = ", ".join(sorted(set(groups[letter])))
            criteria_lines.append(f"هدف التعلم ({ar_letter}): {items}")

        criteria_text = "\n".join(criteria_lines) if criteria_lines else ""

        # Template & output paths
        template_path = resolve_iv_brief_template()
        if not template_path:
            return JSONResponse(
                status_code=404,
                content={"error": "قالب IV غير موجود. ضع iv_assignment_brief_template.docx في uploads/templates/"}
            )

        safe_title = re.sub(r'[^\w\u0600-\u06FF-]', '_', assignment_title)[:40]
        output_filename = f"iv_brief_asgn{assignment_id}_{safe_title}.docx"
        output_path = os.path.join("uploads", "reports", "assessment_records", output_filename)

        await asyncio.to_thread(
            fill_iv_assignment_brief,
            template_path=template_path,
            output_path=output_path,
            assessor_name=assessor_name.strip(),
            iv_name=iv_name.strip(),
            assessment_date=assessment_date.strip(),
            assignment_title=criteria_text,
            unit_title=unit_title,
            criteria_text=criteria_text,
            program_title=program_title,
        )

        return JSONResponse(content={
            "success": True,
            "message": "تم إنشاء ملف IV of Assignment Brief بنجاح",
            "file_url": f"/{output_path}",
            "filename": output_filename,
        })

    except Exception as e:
        print(f"❌ IV assignment brief generation failed: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": f"فشل في إنشاء الملف: {str(e)}"}
        )



# ─── 5) إنشاء IV Assessment Decisions (لكل طالب) ───
@app.post("/api/fill-iv-assessment-decisions/{batch_id}")
async def fill_iv_assessment_decisions_api(
    request: Request,
    batch_id: int,
    selected_submission_ids_json: str = Form("[]"),
    db: Session = Depends(get_db),
):
    """Generate IV Assessment Decisions documents (one per student) for a batch."""
    try:
        user = get_current_user(request, db)
        if not user:
            return JSONResponse(status_code=401, content={"error": "يجب تسجيل الدخول أولاً"})

        import json as _json
        try:
            selected_ids = _json.loads(selected_submission_ids_json)
        except Exception:
            selected_ids = []

        from app.pearson_templates import resolve_iv_decisions_template  # type: ignore

        template_path = resolve_iv_decisions_template()
        if not template_path:
            return JSONResponse(
                status_code=404,
                content={
                    "error": (
                        "قالب BTEC IV Assessment Decisions غير موجود. "
                        "ضع iv_assessment_decisions_template.docx في uploads/templates/"
                    )
                },
            )

        from app.template_filler import fill_batch_iv_assessment_decisions

        generated_files = await fill_batch_iv_assessment_decisions(
            batch_id=batch_id,
            template_path=template_path,
            selected_submission_ids=selected_ids if selected_ids else None,
            db_session=db,
        )

        if not generated_files:
            return JSONResponse(
                status_code=400,
                content={"error": "لم يتم العثور على طلاب مكتملي التصحيح في هذه الدفعة"}
            )

        # If single file, return direct download link
        if len(generated_files) == 1:
            fname = os.path.basename(generated_files[0])
            return JSONResponse(content={
                "success": True,
                "message": "تم إنشاء وثيقة قرارات التقييم بنجاح",
                "file_url": f"/api/download-assessment-record/{fname}",
                "filename": fname,
                "files": [fname],
            })

        # Multiple files → create ZIP
        import zipfile
        import io
        zip_buffer = io.BytesIO()
        zip_name = f"iv_decisions_batch{batch_id}.zip"
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for fpath in generated_files:
                zf.write(fpath, os.path.basename(fpath))
        zip_buffer.seek(0)

        zip_output = os.path.join("uploads", "reports", "assessment_records", zip_name)
        with open(zip_output, 'wb') as f:
            f.write(zip_buffer.read())

        file_list = [os.path.basename(fp) for fp in generated_files]
        return JSONResponse(content={
            "success": True,
            "message": f"تم إنشاء {len(generated_files)} وثيقة قرارات تقييم بنجاح",
            "zip_url": f"/api/download-assessment-record/{zip_name}",
            "files": file_list,
        })

    except Exception as e:
        print(f"❌ IV assessment decisions generation failed: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": f"فشل في إنشاء الملفات: {str(e)}"}
        )



# ─── 6) رفع قالب سجل التقييم ───
@app.post("/api/upload-assessment-template")
async def upload_assessment_template(
    request: Request,
    template_file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """Upload a new BTEC Assessment Record template."""
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    if not template_file.filename or not template_file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="يرجى رفع ملف Word (.docx)")

    template_dir = os.path.join("uploads", "templates")
    os.makedirs(template_dir, exist_ok=True)
    template_path = os.path.join(template_dir, "assessment_record_template.docx")

    content = await template_file.read()
    with open(template_path, "wb") as f:
        f.write(content)

    return JSONResponse(content={
        "success": True,
        "message": "تم رفع قالب سجل التقييم بنجاح"
    })
    """
نقطة النهاية (route) المطلوب إضافتها إلى main.py
بعد المسارات الأخرى الخاصة بسجلات التقييم.

الميزة: إنشاء "سجلات الأدلة" (Evidence Records + LA Evidence) لطلاب دفعة.
ينتج ملفين Word لكل طالب:
  1. Evidance - IT.docx       (سطر لكل معيار)
  2. نموذج ربط أدلة المتعلم    (سطر لكل Learning Aim)

ثم يضمّ كل الملفات في ZIP واحد للتنزيل.
"""

@app.post("/api/fill-evidence-records/{batch_id}")
async def fill_evidence_records_api(
    batch_id: int,
    request: Request,
    selected_submission_ids_json: str = Form("[]"),
    db: Session = Depends(get_db),
):
    """Fill BOTH Pearson evidence templates for students in a batch and return a ZIP."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"error": "يجب تسجيل الدخول أولاً"}, status_code=401)

    try:
        from app.evidence_filler import fill_batch_evidence_records  # type: ignore
        from app.evidence_templates import detect_btec_level_hint, resolve_evidence_templates  # type: ignore
        from app.template_filler import build_unit_title  # type: ignore
        import zipfile
        import json

        batch = db.query(BatchGrading).filter(BatchGrading.id == batch_id).first()
        if not batch:
            return JSONResponse({"error": "الدفعة غير موجودة"}, status_code=404)

        assignment = db.query(Assignment).filter(Assignment.id == batch.assignment_id).first()
        unit_info = build_unit_title(
            (assignment.unit_number or "") if assignment else "",
            (assignment.unit_name or "") if assignment else "",
        )
        level_hint = detect_btec_level_hint(
            assignment.title if assignment else "",
            assignment.unit_name if assignment else "",
            unit_info,
        )
        evidence_template_path, la_template_path, missing_templates = resolve_evidence_templates(level_hint)

        if missing_templates:
            return JSONResponse(
                {
                    "error": (
                        "قوالب Pearson Evidence غير موجودة في uploads/templates/: "
                        + "، ".join(missing_templates)
                        + ". الرجاء رفعها أولاً."
                    )
                },
                status_code=400,
            )

        try:
            selected_submission_ids = json.loads(selected_submission_ids_json or "[]")
            if not isinstance(selected_submission_ids, list):
                selected_submission_ids = []
            selected_submission_ids = [int(x) for x in selected_submission_ids]
        except Exception:
            selected_submission_ids = []

        output_dir = os.path.join("uploads", "reports", "evidence_records", f"batch_{batch_id}")
        os.makedirs(output_dir, exist_ok=True)

        generated_files = await fill_batch_evidence_records(
            batch_id=batch_id,
            evidence_template_path=evidence_template_path,
            la_template_path=la_template_path,
            output_dir=output_dir,
            selected_submission_ids=selected_submission_ids or None,
            db_session=db,
        )

        if not generated_files:
            return JSONResponse(
                {"error": "لا توجد سجلات تقييم لطلاب هذه الدفعة. تأكد من إجراء التصحيح أولاً."},
                status_code=400,
            )

        # Bundle into a ZIP for one-click download
        zip_path = os.path.join("uploads", "reports", "evidence_records", f"batch_{batch_id}_evidence.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fp in generated_files:
                zf.write(fp, arcname=os.path.basename(fp))

        log_activity(
            db, "fill_evidence_records", "export",
            f"إنشاء سجلات الأدلة للدفعة #{batch_id} ({len(generated_files)} ملف)",
            user_id=user_id,
            user_name=_get_user_display(db, user_id),
            user_email=_get_user_email(db, user_id),
            ip_address=_get_client_ip(request),
            user_agent=_get_user_agent(request),
            level="success",
        )

        return JSONResponse({
            "success": True,
            "message": f"تم إنشاء {len(generated_files)} ملف لسجلات الأدلة",
            "zip_url": f"/uploads/reports/evidence_records/batch_{batch_id}_evidence.zip",
            "files_generated": len(generated_files),
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"error": f"فشل في إنشاء سجلات الأدلة: {str(e)}"},
            status_code=500,
        )


@app.get("/api/download-evidence-zip/{batch_id}")
async def download_evidence_zip(batch_id: int, request: Request, db: Session = Depends(get_db)):
    """تنزيل ZIP السجلات إن وُجد."""
    import os
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"error": "يجب تسجيل الدخول"}, status_code=401)

    zip_path = os.path.join("uploads", "reports", "evidence_records", f"batch_{batch_id}_evidence.zip")
    if not os.path.exists(zip_path):
        raise HTTPException(status_code=404, detail="لم يتم إنشاء السجلات بعد")

    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename=f"evidence_records_batch_{batch_id}.zip",
    )


@app.post("/api/fill-evidence-records-for-student/{submission_id}")
async def fill_evidence_records_for_student_api(
    submission_id: int,
    request: Request,
    db: Session = Depends(get_db),
):
    """Fill BOTH evidence templates for a SINGLE student (per-student button)."""
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"error": "يجب تسجيل الدخول أولاً"}, status_code=401)

    try:
        from app.evidence_filler import fill_single_student_evidence_records  # type: ignore
        from app.evidence_templates import detect_btec_level_hint, resolve_evidence_templates  # type: ignore
        from app.template_filler import build_unit_title  # type: ignore
        import zipfile

        submission = db.query(Submission).filter(Submission.id == submission_id).first()
        if not submission:
            return JSONResponse({"error": "التسليم غير موجود"}, status_code=404)

        assignment = db.query(Assignment).filter(Assignment.id == submission.assignment_id).first()
        unit_info = build_unit_title(
            (assignment.unit_number or "") if assignment else "",
            (assignment.unit_name or "") if assignment else "",
        )
        level_hint = detect_btec_level_hint(
            assignment.title if assignment else "",
            assignment.unit_name if assignment else "",
            unit_info,
        )
        evidence_template_path, la_template_path, missing_templates = resolve_evidence_templates(level_hint)

        if missing_templates:
            return JSONResponse(
                {"error": "قوالب Pearson Evidence غير موجودة: " + "، ".join(missing_templates)},
                status_code=400,
            )

        output_dir = os.path.join("uploads", "reports", "evidence_records", "per_student")
        os.makedirs(output_dir, exist_ok=True)

        generated_files = await fill_single_student_evidence_records(
            submission_id=submission_id,
            evidence_template_path=evidence_template_path,
            la_template_path=la_template_path,
            output_dir=output_dir,
            db_session=db,
        )

        if not generated_files:
            return JSONResponse(
                {"error": "لم يتم إنشاء أي ملف. تأكد من وجود نتائج تصحيح للطالب."},
                status_code=400,
            )

        # Bundle into a ZIP per student
        zip_path = os.path.join("uploads", "reports", "evidence_records", f"student_{submission_id}_evidence.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fp in generated_files:
                zf.write(fp, arcname=os.path.basename(fp))

        log_activity(
            db, "fill_evidence_records_student", "export",
            f"إنشاء سجلات أدلة للطالب submission#{submission_id} ({len(generated_files)} ملف)",
            user_id=user_id,
            user_name=_get_user_display(db, user_id),
            user_email=_get_user_email(db, user_id),
            ip_address=_get_client_ip(request),
            user_agent=_get_user_agent(request),
            level="success",
        )

        return JSONResponse({
            "success": True,
            "message": f"تم إنشاء {len(generated_files)} ملف لسجلات الأدلة",
            "zip_url": f"/api/download-evidence-zip-student/{submission_id}",
            "files_generated": len(generated_files),
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": f"فشل في إنشاء سجلات الأدلة: {str(e)}"}, status_code=500)


@app.get("/api/download-evidence-zip-student/{submission_id}")
async def download_evidence_zip_student(submission_id: int, request: Request, db: Session = Depends(get_db)):
    """تنزيل ZIP السجلات لطالب واحد."""
    import os
    user_id = get_current_user_id(request)
    if not user_id:
        return JSONResponse({"error": "يجب تسجيل الدخول"}, status_code=401)

    zip_path = os.path.join("uploads", "reports", "evidence_records", f"student_{submission_id}_evidence.zip")
    if not os.path.exists(zip_path):
        raise HTTPException(status_code=404, detail="لم يتم إنشاء السجلات بعد")

    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename=f"evidence_records_student_{submission_id}.zip",
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "main:app",
        host=os.getenv("HOST", "0.0.0.0"),
        port=int(os.getenv("PORT", "5556")),
        reload=os.getenv("DEBUG", "False").lower() == "true",
        timeout_keep_alive=int(os.getenv("UVICORN_TIMEOUT_KEEP_ALIVE", "300")),
    )
