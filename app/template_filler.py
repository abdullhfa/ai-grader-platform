"""
BTEC Assessment Record Template Auto-Filler
Fills Word (.docx) templates with grading results for each student.
"""
import os
import re
import json as _json_mod
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml import OxmlElement  # type: ignore


def build_unit_title(unit_number: str, unit_name: str) -> str:
    """
    Build a clean unit title string shared across all document templates.
    Strips L2_/L3_ prefix from unit_number so 'L2_4' displays as '4'.
    Examples:
      ('L2_4', 'Introduction to Computer Networking')
        → 'الوحدة 4: Introduction to Computer Networking'
      ('21', 'INTRODUCTION TO AI')
        → 'الوحدة 21: INTRODUCTION TO AI'
      ('', 'Introduction to Programming')
        → 'Introduction to Programming'
    """
    num = (unit_number or "").strip()
    name = (unit_name or "").strip()
    # Strip L2_ / L3_ prefix (case-insensitive)
    clean_num = re.sub(r'^[Ll][23]_', '', num)
    if clean_num and name:
        return f"الوحدة {clean_num}: {name}"
    if name:
        return name
    if clean_num:
        return f"الوحدة {clean_num}"
    return ""


# ═══════════════════════════════════════════════════════════════
# TABLE CELL MAP for "سجل التقييم" (Assessment Record)
# Table 0 layout (30 rows x 6 cols):
#   [1,3]/[1,4-5] = رقم تسجيل الطالب / value
#   [2,3]/[2,4-5] = اسم الطالب / value
#   [3,2]         = عنوان الواجب (assignment title)
#   [3,3]/[3,4-5] = اسم المُقيّم / value
#   [4,2-5]       = رقم الوحدة وعنوانها (unit number + title)
#   [7,0]         = "المعايير المستهدفة" header
#   [7,1]         = "هل تم تحقيق المعيار" header
#   [7,2-5]       = "تعليقات التقييم" header
#   [8-22, 0]     = Criteria codes (A.P1, A.P2, ...)
#   [8-22, 1]     = Achievement (نعم/لا)
#   [8-22, 2-5]   = Assessment comments (merged)
#   [23-24]       = General comments area
#   [25, 3-5]     = Student signature
#   [26, 3-5]     = Date
#   [27, 3-5]     = Assessor signature
#   [28, 3-5]     = Date
#   [29, 3-5]     = Feedback date
# ═══════════════════════════════════════════════════════════════


def _ensure_rtl_paragraph(paragraph, alignment: Optional[str] = None):
    """Ensure a paragraph has RTL (bidi) direction.
    With bidi set, Word defaults to right-aligned — no need for w:jc.
    Use alignment='center' only when explicitly needed."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    # Only set jc if an explicit alignment is requested (e.g. 'center')
    jc = pPr.find(qn('w:jc'))
    if alignment:
        if jc is not None:
            pPr.remove(jc)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        rPr = pPr.find(qn('w:rPr'))
        if rPr is not None:
            rPr.addprevious(jc)
        else:
            pPr.append(jc)


def _ensure_rtl_run(run):
    """Ensure a run has RTL direction."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)


def _set_cell_text(cell, text: str, rtl: bool = True,
                   alignment: Optional[str] = None):
    """Set cell text while preserving the first paragraph's formatting.
    alignment: optional explicit alignment (e.g. 'center')."""
    if not cell.paragraphs:
        cell.text = text
        return
    first_para = cell.paragraphs[0]
    if first_para.runs:
        run = first_para.runs[0]
        run.text = text
        if rtl:
            _ensure_rtl_run(run)
        for extra_run in first_para.runs[1:]:
            extra_run.text = ""
    else:
        first_para.text = text
        if first_para.runs and rtl:
            _ensure_rtl_run(first_para.runs[0])
    if rtl:
        _ensure_rtl_paragraph(first_para, alignment=alignment)
    # Remove extra paragraphs
    for para in cell.paragraphs[1:]:
        p_element = para._element
        p_element.getparent().remove(p_element)


def _set_cell_multiline(cell, text: str, rtl: bool = True):
    """Set cell text with line breaks, preserving formatting."""
    if not text:
        return
    lines = text.split('\n')
    if not cell.paragraphs:
        cell.text = text
        return
    first_para = cell.paragraphs[0]
    # Get formatting from first run
    fmt_size = None
    fmt_name = None
    if first_para.runs:
        fmt_size = first_para.runs[0].font.size
        fmt_name = first_para.runs[0].font.name
    # Clear extra paragraphs
    for para in cell.paragraphs[1:]:
        p_element = para._element
        p_element.getparent().remove(p_element)
    # Set first line
    if first_para.runs:
        first_para.runs[0].text = lines[0]
        if rtl:
            _ensure_rtl_run(first_para.runs[0])
        for extra_run in first_para.runs[1:]:
            extra_run.text = ""
    else:
        first_para.text = lines[0]
        if first_para.runs and rtl:
            _ensure_rtl_run(first_para.runs[0])
    if rtl:
        _ensure_rtl_paragraph(first_para)
    # Add remaining lines
    for line in lines[1:]:
        new_para = cell.add_paragraph(line)
        if rtl:
            _ensure_rtl_paragraph(new_para)
        if new_para.runs:
            if fmt_size:
                new_para.runs[0].font.size = fmt_size
            if fmt_name:
                new_para.runs[0].font.name = fmt_name
            if rtl:
                _ensure_rtl_run(new_para.runs[0])


def _find_criteria_row(table, criteria_code: str) -> Optional[int]:
    """Find the row index for a given criteria code in the table."""
    for ri, row in enumerate(table.rows):
        cell_text = row.cells[0].text.strip()
        if cell_text == criteria_code:
            return ri
    return None


def _delete_row(table, row_index: int):
    """Delete a row from a table by index."""
    row_element = table.rows[row_index]._tr
    table._tbl.remove(row_element)


def _delete_unused_criteria_rows(table, used_criteria: set):
    """Delete criteria rows that are NOT in the used_criteria set.
    Criteria rows are between the header row (المعايير المستهدفة) and
    the general comments row (تعليقات عامة)."""
    rows_to_delete = []
    in_criteria_zone = False

    for ri, row in enumerate(table.rows):
        cell0_text = row.cells[0].text.strip()

        if 'المعايير المستهدفة' in cell0_text:
            in_criteria_zone = True
            continue

        if 'تعليقات عامة' in cell0_text:
            break

        if in_criteria_zone:
            # This is a criteria row - check if it's used
            if cell0_text and cell0_text not in used_criteria:
                rows_to_delete.append(ri)

    # Delete from bottom to top to preserve indices
    for ri in reversed(rows_to_delete):
        _delete_row(table, ri)


def _clear_resubmission_table(doc):
    """Clear/remove the resubmission record table (Table index 3)."""
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        # Remove the entire table element from document
        t3._element.getparent().remove(t3._element)
    # Also remove table 2 (resubmission authorization) if it exists
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        cell0_text = t2.rows[0].cells[0].text.strip() if t2.rows else ""
        if 'إعادة' in cell0_text or 'تسليم' in cell0_text:
            t2._element.getparent().remove(t2._element)


def _find_general_comments_row(table) -> Optional[int]:
    """Find the row index for 'تعليقات عامة' (general comments)."""
    for ri, row in enumerate(table.rows):
        if 'تعليقات عامة' in row.cells[0].text.strip():
            return ri
    return None


def fill_assessment_record(
    template_path: str,
    output_path: str,
    student_name: str,
    student_id: str,
    assessor_name: str,
    assignment_title: str,
    unit_info: str,
    criteria_results: List[Dict],
    general_comments: str,
    assessment_date: str = "",
    submission_date: str = "",
    deadline_date: str = "",
    feedback_date: str = "",
    table_index: int = 0,
    program_title: str = "",
) -> str:
    """
    Fill a BTEC Assessment Record template with student grading data.
    Deletes unused criteria rows and clears the resubmission table.
    All text is written RTL.
    """
    doc = Document(template_path)

    if table_index >= len(doc.tables):
        raise ValueError(f"Template has only {len(doc.tables)} tables, requested index {table_index}")

    table = doc.tables[table_index]

    # ── Fill program title (عنوان البرنامج) ──
    if program_title:
        if len(table.rows) > 1 and len(table.rows[1].cells) > 2:
            _set_cell_text(table.rows[1].cells[2], program_title)
        if len(table.rows) > 2 and len(table.rows[2].cells) > 2:
            _set_cell_text(table.rows[2].cells[2], program_title)

    # ── Fill student info ──
    # Row 1, cols 4-5: Student registration number
    if len(table.rows) > 1 and len(table.rows[1].cells) > 4:
        _set_cell_text(table.rows[1].cells[4], student_id, rtl=False)

    # Row 2, cols 4-5: Student name
    if len(table.rows) > 2 and len(table.rows[2].cells) > 4:
        _set_cell_text(table.rows[2].cells[4], student_name)

    # Row 3, cols 4-5: Assessor name — left blank for manual entry

    # Row 3, col 2: Assignment title
    if assignment_title and len(table.rows) > 3:
        _set_cell_text(table.rows[3].cells[2], assignment_title)

    # Row 4, cols 2-5: Unit info
    if unit_info and len(table.rows) > 4:
        _set_cell_text(table.rows[4].cells[2], unit_info)

    # Row 5, col 2: Submission date (تاريخ تسليم المهمة)
    if submission_date and len(table.rows) > 5 and len(table.rows[5].cells) > 2:
        _set_cell_text(table.rows[5].cells[2], submission_date, rtl=False)

    # Row 5, col 5: Deadline date (الموعد النهائي)
    if deadline_date and len(table.rows) > 5 and len(table.rows[5].cells) > 5:
        _set_cell_text(table.rows[5].cells[5], deadline_date, rtl=False)

    # ── Delete unused criteria rows first, then fill ──
    used_criteria = {c.get("criteria_level", "") for c in criteria_results}
    _delete_unused_criteria_rows(table, used_criteria)

    # ── Fill criteria achievement and comments ──
    for criterion in criteria_results:
        level = criterion.get("criteria_level", "")
        achieved = criterion.get("achieved", False)
        feedback = criterion.get("feedback", "")

        row_idx = _find_criteria_row(table, level)
        if row_idx is None:
            continue

        row = table.rows[row_idx]

        # Column 1: Achievement (نعم/لا) — centered like reference
        if len(row.cells) > 1:
            _set_cell_text(row.cells[1], "نعم" if achieved else "لا",
                           alignment='center')

        # Columns 2-5: Assessment comments (merged cells)
        if len(row.cells) > 2:
            _set_cell_text(row.cells[2], feedback)

    # ── Fill general comments ──
    comments_row = _find_general_comments_row(table)
    if comments_row is not None:
        content_row = comments_row + 1
        if content_row < len(table.rows):
            row = table.rows[content_row]
            _set_cell_multiline(row.cells[0], general_comments)

    # ── Signatures, assessor name, and dates — left blank for manual entry ──

    # ── Clear resubmission table ──
    _clear_resubmission_table(doc)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _remove_tashkeel(text: str) -> str:
    """Remove Arabic diacritics (tashkeel) from text."""
    import re as _re
    # Unicode ranges for Arabic tashkeel
    tashkeel_pattern = _re.compile(r'[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06DC\u06DF-\u06E4\u06E7\u06E8\u06EA-\u06ED]')
    return tashkeel_pattern.sub('', text)


def _build_general_comments_prompt(
    student_name: str,
    grade_level: str,
    criteria_results: List[Dict],
) -> str:
    """Build the AI prompt for general comments in formal Arabic."""
    achieved = [c for c in criteria_results if c.get("achieved")]
    not_achieved = [c for c in criteria_results if not c.get("achieved")]

    achieved_list = ", ".join([c["criteria_level"] for c in achieved]) if achieved else "لا يوجد"
    not_achieved_list = ", ".join([c["criteria_level"] for c in not_achieved]) if not_achieved else "لا يوجد"

    criteria_details = ""
    for c in criteria_results:
        status = "✓ تم تحقيقه" if c.get("achieved") else "✗ لم يتحقق"
        fb = c.get("feedback", "")[:150]
        criteria_details += f"\n- {c['criteria_level']}: {status} - {fb}"

    return f"""أنت مقيّم BTEC معتمد. اكتب تعليقات عامة لسجل التقييم الخاص بالطالب باللغة العربية الفصحى.

بيانات الطالب:
- الاسم: {student_name}
- الدرجة النهائية: {grade_level}
- المعايير المحققة: {achieved_list}
- المعايير غير المحققة: {not_achieved_list}

تفاصيل المعايير:{criteria_details}

اكتب التعليقات بالعربية الفصحى بالتنسيق التالي:
- السطر الأول: علامتك النهائية هي {grade_level}، ثم عبارة تحفيزية مناسبة للدرجة
- السطر الثاني: ثناء عام على مستوى الفهم والأداء والمهارات المُظهرة
- السطر الثالث: وصف لجهد الطالب في حل الواجب وكتابة التقارير
- الأسطر التالية: تلخيص لما قدمه الطالب في كل مجموعة معايير (ماذا شرح، ماذا حلل، ماذا قدم) بأسلوب "لقد شرحت..." أو "وصفت..." أو "حللت..." أو "قدمت..."
{"- أضف سطراً عن المعايير غير المحققة وما يحتاج لتحسينه" if not_achieved else ""}

شروط مهمة:
- اكتب بالعربية الفصحى فقط
- استخدم صيغة المخاطب المذكر (شرحت، حللت، قدمت، وصفت)
- لا تستخدم ترقيم او نقاط او شرطات في بداية الاسطر
- لا تضع عناوين للاسطر
- لا تستخدم تشكيل (حركات) على الكلمات ابدا مثل الفتحة والضمة والكسرة والتنوين والشدة
- اكتب الكلمات بدون اي علامات ضبط
- اجعل التعليق مهني وبناء ومفصل
- كل طالب يجب ان تكتب له تعليقات مختلفة ومخصصة
"""


async def generate_general_comments_via_ai_async(
    student_name: str,
    grade_level: str,
    criteria_results: List[Dict],
    ai_provider=None,
) -> str:
    """Async version of general comments generation using formal Arabic."""
    achieved = [c for c in criteria_results if c.get("achieved")]
    not_achieved = [c for c in criteria_results if not c.get("achieved")]

    prompt = _build_general_comments_prompt(student_name, grade_level, criteria_results)

    if ai_provider:
        try:
            import asyncio
            messages = [
                {"role": "system", "content": "أنت مقيّم BTEC محترف تكتب تعليقات تقييم مخصصة باللغة العربية الفصحى."},
                {"role": "user", "content": prompt},
            ]
            response = await asyncio.to_thread(
                ai_provider.chat_completion,
                messages=messages,
                temperature=0.7
            )
            if isinstance(response, str) and response.strip():
                return _remove_tashkeel(response.strip())
        except Exception as e:
            print(f"⚠️ AI comments generation failed: {e}")

    return _remove_tashkeel(_generate_fallback_comments(student_name, grade_level, achieved, not_achieved))


def _generate_fallback_comments(
    student_name: str,
    grade_level: str,
    achieved: List[Dict],
    not_achieved: List[Dict],
) -> str:
    """Generate deterministic comments in formal Arabic without AI."""
    import hashlib
    name_hash = int(hashlib.md5(student_name.encode()).hexdigest(), 16)

    # Line 1: Grade + motivational phrase
    if not not_achieved:
        motivational_pool = [
            f"علامتك النهائية هي {grade_level} ، أتمنى لك الاستمرار والاجتهاد دائماً",
            f"علامتك النهائية هي {grade_level} ، عمل متميز يعكس مستوى عالٍ من الفهم والإتقان",
            f"علامتك النهائية هي {grade_level} ، إنجاز رائع يدل على جهد واضح ومثابرة",
            f"علامتك النهائية هي {grade_level} ، أداء يستحق الثناء والتقدير",
        ]
    else:
        motivational_pool = [
            f"علامتك النهائية هي {grade_level} ، لديك إمكانيات جيدة وأتوقع منك المزيد من التطور",
            f"علامتك النهائية هي {grade_level} ، بذلت جهداً واضحاً ويمكنك تحسين أدائك في بعض المعايير",
            f"علامتك النهائية هي {grade_level} ، أداء جيد مع وجود فرص للتحسين والتطوير",
            f"علامتك النهائية هي {grade_level} ، عمل مقبول وأتطلع لرؤية تحسن في المعايير المتبقية",
        ]
    line1 = motivational_pool[name_hash % len(motivational_pool)]

    # Line 2: Praise on understanding
    if not not_achieved:
        line2 = "لقد غطيت جميع النقاط المطلوبة وأظهرت أنك تفهم الموضوع بشكل رائع. أنت تفهم بوضوح الموضوع ويمكنك تطبيقه على ممارستك. لقد أظهرت مهارات وكفاءة ووعياً جيداً."
    else:
        line2 = "أظهرت فهماً جيداً في عدة جوانب من الموضوع، ولديك القدرة على تطوير أدائك بشكل أفضل في المعايير المتبقية."

    # Line 3: Effort
    effort_pool = [
        "يظهر عملك جهداً واضحاً ورائعاً أثناء حل الواجب وكتابة التقارير.",
        "يعكس عملك التزاماً ومثابرة في إنجاز المهام المطلوبة.",
    ]
    line3 = effort_pool[name_hash % len(effort_pool)]

    # Lines 4+: Criteria summary grouped by section
    summary_lines = []
    not_achieved_codes = [c["criteria_level"] for c in not_achieved]

    achieved_count = len(achieved)
    total = achieved_count + len(not_achieved)
    summary_lines.append(f"لقد حققت {achieved_count} من أصل {total} معيار مطلوب في هذه المهمة.")

    if not_achieved:
        missing = ", ".join(not_achieved_codes)
        summary_lines.append(f"المعايير التي تحتاج لمراجعة وتحسين: {missing}. يُرجى مراجعة التعليقات التفصيلية لكل معيار.")

    return f"{line1}\n{line2}\n{line3}\n" + "\n".join(summary_lines)


async def fill_batch_assessment_records(
    batch_id: int,
    template_path: str,
    assessor_name: str,
    assessment_date: str = "",
    submission_date: str = "",
    deadline_date: str = "",
    student_ids: Optional[Dict[str, str]] = None,
    selected_submission_ids: Optional[List[int]] = None,
    db_session=None,
) -> List[str]:
    """
    Fill assessment record templates for ALL students in a batch.

    Args:
        batch_id: The batch grading ID
        template_path: Path to the Assessment Record Word template
        assessor_name: Name of the assessor/teacher
        assessment_date: Date string for signatures
        submission_date: Date assignment was given to students
        deadline_date: Deadline date for submission
        student_ids: Dict mapping student_name -> student_id (registration number)
        db_session: SQLAlchemy database session

    Returns:
        List of paths to generated filled documents
    """
    from app.models import (
        BatchGrading, Submission, GradingSummary, GradingResult,
        GradingCriteria, Assignment, SubmissionStatus
    )
    from app.ai_provider import get_global_provider

    if db_session is None:
        from app.database import SessionLocal
        db_session = SessionLocal()
        should_close = True
    else:
        should_close = False

    try:
        # Get batch info
        batch = db_session.query(BatchGrading).filter(BatchGrading.id == batch_id).first()
        if not batch:
            raise ValueError(f"Batch {batch_id} not found")

        # Get assignment info
        assignment = db_session.query(Assignment).filter(Assignment.id == batch.assignment_id).first()
        if not assignment:
            raise ValueError(f"Assignment not found for batch {batch_id}")

        assignment_title = assignment.title or ""
        unit_info = build_unit_title(assignment.unit_number or "", assignment.unit_name or "")

        # Compute program title (same logic as IV endpoint)
        _unit_num_upper = (assignment.unit_number or "").upper()
        _level = 2 if _unit_num_upper.startswith("L2") else 3
        program_title = f"شهادات Pearson BTEC International من المستوى {_level} في تكنولوجيا المعلومات"

        # Build criteria text (learning objectives) to use as assignment title
        _letter_to_ar = {'A': 'أ', 'B': 'ب', 'C': 'ج', 'D': 'د', 'E': 'هـ', 'F': 'و', 'G': 'ز'}
        _groups: dict = {}
        if assignment.unit_criteria_json:
            try:
                for item in _json_mod.loads(assignment.unit_criteria_json):
                    code = (item.get('code') or '').strip()
                    aim = (item.get('learning_aim_ref') or (code[0] if code else '?')).upper()
                    if code:
                        _groups.setdefault(aim, []).append(code)
            except Exception:
                pass
        if not _groups:
            for cr in db_session.query(GradingCriteria).filter(
                GradingCriteria.assignment_id == assignment.id
            ).order_by(GradingCriteria.criteria_level).all():
                lvl = (cr.criteria_level or '').strip()
                letter = lvl[0].upper() if lvl else '?'
                _groups.setdefault(letter, []).append(lvl)
        _lines = []
        for _l in sorted(_groups.keys()):
            _ar = _letter_to_ar.get(_l, _l)
            _items = ', '.join(sorted(set(_groups[_l])))
            _lines.append(f'هدف التعلم ({_ar}): {_items}')
        criteria_title = '\n'.join(_lines) if _lines else assignment_title

        # Get all completed submissions in this batch
        query = db_session.query(Submission).filter(
            Submission.batch_id == batch_id,
            Submission.status == SubmissionStatus.COMPLETED,
        )
        # Filter by selected submission IDs if provided
        # When IDs span multiple batches, bypass batch_id filter and use IDs only
        if selected_submission_ids:
            query = db_session.query(Submission).filter(
                Submission.id.in_(selected_submission_ids),
                Submission.status == SubmissionStatus.COMPLETED,
            )
        submissions = query.all()

        if not submissions:
            print("⚠️ No completed submissions found in this batch")
            return []

        # Get AI provider for generating comments
        try:
            ai_provider = get_global_provider()
        except Exception:
            ai_provider = None

        output_dir = Path("uploads/reports/assessment_records")
        output_dir.mkdir(parents=True, exist_ok=True)

        generated_files = []

        for submission in submissions:
            print(f"📝 Filling assessment record for: {submission.student_name}")

            # Get grading summary
            summary = (
                db_session.query(GradingSummary)
                .filter(GradingSummary.submission_id == submission.id)
                .first()
            )

            if not summary:
                print(f"  ⚠️ No grading summary for {submission.student_name}, skipping")
                continue

            # Get grading results with criteria info
            results = (
                db_session.query(GradingResult, GradingCriteria)
                .join(GradingCriteria, GradingResult.criteria_id == GradingCriteria.id)
                .filter(GradingResult.submission_id == submission.id)
                .all()
            )

            # Build criteria_results list
            criteria_results = []
            for result, criteria in results:
                criteria_results.append({
                    "criteria_level": criteria.criteria_level,
                    "achieved": result.achieved,
                    "feedback": result.feedback or ("تم اجتياز المعيار." if result.achieved else "لم يتم اجتياز المعيار."),
                })

            # Sort criteria P → M → D
            def _sort_key(c):
                lv = c.get('criteria_level', '')
                short = lv.split('.')[-1] if '.' in lv else lv
                _order = {'P': 0, 'M': 1, 'D': 2}
                letter = short[0].upper() if short else 'Z'
                try:
                    num = int(short[1:]) if len(short) > 1 else 0
                except ValueError:
                    num = 99
                return (_order.get(letter, 9), num)

            criteria_results.sort(key=_sort_key)

            # Generate general comments
            grade_level = summary.grade_level or "U"
            # Extract short grade (first letter/word)
            grade_short = grade_level.split(" ")[0] if grade_level else "U"

            general_comments = await generate_general_comments_via_ai_async(
                student_name=submission.student_name,
                grade_level=grade_short,
                criteria_results=criteria_results,
                ai_provider=ai_provider,
            )

            # Student info - use provided student_ids or fallback to DB
            student_name_val = submission.student_name or ""
            student_id_val = ""
            if student_ids and student_name_val in student_ids:
                student_id_val = student_ids[student_name_val]
            elif submission.student_id:
                student_id_val = submission.student_id

            # Generate output filename (unique per submission — one record per student policy)
            safe_name = re.sub(r'[^\w\s\u0600-\u06FF-]', '', student_name_val).strip()
            output_filename = f"assessment_record_sub{submission.id}_{safe_name}.docx"
            output_path = str(output_dir / output_filename)

            # ── One-record-per-student policy ──────────────────────────────
            if os.path.exists(output_path):
                print(f"  ⏭️  Record already exists for {student_name_val}, skipping (sub_id={submission.id})")
                generated_files.append(output_path)   # include existing file in result
                continue
            # ──────────────────────────────────────────────────────────────
            try:
                fill_assessment_record(
                    template_path=template_path,
                    output_path=output_path,
                    student_name=student_name_val,
                    student_id=student_id_val,
                    assessor_name=assessor_name,
                    assignment_title=criteria_title,
                    unit_info=unit_info,
                    criteria_results=criteria_results,
                    general_comments=general_comments,
                    assessment_date=assessment_date,
                    submission_date=submission_date,
                    deadline_date=deadline_date,
                    feedback_date=assessment_date,
                    program_title=program_title,
                )
                generated_files.append(output_path)
                print(f"  ✅ Saved: {output_path}")
            except Exception as e:
                print(f"  ❌ Failed for {student_name_val}: {e}")

        print(f"\n📄 Generated {len(generated_files)} assessment records")
        return generated_files

    finally:
        if should_close:
            db_session.close()


def fill_single_assessment_record_from_data(
    template_path: str,
    output_path: str,
    student_name: str,
    student_id: str,
    assessor_name: str,
    assignment_title: str,
    unit_info: str,
    criteria_results: List[Dict],
    grade_level: str,
    assessment_date: str = "",
    submission_date: str = "",
    deadline_date: str = "",
    ai_provider=None,
    program_title: str = "",
) -> str:
    """
    Convenience function: fill template for one student from raw data (no DB).
    Generates general comments automatically.
    """
    general_comments = _generate_fallback_comments(
        student_name,
        grade_level,
        [c for c in criteria_results if c.get("achieved")],
        [c for c in criteria_results if not c.get("achieved")],
    )

    return fill_assessment_record(
        template_path=template_path,
        output_path=output_path,
        student_name=student_name,
        student_id=student_id,
        assessor_name=assessor_name,
        assignment_title=assignment_title,
        unit_info=unit_info,
        criteria_results=criteria_results,
        general_comments=general_comments,
        assessment_date=assessment_date,
        submission_date=submission_date,
        deadline_date=deadline_date,
        feedback_date=assessment_date,
        program_title=program_title,
    )


# ═══════════════════════════════════════════════════════════════
# IV of Assignment Brief Filler
# ═══════════════════════════════════════════════════════════════

def fill_iv_assignment_brief(
    template_path: str,
    output_path: str,
    assessor_name: str,
    iv_name: str,
    assessment_date: str,
    assignment_title: str,
    unit_title: str,
    criteria_text: str,
    program_title: str = "",
) -> str:
    """
    Fill the Internal Verification of Assignment Brief template.
    Replaces all red-colored (EE0000) placeholder cells with actual data.
    """
    doc = Document(template_path)

    RED = 'EE0000'
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    processed_tcs: set = set()   # track unique tc elements (merged cells)

    def _dered(cell):
        """Change EE0000 runs in this cell to black."""
        for para in cell.paragraphs:
            for run in para.runs:
                rPr = run._r.find(f'{{{ns}}}rPr')
                if rPr is not None:
                    cel = rPr.find(f'{{{ns}}}color')
                    if cel is not None and cel.get(f'{{{ns}}}val', '').upper() == RED:
                        cel.set(f'{{{ns}}}val', '000000')

    def fill_row_value(table, row_idx: int, text: str, col_start: int = 1,
                       rtl: bool = True):
        """Set all unique cells in the row from col_start onwards to *text*."""
        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        for ci in range(col_start, len(row.cells)):
            tc_id = id(row.cells[ci]._tc)
            if tc_id in processed_tcs:
                continue
            processed_tcs.add(tc_id)
            cell = row.cells[ci]
            if rtl:
                _set_cell_multiline(cell, text, rtl=True)
            else:
                _set_cell_text(cell, text, rtl=False)
            _dered(cell)

    def fill_single_cell(table, row_idx: int, col_idx: int, text: str,
                         rtl: bool = True):
        """Set one specific cell."""
        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        if col_idx >= len(row.cells):
            return
        tc_id = id(row.cells[col_idx]._tc)
        if tc_id in processed_tcs:
            return
        processed_tcs.add(tc_id)
        cell = row.cells[col_idx]
        if rtl:
            _set_cell_text(cell, text, rtl=True)
        else:
            _set_cell_text(cell, text, rtl=False)
        _dered(cell)

    def fill_row_value_force(table, row_idx: int, text: str, col_start: int = 1,
                             rtl: bool = True):
        """Write to cells regardless of color (for non-red placeholder rows)."""
        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        seen: set = set()
        for ci in range(col_start, len(row.cells)):
            tc_id = id(row.cells[ci]._tc)
            if tc_id in seen:
                continue
            seen.add(tc_id)
            cell = row.cells[ci]
            if rtl:
                _set_cell_multiline(cell, text, rtl=True)
            else:
                _set_cell_text(cell, text, rtl=False)

    # ── Table 0 ───────────────────────────────────────────────
    t0 = doc.tables[0]
    if program_title:
        fill_row_value_force(t0, 1, program_title)  # عنوان البرنامج
    # Rows 2 (اسم المقيم) and 3 (اسم المدقق الداخلي) — left blank for manual entry
    fill_row_value(t0, 4, unit_title)         # رقم الوحدة وعنوانها
    fill_row_value(t0, 5, assignment_title)   # عنوان موجز المهام
    fill_row_value(t0, 6, criteria_text, rtl=True)  # معايير التقييم

    # ── Table 1 (signatures) — left blank for manual signing ──────────
    # Cells for signatures, names, and dates are intentionally not filled
    # so that the assessor and IV can sign and date manually on paper.

    # ── Final pass: de-red ALL remaining red text in the whole document ──
    # This turns checklist "نعم" / "لا ينطبق" cells from red to black
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _dered(cell)
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.find(f'{{{ns}}}rPr')
            if rPr is not None:
                cel = rPr.find(f'{{{ns}}}color')
                if cel is not None and cel.get(f'{{{ns}}}val', '').upper() == RED:
                    cel.set(f'{{{ns}}}val', '000000')

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════
# IV Assessment Decisions (Single Learner) Filler
# ═══════════════════════════════════════════════════════════════

def fill_iv_assessment_decisions(
    template_path: str,
    output_path: str,
    student_name: str,
    program_title: str,
    unit_title: str,
    assignment_title: str,
    achieved_criteria: List[str],
    general_comments: str = "",
) -> str:
    """
    Fill the IV Assessment Decisions for Single Learner template.
    Generates one document per student showing which criteria were awarded
    and IV checklist answers.
    """
    doc = Document(template_path)

    RED = 'EE0000'
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    if not doc.tables:
        raise ValueError("No tables found in IV Assessment Decisions template")

    t0 = doc.tables[0]
    processed_tcs: set = set()

    def _dered(cell):
        for para in cell.paragraphs:
            for run in para.runs:
                rPr = run._r.find(f'{{{ns}}}rPr')
                if rPr is not None:
                    cel = rPr.find(f'{{{ns}}}color')
                    if cel is not None and cel.get(f'{{{ns}}}val', '').upper() == RED:
                        cel.set(f'{{{ns}}}val', '000000')

    def _set_unique(row_idx: int, col_idx: int, text: str, rtl: bool = True):
        """Set a single cell, skip if already processed (merged)."""
        if row_idx >= len(t0.rows):
            return
        row = t0.rows[row_idx]
        if col_idx >= len(row.cells):
            return
        cell = row.cells[col_idx]
        tc_id = id(cell._tc)
        if tc_id in processed_tcs:
            return
        processed_tcs.add(tc_id)
        _set_cell_text(cell, text, rtl=rtl)
        _dered(cell)

    def _set_row_force(row_idx: int, text: str, col_start: int = 1, rtl: bool = True):
        """Write all unique cells in a row from col_start (ignoring processed set)."""
        if row_idx >= len(t0.rows):
            return
        row = t0.rows[row_idx]
        seen: set = set()
        for ci in range(col_start, len(row.cells)):
            tc_id = id(row.cells[ci]._tc)
            if tc_id in seen:
                continue
            seen.add(tc_id)
            _set_cell_text(row.cells[ci], text, rtl=rtl)
            _dered(row.cells[ci])

    # ── Header rows ────────────────────────────────────────────
    _set_row_force(1, program_title, col_start=2)        # عنوان البرنامج
    _set_row_force(2, unit_title, col_start=2)           # رقم الوحدة والعنوان
    # R3: اسم المقيم / اسم المُدقق الداخلي — BLANK (manual)
    _set_row_force(4, assignment_title, col_start=2)     # عنوان المهمة

    # ── R6: student data row ───────────────────────────────────
    _set_unique(6, 0, student_name)
    _set_unique(6, 1, "التسليم الأول")
    _set_unique(6, 2, "، ".join(achieved_criteria) if achieved_criteria else "")
    _set_unique(6, 3, "نعم")
    # col 4 (wrongly awarded) and col 6 (reason) — BLANK

    # ── R8–R13: checklist answers ──────────────────────────────
    _set_unique(8, 7, "نعم")   # هل أكد الطالب والمقيم مصداقية الأدلة
    _set_unique(9, 7, "لا")    # هل هناك دليل على غش
    _set_unique(11, 7, "نعم")  # مرتبطة بشكل مباشر مع معايير التقييم
    _set_unique(12, 7, "نعم")  # تقدم التبرير المناسب
    _set_unique(13, 7, "نعم")  # التوجيه المناسب لتحسين الأداء

    # ── R15: general comments ──────────────────────────────────
    if general_comments and len(t0.rows) > 15:
        cell15 = t0.rows[15].cells[0]
        _set_cell_multiline(cell15, general_comments, rtl=True)
        _dered(cell15)

    # ── R19-R22: required actions = "لا يوجد" ─────────────────
    for ri in range(19, 23):
        if ri >= len(t0.rows):
            break
        row = t0.rows[ri]
        if row.cells:
            cell = row.cells[0]
            tc_id = id(cell._tc)
            if tc_id not in processed_tcs:
                processed_tcs.add(tc_id)
                _set_cell_text(cell, "لا يوجد")
                _dered(cell)

    # ── Final: de-red ALL remaining red text ───────────────────
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _dered(cell)
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.find(f'{{{ns}}}rPr')
            if rPr is not None:
                cel = rPr.find(f'{{{ns}}}color')
                if cel is not None and cel.get(f'{{{ns}}}val', '').upper() == RED:
                    cel.set(f'{{{ns}}}val', '000000')

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


async def fill_batch_iv_assessment_decisions(
    batch_id: int,
    template_path: str,
    selected_submission_ids: Optional[List[int]] = None,
    db_session=None,
) -> List[str]:
    """
    Generate IV Assessment Decisions documents for all students in a batch.
    Returns a list of paths to the generated .docx files.
    """
    from app.models import (
        BatchGrading, Submission, GradingResult,
        GradingCriteria, Assignment, SubmissionStatus
    )

    if db_session is None:
        from app.database import SessionLocal
        db_session = SessionLocal()
        should_close = True
    else:
        should_close = False

    try:
        batch = db_session.query(BatchGrading).filter(BatchGrading.id == batch_id).first()
        if not batch:
            raise ValueError(f"Batch {batch_id} not found")

        assignment = db_session.query(Assignment).filter(Assignment.id == batch.assignment_id).first()
        if not assignment:
            raise ValueError(f"Assignment not found for batch {batch_id}")

        # Build program title
        _unit_num_upper = (assignment.unit_number or "").upper()
        _level = 2 if _unit_num_upper.startswith("L2") else 3
        program_title = f"شهادات Pearson BTEC International من المستوى {_level} في تكنولوجيا المعلومات"

        unit_title = build_unit_title(assignment.unit_number or "", assignment.unit_name or "")

        # Build criteria title (same as assessment records)
        _letter_to_ar = {'A': 'أ', 'B': 'ب', 'C': 'ج', 'D': 'د', 'E': 'هـ', 'F': 'و', 'G': 'ز'}
        _groups: dict = {}
        if assignment.unit_criteria_json:
            try:
                for item in _json_mod.loads(assignment.unit_criteria_json):
                    code = (item.get('code') or '').strip()
                    aim = (item.get('learning_aim_ref') or (code[0] if code else '?')).upper()
                    if code:
                        _groups.setdefault(aim, []).append(code)
            except Exception:
                pass
        if not _groups:
            for cr in db_session.query(GradingCriteria).filter(
                GradingCriteria.assignment_id == assignment.id
            ).order_by(GradingCriteria.criteria_level).all():
                lvl = (cr.criteria_level or '').strip()
                letter = lvl[0].upper() if lvl else '?'
                _groups.setdefault(letter, []).append(lvl)
        _lines = []
        for _l in sorted(_groups.keys()):
            _ar = _letter_to_ar.get(_l, _l)
            _items = ', '.join(sorted(set(_groups[_l])))
            _lines.append(f'هدف التعلم ({_ar}): {_items}')
        assignment_title = '\n'.join(_lines) if _lines else (assignment.title or "")

        # Query submissions
        if selected_submission_ids:
            submissions = db_session.query(Submission).filter(
                Submission.id.in_(selected_submission_ids),
                Submission.status == SubmissionStatus.COMPLETED,
            ).all()
        else:
            submissions = db_session.query(Submission).filter(
                Submission.batch_id == batch_id,
                Submission.status == SubmissionStatus.COMPLETED,
            ).all()

        if not submissions:
            print("⚠️ No completed submissions found")
            return []

        output_dir = Path("uploads/reports/assessment_records")
        output_dir.mkdir(parents=True, exist_ok=True)

        generated_files: List[str] = []

        for submission in submissions:
            print(f"📝 IV decisions for: {submission.student_name}")

            # Collect achieved criteria codes
            results = (
                db_session.query(GradingResult, GradingCriteria)
                .join(GradingCriteria, GradingResult.criteria_id == GradingCriteria.id)
                .filter(GradingResult.submission_id == submission.id)
                .all()
            )
            achieved_criteria = sorted(
                [cr.criteria_level for res, cr in results if res.achieved and cr.criteria_level],
                key=lambda c: (
                    {'P': 0, 'M': 1, 'D': 2}.get(
                        (c.split('.')[-1][0].upper() if '.' in c else c[0].upper()),
                        9
                    ),
                    int(c.split('.')[-1][1:]) if '.' in c and len(c.split('.')[-1]) > 1
                    else int(c[1:]) if len(c) > 1 else 0
                )
            )

            # Build brief general comments
            student_name_val = submission.student_name or ""
            if achieved_criteria:
                crit_str = "، ".join(achieved_criteria)
                general_comments = (
                    f"تمت مراجعة الأدلة المقدمة من الطالب {student_name_val} "
                    f"وتبين أنها ذات صلة ومناسبة للمعايير المستهدفة. "
                    f"المعايير المحققة هي: {crit_str}. "
                    f"قرار التقييم المقدم صحيح ومتسق مع الأدلة والمعايير المستهدفة."
                )
            else:
                general_comments = (
                    f"تمت مراجعة الأدلة المقدمة من الطالب {student_name_val}. "
                    f"لم يتم تحقيق أي من معايير التقييم المستهدفة."
                )

            safe_name = re.sub(r'[^\w\s\u0600-\u06FF-]', '', student_name_val).strip()
            output_filename = f"iv_decisions_sub{submission.id}_{safe_name}.docx"
            output_path = str(output_dir / output_filename)

            try:
                fill_iv_assessment_decisions(
                    template_path=template_path,
                    output_path=output_path,
                    student_name=student_name_val,
                    program_title=program_title,
                    unit_title=unit_title,
                    assignment_title=assignment_title,
                    achieved_criteria=achieved_criteria,
                    general_comments=general_comments,
                )
                generated_files.append(output_path)
                print(f"  ✅ Saved: {output_path}")
            except Exception as e:
                print(f"  ❌ Failed for {student_name_val}: {e}")

        print(f"\n📄 Generated {len(generated_files)} IV decisions documents")
        return generated_files

    finally:
        if should_close:
            db_session.close()
