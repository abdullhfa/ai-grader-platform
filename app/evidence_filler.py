"""
Evidence Record Template Filler — fills two BTEC evidence templates per student.
"""

from __future__ import annotations

import os
import re
import json as _json_mod
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml import OxmlElement  # type: ignore

from app.template_filler import (  # type: ignore
    _set_cell_text,
    _set_cell_multiline,
    _ensure_rtl_paragraph,
    _ensure_rtl_run,
)


_PAGE_PATTERNS = [
    re.compile(r'(?:صفحة|صفحات|الصفحة|الصفحات)\s*[:\-]?\s*([\d٠-٩\-–,،\s]+)', re.IGNORECASE),
    re.compile(r'\b[Pp]ages?\.?\s*([\d\-–,\s]+)\b', re.IGNORECASE),
    re.compile(r'\bp\.?\s*([\d\-–,\s]+)\b', re.IGNORECASE),
]

_EVIDENCE_TYPE_MAP = [
    (re.compile(r'(?:تقرير|report|writing|document)', re.IGNORECASE), "Report"),
    (re.compile(r'(?:عرض|بوربوينت|powerpoint|\.pptx?\b|presentation|slides?)', re.IGNORECASE), "PPT"),
    (re.compile(r'(?:فيديو|video|recording|\.mp4\b|\.mov\b|\.avi\b)', re.IGNORECASE), "Video"),
    (re.compile(r'(?:صور|صورة|photo|photograph|image|screenshot|لقطة شاشة|لقطات شاشة)', re.IGNORECASE), "Photos"),
    (re.compile(r'(?:منتج نهائي|نموذج نهائي|final product|prototype)', re.IGNORECASE), "Final Product"),
    (re.compile(r'(?:excel|xlsx?\b|spreadsheet|ورقة عمل|جدول بيانات)', re.IGNORECASE), "Excel worksheet"),
    (re.compile(r'(?:observation|ملاحظة|مراقبة)', re.IGNORECASE), "Observation"),
    (re.compile(r'(?:Packet Tracer|محاكاة|simulation)', re.IGNORECASE), "Simulation file"),
    (re.compile(r'(?:كود|code|sketch|github|repository|\.py\b|\.js\b|\.html\b|\.css\b|\.cs\b)', re.IGNORECASE), "Code file"),
]


def _extract_page_refs(text: str) -> str:
    if not text:
        return ""
    for pat in _PAGE_PATTERNS:
        m = pat.search(text)
        if m:
            return m.group(0).strip()
    return ""


def _classify_evidence_type(*texts: str) -> str:
    blob = " ".join(t for t in texts if t)
    if not blob:
        return "Report"
    found = []
    for pat, name in _EVIDENCE_TYPE_MAP:
        if pat.search(blob) and name not in found:
            found.append(name)
    return ", ".join(found) if found else "Report"


# Range of Arabic Unicode codepoints used to detect RTL strings
_ARABIC_RE = re.compile(r"[؀-ۿݐ-ݿࢠ-ࣿﭐ-﷿ﹰ-﻿]")

# Pairs of mirrored characters that Word's bidi auto-mirroring flips visually
# in RTL paragraphs. We pre-swap them so Word's mirroring renders them in the
# user's expected visual position (open on the right, close on the left).
_BRACKET_SWAP = str.maketrans({
    "(": ")",
    ")": "(",
    "[": "]",
    "]": "[",
    "{": "}",
    "}": "{",
    "<": ">",
    ">": "<",
})


def _is_arabic_char(ch):
    """Return True if `ch` is an Arabic letter (Unicode Arabic block)."""
    if not ch:
        return False
    return bool(_ARABIC_RE.match(ch))


def _has_arabic_in_range(text, start, end):
    """Return True if any char in text[start:end] is Arabic."""
    for k in range(max(0, start), min(len(text), end)):
        if _is_arabic_char(text[k]):
            return True
    return False


def _fix_rtl_brackets(text):
    """Swap matching bracket *pairs* whose contents include Arabic letters.

    Word's bidi algorithm auto-mirrors brackets when they sit in RTL context
    (Arabic neighbours). Plain Latin/code segments like '[B.P3]' keep their
    LTR direction and are NOT mirrored, so we leave them alone.

    Strategy: walk the string, push opening brackets onto a stack, and when
    a closing bracket appears, pair it with the matching opener and swap
    BOTH only if the enclosed content (or the immediate Arabic neighbours)
    contain Arabic letters. This keeps balanced visual output."""
    if not text or not _ARABIC_RE.search(text):
        return text

    openers = {"(": ")", "[": "]", "{": "}"}
    closers = {v: k for k, v in openers.items()}

    chars = list(text)
    stack = []  # list of opener positions, with their bracket char
    swap_positions = []  # list of (open_idx, close_idx) pairs to swap

    for i, c in enumerate(chars):
        if c in openers:
            stack.append((i, c))
        elif c in closers and stack:
            open_idx, open_ch = stack[-1]
            # Only pop if the matching pair lines up
            if openers[open_ch] == c:
                stack.pop()
                # STRICT rule: swap only if Arabic appears INSIDE the
                # bracket pair. This matches Word's bidi behavior — pure
                # Latin content like '[B.P3]' or '(P)' is never mirrored
                # by Word, so we leave it alone regardless of neighbours.
                if _has_arabic_in_range(text, open_idx + 1, i):
                    swap_positions.append((open_idx, i))

    # Apply swaps to BOTH halves of each matched pair
    pair_swap = {"(": ")", ")": "(", "[": "]", "]": "[", "{": "}", "}": "{"}
    for open_idx, close_idx in swap_positions:
        chars[open_idx] = pair_swap[chars[open_idx]]
        chars[close_idx] = pair_swap[chars[close_idx]]

    return "".join(chars)


def _build_evidence_for_criterion(criterion):
    task_ref = (criterion.get("criteria_level") or "").strip()
    achieved = bool(criterion.get("achieved", False))
    dm = criterion.get("decision_matrix") or []
    if not isinstance(dm, list):
        dm = []
    evidence_pieces = []
    for item in dm:
        if not isinstance(item, dict):
            continue
        ev = (item.get("evidence") or "").strip()
        met = bool(item.get("met", False))
        if ev and ev != "-" and met:
            evidence_pieces.append(ev)
    description_text = "؛ ".join(evidence_pieces).strip()
    if not description_text:
        feedback = (criterion.get("feedback") or "").strip()
        if feedback and achieved:
            description_text = feedback
    page_ref = _extract_page_refs(description_text)
    if page_ref and not description_text.lstrip().startswith(page_ref):
        page_or_desc = f"{page_ref} — {description_text}"
    else:
        page_or_desc = description_text
    evidence_provided = _classify_evidence_type(description_text, criterion.get("feedback") or "")
    if not achieved:
        return (task_ref, "", "")
    return (task_ref, evidence_provided, page_or_desc.strip())


def _group_by_learning_aim(criteria_results):
    groups = {}
    for cr in criteria_results:
        lvl = (cr.get("criteria_level") or "").strip()
        if "." in lvl:
            prefix = lvl.split(".", 1)[0].upper()
        else:
            prefix = "A"
        groups.setdefault(prefix, []).append(cr)
    return groups


def _aggregate_la_evidence(criteria_in_la):
    all_desc = []
    all_types_blob = []
    for cr in criteria_in_la:
        if not cr.get("achieved"):
            continue
        _, ev_type, page_or_desc = _build_evidence_for_criterion(cr)
        if page_or_desc:
            tag = (cr.get("criteria_level") or "").strip()
            if tag:
                all_desc.append(f"[{tag}] {page_or_desc}")
            else:
                all_desc.append(page_or_desc)
        if ev_type:
            all_types_blob.append(ev_type)
    seen = set()
    unique_types = []
    for t in all_types_blob:
        for sub in [s.strip() for s in t.split(",")]:
            if sub and sub not in seen:
                seen.add(sub)
                unique_types.append(sub)
    description = "\n".join(all_desc).strip()
    if unique_types:
        evidence_provided = ", ".join(unique_types)
    elif description:
        evidence_provided = "Report"
    else:
        evidence_provided = ""
    return (evidence_provided, description)


# ════════════════════════════════════════════════════════════════════
#   DYNAMIC PER-STUDENT ASSESSOR COMMENT GENERATION
# ════════════════════════════════════════════════════════════════════

def _summarise_evidence_types(criteria_results):
    types = []
    seen = set()
    for cr in criteria_results:
        if not cr.get("achieved"):
            continue
        _, ev_type, _ = _build_evidence_for_criterion(cr)
        for sub in [s.strip() for s in (ev_type or "").split(",")]:
            if sub and sub not in seen:
                seen.add(sub)
                types.append(sub)
    return types


def _summarise_la_progress(criteria_results):
    grouped = _group_by_learning_aim(criteria_results)
    out = {}
    for prefix, items in grouped.items():
        ach = sum(1 for c in items if c.get("achieved"))
        out[prefix] = (ach, len(items))
    return out


def _generate_evidence_comment_fallback(student_name, grade_level, criteria_results):
    """Deterministic per-student comment that varies by name/grade/criteria.
    Always different across different students; identical across re-runs for
    the same student so the record is reproducible."""
    name_hash = int(hashlib.md5((student_name or "").encode("utf-8")).hexdigest(), 16)

    achieved = [c for c in criteria_results if c.get("achieved")]
    not_achieved = [c for c in criteria_results if not c.get("achieved")]
    achieved_codes = [str(c.get("criteria_level") or "") for c in achieved]
    not_achieved_codes = [str(c.get("criteria_level") or "") for c in not_achieved]

    la_progress = _summarise_la_progress(criteria_results)
    la_phrases = []
    for prefix in ["A", "B", "C", "D"]:
        if prefix not in la_progress:
            continue
        ach, total = la_progress[prefix]
        if ach == total and total > 0:
            la_phrases.append(f"LA. {prefix} ({ach}/{total} مكتمل)")
        elif ach > 0:
            la_phrases.append(f"LA. {prefix} ({ach}/{total} جزئي)")
        else:
            la_phrases.append(f"LA. {prefix} (لم يتحقق)")

    evidence_types = _summarise_evidence_types(criteria_results)
    evidence_blob = "، ".join(evidence_types) if evidence_types else "تقرير"

    if not_achieved:
        opener_pool = [
            f"قدّم الطالب/ة {student_name} عملاً يعكس جهداً ملموساً، حيث حقّق المعايير: {', '.join(achieved_codes)} عبر تقديم {evidence_blob}.",
            f"اطّلع المُقيِّم على عمل الطالب/ة {student_name} الذي تضمّن {evidence_blob}، وقد استوفى المعايير التالية: {', '.join(achieved_codes)}.",
            f"بناءً على مراجعة المُقيِّم، أنجز الطالب/ة {student_name} المعايير {', '.join(achieved_codes)} مستخدماً {evidence_blob}.",
            f"حقّق الطالب/ة {student_name} ({grade_level}) المعايير {', '.join(achieved_codes)} من خلال {evidence_blob} المُقدَّمة.",
        ]
    else:
        opener_pool = [
            f"قدّم الطالب/ة {student_name} عملاً متكاملاً يستوفي جميع المعايير المطلوبة عبر {evidence_blob} مرفقة بأرقام الصفحات.",
            f"أنجز الطالب/ة {student_name} كافة معايير المهمة بنجاح، مع تقديم {evidence_blob} موثّقة بشكل واضح.",
            f"بناءً على مراجعة المُقيِّم، حقّق الطالب/ة {student_name} المستوى {grade_level} الكامل من خلال {evidence_blob} المرفقة.",
            f"يعكس عمل الطالب/ة {student_name} استيفاءً تامّاً للمعايير المستهدفة عبر {evidence_blob} المنظَّمة بأرقام الصفحات.",
        ]
    line1 = opener_pool[name_hash % len(opener_pool)]

    line2 = ""
    if la_phrases:
        line2 = "تتوزَّع الإنجازات على أهداف التعلم كالآتي: " + "، ".join(la_phrases) + "."

    if not_achieved:
        missing = "، ".join(not_achieved_codes)
        closer_pool = [
            f"يُوصى بمعالجة المعايير غير المكتملة ({missing}) من خلال إضافة الأدلة العملية المطلوبة لتحقيق مستوى أعلى.",
            f"المعايير {missing} تحتاج إلى تقديم أدلة إضافية وفق متطلبات مواصفات BTEC.",
            f"للوصول إلى المستوى التالي، يُنصح بإكمال المعايير {missing} مع توثيق أرقام الصفحات.",
        ]
        line3 = closer_pool[name_hash % len(closer_pool)]
    else:
        closer_pool = [
            f"المستوى المُحقَّق ({grade_level}) يعكس فهماً عميقاً وتطبيقاً عملياً للمعايير المستهدفة.",
            f"العمل المُقدَّم يستحق الإشادة لاتساقه واكتماله، وقد بلغ المستوى ({grade_level}) عن جدارة.",
            f"تمَّت مراجعة الأدلة المُقدَّمة وتأكَّد المُقيِّم من مصداقيتها واتساقها مع متطلبات المعايير ({grade_level}).",
        ]
        line3 = closer_pool[name_hash % len(closer_pool)]

    parts = [line1]
    if line2:
        parts.append(line2)
    parts.append(line3)
    return "\n".join(parts)


async def _generate_evidence_comment_via_ai(student_name, grade_level, criteria_results, ai_provider=None):
    """Try AI first, fall back to deterministic generator if unavailable."""
    if ai_provider is None:
        return _generate_evidence_comment_fallback(student_name, grade_level, criteria_results)

    achieved = [c for c in criteria_results if c.get("achieved")]
    not_achieved = [c for c in criteria_results if not c.get("achieved")]
    evidence_types = _summarise_evidence_types(criteria_results)

    achieved_summary = "\n".join(
        f"- {c.get('criteria_level', '?')}: {(c.get('feedback') or '')[:200]}"
        for c in achieved
    )
    not_achieved_summary = "\n".join(
        f"- {c.get('criteria_level', '?')}: {(c.get('feedback') or '')[:120]}"
        for c in not_achieved
    )

    prompt_lines = [
        "أنت مُقيِّم BTEC محترف. اكتب تعليق المُقيِّم لأسفل ملف \"أدلة الإنجاز\" (Evidence Record) للطالب/ة.",
        "التعليق يجب أن:",
        "- يكون مخصصاً لهذا الطالب بالذات (لا قوالب جاهزة)",
        "- يَذكر أنواع الأدلة التي قدّمها فعلياً",
        "- يُشير إلى المعايير المُحقَّقة وغير المُحقَّقة",
        "- يكون من 3-5 جمل بالعربية الفصحى المختصرة",
        "- لا يحتوي على علامات تنسيق Markdown",
        "",
        "بيانات الطالب:",
        f"- الاسم: {student_name}",
        f"- الدرجة النهائية: {grade_level}",
        f"- أنواع الأدلة المُقدَّمة: {', '.join(evidence_types) if evidence_types else 'لا توجد'}",
        "",
        "المعايير المُحقَّقة:",
        achieved_summary or "(لا يوجد)",
        "",
        "المعايير غير المُحقَّقة:",
        not_achieved_summary or "(لا يوجد)",
        "",
        "اكتب التعليق الآن مباشرةً (بدون مقدمة أو عنوان):",
    ]
    prompt = "\n".join(prompt_lines)

    try:
        import asyncio
        messages = [
            {"role": "system", "content": "أنت مُقيِّم BTEC محترف تكتب تعليقات مُقيِّم مخصّصة لكل طالب باللغة العربية الفصحى المختصرة."},
            {"role": "user", "content": prompt},
        ]
        response = await asyncio.to_thread(
            ai_provider.chat_completion,
            messages=messages,
            temperature=0.7,
        )
        if isinstance(response, str) and response.strip():
            return response.strip()
    except Exception as e:
        print(f"⚠️ AI evidence comment generation failed (using fallback): {e}")

    return _generate_evidence_comment_fallback(student_name, grade_level, criteria_results)


# ════════════════════════════════════════════════════════════════════
#   ASSESSOR-COMMENT INJECTION INTO TEMPLATES
# ════════════════════════════════════════════════════════════════════

def _write_comment_into_table(doc, header_text, comment_text):
    """Find the 'Comments...' row in the first table and write the
    personalised comment into it, replacing the placeholder.

    The 3 cells in the comment row are typically MERGED (they share the
    same _tc element), so we must write to cell 0 only — clearing the
    other "cells" would actually wipe the comment we just wrote. We
    detect this by comparing the underlying _tc identities."""
    if not comment_text or not doc.tables:
        return
    table = doc.tables[0]
    comment_row = None
    for ri in range(len(table.rows) - 1, -1, -1):
        if table.rows[ri].cells[0].text.strip().lower().startswith("comment"):
            comment_row = ri
            break
    if comment_row is None:
        return
    # Apply RTL bracket-swap so Word's auto-mirroring renders the brackets
    # in their visually correct positions for Arabic readers.
    full_text = f"{header_text}\n{_fix_rtl_brackets(comment_text)}"
    cells = table.rows[comment_row].cells
    # Identify unique underlying cells (merged cells share the same _tc).
    seen_ids = {id(cells[0]._tc)}
    _set_cell_multiline(cells[0], full_text, rtl=True)
    for ci in range(1, len(cells)):
        if id(cells[ci]._tc) in seen_ids:
            # Same merged cell — skip, otherwise we'd erase the comment.
            continue
        seen_ids.add(id(cells[ci]._tc))
        _set_cell_text(cells[ci], "", rtl=True)


# ════════════════════════════════════════════════════════════════════
#   TEMPLATE 1: "Evidance - IT.docx"
# ════════════════════════════════════════════════════════════════════

def fill_evidence_record(template_path, output_path, student_name,
                         criteria_results, assessor_comments=""):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Template has no tables — wrong template?")
    table = doc.tables[0]
    n_rows = len(table.rows)
    if n_rows < 3:
        raise ValueError(f"Template has too few rows ({n_rows})")

    def _sort_key(c):
        lv = (c.get("criteria_level") or "").strip()
        short = lv.split(".")[-1] if "." in lv else lv
        order = {"P": 0, "M": 1, "D": 2}
        letter = short[:1].upper() if short else "Z"
        m = re.search(r"\d+", short)
        num = int(m.group()) if m else 99
        prefix = lv.split(".")[0] if "." in lv else ""
        return (prefix, order.get(letter, 9), num)

    sorted_criteria = sorted(criteria_results, key=_sort_key)

    comment_row_idx = n_rows - 1
    for ri in range(n_rows - 1, 0, -1):
        if table.rows[ri].cells[0].text.strip().lower().startswith("comment"):
            comment_row_idx = ri
            break

    max_data_rows = comment_row_idx - 1
    rows_to_fill = sorted_criteria[:max_data_rows]

    for i, cr in enumerate(rows_to_fill, start=1):
        task_ref, ev_provided, page_or_desc = _build_evidence_for_criterion(cr)
        cells = table.rows[i].cells
        _set_cell_text(cells[0], task_ref or "", rtl=False, alignment="center")
        _set_cell_text(cells[1], ev_provided or "", rtl=False)
        # Apply RTL bracket-swap ONCE just before writing — Word will then
        # mirror the brackets back to their visually correct positions.
        _set_cell_multiline(cells[2], _fix_rtl_brackets(page_or_desc or ""), rtl=False)

    for i in range(len(rows_to_fill) + 1, comment_row_idx):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], "", rtl=False)
        _set_cell_text(cells[1], "", rtl=False)
        _set_cell_text(cells[2], "", rtl=False)

    _write_comment_into_table(doc, "Comments to be recorded by the assessor:", assessor_comments)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


_LA_LABELS = ["LA. A", "LA. B", "LA. C", "LA. D"]
_LA_PREFIXES = ["A", "B", "C", "D"]


def fill_la_evidence_record(template_path, output_path, student_name,
                            criteria_results, assessor_comments=""):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Template has no tables")
    table = doc.tables[0]
    if len(table.rows) < 5:
        raise ValueError(f"Template Table 0 has too few rows ({len(table.rows)})")

    grouped = _group_by_learning_aim(criteria_results)
    for la_idx, prefix in enumerate(_LA_PREFIXES):
        row_idx = la_idx + 1
        if row_idx >= len(table.rows):
            break
        crits_in_la = grouped.get(prefix, [])
        if crits_in_la:
            evidence_provided, description = _aggregate_la_evidence(crits_in_la)
        else:
            evidence_provided, description = ("", "")
        cells = table.rows[row_idx].cells
        _set_cell_text(cells[0], _LA_LABELS[la_idx], rtl=False, alignment="center")
        _set_cell_multiline(cells[1], evidence_provided, rtl=False)
        # Apply RTL bracket-swap ONCE just before writing
        _set_cell_multiline(cells[2], _fix_rtl_brackets(description), rtl=False)

    _write_comment_into_table(doc, "Comments for note by the Assessor:", assessor_comments)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ════════════════════════════════════════════════════════════════════
#   BATCH WRAPPER
# ════════════════════════════════════════════════════════════════════

async def fill_batch_evidence_records(
    batch_id,
    evidence_template_path,
    la_template_path,
    output_dir="uploads/reports/evidence_records",
    selected_submission_ids=None,
    db_session=None,
):
    from app.models import (  # type: ignore
        BatchGrading, Submission, GradingSummary, GradingResult, GradingCriteria
    )

    # Try to obtain the global AI provider — if unavailable, comments
    # fall back to the deterministic per-student generator.
    ai_provider = None
    try:
        from app.ai_provider import get_global_provider  # type: ignore
        ai_provider = get_global_provider()
    except Exception:
        ai_provider = None

    if db_session is None:
        from app.database import SessionLocal  # type: ignore
        db_session = SessionLocal()
        should_close = True
    else:
        should_close = False

    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)
    generated_files = []

    try:
        batch = db_session.query(BatchGrading).filter(BatchGrading.id == batch_id).first()
        if not batch:
            raise ValueError(f"Batch {batch_id} not found")

        sub_query = db_session.query(Submission).filter(Submission.batch_id == batch.id)
        if selected_submission_ids:
            sub_query = sub_query.filter(Submission.id.in_(selected_submission_ids))
        submissions = sub_query.all()

        for submission in submissions:
            print(f"📝 Filling evidence records for: {submission.student_name}")

            summary = (
                db_session.query(GradingSummary)
                .filter(GradingSummary.submission_id == submission.id)
                .first()
            )
            grade_level = (summary.grade_level if summary else "U") or "U"

            snapshot = None
            if getattr(submission, "grading_snapshot_json", None):
                try:
                    snapshot = _json_mod.loads(str(submission.grading_snapshot_json))
                except Exception:
                    snapshot = None

            if not snapshot or not snapshot.get("criteria_results"):
                results_db = (
                    db_session.query(GradingResult, GradingCriteria)
                    .join(GradingCriteria, GradingResult.criteria_id == GradingCriteria.id)
                    .filter(GradingResult.submission_id == submission.id)
                    .all()
                )
                if not results_db:
                    print(f"  ⚠️  No grading results for {submission.student_name}, skipping")
                    continue
                criteria_results = []
                for r, c in results_db:
                    fb = r.feedback or ""
                    dm_evidence = fb if fb and r.achieved else "-"
                    criteria_results.append({
                        "criteria_level": c.criteria_level or "",
                        "achieved": bool(r.achieved),
                        "feedback": fb,
                        "decision_matrix": [{
                            "requirement": c.criteria_description or c.criteria_level or "",
                            "met": bool(r.achieved),
                            "evidence": dm_evidence,
                        }],
                    })
            else:
                criteria_results = snapshot.get("criteria_results") or []

            # ── Generate per-student dynamic comment ──
            comment_text = await _generate_evidence_comment_via_ai(
                student_name=submission.student_name or "",
                grade_level=grade_level,
                criteria_results=criteria_results,
                ai_provider=ai_provider,
            )
            print(f"  💬 Comment ({len(comment_text)} chars) generated for {submission.student_name}")

            safe_name = re.sub(r'[^\w\s؀-ۿ-]', '', submission.student_name or '').strip()

            ev1_filename = f"evidence_record_sub{submission.id}_{safe_name}.docx"
            ev1_path = str(out_path / ev1_filename)
            try:
                fill_evidence_record(
                    template_path=evidence_template_path,
                    output_path=ev1_path,
                    student_name=submission.student_name or "",
                    criteria_results=criteria_results,
                    assessor_comments=comment_text,
                )
                generated_files.append(ev1_path)
                print(f"  ✅ Saved evidence record: {ev1_path}")
            except Exception as e:
                print(f"  ❌ Failed evidence record for {submission.student_name}: {e}")

            la_filename = f"la_evidence_sub{submission.id}_{safe_name}.docx"
            la_path = str(out_path / la_filename)
            try:
                fill_la_evidence_record(
                    template_path=la_template_path,
                    output_path=la_path,
                    student_name=submission.student_name or "",
                    criteria_results=criteria_results,
                    assessor_comments=comment_text,
                )
                generated_files.append(la_path)
                print(f"  ✅ Saved LA evidence: {la_path}")
            except Exception as e:
                print(f"  ❌ Failed LA evidence for {submission.student_name}: {e}")

        return generated_files

    finally:
        if should_close:
            db_session.close()


# ════════════════════════════════════════════════════════════════════
#   SINGLE-STUDENT WRAPPER (called from per-student button)
# ════════════════════════════════════════════════════════════════════

async def fill_single_student_evidence_records(
    submission_id,
    evidence_template_path,
    la_template_path,
    output_dir="uploads/reports/evidence_records",
    db_session=None,
):
    """Fill BOTH evidence templates for a SINGLE student (by submission_id).
    Returns the list of generated file paths (length 0–2)."""
    from app.models import (  # type: ignore
        Submission, GradingSummary, GradingResult, GradingCriteria
    )

    ai_provider = None
    try:
        from app.ai_provider import get_global_provider  # type: ignore
        ai_provider = get_global_provider()
    except Exception:
        ai_provider = None

    if db_session is None:
        from app.database import SessionLocal  # type: ignore
        db_session = SessionLocal()
        should_close = True
    else:
        should_close = False

    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)
    generated_files = []

    try:
        submission = db_session.query(Submission).filter(Submission.id == submission_id).first()
        if not submission:
            raise ValueError(f"Submission {submission_id} not found")

        summary = (
            db_session.query(GradingSummary)
            .filter(GradingSummary.submission_id == submission.id)
            .first()
        )
        grade_level = (summary.grade_level if summary else "U") or "U"

        snapshot = None
        if getattr(submission, "grading_snapshot_json", None):
            try:
                snapshot = _json_mod.loads(str(submission.grading_snapshot_json))
            except Exception:
                snapshot = None

        if not snapshot or not snapshot.get("criteria_results"):
            results_db = (
                db_session.query(GradingResult, GradingCriteria)
                .join(GradingCriteria, GradingResult.criteria_id == GradingCriteria.id)
                .filter(GradingResult.submission_id == submission.id)
                .all()
            )
            if not results_db:
                raise ValueError("لا توجد نتائج تصحيح لهذا الطالب — صحّح الواجب أولاً.")
            criteria_results = []
            for r, c in results_db:
                fb = r.feedback or ""
                dm_evidence = fb if fb and r.achieved else "-"
                criteria_results.append({
                    "criteria_level": c.criteria_level or "",
                    "achieved": bool(r.achieved),
                    "feedback": fb,
                    "decision_matrix": [{
                        "requirement": c.criteria_description or c.criteria_level or "",
                        "met": bool(r.achieved),
                        "evidence": dm_evidence,
                    }],
                })
        else:
            criteria_results = snapshot.get("criteria_results") or []

        comment_text = await _generate_evidence_comment_via_ai(
            student_name=submission.student_name or "",
            grade_level=grade_level,
            criteria_results=criteria_results,
            ai_provider=ai_provider,
        )

        safe_name = re.sub(r'[^\w\s؀-ۿ-]', '', submission.student_name or '').strip()

        ev1_filename = f"evidence_record_sub{submission.id}_{safe_name}.docx"
        ev1_path = str(out_path / ev1_filename)
        fill_evidence_record(
            template_path=evidence_template_path,
            output_path=ev1_path,
            student_name=submission.student_name or "",
            criteria_results=criteria_results,
            assessor_comments=comment_text,
        )
        generated_files.append(ev1_path)

        la_filename = f"la_evidence_sub{submission.id}_{safe_name}.docx"
        la_path = str(out_path / la_filename)
        fill_la_evidence_record(
            template_path=la_template_path,
            output_path=la_path,
            student_name=submission.student_name or "",
            criteria_results=criteria_results,
            assessor_comments=comment_text,
        )
        generated_files.append(la_path)

        return generated_files

    finally:
        if should_close:
            db_session.close()
