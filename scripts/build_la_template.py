"""
One-time script: generate the missing Arabic LA-Evidence template at
uploads/templates/نموذج ربط أدلة المتعلم بأهداف التعلّم.docx

This is NOT the official Pearson template — it is a structurally valid
substitute so the evidence_filler code can populate it. If the institution
has the official template, drop it in uploads/templates/ with the same
filename and it will be used instead.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _shade(cell, hex_color):
    """Add cell background shading."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_rtl(paragraph):
    """Make paragraph RTL (Arabic)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def build_la_evidence_template(output_path):
    doc = Document()

    # Page setup — narrow margins for a single-page form
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ─── Title ───
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("نموذج ربط أدلة المتعلم بأهداف التعلّم")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    _set_rtl(title)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Learner Evidence to Learning Aim Mapping")
    sr.font.size = Pt(11)
    sr.italic = True
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ─── Header info as a paragraph (NOT a table) so the LA table stays at tables[0]
    # The evidence_filler.py code does: table = doc.tables[0] — so the first table
    # MUST be the LA mapping table.
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ip = info_para.add_run("اسم المتعلم: __________________________   |   اسم الواجب: __________________________   |   التاريخ: ____________")
    ip.font.size = Pt(10)
    ip.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _set_rtl(info_para)

    doc.add_paragraph()

    # ─── Main LA evidence mapping table (the one the code fills) ───
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Column widths
    widths = [Cm(2.5), Cm(7.0), Cm(7.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # Header row (row 0)
    hdr = table.rows[0].cells
    hdr[0].text = "هدف التعلّم\nLearning Aim"
    hdr[1].text = "الأدلة المقدمة\nEvidence Provided"
    hdr[2].text = "وصف ربط الدليل بالمعيار\nDescription"

    for c in hdr:
        _shade(c, "1F497D")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(11)

    # Rows 1-4: one per Learning Aim. The code fills these.
    # Pre-populate column 0 with LA labels — the filler will overwrite these
    # but it's nice to have them visible if a teacher opens the empty template.
    la_labels = ["LA. A", "LA. B", "LA. C", "LA. D"]
    for i, label in enumerate(la_labels):
        row = table.rows[i + 1]
        cell0 = row.cells[0]
        cell0.text = label
        cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
        # Empty placeholder cells for the filler to populate
        for j in (1, 2):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Empty paragraph so cell isn't collapsed
            cell.text = ""

    doc.add_paragraph()

    # ─── Assessor comments section (the code locates this by exact heading text) ───
    h = doc.add_paragraph()
    hr = h.add_run("Comments for note by the Assessor:")
    hr.bold = True
    hr.font.size = Pt(11)
    hr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    comments_table = doc.add_table(rows=1, cols=1)
    comments_table.style = "Light Grid Accent 1"
    comments_cell = comments_table.rows[0].cells[0]
    # Provide ~8 lines of vertical space
    for _ in range(8):
        comments_cell.add_paragraph("")

    doc.add_paragraph()

    # ─── Footer / signatures ───
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.rows[0].cells[0].text = "توقيع المُقيِّم / Assessor Signature"
    sig_table.rows[0].cells[1].text = "التاريخ / Date"
    sig_table.rows[1].cells[0].text = ""
    sig_table.rows[1].cells[1].text = ""
    for c in sig_table.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[OK] Template created at: {output_path}")


if __name__ == "__main__":
    out = "uploads/templates/نموذج ربط أدلة المتعلم بأهداف التعلّم.docx"
    build_la_evidence_template(out)
