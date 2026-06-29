import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Helpers for RTL and Styling
def set_p_rtl(p):
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def set_run_rtl(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=None):
    run.font.rtl = True
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    rPr = run._r.get_or_add_rPr()
    # Complex Script font configuration
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    run.font.size = Pt(size_pt)

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_rtl(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    # Set RTL cell layout if needed (Word auto-handles paragraph alignment mostly)
    # The best way is to ensure all paragraphs inside are RTL

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

# Inline style parsing
def parse_inline_styles(text):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    segments = []
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            segments.append((part[2:-2], 'bold', None))
        elif part.startswith('*') and part.endswith('*'):
            segments.append((part[1:-1], 'italic', None))
        elif part.startswith('`') and part.endswith('`'):
            segments.append((part[1:-1], 'code', None))
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                segments.append((match.group(1), 'link', match.group(2)))
            else:
                segments.append((part, 'normal', None))
        else:
            segments.append((part, 'normal', None))
    return segments

def add_runs_to_paragraph(p, text, is_arabic=True, font_size=11, force_color=None, is_code=False):
    segments = parse_inline_styles(text)
    for seg_text, style, url in segments:
        run = p.add_run(seg_text)
        
        # Color resolution
        color = force_color
        if style == 'link' and not color:
            color = RGBColor(37, 99, 235)  # Blue
            
        bold = (style == 'bold')
        italic = (style == 'italic')
        
        # Set RTL and fonts
        font_name = "Courier New" if (style == 'code' or is_code) else "Arial"
        size = font_size - 1 if (style == 'code' or is_code) else font_size
        
        set_run_rtl(run, font_name=font_name, size_pt=size, bold=bold, italic=italic, color_rgb=color)
        
        if style == 'code':
            # Add a slight highlight or shading if supported, else just monospace
            pass

def create_callout_paragraph(doc, alert_type, text_content):
    # Set colors based on alert type
    # Warning: Yellow/Orange, Caution: Red, Note: Blue, Important: Purple
    colors = {
        'WARNING': ('FEF3C7', RGBColor(180, 83, 9)),      # Background, Text
        'CAUTION': ('FEE2E2', RGBColor(185, 28, 28)),
        'NOTE': ('EFF6FF', RGBColor(29, 78, 216)),
        'IMPORTANT': ('F5F3FF', RGBColor(109, 40, 217)),
        'DEFAULT': ('F3F4F6', RGBColor(75, 85, 99))
    }
    
    bg_color, text_color = colors.get(alert_type.upper(), colors['DEFAULT'])
    
    # We can simulate a blockquote by using a table with 1 cell to have borders and background
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=200, right=200)
    
    # Set left border (or right border since it's RTL, let's set left and right to be safe/symmetrical or left)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{bg_color}"/>'
        f'<w:right w:val="single" w:sz="24" w:space="0" w:color="{bg_color}"/>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    set_p_rtl(p)
    
    # Add Alert tag
    run_tag = p.add_run(f"【 {alert_type} 】 ")
    set_run_rtl(run_tag, font_name="Arial", size_pt=11, bold=True, color_rgb=text_color)
    
    # Add text
    add_runs_to_paragraph(p, text_content, font_size=11, force_color=text_color)

def parse_markdown_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    
    doc = Document()
    
    # Configure document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    code_lang = ""
    
    in_table = False
    table_rows = []
    
    i = 0
    num_lines = len(lines)
    
    while i < num_lines:
        line = lines[i]
        stripped = line.strip()
        
        # 1. Handle Code Blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block, create a table cell for code block to look pretty
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F9FAFB")
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                
                # thin borders
                tcPr = cell._tc.get_or_add_tcPr()
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="single" w:sz="4" w:color="E5E7EB"/>'
                    f'<w:bottom w:val="single" w:sz="4" w:color="E5E7EB"/>'
                    f'<w:left w:val="single" w:sz="4" w:color="E5E7EB"/>'
                    f'<w:right w:val="single" w:sz="4" w:color="E5E7EB"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(borders)
                
                p = cell.paragraphs[0]
                # Code blocks are left aligned
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.15
                
                code_text = "\n".join(code_content)
                run = p.add_run(code_text)
                set_run_rtl(run, font_name="Courier New", size_pt=9.5, color_rgb=RGBColor(55, 65, 81))
                run.font.rtl = False # code is LTR usually
                
                # Add spacing after
                p_after = doc.add_paragraph()
                p_after.paragraph_format.space_before = Pt(0)
                p_after.paragraph_format.space_after = Pt(6)
                
                in_code_block = False
                code_content = []
            else:
                in_code_block = True
                code_lang = stripped[3:].strip()
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line.rstrip('\r\n'))
            i += 1
            continue
            
        # 2. Handle Tables
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            # Parse row cells
            # Split by | but ignore leading/trailing empty elements
            row_cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(row_cells)
            i += 1
            continue
        elif in_table:
            # End of table
            # Process table_rows
            if len(table_rows) > 0:
                # Filter out the separator row (contains ---)
                filtered_rows = []
                for r in table_rows:
                    if all(re.match(r'^:?-+:?$', c) for c in r if c):
                        continue
                    filtered_rows.append(r)
                
                if len(filtered_rows) > 0:
                    num_cols = max(len(r) for r in filtered_rows)
                    num_rows = len(filtered_rows)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(table)
                    
                    for row_idx, r_data in enumerate(filtered_rows):
                        row = table.rows[row_idx]
                        is_header = (row_idx == 0)
                        
                        bg_color = "1B365D" if is_header else ("F9FAFB" if row_idx % 2 == 1 else "FFFFFF")
                        text_color = RGBColor(255, 255, 255) if is_header else RGBColor(31, 41, 55)
                        
                        for col_idx, cell_value in enumerate(r_data):
                            if col_idx >= num_cols:
                                break
                            cell = row.cells[col_idx]
                            set_cell_background(cell, bg_color)
                            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                            
                            p = cell.paragraphs[0]
                            set_p_rtl(p)
                            p.paragraph_format.space_after = Pt(0)
                            
                            # Bold headers
                            add_runs_to_paragraph(p, cell_value, font_size=10, force_color=text_color)
                            if is_header:
                                for run in p.runs:
                                    run.bold = True
                                    
                    # Add spacing paragraph after table
                    p_after = doc.add_paragraph()
                    p_after.paragraph_format.space_after = Pt(12)
                    
            in_table = False
            table_rows = []
            # continue parsing current line (which is not part of the table)
            
        # 3. Handle Alert Callouts (GitHub format: > [!WARNING] followed by content)
        if stripped.startswith(">"):
            # Check if this is an alert header
            alert_match = re.match(r'^>\s*\[!(WARNING|CAUTION|NOTE|IMPORTANT)\]\s*(.*)$', stripped)
            if alert_match:
                alert_type = alert_match.group(1)
                alert_text = alert_match.group(2)
                
                # Check if the next line(s) are part of the blockquote
                blockquote_lines = []
                if alert_text:
                    blockquote_lines.append(alert_text)
                
                next_i = i + 1
                while next_i < num_lines and lines[next_i].strip().startswith(">"):
                    next_line_stripped = lines[next_i].strip()[1:].strip()
                    # If it's another alert header, stop or continue
                    if re.match(r'^\[!(WARNING|CAUTION|NOTE|IMPORTANT)\]', next_line_stripped):
                        break
                    blockquote_lines.append(next_line_stripped)
                    next_i += 1
                
                i = next_i
                combined_text = " ".join(blockquote_lines)
                create_callout_paragraph(doc, alert_type, combined_text)
                continue
            
            # Simple blockquote (no alert tag)
            quote_text = stripped[1:].strip()
            next_i = i + 1
            while next_i < num_lines and lines[next_i].strip().startswith(">") and not re.match(r'^>\s*\[!', lines[next_i].strip()):
                quote_text += " " + lines[next_i].strip()[1:].strip()
                next_i += 1
            i = next_i
            create_callout_paragraph(doc, "NOTE", quote_text)
            continue
            
        # 4. Handle Headers
        if stripped.startswith("#"):
            match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if match:
                level = len(match.group(1))
                title_text = match.group(2)
                
                p = doc.add_paragraph()
                set_p_rtl(p)
                p.paragraph_format.keep_with_next = True
                
                if level == 1:
                    p.paragraph_format.space_before = Pt(24)
                    p.paragraph_format.space_after = Pt(12)
                    add_runs_to_paragraph(p, title_text, font_size=18, force_color=RGBColor(27, 54, 93))
                    for run in p.runs:
                        run.bold = True
                elif level == 2:
                    p.paragraph_format.space_before = Pt(18)
                    p.paragraph_format.space_after = Pt(8)
                    add_runs_to_paragraph(p, title_text, font_size=14, force_color=RGBColor(37, 99, 235))
                    for run in p.runs:
                        run.bold = True
                elif level == 3:
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(6)
                    add_runs_to_paragraph(p, title_text, font_size=12, force_color=RGBColor(31, 41, 55))
                    for run in p.runs:
                        run.bold = True
                else:
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                    add_runs_to_paragraph(p, title_text, font_size=11, force_color=RGBColor(75, 85, 99))
                    for run in p.runs:
                        run.bold = True
                i += 1
                continue
                
        # 5. Handle Horizontal Rules
        if stripped in ["---", "***", "___"]:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E5E7EB"/></w:pBdr>')
            pPr.append(pBdr)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
            
        # 6. Handle Empty Lines
        if not stripped:
            i += 1
            continue
            
        # 7. Handle Lists (Unordered & Ordered)
        list_match = re.match(r'^([\-\*\+])\s+(.*)$', stripped)
        num_list_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        
        if list_match:
            bullet_char = list_match.group(1)
            list_text = list_match.group(2)
            
            p = doc.add_paragraph()
            set_p_rtl(p)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.25)
            
            # Simulated bullet marker (using RTL bullet char)
            run_bullet = p.add_run("•  ")
            set_run_rtl(run_bullet, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(37, 99, 235))
            
            add_runs_to_paragraph(p, list_text, font_size=11)
            i += 1
            continue
            
        elif num_list_match:
            num_str = num_list_match.group(1)
            list_text = num_list_match.group(2)
            
            p = doc.add_paragraph()
            set_p_rtl(p)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.25)
            
            run_num = p.add_run(f"{num_str}.  ")
            set_run_rtl(run_num, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(37, 99, 235))
            
            add_runs_to_paragraph(p, list_text, font_size=11)
            i += 1
            continue
            
        # 8. Regular Paragraph
        p = doc.add_paragraph()
        set_p_rtl(p)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        
        add_runs_to_paragraph(p, stripped, font_size=11)
        i += 1

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}!")

if __name__ == "__main__":
    base_dir = r"C:\Users\aalsa\.gemini\antigravity-ide\brain\63b19ed6-10e8-4de2-9d8e-3361ebaf5cfb"
    
    md1 = os.path.join(base_dir, "audit_report.md")
    docx1 = os.path.join(base_dir, "audit_report.docx")
    parse_markdown_to_docx(md1, docx1)
    
    md2 = os.path.join(base_dir, "audit_report_part2.md")
    docx2 = os.path.join(base_dir, "audit_report_part2.docx")
    parse_markdown_to_docx(md2, docx2)
